#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sincronizza offerte dalle cartelle 2025 e 2026
Associa i file alle opportunità esistenti con matching migliorato dei nomi account
"""
import os
import sys
import re
from datetime import datetime
from pathlib import Path
from difflib import SequenceMatcher

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from crm_app_completo import db, app, Account, Opportunity, User, OfferDocument, extract_offer_amount_from_file

def extract_address_from_pdf(pdf_path):
    """Estrae l'indirizzo dal PDF dell'offerta"""
    try:
        import pdfplumber
        import re
        
        with pdfplumber.open(pdf_path) as pdf:
            all_text = ""
            # Leggi le prime 3 pagine (di solito l'indirizzo è all'inizio)
            for page in pdf.pages[:3]:
                all_text += page.extract_text() or ""
            
            # Pattern per indirizzi italiani
            # Cerca: Via/Nome Via, Numero Civico, CAP Città (Provincia)
            address_patterns = [
                r'(Via|Viale|Vicolo|Piazza|Corso|Largo|Piazzale|Strada|Contrada)\s+[A-Za-z0-9\s\',\.]+,\s*\d{5}\s+[A-Za-z\s]+(?:\([A-Z]{2}\))?',
                r'[A-Za-z0-9\s\',\.]+,\s*\d{5}\s+[A-Za-z\s]+(?:\([A-Z]{2}\))?',
                r'\d{5}\s+[A-Za-z\s]+(?:\([A-Z]{2}\))?',
            ]
            
            for pattern in address_patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE | re.MULTILINE)
                if matches:
                    # Prendi il primo match che sembra un indirizzo completo
                    for match in matches:
                        match_clean = match.strip()
                        # Verifica che contenga almeno CAP e città
                        if re.search(r'\d{5}', match_clean) and len(match_clean) > 10:
                            return match_clean
            
            return None
    except Exception as e:
        return None

def extract_address_from_docx(docx_path):
    """Estrae l'indirizzo dal DOCX dell'offerta"""
    try:
        from docx import Document
        import re
        
        doc = Document(docx_path)
        all_text = "\n".join([para.text for para in doc.paragraphs[:20]])  # Prime 20 righe
        
        # Stessi pattern del PDF
        address_patterns = [
            r'(Via|Viale|Vicolo|Piazza|Corso|Largo|Piazzale|Strada|Contrada)\s+[A-Za-z0-9\s\',\.]+,\s*\d{5}\s+[A-Za-z\s]+(?:\([A-Z]{2}\))?',
            r'[A-Za-z0-9\s\',\.]+,\s*\d{5}\s+[A-Za-z\s]+(?:\([A-Z]{2}\))?',
            r'\d{5}\s+[A-Za-z\s]+(?:\([A-Z]{2}\))?',
        ]
        
        for pattern in address_patterns:
            matches = re.findall(pattern, all_text, re.IGNORECASE | re.MULTILINE)
            if matches:
                for match in matches:
                    match_clean = match.strip()
                    if re.search(r'\d{5}', match_clean) and len(match_clean) > 10:
                        return match_clean
        
        return None
    except Exception as e:
        return None

# Cartelle da sincronizzare
import platform
if platform.system() == 'Windows':
    OFFERS_2025 = r"K:\OFFERTE - K\OFFERTE 2025 - K"
    OFFERS_2026 = r"K:\OFFERTE - K\OFFERTE 2026 - K"
else:
    OFFERS_2025 = os.environ.get('OFFERS_2025_FOLDER', '/mnt/k/OFFERTE - K/OFFERTE 2025 - K')
    OFFERS_2026 = os.environ.get('OFFERS_2026_FOLDER', '/mnt/k/OFFERTE - K/OFFERTE 2026 - K')

def similarity(a, b):
    """Calcola similarità tra due stringhe (0-1)"""
    if not a or not b:
        return 0.0
    # Normalizza: rimuovi spazi extra, lowercase, rimuovi caratteri speciali comuni
    a_norm = re.sub(r'[^\w\s]', '', a.lower().strip())
    b_norm = re.sub(r'[^\w\s]', '', b.lower().strip())
    return SequenceMatcher(None, a_norm, b_norm).ratio()

def clean_account_name(name):
    """Pulisce il nome account rimuovendo caratteri strani e normalizzando"""
    if not name:
        return ""
    
    # Rimuovi prefissi ~$ (file temporanei Windows)
    cleaned = re.sub(r'^~\$\s*', '', name)
    
    # Rimuovi caratteri speciali non standard (mantieni solo lettere, numeri, spazi, trattini, apostrofi, parentesi, virgole)
    cleaned = re.sub(r'[^\w\s\-\.\,\'\(\)]', '', cleaned)
    
    # Rimuovi spazi multipli
    cleaned = re.sub(r'\s+', ' ', cleaned)
    
    # Rimuovi spazi all'inizio e alla fine
    cleaned = cleaned.strip()
    
    # Capitalizza correttamente (prima lettera di ogni parola, ma mantieni maiuscole esistenti per sigle)
    words = cleaned.split()
    cleaned_words = []
    for word in words:
        # Se è tutto maiuscolo o contiene numeri, mantieni com'è (probabilmente una sigla)
        if word.isupper() or any(c.isdigit() for c in word):
            cleaned_words.append(word)
        else:
            # Capitalizza solo la prima lettera
            cleaned_words.append(word.capitalize())
    
    cleaned = ' '.join(cleaned_words)
    
    return cleaned

def normalize_account_name(name):
    """Normalizza il nome account per il matching"""
    if not name:
        return ""
    # Rimuovi caratteri speciali comuni, spazi multipli, lowercase
    normalized = re.sub(r'[^\w\s]', ' ', name.lower())
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    # Rimuovi parole comuni che non aiutano nel matching
    stop_words = ['srl', 'spa', 'sas', 's.n.c.', 's.n.c', 'snc', 'ltd', 'limited', 'inc', 'corp', 'corporation']
    words = normalized.split()
    words = [w for w in words if w not in stop_words]
    return ' '.join(words)

def find_best_account_match(client_name, accounts):
    """Trova il miglior match account usando fuzzy matching"""
    if not client_name:
        return None
    
    client_normalized = normalize_account_name(client_name)
    best_match = None
    best_score = 0.0
    
    for account in accounts:
        account_normalized = normalize_account_name(account.name)
        
        # Calcola similarità
        score = similarity(client_normalized, account_normalized)
        
        # Bonus se il nome contiene parole chiave del cliente
        client_words = set(client_normalized.split())
        account_words = set(account_normalized.split())
        if client_words and account_words:
            common_words = client_words.intersection(account_words)
            if common_words:
                # Aggiungi bonus per parole comuni (max 0.2)
                word_bonus = min(0.2, len(common_words) * 0.05)
                score += word_bonus
        
        if score > best_score:
            best_score = score
            best_match = account
    
    # Soglia minima di similarità: 0.6
    return best_match if best_score >= 0.6 else None

def extract_offer_info(filename):
    """Estrae informazioni dal nome del file offerta"""
    name = os.path.splitext(filename)[0]
    
    # Pattern 1: K####-YY - CLIENTE - PRODOTTO - DD-MM-YYYY
    k_pattern = re.compile(r'K(\d{4})-(\d{2})\s*-\s*(.+?)\s*-\s*(.+?)\s*-\s*(\d{2}-\d{2}-\d{4})', re.IGNORECASE)
    match = k_pattern.search(name)
    if match:
        offer_number = f"K{match.group(1)}-{match.group(2)}"
        client_name = match.group(3).strip()
        product = match.group(4).strip()
        date_str = match.group(5)
        return {
            'offer_number': offer_number,
            'client_name': client_name,
            'product': product,
            'date': date_str,
            'year': match.group(2),
            'filename': filename
        }
    
    # Pattern 2: K####-YY - CLIENTE (senza prodotto)
    k_pattern2 = re.compile(r'K(\d{4})-(\d{2})\s*-\s*(.+?)(?:\s*-\s*\d{2}-\d{2}-\d{4})?$', re.IGNORECASE)
    match = k_pattern2.search(name)
    if match:
        offer_number = f"K{match.group(1)}-{match.group(2)}"
        client_name = match.group(3).strip()
        # Rimuovi eventuale data finale
        client_name = re.sub(r'\s*-\s*\d{2}-\d{2}-\d{4}$', '', client_name).strip()
        return {
            'offer_number': offer_number,
            'client_name': client_name,
            'year': match.group(2),
            'filename': filename
        }
    
    # Pattern 3: K####-YY (solo numero)
    k_simple = re.compile(r'K(\d{4})-(\d{2})', re.IGNORECASE)
    match = k_simple.search(name)
    if match:
        offer_number = f"K{match.group(1)}-{match.group(2)}"
        # Cerca nome cliente dopo il numero
        rest = name[match.end():].strip()
        if rest.startswith('-'):
            rest = rest[1:].strip()
        parts = rest.split(' - ')
        client_name = parts[0].strip() if parts else None
        return {
            'offer_number': offer_number,
            'client_name': client_name,
            'year': match.group(2),
            'filename': filename
        }
    
    # Fallback: cerca qualsiasi pattern con numero
    numbers = re.findall(r'\d{4,}', name)
    if numbers:
        return {
            'offer_number': f"K{numbers[0]}",
            'client_name': name.replace(f"K{numbers[0]}", '').replace('-', ' ').strip(),
            'year': str(datetime.now().year)[-2:],
            'filename': filename
        }
    
    return {
        'offer_number': None,
        'client_name': name.replace('_', ' ').replace('-', ' ').title(),
        'year': str(datetime.now().year)[-2:],
        'filename': filename
    }

def sync_offers_from_folder(folder_path, year):
    """Sincronizza offerte da una cartella specifica"""
    if not os.path.exists(folder_path):
        print(f"⚠️  Cartella {year} non trovata: {folder_path}")
        return {
            'accounts_created': 0,
            'opportunities_created': 0,
            'offers_created': 0,
            'offers_updated': 0,
            'files_processed': 0
        }
    
    print(f"\n📂 Sincronizzazione cartella {year}: {folder_path}")
    
    with app.app_context():
        admin_user = User.query.filter_by(role='admin').first()
        if not admin_user:
            print("❌ Utente admin non trovato")
            return {}
        
        # Carica tutti gli account per il matching
        all_accounts = Account.query.all()
        print(f"📊 Account disponibili per matching: {len(all_accounts)}")
        
        offer_extensions = ['.docx', '.pdf', '.doc']
        all_files = []
        
        # Raccogli tutti i file offerte
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if any(file.lower().endswith(ext) for ext in offer_extensions):
                    all_files.append(os.path.join(root, file))
        
        print(f"📄 Trovati {len(all_files)} file offerte")
        
        stats = {
            'accounts_created': 0,
            'opportunities_created': 0,
            'offers_created': 0,
            'offers_updated': 0,
            'files_processed': 0
        }
        
        for file_path in all_files:
            try:
                filename = os.path.basename(file_path)
                folder_path_file = os.path.dirname(file_path)
                offer_info = extract_offer_info(filename)
                
                client_name = offer_info.get('client_name')
                offer_number = offer_info.get('offer_number')
                
                if not client_name or client_name == 'Cliente Sconosciuto':
                    continue
                
                stats['files_processed'] += 1
                
                # Trova account con fuzzy matching
                account = find_best_account_match(client_name, all_accounts)
                
                if not account:
                    # Crea nuovo account con nome pulito
                    clean_name = clean_account_name(client_name)
                    image_url = f"https://ui-avatars.com/api/?name={clean_name.replace(' ', '+')}&size=128&background=667eea&color=fff"
                    account = Account(
                        name=clean_name,
                        owner_id=admin_user.id,
                        image_url=image_url
                    )
                    db.session.add(account)
                    db.session.flush()
                    all_accounts.append(account)  # Aggiungi alla lista per matching futuri
                    stats['accounts_created'] += 1
                    print(f"  ✅ Account creato: {clean_name}")
                else:
                    # Pulisci il nome dell'account esistente se necessario
                    if account.name != clean_account_name(account.name):
                        old_name = account.name
                        account.name = clean_account_name(account.name)
                        print(f"  🔧 Nome account corretto: '{old_name}' -> '{account.name}'")
                    
                    print(f"  🔗 Account trovato: {account.name} (similarità: {similarity(normalize_account_name(client_name), normalize_account_name(account.name)):.2f})")
                
                # Trova o crea opportunità
                opp_name = f"Offerta {offer_number} - {account.name}" if offer_number else f"Offerta - {account.name}"
                
                # Cerca opportunità esistente per numero offerta PRIMA (più preciso)
                existing_opp = None
                
                if offer_number:
                    # Cerca per numero offerta nelle offerte collegate
                    existing_offer_doc = OfferDocument.query.filter_by(number=offer_number).first()
                    if existing_offer_doc and existing_offer_doc.opportunity_id:
                        existing_opp = Opportunity.query.get(existing_offer_doc.opportunity_id)
                        # Verifica che l'account sia corretto
                        if existing_opp and existing_opp.account_id != account.id:
                            print(f"  ⚠️  Offerta {offer_number} già collegata a account diverso, aggiorno account")
                            existing_opp.account_id = account.id
                
                # Se non trovata per numero offerta, cerca per account_id
                if not existing_opp:
                    existing_opp = Opportunity.query.filter(
                        Opportunity.account_id == account.id
                    ).order_by(Opportunity.created_at.desc()).first()
                
                # Se non trovata, cerca per numero offerta nel nome (fallback)
                if not existing_opp and offer_number:
                    # Cerca opportunità che hanno offerte con questo numero
                    from sqlalchemy import and_
                    existing_offers = OfferDocument.query.filter_by(number=offer_number).all()
                    if existing_offers:
                        for off in existing_offers:
                            if off.opportunity_id:
                                existing_opp = Opportunity.query.get(off.opportunity_id)
                                if existing_opp:
                                    # Se l'opportunità non è collegata all'account corretto, aggiornala
                                    if existing_opp.account_id != account.id:
                                        print(f"  🔗 Collegando opportunità esistente '{existing_opp.name}' all'account '{account.name}'")
                                        existing_opp.account_id = account.id
                                    break
                
                if not existing_opp:
                    # Crea nuova opportunità
                    import platform
                    if platform.system() == 'Windows':
                        base_dir = r"K:\OFFERTE - K\OPPORTUNITA"
                    else:
                        base_dir = os.environ.get('OPPORTUNITIES_FOLDER', '/mnt/k/OFFERTE - K/OPPORTUNITA')
                    
                    if not os.path.exists(base_dir):
                        base_dir = os.path.join(os.path.dirname(__file__), 'opportunita')
                    
                    os.makedirs(base_dir, exist_ok=True)
                    opp_name_clean = "".join(c for c in opp_name if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
                    folder_name_opp = f"OPP-{datetime.now().strftime('%Y%m%d')}-{opp_name_clean}"
                    folder_path_opp = os.path.join(base_dir, folder_name_opp)
                    
                    if not os.path.exists(folder_path_opp):
                        os.makedirs(folder_path_opp)
                    
                    existing_opp = Opportunity(
                        name=opp_name,
                        stage='Qualification',
                        amount=0.0,
                        probability=10,
                        account_id=account.id,
                        folder_path=folder_path_opp,
                        owner_id=admin_user.id,
                        description=f"Offerta sincronizzata da cartella {year}: {filename}"
                    )
                    db.session.add(existing_opp)
                    db.session.flush()
                    stats['opportunities_created'] += 1
                    print(f"  ✅ Opportunità creata: {opp_name}")
                
                # Cerca PDF e DOCX nella stessa cartella
                pdf_path = None
                docx_path = None
                
                if filename.lower().endswith('.pdf'):
                    pdf_path = file_path
                else:
                    base_name = os.path.splitext(filename)[0]
                    pdf_candidate = os.path.join(folder_path_file, base_name + '.pdf')
                    if os.path.exists(pdf_candidate):
                        pdf_path = pdf_candidate
                
                if filename.lower().endswith('.docx'):
                    docx_path = file_path
                else:
                    base_name = os.path.splitext(filename)[0]
                    docx_candidate = os.path.join(folder_path_file, base_name + '.docx')
                    if os.path.exists(docx_candidate):
                        docx_path = docx_candidate
                
                # Estrai valore e indirizzo dal PDF - cerca il numero più alto
                offer_amount = 0.0
                extracted_address = None
                
                # Prova prima il PDF
                if pdf_path and os.path.exists(pdf_path):
                    try:
                        offer_amount = extract_offer_amount_from_file(pdf_path)
                        extracted_address = extract_address_from_pdf(pdf_path)
                        if offer_amount > 0:
                            print(f"  💰 Valore estratto da PDF: €{offer_amount:,.2f}")
                    except Exception as e:
                        print(f"  ⚠️  Errore estrazione PDF: {str(e)}")
                
                # Se non trovato nel PDF, prova nel DOCX
                if offer_amount == 0 and docx_path and os.path.exists(docx_path):
                    try:
                        offer_amount = extract_offer_amount_from_file(docx_path)
                        if offer_amount > 0:
                            print(f"  💰 Valore estratto da DOCX: €{offer_amount:,.2f}")
                    except Exception as e:
                        print(f"  ⚠️  Errore estrazione DOCX: {str(e)}")
                
                if not extracted_address and docx_path and os.path.exists(docx_path):
                    try:
                        extracted_address = extract_address_from_docx(docx_path)
                    except:
                        pass
                
                # Aggiorna indirizzo account se trovato e non presente
                if extracted_address and account:
                    if not account.billing_address and not account.shipping_address:
                        account.billing_address = extracted_address
                        account.shipping_address = extracted_address
                        print(f"  📍 Indirizzo aggiunto all'account: {extracted_address[:50]}...")
                    elif account.billing_address != extracted_address and account.shipping_address != extracted_address:
                        # Se l'indirizzo è diverso, aggiorna solo se quello esistente è vuoto o molto corto
                        if len(account.billing_address or '') < 20:
                            account.billing_address = extracted_address
                            print(f"  📍 Indirizzo billing aggiornato")
                        if len(account.shipping_address or '') < 20:
                            account.shipping_address = extracted_address
                            print(f"  📍 Indirizzo shipping aggiornato")
                
                # Crea o aggiorna OfferDocument
                offer_number_str = offer_number if offer_number else filename.replace('.pdf', '').replace('.docx', '').replace('.doc', '')
                
                existing_offer = OfferDocument.query.filter_by(number=offer_number_str).first()
                
                if not existing_offer:
                    offer_doc = OfferDocument(
                        number=offer_number_str,
                        status='draft',
                        amount=offer_amount,
                        description=f"Offerta sincronizzata da {year}: {client_name}",
                        pdf_path=pdf_path if pdf_path and os.path.exists(pdf_path) else None,
                        docx_path=docx_path if docx_path and os.path.exists(docx_path) else None,
                        folder_path=folder_path_file,
                        opportunity_id=existing_opp.id,
                        owner_id=admin_user.id
                    )
                    db.session.add(offer_doc)
                    stats['offers_created'] += 1
                else:
                    # Aggiorna offerta esistente
                    if offer_amount > 0 and (not existing_offer.amount or existing_offer.amount == 0):
                        existing_offer.amount = offer_amount
                    if not existing_offer.pdf_path and pdf_path and os.path.exists(pdf_path):
                        existing_offer.pdf_path = pdf_path
                    if not existing_offer.docx_path and docx_path and os.path.exists(docx_path):
                        existing_offer.docx_path = docx_path
                    if not existing_offer.opportunity_id:
                        existing_offer.opportunity_id = existing_opp.id
                    if not existing_offer.folder_path:
                        existing_offer.folder_path = folder_path_file
                    stats['offers_updated'] += 1
                
                # Commit ogni 50 file
                if stats['files_processed'] % 50 == 0:
                    db.session.commit()
                    
            except Exception as e:
                print(f"  ❌ Errore processando {filename}: {e}")
                db.session.rollback()
                continue
        
        db.session.commit()
        return stats

def sync_all_offers():
    """Sincronizza offerte da entrambe le cartelle"""
    print("=" * 80)
    print("SINCRONIZZAZIONE OFFERTE 2025 E 2026")
    print("=" * 80)
    
    stats_2025 = sync_offers_from_folder(OFFERS_2025, "2025")
    stats_2026 = sync_offers_from_folder(OFFERS_2026, "2026")
    
    print("\n" + "=" * 80)
    print("RIEPILOGO SINCRONIZZAZIONE")
    print("=" * 80)
    print(f"\n📊 Cartella 2025:")
    print(f"   File processati: {stats_2025.get('files_processed', 0)}")
    print(f"   Account creati: {stats_2025.get('accounts_created', 0)}")
    print(f"   Opportunità create: {stats_2025.get('opportunities_created', 0)}")
    print(f"   Offerte create: {stats_2025.get('offers_created', 0)}")
    print(f"   Offerte aggiornate: {stats_2025.get('offers_updated', 0)}")
    
    print(f"\n📊 Cartella 2026:")
    print(f"   File processati: {stats_2026.get('files_processed', 0)}")
    print(f"   Account creati: {stats_2026.get('accounts_created', 0)}")
    print(f"   Opportunità create: {stats_2026.get('opportunities_created', 0)}")
    print(f"   Offerte create: {stats_2026.get('offers_created', 0)}")
    print(f"   Offerte aggiornate: {stats_2026.get('offers_updated', 0)}")
    
    print("\n✅ Sincronizzazione completata!")

if __name__ == '__main__':
    sync_all_offers()
