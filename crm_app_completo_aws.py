#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MekoCRM - Sistema CRM completo stile Salesforce
"""
import sys
import os
from pathlib import Path

# Carica variabili d'ambiente da file .env se esiste
def load_env_file():
    """Carica variabili d'ambiente da file .env"""
    env_file = Path('.env')
    if env_file.exists():
        try:
            # Prova con python-dotenv se disponibile
            try:
                from dotenv import load_dotenv
                load_dotenv()
                return
            except ImportError:
                pass
            
            # Fallback: leggi manualmente il file .env
            with open(env_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        key, value = line.split('=', 1)
                        os.environ[key.strip()] = value.strip()
        except Exception as e:
            print(f"Warning: Errore caricamento .env: {e}")

load_env_file()

if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date, timedelta
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy import func
import json
import os
import shutil
from pathlib import Path

app = Flask(__name__)
app.config['SECRET_KEY'] = 'mekocrm-2025-secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///mekocrm.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)


# ==================== MODELLI DATABASE SALESFORCE-LIKE ====================

class TimestampMixin:
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class User(TimestampMixin, db.Model):
    __tablename__ = "users"
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    first_name = db.Column(db.String(100))
    last_name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    role = db.Column(db.String(20), default="commerciale")
    
    leads = db.relationship("Lead", backref="owner", lazy=True)
    accounts = db.relationship("Account", backref="owner", lazy=True)
    contacts = db.relationship("Contact", backref="owner", lazy=True)
    opportunities = db.relationship("Opportunity", backref="owner", lazy=True)
    offers = db.relationship("OfferDocument", backref="owner", lazy=True)


class Lead(TimestampMixin, db.Model):
    __tablename__ = "leads"
    id = db.Column(db.Integer, primary_key=True)
    company_name = db.Column(db.String(200), nullable=False)
    contact_first_name = db.Column(db.String(100))
    contact_last_name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    vat_number = db.Column(db.String(20))
    address = db.Column(db.Text)
    industry = db.Column(db.String(100))
    interest = db.Column(db.String(50))  # MIR, UR, Unitree, ROEQ, Servizi
    source = db.Column(db.String(50))
    rating = db.Column(db.String(20), default="Warm")
    status = db.Column(db.String(20), default="New")  # New, Contacted, Qualified, Converted, Lost
    notes = db.Column(db.Text)
    last_contact_at = db.Column(db.DateTime)
    
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    converted_account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"))
    converted_contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"))
    converted_opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"))


class Account(TimestampMixin, db.Model):
    __tablename__ = "accounts"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), unique=True, nullable=False)
    industry = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    website = db.Column(db.String(200))
    billing_address = db.Column(db.Text)
    shipping_address = db.Column(db.Text)
    description = db.Column(db.Text)
    image_url = db.Column(db.String(500))  # URL immagine account (fittizia)
    
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    contacts = db.relationship("Contact", backref="account", lazy=True)
    opportunities = db.relationship("Opportunity", backref="account", lazy=True)


class Contact(TimestampMixin, db.Model):
    __tablename__ = "contacts"
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(100), nullable=False)
    last_name = db.Column(db.String(100), nullable=False)
    title = db.Column(db.String(100))
    email = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    department = db.Column(db.String(100))
    description = db.Column(db.Text)
    
    account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"))
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)


class Opportunity(TimestampMixin, db.Model):
    __tablename__ = "opportunities"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    stage = db.Column(db.String(40), default="Qualification")
    amount = db.Column(db.Float, default=0.0)
    probability = db.Column(db.Integer, default=10)
    close_date = db.Column(db.Date)
    opportunity_type = db.Column(db.String(50))
    description = db.Column(db.Text)
    supplier_category = db.Column(db.String(50))  # Universal Robots, Unitree, MiR, Altri Fornitori, Servizi
    folder_path = db.Column(db.String(400))  # Cartella per documenti dell'opportunità
    
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"))
    contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"))
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"))
    
    line_items = db.relationship("OpportunityLineItem", backref="opportunity", lazy=True, cascade="all, delete-orphan")
    offers = db.relationship("OfferDocument", backref="opportunity", lazy=True)
    chatter_posts = db.relationship("OpportunityChatter", backref="opportunity", lazy=True, cascade="all, delete-orphan", order_by="OpportunityChatter.created_at.desc()")
    tasks = db.relationship("OpportunityTask", backref="opportunity", lazy=True, cascade="all, delete-orphan", order_by="OpportunityTask.created_at.desc()")


class OpportunityChatter(TimestampMixin, db.Model):
    __tablename__ = "opportunity_chatter"
    id = db.Column(db.Integer, primary_key=True)
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    message = db.Column(db.Text, nullable=False)
    
    user = db.relationship("User", backref="chatter_posts")


class OpportunityTask(TimestampMixin, db.Model):
    __tablename__ = "opportunity_tasks"
    id = db.Column(db.Integer, primary_key=True)
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=True)  # NULL = task globale
    account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"), nullable=True)  # Collegamento diretto ad account
    task_type = db.Column(db.String(50), nullable=False)  # 'send_offer', 'call', 'recall', 'other'
    description = db.Column(db.Text)
    assigned_to_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)  # Chi deve farlo
    due_date = db.Column(db.Date)
    is_completed = db.Column(db.Boolean, default=False)
    completed_at = db.Column(db.DateTime)
    completed_by_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    
    account = db.relationship("Account", backref="tasks")
    
    assigned_to = db.relationship("User", foreign_keys=[assigned_to_id], backref="assigned_tasks")
    completed_by = db.relationship("User", foreign_keys=[completed_by_id])


class OpportunityLineItem(TimestampMixin, db.Model):
    __tablename__ = "opportunity_line_items"
    id = db.Column(db.Integer, primary_key=True)
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=False)
    product_id = db.Column(db.Integer, db.ForeignKey("products.id"), nullable=False)
    quantity = db.Column(db.Integer, default=1)
    unit_price = db.Column(db.Float, nullable=False)
    discount = db.Column(db.Float, default=0.0)
    total_price = db.Column(db.Float, nullable=False)
    
    product = db.relationship("Product", backref="line_items")


class Product(TimestampMixin, db.Model):
    __tablename__ = "products"
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(80), unique=True, nullable=False)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    category = db.Column(db.String(80))  # MIR, UR, Unitree, ROEQ, Servizi
    family = db.Column(db.String(80))
    is_active = db.Column(db.Boolean, default=True)
    
    prices = db.relationship("PriceBookEntry", backref="product", lazy=True, cascade="all, delete-orphan")


class PriceBookEntry(TimestampMixin, db.Model):
    __tablename__ = "pricebook_entries"
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey("products.id"), nullable=False)
    price_level = db.Column(db.String(80), nullable=False)  # bronze, silver, gold, ur_acquisto, ur_si, mir_distributor, etc.
    currency = db.Column(db.String(8), default="EUR")
    amount = db.Column(db.Float, nullable=False)
    source = db.Column(db.String(80), default="official")
    effective_date = db.Column(db.Date, default=date.today)
    is_override = db.Column(db.Boolean, default=False)
    is_active = db.Column(db.Boolean, default=True)


class OfferDocument(TimestampMixin, db.Model):
    __tablename__ = "offers"
    id = db.Column(db.Integer, primary_key=True)
    number = db.Column(db.String(32), unique=True, nullable=False)
    status = db.Column(db.String(20), default="draft")  # draft, sent, accepted, rejected, expired
    amount = db.Column(db.Float, default=0.0)
    description = db.Column(db.Text)
    docx_path = db.Column(db.String(400))
    pdf_path = db.Column(db.String(400))
    folder_path = db.Column(db.String(400))
    generated_from_nlp = db.Column(db.Boolean, default=False)
    
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"))
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"))
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"))


# ==================== MODELLI EMAIL MARKETING ====================

class EmailList(TimestampMixin, db.Model):
    """Lista email per campagne marketing"""
    __tablename__ = "email_lists"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    is_active = db.Column(db.Boolean, default=True)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    
    subscriptions = db.relationship("EmailListSubscription", backref="email_list", lazy=True, cascade="all, delete-orphan")
    campaigns = db.relationship("EmailCampaign", backref="email_list", lazy=True)


class EmailListSubscription(TimestampMixin, db.Model):
    """Iscrizioni a liste email"""
    __tablename__ = "email_list_subscriptions"
    id = db.Column(db.Integer, primary_key=True)
    list_id = db.Column(db.Integer, db.ForeignKey("email_lists.id"), nullable=False)
    contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"), nullable=True)
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"), nullable=True)
    email = db.Column(db.String(200), nullable=False)
    first_name = db.Column(db.String(100))
    last_name = db.Column(db.String(100))
    status = db.Column(db.String(20), default="subscribed")  # subscribed, unsubscribed, bounced
    subscribed_at = db.Column(db.DateTime, default=datetime.utcnow)
    unsubscribed_at = db.Column(db.DateTime)


class EmailTemplate(TimestampMixin, db.Model):
    """Template email HTML"""
    __tablename__ = "email_templates"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    subject = db.Column(db.String(500), nullable=False)
    html_content = db.Column(db.Text, nullable=False)
    text_content = db.Column(db.Text)
    category = db.Column(db.String(100))  # newsletter, promotional, transactional, etc.
    is_active = db.Column(db.Boolean, default=True)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    
    campaigns = db.relationship("EmailCampaign", backref="template", lazy=True)


class EmailCampaign(TimestampMixin, db.Model):
    """Campagna email"""
    __tablename__ = "email_campaigns"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    subject = db.Column(db.String(500), nullable=False)
    html_content = db.Column(db.Text, nullable=False)
    text_content = db.Column(db.Text)
    status = db.Column(db.String(20), default="draft")  # draft, scheduled, sending, sent, paused, cancelled
    list_id = db.Column(db.Integer, db.ForeignKey("email_lists.id"), nullable=False)
    template_id = db.Column(db.Integer, db.ForeignKey("email_templates.id"), nullable=True)
    scheduled_at = db.Column(db.DateTime)
    sent_at = db.Column(db.DateTime)
    total_recipients = db.Column(db.Integer, default=0)
    total_sent = db.Column(db.Integer, default=0)
    total_delivered = db.Column(db.Integer, default=0)
    total_opened = db.Column(db.Integer, default=0)
    total_clicked = db.Column(db.Integer, default=0)
    total_bounced = db.Column(db.Integer, default=0)
    total_unsubscribed = db.Column(db.Integer, default=0)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    
    sends = db.relationship("EmailCampaignSend", backref="campaign", lazy=True, cascade="all, delete-orphan")


class EmailCampaignSend(TimestampMixin, db.Model):
    """Singola email inviata in una campagna"""
    __tablename__ = "email_campaign_sends"
    id = db.Column(db.Integer, primary_key=True)
    campaign_id = db.Column(db.Integer, db.ForeignKey("email_campaigns.id"), nullable=False)
    email = db.Column(db.String(200), nullable=False)
    contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"), nullable=True)
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"), nullable=True)


class Ticket(TimestampMixin, db.Model):
    """Ticket di supporto/richiesta clienti"""
    __tablename__ = "tickets"
    id = db.Column(db.Integer, primary_key=True)
    ticket_number = db.Column(db.String(20), unique=True, nullable=False)  # TKT-2025-001
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(20), default="Aperto")  # Aperto, In Lavorazione, In Attesa Cliente, Chiuso
    priority = db.Column(db.String(20), default="Media")  # Bassa, Media, Alta, Urgente
    category = db.Column(db.String(50))  # Supporto, Richiesta, Bug, Altro
    product_category = db.Column(db.String(100))  # MiR, Universal Robots, Unitree, ecc.
    
    customer_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)  # Cliente che ha creato il ticket
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)  # Operatore che ha preso ownership
    
    customer = db.relationship("User", foreign_keys=[customer_id], backref="tickets_as_customer")
    owner = db.relationship("User", foreign_keys=[owner_id], backref="tickets_as_owner")
    comments = db.relationship("TicketComment", backref="ticket", lazy=True, cascade="all, delete-orphan", order_by="TicketComment.created_at.asc()")


class TicketComment(TimestampMixin, db.Model):
    """Commenti su ticket"""
    __tablename__ = "ticket_comments"
    id = db.Column(db.Integer, primary_key=True)
    ticket_id = db.Column(db.Integer, db.ForeignKey("tickets.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    message = db.Column(db.Text, nullable=False)
    is_internal = db.Column(db.Boolean, default=False)  # True = solo operatori vedono
    
    user = db.relationship("User", backref="ticket_comments")


class TicketAttachment(TimestampMixin, db.Model):
    """Allegati ai ticket (immagini, video)"""
    __tablename__ = "ticket_attachments"
    id = db.Column(db.Integer, primary_key=True)
    ticket_id = db.Column(db.Integer, db.ForeignKey("tickets.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_type = db.Column(db.String(50))  # image, video
    file_size = db.Column(db.Integer)  # in bytes
    mime_type = db.Column(db.String(100))
    is_compressed = db.Column(db.Boolean, default=False)
    
    user = db.relationship("User", backref="ticket_attachments")
    ticket = db.relationship("Ticket", backref="attachments")


class Notification(TimestampMixin, db.Model):
    """Notifiche per gli utenti"""
    __tablename__ = "notifications"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    message = db.Column(db.Text, nullable=False)
    type = db.Column(db.String(50), default="info")  # info, success, warning, error
    is_read = db.Column(db.Boolean, default=False)
    related_ticket_id = db.Column(db.Integer, db.ForeignKey("tickets.id"), nullable=True)
    related_opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=True)
    
    user = db.relationship("User", backref="notifications")
    ticket = db.relationship("Ticket", backref="notifications")
    opportunity = db.relationship("Opportunity", backref="notifications")


def extract_offer_amount(offer):
    """Estrae il valore da un'offerta (dal campo amount o dal PDF se disponibile)"""
    if offer.amount and offer.amount > 0:
        return offer.amount
    
    # Se non c'è amount, prova a estrarre dal PDF
    if offer.pdf_path and os.path.exists(offer.pdf_path):
        return extract_offer_amount_from_file(offer.pdf_path)
    
    return 0.0


def extract_offer_amount_from_file(file_path):
    """Estrae il valore da un file PDF"""
    if not file_path or not file_path.endswith('.pdf') or not os.path.exists(file_path):
        return 0.0
    
    try:
        import pdfplumber
        import re
        with pdfplumber.open(file_path) as pdf:
            all_text = ""
            for page in pdf.pages:
                all_text += page.extract_text() or ""
            
            # Cerca "TOTALE" o "TOTALE FORNITURA"
            patterns = [
                r'TOTALE\s*(?:FORNITURA)?\s*[:\s]*€?\s*([\d.,]+)',
                r'TOTALE\s*[:\s]*€?\s*([\d.,]+)',
            ]
            
            max_amount = 0.0
            for pattern in patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.replace('.', '').replace(',', '.')
                    try:
                        amount = float(amount_str)
                        if amount > max_amount:
                            max_amount = amount
                    except ValueError:
                        continue
            
            if max_amount > 0:
                return max_amount
            
            # Cerca il numero più alto
            numbers = re.findall(r'€\s*([\d.,]+)', all_text)
            amounts = []
            for num in numbers:
                try:
                    amount_str = num.replace('.', '').replace(',', '.')
                    amount = float(amount_str)
                    if amount > 100:
                        amounts.append(amount)
                except ValueError:
                    continue
            
            if amounts:
                return max(amounts)
    except Exception:
        pass
    
    return 0.0


# ==================== ROUTES ====================

@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('app_route'))
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        try:
            data = request.json
            if not data:
                return jsonify({'success': False, 'error': 'Dati mancanti'}), 400
            
            # Prova prima con username, poi con email
            user = User.query.filter_by(username=data.get('username')).first()
            if not user:
                user = User.query.filter_by(email=data.get('username')).first()
            
            if user and check_password_hash(user.password, data.get('password', '')):
                session['user_id'] = user.id
                session['username'] = user.username
                session['user_role'] = user.role
                return jsonify({'success': True, 'role': user.role})
            return jsonify({'success': False, 'error': 'Credenziali errate'}), 401
        except Exception as e:
            return jsonify({'success': False, 'error': f'Errore: {str(e)}'}), 500
    return render_template('login.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    """Registrazione clienti"""
    if request.method == 'POST':
        try:
            data = request.json
            if not data:
                return jsonify({'success': False, 'error': 'Dati mancanti'}), 400
            
            # Verifica campi obbligatori
            if not data.get('email') or not data.get('password') or not data.get('first_name'):
                return jsonify({'success': False, 'error': 'Email, password e nome sono obbligatori'}), 400
            
            # Verifica se email già esistente
            existing = User.query.filter_by(email=data.get('email')).first()
            if existing:
                return jsonify({'success': False, 'error': 'Email già registrata'}), 400
            
            # Verifica se username già esistente
            username = data.get('username') or data.get('email').split('@')[0]
            existing_username = User.query.filter_by(username=username).first()
            if existing_username:
                # Aggiungi numero se username già esiste
                counter = 1
                while User.query.filter_by(username=f"{username}{counter}").first():
                    counter += 1
                username = f"{username}{counter}"
            
            # Crea nuovo utente cliente
            user = User(
                username=username,
                password=generate_password_hash(data.get('password')),
                email=data.get('email'),
                first_name=data.get('first_name'),
                last_name=data.get('last_name', ''),
                role='cliente'  # Ruolo cliente per distinguerli dagli operatori
            )
            
            db.session.add(user)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Registrazione completata! Ora puoi accedere.',
                'username': username
            })
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': f'Errore: {str(e)}'}), 500
    
    # GET: mostra pagina registrazione
    return render_template('register.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/dashboard')
def dashboard():
    """Redirect a /app per compatibilità"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('app_route'))


@app.route('/app')
def app_route():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    user = User.query.get(session['user_id'])
    if user and user.role == 'cliente':
        # Cliente vede solo la vista ticket
        return render_template('app_cliente.html')
    else:
        # Operatori e admin vedono il CRM completo
        return render_template('app_mekocrm_completo.html')


@app.route('/api/user/info')
def user_info():
    """Ottieni informazioni utente corrente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user:
        return jsonify({'error': 'Utente non trovato'}), 404
    
    return jsonify({
        'user': {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'role': user.role
        }
    })


@app.route('/favicon.ico')
def favicon():
    """Route per favicon per evitare errori 404"""
    return '', 204  # No Content


# ==================== API ROUTES ====================

def extract_offer_amount(offer):
    """Estrae il valore dell'offerta dal PDF o DOCX se non presente nel database"""
    if offer.amount and offer.amount > 0:
        return offer.amount
    
    # Prova a estrarre dal PDF
    if offer.pdf_path and os.path.exists(offer.pdf_path):
        try:
            import pdfplumber
            import re
            with pdfplumber.open(offer.pdf_path) as pdf:
                all_text = ""
                for page in pdf.pages:
                    all_text += page.extract_text() or ""
                
                # Cerca "TOTALE" o "TOTALE FORNITURA" seguito da un numero
                patterns = [
                    r'TOTALE\s*(?:FORNITURA)?\s*[:\s]*€?\s*([\d.,]+)',
                    r'TOTALE\s*[:\s]*€?\s*([\d.,]+)',
                    r'€\s*([\d.,]+)\s*(?:TOTALE|TOTALE FORNITURA)',
                ]
                
                max_amount = 0.0
                for pattern in patterns:
                    matches = re.findall(pattern, all_text, re.IGNORECASE)
                    for match in matches:
                        # Converti il numero (gestisce sia . che , come separatore decimale)
                        amount_str = match.replace('.', '').replace(',', '.')
                        try:
                            amount = float(amount_str)
                            if amount > max_amount:
                                max_amount = amount
                        except ValueError:
                            continue
                
                if max_amount > 0:
                    return max_amount
                
                # Se non trova "TOTALE", cerca il numero più alto nel documento
                numbers = re.findall(r'€\s*([\d.,]+)', all_text)
                amounts = []
                for num in numbers:
                    try:
                        amount_str = num.replace('.', '').replace(',', '.')
                        amount = float(amount_str)
                        if amount > 100:  # Filtra numeri troppo piccoli (probabilmente quantità)
                            amounts.append(amount)
                    except ValueError:
                        continue
                
                if amounts:
                    return max(amounts)
        except Exception as e:
            pass  # Ignora errori di estrazione
    
    # Prova a estrarre dal DOCX
    if offer.docx_path and os.path.exists(offer.docx_path):
        try:
            from docx import Document
            import re
            doc = Document(offer.docx_path)
            all_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Cerca "TOTALE" o "TOTALE FORNITURA"
            patterns = [
                r'TOTALE\s*(?:FORNITURA)?\s*[:\s]*€?\s*([\d.,]+)',
                r'TOTALE\s*[:\s]*€?\s*([\d.,]+)',
            ]
            
            max_amount = 0.0
            for pattern in patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.replace('.', '').replace(',', '.')
                    try:
                        amount = float(amount_str)
                        if amount > max_amount:
                            max_amount = amount
                    except ValueError:
                        continue
            
            if max_amount > 0:
                return max_amount
            
            # Cerca il numero più alto
            numbers = re.findall(r'€\s*([\d.,]+)', all_text)
            amounts = []
            for num in numbers:
                try:
                    amount_str = num.replace('.', '').replace(',', '.')
                    amount = float(amount_str)
                    if amount > 100:
                        amounts.append(amount)
                except ValueError:
                    continue
            
            if amounts:
                return max(amounts)
        except Exception as e:
            pass  # Ignora errori di estrazione
    
    return 0.0


@app.route('/api/user/me', methods=['GET'])
def get_current_user():
    """Ottieni informazioni sull'utente corrente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user:
        return jsonify({'error': 'Utente non trovato'}), 404
    
    return jsonify({
        'user': {
            'id': user.id,
            'username': user.username,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email,
            'role': user.role
        }
    })


@app.route('/api/dashboard/summary')
def dashboard_summary():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    # Metriche base per l'utente corrente
    # Leads e Accounts condivisi - conta tutti
    leads_count = Lead.query.count()
    accounts_count = Account.query.count()
    opps = Opportunity.query.filter_by(owner_id=user_id).all()
    opportunities_count = len(opps)
    
    # Calcola pipeline includendo valori offerte
    pipeline = 0.0
    for opp in opps:
        if opp.stage not in ['Closed Won', 'Closed Lost']:
            if opp.amount and opp.amount > 0:
                pipeline += opp.amount
            else:
                opp_value = sum(
                    (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                    for o in opp.offers
                )
                if opp_value > 0:
                    pipeline += opp_value
    
    # Calcola chiusi includendo valori offerte
    closed_won = 0.0
    for opp in opps:
        if opp.stage == 'Closed Won':
            if opp.amount and opp.amount > 0:
                closed_won += opp.amount
            else:
                opp_value = sum(
                    (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                    for o in opp.offers
                )
                if opp_value > 0:
                    closed_won += opp_value
    
    closed_won_count = len([o for o in opps if o.stage == 'Closed Won'])
    
    # Metriche globali (tutte le opportunità - leads UR condivisi)
    all_opps = Opportunity.query.all()
    all_opportunities_count = len(all_opps)
    
    # Calcola pipeline globale includendo valori offerte
    all_pipeline = 0.0
    for opp in all_opps:
        if opp.stage not in ['Closed Won', 'Closed Lost']:
            if opp.amount and opp.amount > 0:
                all_pipeline += opp.amount
            else:
                opp_value = sum(
                    (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                    for o in opp.offers
                )
                if opp_value > 0:
                    all_pipeline += opp_value
    
    # Calcola chiusi globali includendo valori offerte
    all_closed_won = 0.0
    for opp in all_opps:
        if opp.stage == 'Closed Won':
            if opp.amount and opp.amount > 0:
                all_closed_won += opp.amount
            else:
                opp_value = sum(
                    (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                    for o in opp.offers
                )
                if opp_value > 0:
                    all_closed_won += opp_value
    
    all_closed_count = len([o for o in all_opps if o.stage == 'Closed Won'])
    
    # Leaderboard - metriche per tutti gli utenti commerciali
    commercial_users = User.query.filter_by(role='commerciale').all()
    leaderboard = []
    for user in commercial_users:
        user_opps = Opportunity.query.filter_by(owner_id=user.id).all()
        
        # Calcola pipeline includendo valori offerte
        user_pipeline = 0.0
        for opp in user_opps:
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                # Usa amount opportunità o somma delle offerte
                if opp.amount and opp.amount > 0:
                    user_pipeline += opp.amount
                else:
                    # Calcola dalla somma delle offerte
                    opp_value = sum(
                        (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                        for o in opp.offers
                    )
                    if opp_value > 0:
                        user_pipeline += opp_value
        
        # Calcola chiusi includendo valori offerte
        user_closed_won = 0.0
        for opp in user_opps:
            if opp.stage == 'Closed Won':
                if opp.amount and opp.amount > 0:
                    user_closed_won += opp.amount
                else:
                    opp_value = sum(
                        (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                        for o in opp.offers
                    )
                    if opp_value > 0:
                        user_closed_won += opp_value
        
        user_closed_count = len([o for o in user_opps if o.stage == 'Closed Won'])
        user_opps_count = len(user_opps)
        
        leaderboard.append({
            'user_id': user.id,
            'username': user.username,
            'name': f"{user.first_name} {user.last_name}",
            'opportunities': user_opps_count,
            'pipeline': user_pipeline,
            'closed_won': user_closed_won,
            'closed_count': user_closed_count
        })
    
    # Ordina per pipeline (decrescente)
    leaderboard.sort(key=lambda x: x['pipeline'], reverse=True)
    
    # Statistiche per mese/giorno (ultimi 30 giorni)
    from datetime import timedelta
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    
    # Opportunità chiuse per giorno (ultimi 30 giorni)
    daily_closed = {}
    daily_pipeline = {}
    daily_offers = {}  # Valori offerte per giorno
    monthly_stats = {}
    monthly_offers = {}  # Valori offerte per mese
    
    for opp in all_opps:
        if opp.created_at and opp.created_at >= thirty_days_ago:
            day_key = opp.created_at.date().isoformat()
            month_key = opp.created_at.strftime('%Y-%m')
            
            # Pipeline giornaliera
            if day_key not in daily_pipeline:
                daily_pipeline[day_key] = {'count': 0, 'amount': 0}
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                daily_pipeline[day_key]['count'] += 1
                daily_pipeline[day_key]['amount'] += opp.amount
            
            # Chiusi giornalieri
            if opp.stage == 'Closed Won':
                closed_day = opp.close_date.isoformat() if opp.close_date else day_key
                if closed_day not in daily_closed:
                    daily_closed[closed_day] = {'count': 0, 'amount': 0}
                daily_closed[closed_day]['count'] += 1
                daily_closed[closed_day]['amount'] += opp.amount
            
            # Statistiche mensili
            if month_key not in monthly_stats:
                monthly_stats[month_key] = {'pipeline': 0, 'closed': 0, 'count': 0}
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                monthly_stats[month_key]['pipeline'] += opp.amount
            elif opp.stage == 'Closed Won':
                monthly_stats[month_key]['closed'] += opp.amount
                monthly_stats[month_key]['count'] += 1
        
        # Calcola valori offerte per questa opportunità
        opp_offers_value = 0.0
        for offer in opp.offers:
            offer_amount = offer.amount if offer.amount and offer.amount > 0 else extract_offer_amount(offer)
            opp_offers_value += offer_amount
        
        if opp_offers_value > 0:
            # Aggiungi ai totali giornalieri
            if opp.created_at and opp.created_at >= thirty_days_ago:
                day_key = opp.created_at.date().isoformat()
                if day_key not in daily_offers:
                    daily_offers[day_key] = {'count': 0, 'amount': 0.0}
                daily_offers[day_key]['count'] += len([o for o in opp.offers if (o.amount and o.amount > 0) or extract_offer_amount(o) > 0])
                daily_offers[day_key]['amount'] += opp_offers_value
                
                # Aggiungi ai totali mensili
                month_key = opp.created_at.strftime('%Y-%m')
                if month_key not in monthly_offers:
                    monthly_offers[month_key] = {'count': 0, 'amount': 0.0}
                monthly_offers[month_key]['count'] += len([o for o in opp.offers if (o.amount and o.amount > 0) or extract_offer_amount(o) > 0])
                monthly_offers[month_key]['amount'] += opp_offers_value
    
    # Ultima sincronizzazione
    last_sync = PriceBookEntry.query.order_by(PriceBookEntry.created_at.desc()).first()
    last_sync_date = last_sync.created_at if last_sync else None
    
    # Calcola valore totale offerte per tutte le opportunità
    total_offers_value = 0.0
    for opp in all_opps:
        for offer in opp.offers:
            offer_amount = offer.amount if offer.amount and offer.amount > 0 else extract_offer_amount(offer)
            total_offers_value += offer_amount
    
    return jsonify({
        'metrics': {
            'leads': leads_count,
            'accounts': accounts_count,
            'opportunities': opportunities_count,
            'pipeline': pipeline,
            'closed_won': closed_won,
            'closed_count': closed_won_count
        },
        'global_metrics': {
            'opportunities': all_opportunities_count,
            'pipeline': all_pipeline,
            'closed_won': all_closed_won,
            'closed_count': all_closed_count,
            'total_offers_value': total_offers_value
        },
        'leaderboard': leaderboard,
        'daily_pipeline': daily_pipeline,
        'daily_closed': daily_closed,
        'daily_offers': daily_offers,
        'monthly_stats': monthly_stats,
        'monthly_offers': monthly_offers,
        'last_sync': last_sync_date.isoformat() if last_sync_date else None
    })


@app.route('/api/leads', methods=['GET', 'POST'])
def leads():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Leads condivisi per tutti (come opportunità UR)
        leads_list = Lead.query.order_by(Lead.created_at.desc()).all()
        return jsonify({
            'leads': [{
                'id': l.id,
                'company_name': l.company_name,
                'contact_first_name': l.contact_first_name,
                'contact_last_name': l.contact_last_name,
                'email': l.email,
                'phone': l.phone,
                'interest': l.interest,
                'source': l.source,
                'status': l.status,
                'rating': l.rating,
                'owner_id': l.owner_id,
                'owner_name': f"{l.owner.first_name} {l.owner.last_name}" if l.owner else 'Unknown',
                'is_my_lead': l.owner_id == user_id,
                'created_at': l.created_at.isoformat() if l.created_at else None,
                'last_contact_at': l.last_contact_at.isoformat() if l.last_contact_at else None
            } for l in leads_list]
        })
    
    elif request.method == 'POST':
        data = request.json
        lead = Lead(
            company_name=data['company_name'],
            contact_first_name=data.get('contact_first_name'),
            contact_last_name=data.get('contact_last_name'),
            email=data.get('email'),
            phone=data.get('phone'),
            vat_number=data.get('vat_number'),
            address=data.get('address'),
            industry=data.get('industry'),
            interest=data.get('interest'),
            source=data.get('source'),
            rating=data.get('rating', 'Warm'),
            notes=data.get('notes'),
            owner_id=user_id
        )
        db.session.add(lead)
        db.session.commit()
        return jsonify({'success': True, 'lead_id': lead.id})


@app.route('/api/leads/<int:lead_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_lead(lead_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead = Lead.query.get_or_404(lead_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and lead.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        db.session.delete(lead)
        db.session.commit()
        return jsonify({'success': True})
    
    if lead.owner_id != session['user_id']:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        return jsonify({
            'lead': {
                'id': lead.id,
                'company_name': lead.company_name,
                'contact_first_name': lead.contact_first_name,
                'contact_last_name': lead.contact_last_name,
                'email': lead.email,
                'phone': lead.phone,
                'vat_number': lead.vat_number,
                'address': lead.address,
                'industry': lead.industry,
                'interest': lead.interest,
                'source': lead.source,
                'rating': lead.rating,
                'status': lead.status,
                'notes': lead.notes,
                'created_at': lead.created_at.isoformat() if lead.created_at else None
            }
        })
    
    # PUT method
    data = request.json
    if 'company_name' in data:
        lead.company_name = data['company_name']
    if 'contact_first_name' in data:
        lead.contact_first_name = data['contact_first_name']
    if 'contact_last_name' in data:
        lead.contact_last_name = data['contact_last_name']
    if 'email' in data:
        lead.email = data['email']
    if 'phone' in data:
        lead.phone = data['phone']
    if 'vat_number' in data:
        lead.vat_number = data['vat_number']
    if 'address' in data:
        lead.address = data['address']
    if 'industry' in data:
        lead.industry = data['industry']
    if 'interest' in data:
        lead.interest = data['interest']
    if 'source' in data:
        lead.source = data['source']
    if 'rating' in data:
        lead.rating = data['rating']
    if 'status' in data:
        lead.status = data['status']
    if 'notes' in data:
        lead.notes = data['notes']
    if 'created_at' in data:
        if data['created_at']:
            lead.created_at = datetime.fromisoformat(data['created_at'].replace('Z', '+00:00'))
        else:
            lead.created_at = datetime.utcnow()
    if 'last_contact_at' in data:
        if data['last_contact_at']:
            lead.last_contact_at = datetime.fromisoformat(data['last_contact_at'].replace('Z', '+00:00'))
        else:
            lead.last_contact_at = None
    
    db.session.commit()
    return jsonify({'success': True, 'lead_id': lead.id})


@app.route('/api/leads/<int:lead_id>/convert', methods=['POST'])
def convert_lead(lead_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead = Lead.query.get_or_404(lead_id)
    # Tutti possono convertire leads (leads condivisi)
    
    # Crea Account
    account = Account(
        name=lead.company_name,
        industry=lead.industry,
        phone=lead.phone,
        billing_address=lead.address,
        owner_id=lead.owner_id
    )
    db.session.add(account)
    db.session.flush()
    
    # Crea Contact
    contact = Contact(
        first_name=lead.contact_first_name or 'Contact',
        last_name=lead.contact_last_name or 'Unknown',
        email=lead.email,
        phone=lead.phone,
        account_id=account.id,
        owner_id=lead.owner_id
    )
    db.session.add(contact)
    db.session.flush()
    
    # Crea Opportunity
    opportunity = Opportunity(
        name=f"Opportunity {lead.company_name}",
        stage="Qualification",
        account_id=account.id,
        contact_id=contact.id,
        lead_id=lead.id,
        owner_id=lead.owner_id
    )
    db.session.add(opportunity)
    db.session.flush()
    
    # Aggiorna Lead
    lead.status = 'Converted'
    lead.converted_account_id = account.id
    lead.converted_contact_id = contact.id
    lead.converted_opportunity_id = opportunity.id
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/accounts', methods=['GET', 'POST'])
def accounts():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Accounts condivisi per tutti (come opportunità UR)
        accounts_list = Account.query.order_by(Account.created_at.desc()).all()
        return jsonify({
            'accounts': [{
                'id': a.id,
                'name': a.name,
                'industry': a.industry,
                'phone': a.phone,
                'website': a.website,
                'billing_address': a.billing_address,
                'shipping_address': a.shipping_address,
                'description': a.description,
                'image_url': a.image_url,
                'owner_id': a.owner_id,
                'owner_name': f"{a.owner.first_name} {a.owner.last_name}" if a.owner else 'Unknown',
                'is_my_account': a.owner_id == user_id,
                'created_at': a.created_at.isoformat() if a.created_at else None,
                'opportunities': [{
                    'id': o.id, 
                    'name': o.name, 
                    'stage': o.stage, 
                    'amount': o.amount if o.amount and o.amount > 0 else sum(
                        (off.amount if off.amount and off.amount > 0 else extract_offer_amount(off))
                        for off in o.offers
                    ),
                    'created_at': o.created_at.isoformat() if o.created_at else None
                } for o in a.opportunities]
            } for a in accounts_list]
        })
    
    elif request.method == 'POST':
        data = request.json
        account = Account(
            name=data['name'],
            industry=data.get('industry'),
            phone=data.get('phone'),
            website=data.get('website'),
            billing_address=data.get('billing_address'),
            shipping_address=data.get('shipping_address'),
            description=data.get('description'),
            owner_id=user_id
        )
        db.session.add(account)
        db.session.commit()
        return jsonify({'success': True, 'account_id': account.id})


@app.route('/api/accounts/<int:account_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_account(account_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    account = Account.query.get_or_404(account_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and account.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        # Verifica se ci sono opportunità collegate
        opp_count = Opportunity.query.filter_by(account_id=account_id).count()
        if opp_count > 0:
            return jsonify({
                'error': f'Impossibile eliminare: ci sono {opp_count} opportunità collegate',
                'opportunities_count': opp_count
            }), 400
        db.session.delete(account)
        db.session.commit()
        return jsonify({'success': True})
    
    # Rimuovo controllo owner - tutti possono vedere tutti gli accounts
    
    if request.method == 'GET':
        return jsonify({
            'account': {
                'id': account.id,
                'name': account.name,
                'industry': account.industry,
                'phone': account.phone,
                'website': account.website,
                'billing_address': account.billing_address,
                'shipping_address': account.shipping_address,
                'description': account.description
            }
        })
    
    # PUT method
    data = request.json
    if 'name' in data:
        account.name = data['name']
    if 'industry' in data:
        account.industry = data['industry']
    if 'phone' in data:
        account.phone = data['phone']
    if 'website' in data:
        account.website = data['website']
    if 'billing_address' in data:
        account.billing_address = data['billing_address']
    if 'shipping_address' in data:
        account.shipping_address = data['shipping_address']
    if 'description' in data:
        account.description = data['description']
    if 'created_at' in data:
        if data['created_at']:
            account.created_at = datetime.fromisoformat(data['created_at'].replace('Z', '+00:00'))
        else:
            account.created_at = datetime.utcnow()
    
    db.session.commit()
    return jsonify({'success': True, 'account_id': account.id})


@app.route('/api/opportunities', methods=['GET', 'POST'])
def opportunities():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Mostra tutte le opportunità (leads UR condivisi per tutti)
        # Filtri opzionali
        category_filter = request.args.get('category')
        account_id_filter = request.args.get('account_id')
        limit = request.args.get('limit', type=int)
        sort = request.args.get('sort', 'created_at')  # created_at, amount, stage
        
        query = Opportunity.query
        
        if category_filter:
            query = query.filter_by(supplier_category=category_filter)
        if account_id_filter:
            query = query.filter_by(account_id=account_id_filter)
        
        # Ordinamento
        if sort == 'amount':
            query = query.order_by(Opportunity.amount.desc())
        elif sort == 'stage':
            query = query.order_by(Opportunity.stage)
        else:  # default: created_at
            query = query.order_by(Opportunity.created_at.desc())
        
        # Limite
        if limit:
            opps = query.limit(limit).all()
        else:
            opps = query.all()
        
        return jsonify({
            'opportunities': [{
                'id': o.id,
                'name': o.name,
                'stage': o.stage,
                'amount': o.amount if o.amount and o.amount > 0 else sum(
                    (off.amount if off.amount and off.amount > 0 else extract_offer_amount(off))
                    for off in o.offers
                ),
                'created_at': o.created_at.isoformat() if o.created_at else None,
                'probability': o.probability,
                'close_date': o.close_date.isoformat() if o.close_date else None,
                'opportunity_type': o.opportunity_type,
                'description': o.description,
                'supplier_category': o.supplier_category,
                'account_id': o.account_id,
                'account_name': o.account.name if o.account else None,
                'account_image': o.account.image_url if o.account and o.account.image_url else None,
                'contact_id': o.contact_id,
                'owner_id': o.owner_id,
                'owner_name': f"{o.owner.first_name} {o.owner.last_name}" if o.owner else 'Unknown',
                'is_my_opportunity': o.owner_id == user_id,
                'folder_path': o.folder_path,
                'offers': [{
                    'id': off.id,
                    'number': off.number,
                    'amount': off.amount if off.amount and off.amount > 0 else extract_offer_amount(off),
                    'pdf_path': off.pdf_path,
                    'docx_path': off.docx_path,
                    'status': off.status
                } for off in o.offers]
            } for o in opps]
        })
    
    elif request.method == 'POST':
        data = request.json
        close_date = None
        if data.get('close_date'):
            close_date = datetime.fromisoformat(data['close_date'].replace('Z', '+00:00')).date()
        
        # Crea cartella per l'opportunità
        import platform
        if platform.system() == 'Windows':
            base_dir = r"K:\OFFERTE - K\OPPORTUNITA"
        else:
            base_dir = os.environ.get('OPPORTUNITIES_FOLDER', '/mnt/k/OFFERTE - K/OPPORTUNITA')
        
        if not os.path.exists(base_dir):
            base_dir = os.path.join(os.path.dirname(__file__), 'opportunita')
        
        os.makedirs(base_dir, exist_ok=True)
        
        # Nome cartella: OPP-{ID}-{NOME_OPPORTUNITA}
        opp_name_clean = "".join(c for c in data['name'] if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
        folder_name = f"OPP-{datetime.now().strftime('%Y%m%d')}-{opp_name_clean}"
        folder_path = os.path.join(base_dir, folder_name)
        
        # Crea la cartella
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        opp = Opportunity(
            name=data['name'],
            stage=data.get('stage', 'Qualification'),
            amount=float(data.get('amount', 0)),
            probability=int(data.get('probability', 10)),
            close_date=close_date,
            opportunity_type=data.get('opportunity_type'),
            description=data.get('description'),
            account_id=data.get('account_id'),
            contact_id=data.get('contact_id'),
            folder_path=folder_path,
            owner_id=user_id
        )
        db.session.add(opp)
        db.session.commit()
        return jsonify({'success': True, 'opportunity_id': opp.id, 'folder_path': folder_path})


@app.route('/api/opportunities/<int:opportunity_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_opportunity(opportunity_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = Opportunity.query.get_or_404(opportunity_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and opp.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        db.session.delete(opp)
        db.session.commit()
        return jsonify({'success': True})
    
    # Rimuovo il controllo owner - tutti possono vedere tutte le opportunità
    
    if request.method == 'GET':
        return jsonify({
            'opportunity': {
                'id': opp.id,
                'name': opp.name,
                'stage': opp.stage,
                'amount': opp.amount,
                'probability': opp.probability,
                'close_date': opp.close_date.isoformat() if opp.close_date else None,
                'opportunity_type': opp.opportunity_type,
                'description': opp.description,
                'supplier_category': opp.supplier_category,
                'account_id': opp.account_id,
                'account_name': opp.account.name if opp.account else None,
                'account_image': opp.account.image_url if opp.account and opp.account.image_url else None,
                'contact_id': opp.contact_id,
                'owner_id': opp.owner_id,
                'owner_name': f"{opp.owner.first_name} {opp.owner.last_name}" if opp.owner else 'Unknown',
                'folder_path': opp.folder_path,
                'offers': [{
                    'id': off.id,
                    'number': off.number,
                    'amount': off.amount if off.amount and off.amount > 0 else extract_offer_amount(off),
                    'pdf_path': off.pdf_path,
                    'docx_path': off.docx_path,
                    'status': off.status,
                    'description': off.description
                } for off in opp.offers]
            }
        })
    
    # PUT method
    data = request.json
    if 'name' in data:
        opp.name = data['name']
    if 'stage' in data:
        opp.stage = data['stage']
    if 'amount' in data:
        opp.amount = float(data['amount'])
    if 'probability' in data:
        opp.probability = int(data['probability'])
    if 'close_date' in data:
        if data['close_date']:
            opp.close_date = datetime.fromisoformat(data['close_date'].replace('Z', '+00:00')).date()
        else:
            opp.close_date = None
    if 'created_at' in data:
        if data['created_at']:
            opp.created_at = datetime.fromisoformat(data['created_at'].replace('Z', '+00:00'))
        else:
            opp.created_at = datetime.utcnow()
    if 'opportunity_type' in data:
        opp.opportunity_type = data['opportunity_type']
    if 'description' in data:
        opp.description = data['description']
    if 'account_id' in data:
        opp.account_id = data['account_id']
    if 'contact_id' in data:
        opp.contact_id = data['contact_id']
    if 'supplier_category' in data:
        opp.supplier_category = data['supplier_category']
    if 'owner_id' in data:
        # Permetti cambio owner
        opp.owner_id = data['owner_id']
    
    db.session.commit()
    return jsonify({'success': True, 'opportunity_id': opp.id})


@app.route('/api/leads/<int:lead_id>/take-ownership', methods=['POST'])
def take_lead_ownership(lead_id):
    """Prendi ownership di un lead"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead = Lead.query.get_or_404(lead_id)
    user_id = session['user_id']
    
    lead.owner_id = user_id
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': lead.owner_id,
        'owner_name': f"{lead.owner.first_name} {lead.owner.last_name}" if lead.owner else 'Unknown'
    })


@app.route('/api/accounts/<int:account_id>/take-ownership', methods=['POST'])
def take_account_ownership(account_id):
    """Prendi ownership di un account"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    account = Account.query.get_or_404(account_id)
    user_id = session['user_id']
    
    account.owner_id = user_id
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': account.owner_id,
        'owner_name': f"{account.owner.first_name} {account.owner.last_name}" if account.owner else 'Unknown'
    })


@app.route('/api/contacts', methods=['GET', 'POST'])
def contacts():
    """Gestisci contatti"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Contatti condivisi per tutti
        contacts_list = Contact.query.order_by(Contact.created_at.desc()).all()
        return jsonify({
            'contacts': [{
                'id': c.id,
                'first_name': c.first_name,
                'last_name': c.last_name,
                'title': c.title,
                'email': c.email,
                'phone': c.phone,
                'department': c.department,
                'description': c.description,
                'account_id': c.account_id,
                'account_name': c.account.name if c.account else None,
                'owner_id': c.owner_id,
                'owner_name': f"{c.owner.first_name} {c.owner.last_name}" if c.owner else 'Unknown'
            } for c in contacts_list]
        })
    
    elif request.method == 'POST':
        data = request.json
        contact = Contact(
            first_name=data['first_name'],
            last_name=data['last_name'],
            title=data.get('title'),
            email=data.get('email'),
            phone=data.get('phone'),
            department=data.get('department'),
            description=data.get('description'),
            account_id=data.get('account_id') if data.get('account_id') else None,
            owner_id=user_id
        )
        db.session.add(contact)
        db.session.commit()
        return jsonify({'success': True, 'contact_id': contact.id})


@app.route('/api/contacts/<int:contact_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_contact(contact_id):
    """Ottieni, aggiorna o elimina un contatto"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    contact = Contact.query.get_or_404(contact_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and contact.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        db.session.delete(contact)
        db.session.commit()
        return jsonify({'success': True})
    
    if request.method == 'GET':
        return jsonify({
            'contact': {
                'id': contact.id,
                'first_name': contact.first_name,
                'last_name': contact.last_name,
                'title': contact.title,
                'email': contact.email,
                'phone': contact.phone,
                'department': contact.department,
                'description': contact.description,
                'account_id': contact.account_id,
                'account_name': contact.account.name if contact.account else None,
                'owner_id': contact.owner_id
            }
        })
    
    # PUT method
    data = request.json
    if 'first_name' in data:
        contact.first_name = data['first_name']
    if 'last_name' in data:
        contact.last_name = data['last_name']
    if 'title' in data:
        contact.title = data['title']
    if 'email' in data:
        contact.email = data['email']
    if 'phone' in data:
        contact.phone = data['phone']
    if 'department' in data:
        contact.department = data['department']
    if 'description' in data:
        contact.description = data['description']
    if 'account_id' in data:
        contact.account_id = data['account_id'] if data['account_id'] else None
    
    db.session.commit()
    return jsonify({'success': True, 'contact_id': contact.id})


@app.route('/api/accounts/search', methods=['GET'])
def search_accounts():
    """Cerca accounts per autocomplete"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    if len(query) < 2:
        return jsonify({'accounts': []})
    
    accounts = Account.query.filter(
        Account.name.ilike(f'%{query}%')
    ).limit(10).all()
    
    return jsonify({
        'accounts': [{
            'id': a.id,
            'name': a.name,
            'industry': a.industry
        } for a in accounts]
    })


@app.route('/api/opportunities/search', methods=['GET'])
def search_opportunities():
    """Cerca opportunità per autocomplete"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    account_id = request.args.get('account_id')
    
    if len(query) < 2:
        return jsonify({'opportunities': []})
    
    opps_query = Opportunity.query.filter(
        Opportunity.name.ilike(f'%{query}%')
    )
    
    if account_id:
        opps_query = opps_query.filter(Opportunity.account_id == account_id)
    
    opps = opps_query.limit(10).all()
    
    return jsonify({
        'opportunities': [{
            'id': o.id,
            'name': o.name,
            'account_id': o.account_id,
            'account_name': o.account.name if o.account else None,
            'stage': o.stage
        } for o in opps]
    })


@app.route('/api/products/search', methods=['GET'])
def search_products():
    """Cerca prodotti per autocomplete"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    if len(query) < 2:
        return jsonify({'products': []})
    
    products = Product.query.filter(
        Product.is_active == True
    ).filter(
        (Product.name.ilike(f'%{query}%')) | 
        (Product.code.ilike(f'%{query}%'))
    ).limit(20).all()
    
    return jsonify({
        'products': [{
            'id': p.id,
            'code': p.code,
            'name': p.name,
            'category': p.category
        } for p in products]
    })


@app.route('/api/opportunities/<int:opportunity_id>/chatter', methods=['GET', 'POST'])
def opportunity_chatter(opportunity_id):
    """API per chatter delle opportunità"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = Opportunity.query.get_or_404(opportunity_id)
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Ottieni tutti i post del chatter per questa opportunità
        posts = OpportunityChatter.query.filter_by(opportunity_id=opportunity_id).order_by(OpportunityChatter.created_at.desc()).all()
        return jsonify({
            'posts': [{
                'id': p.id,
                'message': p.message,
                'user_id': p.user_id,
                'user_name': f"{p.user.first_name} {p.user.last_name}" if p.user else 'Unknown',
                'created_at': p.created_at.isoformat() if p.created_at else None
            } for p in posts]
        })
    
    elif request.method == 'POST':
        # Crea un nuovo post
        data = request.json
        if not data.get('message'):
            return jsonify({'error': 'Messaggio richiesto'}), 400
        
        post = OpportunityChatter(
            opportunity_id=opportunity_id,
            user_id=user_id,
            message=data['message']
        )
        db.session.add(post)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'post': {
                'id': post.id,
                'message': post.message,
                'user_id': post.user_id,
                'user_name': f"{post.user.first_name} {post.user.last_name}" if post.user else 'Unknown',
                'created_at': post.created_at.isoformat() if post.created_at else None
            }
        })


@app.route('/api/opportunities/<int:opportunity_id>/tasks', methods=['GET', 'POST'])
def opportunity_tasks(opportunity_id):
    """API per tasks delle opportunità"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = Opportunity.query.get_or_404(opportunity_id)
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Ottieni tutti i task per questa opportunità
        tasks = OpportunityTask.query.filter_by(opportunity_id=opportunity_id).order_by(OpportunityTask.created_at.desc()).all()
        return jsonify({
            'tasks': [{
                'id': t.id,
                'task_type': t.task_type,
                'description': t.description,
                'assigned_to_id': t.assigned_to_id,
                'assigned_to_name': f"{t.assigned_to.first_name} {t.assigned_to.last_name}" if t.assigned_to else None,
                'due_date': t.due_date.isoformat() if t.due_date else None,
                'is_completed': t.is_completed,
                'completed_at': t.completed_at.isoformat() if t.completed_at else None,
                'completed_by_name': f"{t.completed_by.first_name} {t.completed_by.last_name}" if t.completed_by else None,
                'created_at': t.created_at.isoformat() if t.created_at else None
            } for t in tasks]
        })
    
    elif request.method == 'POST':
        # Crea un nuovo task
        data = request.json
        if not data.get('task_type'):
            return jsonify({'error': 'Tipo task richiesto'}), 400
        
        due_date = None
        if data.get('due_date'):
            due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
        
        task = OpportunityTask(
            opportunity_id=opportunity_id,
            task_type=data['task_type'],
            description=data.get('description', ''),
            assigned_to_id=data.get('assigned_to_id'),
            due_date=due_date
        )
        db.session.add(task)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'task': {
                'id': task.id,
                'task_type': task.task_type,
                'description': task.description,
                'assigned_to_id': task.assigned_to_id,
                'assigned_to_name': f"{task.assigned_to.first_name} {task.assigned_to.last_name}" if task.assigned_to else None,
                'due_date': task.due_date.isoformat() if task.due_date else None,
                'is_completed': task.is_completed
            }
        })


@app.route('/api/opportunities/<int:opportunity_id>/tasks/<int:task_id>', methods=['PUT', 'DELETE'])
def update_or_delete_task(opportunity_id, task_id):
    """Aggiorna o elimina un task"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    task = OpportunityTask.query.filter_by(id=task_id, opportunity_id=opportunity_id).first_or_404()
    user_id = session['user_id']
    
    if request.method == 'PUT':
        data = request.json
        if 'is_completed' in data:
            task.is_completed = data['is_completed']
            if data['is_completed']:
                task.completed_at = datetime.utcnow()
                task.completed_by_id = user_id
            else:
                task.completed_at = None
                task.completed_by_id = None
        if 'assigned_to_id' in data:
            task.assigned_to_id = data['assigned_to_id']
        if 'description' in data:
            task.description = data['description']
        if 'due_date' in data:
            if data['due_date']:
                task.due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
            else:
                task.due_date = None
        
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        db.session.delete(task)
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/my-tasks', methods=['GET'])
def get_my_tasks():
    """Ottieni tutte le task dell'utente corrente (globali e legate a opportunità)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    # Task globali (opportunity_id IS NULL) assegnate all'utente
    global_tasks = OpportunityTask.query.filter(
        OpportunityTask.opportunity_id.is_(None),
        OpportunityTask.assigned_to_id == user_id
    ).order_by(OpportunityTask.due_date.asc(), OpportunityTask.created_at.desc()).all()
    
    # Task legate a opportunità assegnate all'utente
    opp_tasks = OpportunityTask.query.filter(
        OpportunityTask.opportunity_id.isnot(None),
        OpportunityTask.assigned_to_id == user_id
    ).order_by(OpportunityTask.due_date.asc(), OpportunityTask.created_at.desc()).all()
    
    def task_to_dict(t):
        return {
            'id': t.id,
            'task_type': t.task_type,
            'description': t.description,
            'assigned_to_id': t.assigned_to_id,
            'assigned_to_name': f"{t.assigned_to.first_name} {t.assigned_to.last_name}" if t.assigned_to else None,
            'due_date': t.due_date.isoformat() if t.due_date else None,
            'is_completed': t.is_completed,
            'completed_at': t.completed_at.isoformat() if t.completed_at else None,
            'completed_by_name': f"{t.completed_by.first_name} {t.completed_by.last_name}" if t.completed_by else None,
            'created_at': t.created_at.isoformat() if t.created_at else None,
            'opportunity_id': t.opportunity_id,
            'opportunity_name': t.opportunity.name if t.opportunity else None,
            'account_id': t.account_id,
            'account_name': (t.account.name if t.account else None) or (t.opportunity.account.name if t.opportunity and t.opportunity.account else None),
            'account_image_url': (t.account.image_url if t.account and t.account.image_url else None) or (t.opportunity.account.image_url if t.opportunity and t.opportunity.account and t.opportunity.account.image_url else None),
            'opportunity_offers': [{'id': o.id, 'number': o.number, 'pdf_path': o.pdf_path, 'docx_path': o.docx_path, 'amount': (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))} for o in t.opportunity.offers] if t.opportunity and t.opportunity.offers else []
        }
    
    return jsonify({
        'global_tasks': [task_to_dict(t) for t in global_tasks],
        'opportunity_tasks': [task_to_dict(t) for t in opp_tasks],
        'total': len(global_tasks) + len(opp_tasks)
    })


@app.route('/api/tasks', methods=['POST'])
def create_global_task():
    """Crea una task globale (non legata a opportunità)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    if not data.get('task_type'):
        return jsonify({'error': 'Tipo task richiesto'}), 400
    
    due_date = None
    if data.get('due_date'):
        due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
    
    task = OpportunityTask(
        opportunity_id=data.get('opportunity_id'),  # Può essere None per task globale
        account_id=data.get('account_id'),  # Collegamento diretto ad account
        task_type=data['task_type'],
        description=data.get('description', ''),
        assigned_to_id=data.get('assigned_to_id'),
        due_date=due_date
    )
    db.session.add(task)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'task': {
            'id': task.id,
            'task_type': task.task_type,
            'description': task.description,
            'assigned_to_id': task.assigned_to_id,
            'assigned_to_name': f"{task.assigned_to.first_name} {task.assigned_to.last_name}" if task.assigned_to else None,
            'due_date': task.due_date.isoformat() if task.due_date else None,
            'is_completed': task.is_completed
        }
    })


@app.route('/api/tasks/all', methods=['GET'])
def get_all_tasks():
    """Ottieni tutte le task di tutti gli operatori (per vista per operatore)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Solo admin può vedere tutte le attività
    if session.get('user_role') != 'admin':
        return jsonify({'error': 'Non autorizzato'}), 403
    
    all_tasks = OpportunityTask.query.order_by(
        OpportunityTask.assigned_to_id,
        OpportunityTask.due_date.asc(),
        OpportunityTask.created_at.desc()
    ).all()
    
    def task_to_dict(t):
        return {
            'id': t.id,
            'task_type': t.task_type,
            'description': t.description,
            'assigned_to_id': t.assigned_to_id,
            'assigned_to_name': f"{t.assigned_to.first_name} {t.assigned_to.last_name}" if t.assigned_to else 'Non assegnato',
            'due_date': t.due_date.isoformat() if t.due_date else None,
            'is_completed': t.is_completed,
            'completed_at': t.completed_at.isoformat() if t.completed_at else None,
            'opportunity_id': t.opportunity_id,
            'opportunity_name': t.opportunity.name if t.opportunity else None,
            'account_id': t.account_id,
            'account_name': (t.account.name if t.account else None) or (t.opportunity.account.name if t.opportunity and t.opportunity.account else None),
            'account_image_url': (t.account.image_url if t.account and t.account.image_url else None) or (t.opportunity.account.image_url if t.opportunity and t.opportunity.account and t.opportunity.account.image_url else None),
            'opportunity_offers': [{'id': o.id, 'number': o.number, 'pdf_path': o.pdf_path, 'docx_path': o.docx_path, 'amount': (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))} for o in t.opportunity.offers] if t.opportunity and t.opportunity.offers else []
        }
    
    return jsonify({
        'tasks': [task_to_dict(t) for t in all_tasks]
    })


@app.route('/api/tasks/<int:task_id>', methods=['PUT', 'DELETE'])
def update_or_delete_global_task(task_id):
    """Aggiorna o elimina una task (globale o legata a opportunità)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    task = OpportunityTask.query.get_or_404(task_id)
    user_id = session['user_id']
    
    if request.method == 'PUT':
        data = request.json
        if 'is_completed' in data:
            task.is_completed = data['is_completed']
            if data['is_completed']:
                task.completed_at = datetime.utcnow()
                task.completed_by_id = user_id
            else:
                task.completed_at = None
                task.completed_by_id = None
        if 'assigned_to_id' in data:
            task.assigned_to_id = data['assigned_to_id']
        if 'description' in data:
            task.description = data['description']
        if 'due_date' in data:
            if data['due_date']:
                task.due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
            else:
                task.due_date = None
        
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        db.session.delete(task)
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/opportunities/<int:opportunity_id>/take-ownership', methods=['POST'])
def take_opportunity_ownership(opportunity_id):
    """Prendi ownership di un'opportunità"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = Opportunity.query.get_or_404(opportunity_id)
    user_id = session['user_id']
    
    opp.owner_id = user_id
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': opp.owner_id,
        'owner_name': f"{opp.owner.first_name} {opp.owner.last_name}" if opp.owner else 'Unknown'
    })


@app.route('/api/users/commerciali', methods=['GET'])
def get_commercial_users():
    """Ottieni lista utenti commerciali"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    users = User.query.filter_by(role='commerciale').all()
    return jsonify({
        'users': [{
            'id': u.id,
            'username': u.username,
            'name': f"{u.first_name} {u.last_name}",
            'email': u.email
        } for u in users]
    })


@app.route('/api/products', methods=['GET'])
def products():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    products_list = Product.query.filter_by(is_active=True).all()
    result = []
    
    for p in products_list:
        prices_dict = {}
        for price_entry in p.prices:
            if price_entry.is_active:
                prices_dict[price_entry.price_level] = price_entry.amount
        
        # Compatibilità bronze/silver/gold
        primary_prices = {}
        if 'bronze' in prices_dict:
            primary_prices['bronze'] = prices_dict['bronze']
        if 'silver' in prices_dict:
            primary_prices['silver'] = prices_dict['silver']
        if 'gold' in prices_dict:
            primary_prices['gold'] = prices_dict['gold']
        
        result.append({
            'id': p.id,
            'code': p.code,
            'name': p.name,
            'description': p.description,
            'category': p.category,
            'family': p.family,
            'primary_prices': primary_prices,
            'prices': [{
                'price_level': pe.price_level,
                'amount': pe.amount,
                'currency': pe.currency,
                'source': pe.source,
                'is_override': pe.is_override
            } for pe in p.prices if pe.is_active]
        })
    
    return jsonify({'products': result})


@app.route('/api/products/<int:product_id>/prices', methods=['PUT'])
def update_product_price(product_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    product = Product.query.get_or_404(product_id)
    
    price_level = data['price_level']
    amount = float(data['amount'])
    
    # Cerca entry esistente o crea nuova
    entry = PriceBookEntry.query.filter_by(
        product_id=product_id,
        price_level=price_level
    ).first()
    
    if entry:
        entry.amount = amount
        entry.is_override = True
        entry.is_active = True
    else:
        entry = PriceBookEntry(
            product_id=product_id,
            price_level=price_level,
            amount=amount,
            currency='EUR',
            source='manual',
            is_override=True,
            is_active=True
        )
        db.session.add(entry)
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/pricebook/sync', methods=['POST'])
def sync_pricebook():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        from price_loader import load_price_catalog
        
        catalog = load_price_catalog()
        
        # Sincronizza nel database
        for code, entry in catalog.items():
            # Crea o aggiorna prodotto
            product = Product.query.filter_by(code=code).first()
            if not product:
                product = Product(
                    code=code,
                    name=entry.name,
                    description=entry.description,
                    category=entry.category,
                    family=entry.family,
                    is_active=True
                )
                db.session.add(product)
                db.session.flush()
            
            # Aggiorna prezzi
            for level, amount in entry.price_levels.items():
                price_entry = PriceBookEntry.query.filter_by(
                    product_id=product.id,
                    price_level=level
                ).first()
                
                if price_entry:
                    if not price_entry.is_override:  # Non sovrascrivere override manuali
                        price_entry.amount = amount
                        price_entry.source = entry.sources.get(level, 'official')
                        price_entry.is_active = True
                else:
                    price_entry = PriceBookEntry(
                        product_id=product.id,
                        price_level=level,
                        amount=amount,
                        currency='EUR',
                        source=entry.sources.get(level, 'official'),
                        is_active=True
                    )
                    db.session.add(price_entry)
        
        db.session.commit()
        return jsonify({'success': True, 'message': f'Sincronizzati {len(catalog)} prodotti'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/offers', methods=['GET'])
def offers():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Mostra tutte le offerte (condivise)
    offers_list = OfferDocument.query.order_by(OfferDocument.created_at.desc()).limit(100).all()
    
    offers_data = []
    for o in offers_list:
        # Determina nome cliente
        client_name = 'N/A'
        if o.opportunity and o.opportunity.account:
            client_name = o.opportunity.account.name
        elif o.owner:
            client_name = f"{o.owner.first_name} {o.owner.last_name}"
        
        offers_data.append({
            'id': o.id,
            'number': o.number,
            'status': o.status,
            'amount': o.amount,
            'description': o.description,
            'client': client_name,
            'opportunity_id': o.opportunity_id,
            'opportunity_name': o.opportunity.name if o.opportunity else None,
            'docx_path': o.docx_path,
            'pdf_path': o.pdf_path,
            'created_at': o.created_at.isoformat() if o.created_at else None
        })
    
    return jsonify({'offers': offers_data})


@app.route('/api/offers/<int:offer_id>/pdf', methods=['GET'])
def download_offer_pdf(offer_id):
    """Apri il PDF di un'offerta nel browser"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    offer = OfferDocument.query.get_or_404(offer_id)
    
    if not offer.pdf_path or not os.path.exists(offer.pdf_path):
        return jsonify({'error': 'PDF non trovato'}), 404
    
    # Apri nel browser invece di forzare il download
    return send_file(offer.pdf_path, as_attachment=False, mimetype='application/pdf')


@app.route('/api/offers/<int:offer_id>/docx', methods=['GET'])
def download_offer_docx(offer_id):
    """Scarica il DOCX di un'offerta"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    offer = OfferDocument.query.get_or_404(offer_id)
    
    if not offer.docx_path or not os.path.exists(offer.docx_path):
        return jsonify({'error': 'DOCX non trovato'}), 404
    
    return send_file(offer.docx_path, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/offers/generate', methods=['POST'])
def generate_offer():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        data = request.json
        user_id = session['user_id']
        
        # Estrai dati
        request_text = data.get('request', '').strip()
        account_id = data.get('account_id')
        opportunity_id = data.get('opportunity_id')
        client_name = data.get('client_name', '').strip()
        
        # Se account_id o opportunity_id sono forniti, usa quelli
        if account_id:
            account = Account.query.get(account_id)
            if account:
                client_name = account.name
        
        if not client_name and not account_id:
            return jsonify({'error': 'Specifica un account o un nome cliente'}), 400
        
        if not request_text:
            return jsonify({'error': 'Specifica i prodotti richiesti'}), 400
        
        # Importa generatore
        try:
            from offer_generator import OfferGenerator
            import tempfile
            import platform
            
            # Determina directory offerte
            if platform.system() == 'Windows':
                offers_dir = r"K:\OFFERTE - K\OFFERTE 2025 - K"
            else:
                offers_dir = os.environ.get('OFFERS_FOLDER', '/mnt/k/OFFERTE - K/OFFERTE 2025 - K')
            
            # Fallback se K:\ non disponibile
            if not os.path.exists(offers_dir):
                offers_dir = os.path.join(os.path.dirname(__file__), 'offerte_generate')
                os.makedirs(offers_dir, exist_ok=True)
            
            generator = OfferGenerator(offers_directory=offers_dir)
            
            # Carica listini se disponibili
            pricelist_dir = os.path.join(os.path.dirname(__file__), 'listini')
            if os.path.exists(pricelist_dir):
                pdf_files = [os.path.join(pricelist_dir, f) for f in os.listdir(pricelist_dir) if f.endswith('.pdf')]
                if pdf_files:
                    generator.load_pricelists(pdf_files)
            
            # Crea offerta
            try:
                doc_handler, offer_info = generator.create_offer(
                    request=request_text,
                    company_name=client_name,
                    search_company_online=False
                )
            except Exception as e:
                import traceback
                error_trace = traceback.format_exc()
                logger.error(f"Errore nella creazione offerta: {e}\n{error_trace}")
                return jsonify({
                    'error': f'Errore nella generazione offerta: {str(e)}',
                    'trace': error_trace
                }), 500
            
            if not doc_handler:
                # Verifica se ci sono prodotti disponibili
                available_products = generator.parser.get_all_products()
                if not available_products:
                    return jsonify({
                        'error': 'Nessun prodotto disponibile. Carica prima i listini PDF nella sezione Listini.',
                        'hint': 'Vai su "Listini" e clicca "Sincronizza Listini" per caricare i prodotti.'
                    }), 400
                return jsonify({
                    'error': 'Nessun prodotto trovato nella richiesta. Verifica che i nomi dei prodotti siano corretti.',
                    'hint': f'Prodotti disponibili: {len(available_products)}. Prova a cercare prodotti nella sezione Listini.'
                }), 400
            
            # Salva offerta
            docx_path, pdf_path = generator.save_offer(doc_handler, offer_info, generate_pdf=True)
            
            if not docx_path:
                return jsonify({'error': 'Errore nel salvataggio offerta'}), 500
            
            # Salva offerta nel database
            offer_doc = OfferDocument(
                number=offer_info['offer_number'],
                status='draft',
                amount=offer_info.get('total', 0.0),
                description=f"Prodotti: {', '.join([p.get('name', '') for p in offer_info.get('products', [])[:3]])}",
                docx_path=docx_path,
                pdf_path=pdf_path,
                folder_path=os.path.dirname(docx_path) if docx_path else None,
                generated_from_nlp=True,
                opportunity_id=opportunity_id if opportunity_id else None,
                owner_id=user_id
            )
            db.session.add(offer_doc)
            
            # Se opportunity_id è fornito, aggiorna l'opportunità
            if opportunity_id:
                opp = Opportunity.query.get(opportunity_id)
                if opp:
                    opp.description = f"Offerta generata: {offer_info['offer_number']}\n{opp.description or ''}"
                    # Aggiorna anche l'amount se non è impostato
                    if opp.amount == 0.0 and offer_info.get('total', 0) > 0:
                        opp.amount = offer_info['total']
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'offer_number': offer_info['offer_number'],
                'client': offer_info['client_name'],
                'total': offer_info['total'],
                'products_count': len(offer_info['products']),
                'docx_path': docx_path,
                'pdf_path': pdf_path,
                'offer_id': offer_doc.id
            })
            
        except ImportError as e:
            return jsonify({'error': f'Modulo generatore non disponibile: {e}'}), 500
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Errore generazione offerta: {str(e)}'}), 500
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/offers/import-past', methods=['POST'])
def import_past_offers():
    """Importa offerte passate dalla cartella"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        from analyze_past_offers import analyze_offers_folder, import_offers_to_crm
        offers_data = analyze_offers_folder()
        imported = import_offers_to_crm(offers_data, db.session, session['user_id'])
        return jsonify({
            'success': True,
            'imported': imported,
            'total_found': len(offers_data)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ur/sync', methods=['POST'])
def ur_sync_route():
    """Endpoint per sincronizzazione dati CRM dal portale UR con Selenium"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session.get('user_id', 1)
    
    try:
        # Usa la funzione di sincronizzazione con logging completo
        from ur_portal_integration import sync_ur_portal
        
        # Esegui sincronizzazione (include già tutto il logging)
        stats = sync_ur_portal(db.session, user_id)
        
        if stats.get('errors'):
            return jsonify({
                'success': False,
                'error': '\n'.join(stats['errors']),
                'stats': stats,
                'message': f"Sincronizzazione completata con {len(stats['errors'])} errori"
            }), 500
        
        return jsonify({
            'success': True,
            'message': 'Sincronizzazione completata con successo',
            'stats': stats
        })
        
    except ImportError as e:
        return jsonify({
            'success': False,
            'error': f'Moduli integrazione UR non trovati: {str(e)}. Installa: pip install selenium webdriver-manager'
        }), 500
    except ConnectionError as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'Errore connessione al portale UR: {str(e)}. Verifica la connessione internet e che il dominio sia raggiungibile.',
            'error_type': 'ConnectionError'
        }), 500
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'Errore sincronizzazione UR: {str(e)}',
            'error_type': type(e).__name__
        }), 500


@app.route('/api/offers/sync', methods=['POST'])
def sync_offers_route():
    """Endpoint per sincronizzare offerte dalla cartella Windows"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        # Importa e esegui sync
        from sync_offers_from_folder import sync_offers
        import logging
        import io
        import sys
        
        # Cattura output
        log_capture = io.StringIO()
        handler = logging.StreamHandler(log_capture)
        handler.setLevel(logging.INFO)
        
        # Esegui sync
        result_msg = []
        try:
            sync_offers()
            result_msg.append("✅ Sincronizzazione completata")
            
            # Dopo la sincronizzazione, aggiorna tutti i valori delle offerte
            update_all_offer_values()
            result_msg.append("✅ Valori offerte aggiornati")
        except Exception as sync_error:
            error_msg = f"❌ Errore durante sync: {str(sync_error)}"
            result_msg.append(error_msg)
            import traceback
            result_msg.append(traceback.format_exc())
            raise
        
        return jsonify({
            'success': True,
            'message': '\n'.join(result_msg) if result_msg else 'Sincronizzazione offerte completata. Opportunità aggiornate.'
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Errore sync offerte: {error_trace}")
        return jsonify({
            'success': False,
            'error': str(e),
            'trace': error_trace
        }), 500


def update_all_offer_values():
    """Aggiorna tutti i valori delle offerte estraendoli dai PDF"""
    # Aggiorna tutte le offerte con PDF disponibile
    all_offers = OfferDocument.query.filter(OfferDocument.pdf_path != None).all()
    
    updated_count = 0
    for offer in all_offers:
        if offer.pdf_path and os.path.exists(offer.pdf_path):
            try:
                amount = extract_offer_amount_from_file(offer.pdf_path)
                if amount > 0:
                    # Aggiorna sempre se il valore estratto è maggiore o se non c'è valore
                    old_amount = offer.amount or 0.0
                    if abs(amount - old_amount) > 0.01:
                        offer.amount = amount
                        updated_count += 1
                        if old_amount > 0:
                            print(f"  💰 Valore aggiornato per {offer.number}: €{old_amount:,.2f} → €{amount:,.2f}")
                        else:
                            print(f"  💰 Valore estratto per {offer.number}: €{amount:,.2f}")
            except Exception as e:
                print(f"  ⚠️  Errore estrazione valore per {offer.number}: {e}")
    
    # Aggiorna anche i valori delle opportunità
    opportunities = Opportunity.query.all()
    opp_updated = 0
    for opp in opportunities:
        total_value = 0.0
        for offer in opp.offers:
            if offer.pdf_path and os.path.exists(offer.pdf_path):
                amount = offer.amount if offer.amount and offer.amount > 0 else extract_offer_amount_from_file(offer.pdf_path)
                total_value += amount
            elif offer.amount and offer.amount > 0:
                total_value += offer.amount
        
        if total_value > 0:
            old_amount = opp.amount or 0.0
            if abs(old_amount - total_value) > 0.01:
                opp.amount = total_value
                opp_updated += 1
                print(f"  💰 Valore opportunità {opp.name}: €{old_amount:,.2f} → €{total_value:,.2f}")
    
    db.session.commit()
    print(f"✅ Aggiornati {updated_count} valori offerte e {opp_updated} valori opportunità")


@app.route('/api/offers/update-values', methods=['POST'])
def update_offer_values_route():
    """Endpoint per aggiornare tutti i valori delle offerte dai PDF"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        update_all_offer_values()
        return jsonify({
            'success': True,
            'message': 'Valori offerte aggiornati con successo'
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Errore aggiornamento valori: {error_trace}")
        return jsonify({
            'success': False,
            'error': str(e),
            'trace': error_trace
        }), 500


@app.route('/api/ur/logs', methods=['GET'])
def ur_sync_logs():
    """Endpoint per leggere i log della sincronizzazione UR"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        log_file = 'ur_portal_sync.log'
        lines = request.args.get('lines', 100, type=int)  # Ultime N righe
        
        if not os.path.exists(log_file):
            return jsonify({
                'success': True,
                'logs': [],
                'message': 'File log non ancora creato'
            })
        
        with open(log_file, 'r', encoding='utf-8') as f:
            all_lines = f.readlines()
            # Prendi le ultime N righe
            recent_lines = all_lines[-lines:] if len(all_lines) > lines else all_lines
        
        return jsonify({
            'success': True,
            'logs': [line.strip() for line in recent_lines],
            'total_lines': len(all_lines),
            'returned_lines': len(recent_lines)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Errore lettura log: {str(e)}'
        }), 500


@app.route('/api/analytics/pipeline', methods=['GET'])
def pipeline_analytics():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    period = request.args.get('period', 'month')  # day, week, month
    
    # Calcola data di inizio (ultimi 6 mesi)
    from datetime import datetime, timedelta
    end_date = datetime.utcnow()
    
    # Calcola start_date (6 mesi fa) - usa logica manuale per evitare dipendenze
    start_date = end_date
    for _ in range(6):
        # Sottrai un mese (approssimato a 30 giorni)
        start_date = start_date - timedelta(days=30)
    
    # Query opportunità degli ultimi 6 mesi
    opps = Opportunity.query.filter(
        Opportunity.owner_id == user_id,
        Opportunity.created_at >= start_date
    ).all()
    
    # Genera tutti i periodi degli ultimi 6 mesi
    data = {}
    current = start_date
    
    while current <= end_date:
        if period == 'day':
            period_key = current.strftime('%Y-%m-%d')
            current += timedelta(days=1)
        elif period == 'week':
            period_key = current.strftime('%Y-W%W')
            current += timedelta(weeks=1)
        else:  # month
            period_key = current.strftime('%Y-%m')
            # Avanza di un mese (approssimato)
            if current.month == 12:
                current = current.replace(year=current.year + 1, month=1)
            else:
                current = current.replace(month=current.month + 1)
        
        data[period_key] = {
            'period': period_key,
            'MIR': {'open': 0, 'closed': 0},
            'UR': {'open': 0, 'closed': 0},
            'Unitree': {'open': 0, 'closed': 0},
            'ROEQ': {'open': 0, 'closed': 0},
            'Servizi': {'open': 0, 'closed': 0},
            'Altro': {'open': 0, 'closed': 0}
        }
    
    # Popola con i dati reali
    for opp in opps:
        # Determina periodo
        if period == 'day':
            period_key = opp.created_at.strftime('%Y-%m-%d')
        elif period == 'week':
            period_key = opp.created_at.strftime('%Y-W%W')
        else:
            period_key = opp.created_at.strftime('%Y-%m')
        
        if period_key not in data:
            continue  # Skip se fuori range
        
        # Determina categoria dai prodotti o supplier_category
        category = 'Altro'
        if opp.supplier_category:
            # Mappa supplier_category alle categorie
            if 'UR' in opp.supplier_category or 'Universal Robot' in opp.supplier_category:
                category = 'UR'
            elif 'MiR' in opp.supplier_category or 'MIR' in opp.supplier_category:
                category = 'MIR'
            elif 'Unitree' in opp.supplier_category:
                category = 'Unitree'
            else:
                category = opp.supplier_category
        elif opp.line_items:
            # Prendi categoria dal primo prodotto
            first_product = opp.line_items[0].product
            if first_product:
                category = first_product.category or 'Altro'
        
        # Normalizza categoria
        if category not in data[period_key]:
            category = 'Altro'
        
        # Calcola valore (usa amount o somma offerte)
        opp_value = opp.amount if opp.amount and opp.amount > 0 else 0
        if opp_value == 0 and opp.offers:
            opp_value = sum(
                (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                for o in opp.offers
            )
        
        # Determina se chiusa
        is_closed = opp.stage in ['Closed Won', 'Closed Lost']
        
        if is_closed:
            data[period_key][category]['closed'] += opp_value
        else:
            data[period_key][category]['open'] += opp_value
    
    # Ordina per periodo
    sorted_data = sorted(data.values(), key=lambda x: x['period'])
    
    return jsonify({'data': sorted_data})


# ==================== EMAIL MARKETING API ROUTES ====================

@app.route('/api/email/lists', methods=['GET', 'POST'])
def email_lists():
    """Gestisci liste email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        lists = EmailList.query.filter_by(is_active=True).order_by(EmailList.created_at.desc()).all()
        return jsonify({
            'lists': [{
                'id': l.id,
                'name': l.name,
                'description': l.description,
                'subscriber_count': len([s for s in l.subscriptions if s.status == 'subscribed']),
                'subscribers_count': len([s for s in l.subscriptions if s.status == 'subscribed']),  # Mantenuto per retrocompatibilità
                'owner_id': l.owner_id,
                'created_at': l.created_at.isoformat() if l.created_at else None
            } for l in lists]
        })
    
    elif request.method == 'POST':
        data = request.json
        email_list = EmailList(
            name=data['name'],
            description=data.get('description'),
            owner_id=user_id
        )
        db.session.add(email_list)
        db.session.commit()
        return jsonify({'success': True, 'list_id': email_list.id})


@app.route('/api/email/lists/<int:list_id>', methods=['GET', 'PUT', 'DELETE'])
def email_list_detail(list_id):
    """Dettagli, aggiorna o elimina lista"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    email_list = EmailList.query.get_or_404(list_id)
    
    if request.method == 'GET':
        subscribers = EmailListSubscription.query.filter_by(
            list_id=list_id,
            status='subscribed'
        ).all()
        
        return jsonify({
            'list': {
                'id': email_list.id,
                'name': email_list.name,
                'description': email_list.description,
                'subscribers': [{
                    'id': s.id,
                    'email': s.email,
                    'first_name': s.first_name,
                    'last_name': s.last_name,
                    'contact_id': s.contact_id,
                    'lead_id': s.lead_id
                } for s in subscribers],
                'subscriber_count': len(subscribers),
                'subscribers_count': len(subscribers)  # Mantenuto per retrocompatibilità
            }
        })
    
    elif request.method == 'PUT':
        data = request.json
        if 'name' in data:
            email_list.name = data['name']
        if 'description' in data:
            email_list.description = data['description']
        if 'is_active' in data:
            email_list.is_active = data['is_active']
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        email_list.is_active = False
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/email/lists/<int:list_id>/subscribe', methods=['POST'])
def subscribe_to_list(list_id):
    """Iscrivi contatti/leads a una lista"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    email_list = EmailList.query.get_or_404(list_id)
    data = request.json
    
    contact_ids = data.get('contact_ids', [])
    lead_ids = data.get('lead_ids', [])
    emails = data.get('emails', [])  # Email manuali
    
    subscribed = []
    
    # Iscrivi contatti
    for contact_id in contact_ids:
        contact = Contact.query.get(contact_id)
        if contact and contact.email:
            sub = EmailListSubscription.query.filter_by(
                list_id=list_id,
                contact_id=contact_id
            ).first()
            
            if not sub:
                sub = EmailListSubscription(
                    list_id=list_id,
                    contact_id=contact_id,
                    email=contact.email,
                    first_name=contact.first_name,
                    last_name=contact.last_name,
                    status='subscribed'
                )
                db.session.add(sub)
                subscribed.append(contact.email)
    
    # Iscrivi leads
    for lead_id in lead_ids:
        lead = Lead.query.get(lead_id)
        if lead and lead.email:
            sub = EmailListSubscription.query.filter_by(
                list_id=list_id,
                lead_id=lead_id
            ).first()
            
            if not sub:
                sub = EmailListSubscription(
                    list_id=list_id,
                    lead_id=lead_id,
                    email=lead.email,
                    first_name=lead.contact_first_name,
                    last_name=lead.contact_last_name,
                    status='subscribed'
                )
                db.session.add(sub)
                subscribed.append(lead.email)
    
    # Iscrivi email manuali
    for email_data in emails:
        if isinstance(email_data, str):
            email = email_data
            name = ""
        else:
            email = email_data.get('email')
            name = email_data.get('name', '')
        
        if email:
            # Controlla se esiste già (per email o per contatto/lead collegato)
            sub = EmailListSubscription.query.filter_by(
                list_id=list_id,
                email=email,
                status='subscribed'
            ).first()
            
            if not sub:
                first_name, last_name = (name.split(' ', 1) + [''])[:2] if name else ('', '')
                sub = EmailListSubscription(
                    list_id=list_id,
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    status='subscribed'
                )
                db.session.add(sub)
                subscribed.append(email)
            else:
                # Se esiste già, aggiorna nome se fornito
                if name and (not sub.first_name or not sub.last_name):
                    first_name, last_name = (name.split(' ', 1) + [''])[:2]
                    sub.first_name = sub.first_name or first_name
                    sub.last_name = sub.last_name or last_name
                    db.session.commit()
    
    db.session.commit()
    return jsonify({'success': True, 'subscribed': subscribed, 'count': len(subscribed)})


@app.route('/api/email/templates', methods=['GET', 'POST'])
def email_templates():
    """Gestisci template email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        templates = EmailTemplate.query.filter_by(is_active=True).order_by(EmailTemplate.created_at.desc()).all()
        return jsonify({
            'templates': [{
                'id': t.id,
                'name': t.name,
                'subject': t.subject,
                'category': t.category,
                'template_type': t.category,  # Aggiunto per retrocompatibilità con frontend
                'owner_id': t.owner_id,
                'created_at': t.created_at.isoformat() if t.created_at else None
            } for t in templates]
        })
    
    elif request.method == 'POST':
        data = request.json
        template = EmailTemplate(
            name=data['name'],
            subject=data['subject'],
            html_content=data['html_content'],
            text_content=data.get('text_content'),
            category=data.get('category', 'general'),
            owner_id=user_id
        )
        db.session.add(template)
        db.session.commit()
        return jsonify({'success': True, 'template_id': template.id})


@app.route('/api/email/templates/<int:template_id>', methods=['GET', 'PUT', 'DELETE'])
def email_template_detail(template_id):
    """Dettagli, aggiorna o elimina template"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    template = EmailTemplate.query.get_or_404(template_id)
    
    if request.method == 'GET':
        return jsonify({
            'template': {
                'id': template.id,
                'name': template.name,
                'subject': template.subject,
                'html_content': template.html_content,
                'text_content': template.text_content,
                'category': template.category
            }
        })
    
    elif request.method == 'PUT':
        data = request.json
        if 'name' in data:
            template.name = data['name']
        if 'subject' in data:
            template.subject = data['subject']
        if 'html_content' in data:
            template.html_content = data['html_content']
        if 'text_content' in data:
            template.text_content = data['text_content']
        if 'category' in data:
            template.category = data['category']
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        template.is_active = False
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/email/campaigns', methods=['GET', 'POST'])
def email_campaigns():
    """Gestisci campagne email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        campaigns = EmailCampaign.query.order_by(EmailCampaign.created_at.desc()).limit(100).all()
        result_campaigns = []
        for c in campaigns:
            # Calcola numero destinatari dalla lista (sempre aggiornato)
            subscriber_count = 0
            if c.list_id:
                subscriber_count = EmailListSubscription.query.filter_by(
                    list_id=c.list_id,
                    status='subscribed'
                ).count()
            
            # Usa sempre il conteggio aggiornato degli iscritti per campagne non ancora inviate
            # Per campagne già inviate, mostra il numero effettivamente inviato
            if c.status == 'sent' and c.total_recipients > 0:
                recipients = c.total_recipients  # Numero effettivamente inviato
            else:
                recipients = subscriber_count  # Conteggio aggiornato dalla lista
            
            result_campaigns.append({
                'id': c.id,
                'name': c.name,
                'subject': c.subject,
                'status': c.status,
                'list_id': c.list_id,
                'list_name': c.email_list.name if c.email_list else None,
                'total_recipients': recipients,
                'list_subscriber_count': subscriber_count,
                'total_sent': c.total_sent,
                'total_delivered': c.total_delivered,
                'total_opened': c.total_opened,
                'total_clicked': c.total_clicked,
                'scheduled_at': c.scheduled_at.isoformat() if c.scheduled_at else None,
                'sent_at': c.sent_at.isoformat() if c.sent_at else None,
                'created_at': c.created_at.isoformat() if c.created_at else None
            })
        
        return jsonify({'campaigns': result_campaigns})
    
    elif request.method == 'POST':
        data = request.json
        scheduled_at = None
        if data.get('scheduled_at'):
            scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00'))
        
        # Supporta sia list_id (singolo) che list_ids (array)
        list_ids = []
        if 'list_ids' in data and isinstance(data['list_ids'], list):
            list_ids = data['list_ids']
        elif 'list_id' in data:
            list_ids = [data['list_id']]
        else:
            return jsonify({'success': False, 'error': 'list_id o list_ids richiesto'}), 400
        
        # Se più liste, crea una campagna unica che invierà a tutte le liste (con deduplicazione)
        # Se una sola lista, crea campagna normale
        if len(list_ids) == 1:
            campaign = EmailCampaign(
                name=data['name'],
                subject=data['subject'],
                html_content=data['html_content'],
                text_content=data.get('text_content'),
                list_id=list_ids[0],
                template_id=data.get('template_id'),
                scheduled_at=scheduled_at,
                status='draft',
                owner_id=user_id
            )
            db.session.add(campaign)
            db.session.commit()
            return jsonify({'success': True, 'campaign_id': campaign.id})
        else:
            # Per più liste, salva tutte le liste come JSON nel campo text_content (se vuoto) o in un commento HTML
            import json
            list_ids_json = json.dumps(list_ids)
            
            # Salva le liste come commento JSON nell'HTML (non visibile nell'email)
            html_with_metadata = data['html_content'] + f'\n<!-- CAMPAIGN_LIST_IDS: {list_ids_json} -->'
            
            campaign = EmailCampaign(
                name=f"{data['name']} ({len(list_ids)} liste)",
                subject=data['subject'],
                html_content=html_with_metadata,
                text_content=data.get('text_content') or f'LIST_IDS:{list_ids_json}',  # Backup nel text_content
                list_id=list_ids[0],  # Prima lista come riferimento
                template_id=data.get('template_id'),
                scheduled_at=scheduled_at,
                status='draft',
                owner_id=user_id
            )
            db.session.add(campaign)
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'campaign_id': campaign.id,
                'list_ids': list_ids,
                'message': f'Campagna creata per {len(list_ids)} liste. I destinatari verranno deduplicati all\'invio.'
            })


@app.route('/api/email/campaigns/<int:campaign_id>', methods=['GET', 'PUT', 'DELETE'])
def email_campaign_detail(campaign_id):
    """Dettagli, aggiorna o elimina campagna"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    campaign = EmailCampaign.query.get_or_404(campaign_id)
    
    if request.method == 'GET':
        return jsonify({
            'campaign': {
                'id': campaign.id,
                'name': campaign.name,
                'subject': campaign.subject,
                'html_content': campaign.html_content,
                'text_content': campaign.text_content,
                'status': campaign.status,
                'list_id': campaign.list_id,
                'template_id': campaign.template_id,
                'scheduled_at': campaign.scheduled_at.isoformat() if campaign.scheduled_at else None,
                'sent_at': campaign.sent_at.isoformat() if campaign.sent_at else None,
                'total_recipients': campaign.total_recipients,
                'total_sent': campaign.total_sent,
                'total_delivered': campaign.total_delivered,
                'total_opened': campaign.total_opened,
                'total_clicked': campaign.total_clicked,
                'total_bounced': campaign.total_bounced,
                'total_unsubscribed': campaign.total_unsubscribed
            }
        })
    
    elif request.method == 'PUT':
        data = request.json
        if 'name' in data:
            campaign.name = data['name']
        if 'subject' in data:
            campaign.subject = data['subject']
        if 'html_content' in data:
            campaign.html_content = data['html_content']
        if 'text_content' in data:
            campaign.text_content = data['text_content']
        if 'list_id' in data:
            campaign.list_id = data['list_id']
        if 'template_id' in data:
            campaign.template_id = data['template_id']
        if 'status' in data:
            campaign.status = data['status']
        if 'scheduled_at' in data:
            campaign.scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00')) if data['scheduled_at'] else None
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        db.session.delete(campaign)
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/email/campaigns/<int:campaign_id>/send', methods=['POST'])
def send_email_campaign(campaign_id):
    """Invia una campagna email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        campaign = EmailCampaign.query.get_or_404(campaign_id)
        data = request.json or {}
        
        # Determina le liste da usare
        # Se vengono passate liste nell'endpoint, usale (per supportare più liste)
        # Altrimenti, recupera le liste salvate nella campagna (dal commento HTML o text_content)
        if 'list_ids' in data and isinstance(data['list_ids'], list) and len(data['list_ids']) > 0:
            list_ids = data['list_ids']
        else:
            # Prova a recuperare le liste dal commento HTML o text_content
            import json
            import re
            list_ids = [campaign.list_id]  # Default: solo la lista principale
            
            # Cerca nel commento HTML
            html_match = re.search(r'<!--\s*CAMPAIGN_LIST_IDS:\s*(\[.*?\])\s*-->', campaign.html_content)
            if html_match:
                try:
                    list_ids = json.loads(html_match.group(1))
                except:
                    pass
            
            # Se non trovato, cerca nel text_content
            if len(list_ids) == 1 and campaign.text_content and campaign.text_content.startswith('LIST_IDS:'):
                try:
                    list_ids_json = campaign.text_content.replace('LIST_IDS:', '').strip()
                    list_ids = json.loads(list_ids_json)
                except:
                    pass
        
        # Se il nome contiene "(X liste)", significa che ci sono più liste
        # Per ora, usiamo solo la lista principale, ma in futuro possiamo salvare le liste in un campo JSON
        # TODO: Aggiungere campo list_ids JSON al modello EmailCampaign
        
        # Prepara destinatari da tutte le liste (con deduplicazione)
        subscriptions = EmailListSubscription.query.filter(
            EmailListSubscription.list_id.in_(list_ids),
            EmailListSubscription.status == 'subscribed'
        ).all()
        
        if not subscriptions:
            return jsonify({'error': 'Nessun destinatario trovato nelle liste'}), 400
        
        # Deduplica per email (case-insensitive) mantenendo la prima occorrenza
        seen_emails = set()
        recipients = []
        for sub in subscriptions:
            email_lower = sub.email.lower().strip() if sub.email else None
            if not email_lower or email_lower in seen_emails:
                continue  # Salta duplicati
            
            seen_emails.add(email_lower)
            
            # Ottieni nome normalizzato
            first_name, last_name, full_name = get_recipient_name(sub.email)
            
            recipient = {
                'email': sub.email,
                'name': full_name,
                'first_name': first_name,
                'last_name': last_name,
                'variables': {}
            }
            recipients.append(recipient)
        
        # Inizializza Mailgun
        try:
            from mailgun_service import MailgunService
            mailgun = MailgunService()
        except Exception as e:
            return jsonify({'error': f'Errore configurazione Mailgun: {str(e)}'}), 500
        
        # Aggiorna campagna
        campaign.status = 'sending'
        campaign.total_recipients = len(recipients)
        db.session.commit()
        
        # Rimuovi commenti HTML (come CAMPAIGN_LIST_IDS) prima di inviare
        import re
        html_clean = re.sub(r'<!--\s*CAMPAIGN_LIST_IDS:.*?-->\s*', '', campaign.html_content, flags=re.DOTALL)
        text_clean = campaign.text_content
        if text_clean and text_clean.startswith('LIST_IDS:'):
            text_clean = None  # Rimuovi se era solo metadata
        
        # Aggiungi firma email a tutti i contenuti
        from email_signature import add_signature_to_email
        
        html_with_signature, text_with_signature = add_signature_to_email(
            html_content=html_clean,
            text_content=text_clean
        )
        
        # Invia email
        campaign_data = {
            'name': campaign.name,
            'subject': campaign.subject,
            'html_content': html_with_signature,
            'text_content': text_with_signature
        }
        
        stats = mailgun.send_campaign_email(campaign_data, recipients)
        
        # Aggiorna statistiche campagna
        campaign.total_sent = stats['sent']
        campaign.status = 'sent' if stats['failed'] == 0 else 'partial'
        campaign.sent_at = datetime.utcnow()
        db.session.commit()
        
        # Salva dettagli invii
        for result in stats['results']:
            recipient = result.get('recipient', {})
            send = EmailCampaignSend(
                campaign_id=campaign_id,
                email=recipient.get('email', ''),
                status='sent' if result.get('success') else 'failed',
                mailgun_message_id=result.get('message_id'),
                sent_at=datetime.utcnow() if result.get('success') else None,
                error_message=result.get('error') if not result.get('success') else None
            )
            db.session.add(send)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'stats': {
                'total': stats['total'],
                'sent': stats['sent'],
                'failed': stats['failed']
            }
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        if campaign:
            campaign.status = 'failed'
            db.session.commit()
        return jsonify({'error': str(e)}), 500


@app.route('/api/email/campaigns/<int:campaign_id>/recipients', methods=['GET'])
def campaign_recipients(campaign_id):
    """Lista destinatari di una campagna"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    campaign = EmailCampaign.query.get_or_404(campaign_id)
    
    # Ottieni tutti i destinatari dalla lista
    subscriptions = EmailListSubscription.query.filter_by(
        list_id=campaign.list_id,
        status='subscribed'
    ).all()
    
    recipients = []
    for sub in subscriptions:
        recipients.append({
            'email': sub.email,
            'first_name': sub.first_name or '',
            'last_name': sub.last_name or '',
            'name': f"{sub.first_name or ''} {sub.last_name or ''}".strip() or sub.email,
            'subscribed_at': sub.subscribed_at.isoformat() if sub.subscribed_at else None
        })
    
    return jsonify({
        'success': True,
        'campaign_id': campaign_id,
        'campaign_name': campaign.name,
        'total': len(recipients),
        'recipients': recipients
    })


def normalize_name(name_str):
    """Normalizza un nome: rimuove spazi extra, capitalizza correttamente"""
    if not name_str:
        return ""
    # Rimuovi spazi multipli e trim
    name_str = " ".join(name_str.split())
    # Capitalizza ogni parola (gestisce anche apostrofi e trattini)
    parts = []
    for part in name_str.split():
        if part:
            # Gestisci apostrofi e trattini
            if "'" in part:
                subparts = part.split("'")
                parts.append("'".join([subparts[0].capitalize()] + [s.lower() for s in subparts[1:]]))
            elif "-" in part:
                subparts = part.split("-")
                parts.append("-".join([s.capitalize() for s in subparts]))
            else:
                parts.append(part.capitalize())
    return " ".join(parts)


def extract_name_from_email(email):
    """Estrae e formatta un nome dall'indirizzo email"""
    if not email or '@' not in email:
        return "", "", ""
    
    email_local = email.split('@')[0]
    
    # Prova diversi pattern comuni
    # Pattern 1: nome.cognome
    if '.' in email_local:
        parts = email_local.split('.')
        if len(parts) >= 2:
            first = normalize_name(parts[0])
            last = normalize_name('.'.join(parts[1:]))
            full = f"{first} {last}".strip()
            return first, last, full
    
    # Pattern 2: nome-cognome
    if '-' in email_local:
        parts = email_local.split('-')
        if len(parts) >= 2:
            first = normalize_name(parts[0])
            last = normalize_name('-'.join(parts[1:]))
            full = f"{first} {last}".strip()
            return first, last, full
    
    # Pattern 3: nomecognome (senza separatori)
    if len(email_local) > 3:
        # Prova a dividere a metà (euristico)
        mid = len(email_local) // 2
        first = normalize_name(email_local[:mid])
        last = normalize_name(email_local[mid:])
        full = f"{first} {last}".strip()
        return first, last, full
    
    # Fallback: usa tutto come nome
    name = normalize_name(email_local)
    return name, "", name


def get_recipient_name(email):
    """Ottiene il nome del destinatario dal database o dall'email, normalizzato"""
    first_name = ""
    last_name = ""
    full_name = ""
    
    # Cerca prima nelle subscription (priorità più alta per email marketing)
    subscription = EmailListSubscription.query.filter_by(email=email).first()
    if subscription:
        first_name = normalize_name(subscription.first_name or "")
        last_name = normalize_name(subscription.last_name or "")
        if first_name or last_name:
            full_name = f"{first_name} {last_name}".strip()
    
    # Se non trovato, cerca nei contatti
    if not full_name:
        contact = Contact.query.filter_by(email=email).first()
        if contact:
            first_name = normalize_name(contact.first_name or "")
            last_name = normalize_name(contact.last_name or "")
            if first_name or last_name:
                full_name = f"{first_name} {last_name}".strip()
    
    # Se non trovato, cerca nei lead
    if not full_name:
        lead = Lead.query.filter_by(email=email).first()
        if lead:
            first_name = normalize_name(lead.contact_first_name or "")
            last_name = normalize_name(lead.contact_last_name or "")
            if first_name or last_name:
                full_name = f"{first_name} {last_name}".strip()
    
    # Se ancora non trovato, estrai dall'email
    if not full_name:
        first_name, last_name, full_name = extract_name_from_email(email)
    
    # Fallback finale
    if not full_name:
        full_name = "Cliente"
        first_name = "Cliente"
    
    return first_name, last_name, full_name


@app.route('/api/email/analyze-outlook', methods=['POST'])
def analyze_outlook_emails():
    """Analizza email Outlook con Claude API e aggiunge attività al CRM"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    add_tasks_to_crm = data.get('add_tasks_to_crm', True)
    user_id = session['user_id']
    
    try:
        import requests
        import re
        import time
        from datetime import timedelta
        
        # Claude API Key - usa variabile d'ambiente per sicurezza
        CLAUDE_API_KEY = os.getenv('ANTHROPIC_API_KEY', 'YOUR_ANTHROPIC_API_KEY_HERE')
        CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"
        
        # Microsoft Graph API - Device Code Flow
        # Usa un'app pubblica Microsoft per autenticazione
        CLIENT_ID = "14d82eec-204b-4c2f-b7e8-296a70dab67e"  # Microsoft Graph Explorer App ID (pubblico)
        TENANT_ID = "common"
        SCOPE = "https://graph.microsoft.com/Mail.Read"
        
        # Step 1: Ottieni device code
        device_url = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/devicecode"
        device_data = {
            "client_id": CLIENT_ID,
            "scope": SCOPE
        }
        
        device_response = requests.post(device_url, data=device_data)
        if device_response.status_code != 200:
            return jsonify({
                'error': 'Errore autenticazione Microsoft',
                'details': 'Impossibile ottenere device code. Usa il metodo alternativo con email/password.'
            }), 500
        
        device_info = device_response.json()
        
        # Ritorna il device code al frontend per mostrare all'utente
        return jsonify({
            'success': False,
            'auth_required': True,
            'verification_uri': device_info['verification_uri'],
            'user_code': device_info['user_code'],
            'device_code': device_info['device_code'],
            'expires_in': device_info['expires_in'],
            'interval': device_info['interval'],
            'message': f"Apri {device_info['verification_uri']} e inserisci il codice: {device_info['user_code']}"
        })
        
    except Exception as e:
        return jsonify({'error': f'Errore inizializzazione: {str(e)}'}), 500


@app.route('/api/email/analyze-outlook-complete', methods=['POST'])
def analyze_outlook_emails_complete():
    """Completa l'analisi email dopo l'autenticazione"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    device_code = data.get('device_code')
    add_tasks_to_crm = data.get('add_tasks_to_crm', True)
    user_id = session['user_id']
    
    if not device_code:
        return jsonify({'error': 'Device code richiesto'}), 400
    
    try:
        import requests
        import re
        import time
        from datetime import timedelta
        
        # Claude API Key - usa variabile d'ambiente per sicurezza
        CLAUDE_API_KEY = os.getenv('ANTHROPIC_API_KEY', 'YOUR_ANTHROPIC_API_KEY_HERE')
        CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"
        
        # Microsoft Graph API
        CLIENT_ID = "14d82eec-204b-4c2f-b7e8-296a70dab67e"
        TENANT_ID = "common"
        
        # Step 2: Ottieni access token
        token_url = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"
        token_data = {
            "grant_type": "urn:ietf:params:oauth:grant-type:device_code",
            "client_id": CLIENT_ID,
            "device_code": device_code
        }
        
        # Prova a ottenere il token (l'utente deve aver autorizzato)
        token_response = requests.post(token_url, data=token_data)
        token_result = token_response.json()
        
        if "access_token" not in token_result:
            if token_result.get("error") == "authorization_pending":
                return jsonify({
                    'success': False,
                    'auth_pending': True,
                    'message': 'Attendi che autorizzi l\'app nel browser...'
                })
            else:
                return jsonify({
                    'error': 'Errore autenticazione',
                    'details': token_result.get('error_description', token_result.get('error'))
                }), 500
        
        access_token = token_result["access_token"]
        
        # Step 3: Leggi email da Microsoft Graph
        graph_url = "https://graph.microsoft.com/v1.0/me/messages"
        filter_date = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%dT%H:%M:%SZ')
        
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        params = {
            "$filter": f"receivedDateTime ge {filter_date}",
            "$orderby": "receivedDateTime desc",
            "$top": 50,
            "$select": "subject,from,receivedDateTime,bodyPreview,body"
        }
        
        graph_response = requests.get(graph_url, headers=headers, params=params)
        
        if graph_response.status_code != 200:
            return jsonify({
                'error': f'Errore lettura email: {graph_response.status_code}',
                'details': graph_response.text
            }), 500
        
        graph_data = graph_response.json()
        emails = []
        
        for msg in graph_data.get('value', []):
            body_content = msg.get('bodyPreview', '') or msg.get('body', {}).get('content', '')
            # Rimuovi HTML se presente
            body_content = re.sub(r'<[^>]+>', '', body_content)
            
            emails.append({
                'subject': msg.get('subject', '(Nessun oggetto)'),
                'from': msg.get('from', {}).get('emailAddress', {}).get('address', ''),
                'date': msg.get('receivedDateTime', ''),
                'body': body_content[:5000]
            })
        
        if not emails:
            return jsonify({
                'success': True,
                'summary': 'Nessuna email trovata negli ultimi 7 giorni',
                'tasks': []
            })
        
        # Analizza con Claude
        email_text = ""
        for i, email_data in enumerate(emails, 1):
            email_text += f"\n--- Email {i} ---\n"
            email_text += f"Da: {email_data['from']}\n"
            email_text += f"Oggetto: {email_data['subject']}\n"
            email_text += f"Data: {email_data['date']}\n"
            email_text += f"Contenuto: {email_data['body']}\n\n"
        
        prompt = f"""Analizza le seguenti email e fornisci:
1. Un riassunto generale di tutte le email
2. Una lista di attività da fare basate sulle email

Per ogni attività, indica:
- Tipo attività (send_offer, call, recall, other)
- Descrizione chiara dell'attività
- Data scadenza se presente o suggerita (formato YYYY-MM-DD o null)
- Priorità se evidente (high, medium, low)

Rispondi SOLO in formato JSON valido con questa struttura:
{{
    "summary": "Riassunto generale delle email",
    "tasks": [
        {{
            "task_type": "send_offer|call|recall|other",
            "description": "Descrizione dettagliata dell'attività",
            "due_date": "YYYY-MM-DD o null",
            "priority": "high|medium|low"
        }}
    ]
}}

Email da analizzare:
{email_text}
"""
        
        response = requests.post(
            CLAUDE_API_URL,
            headers={
                "x-api-key": CLAUDE_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json"
            },
            json={
                "model": "claude-3-5-sonnet-20241022",
                "max_tokens": 4000,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=120
        )
        
        if response.status_code != 200:
            return jsonify({'error': f'Errore API Claude: {response.status_code}'}), 500
        
        result = response.json()
        content = result.get('content', [])
        if not content or len(content) == 0:
            return jsonify({'error': 'Risposta Claude vuota'}), 500
        
        text_response = content[0].get('text', '')
        
        # Estrai JSON dalla risposta
        json_match = re.search(r'\{.*\}', text_response, re.DOTALL)
        if not json_match:
            return jsonify({
                'success': True,
                'summary': text_response,
                'tasks': []
            })
        
        try:
            analysis = json.loads(json_match.group())
        except json.JSONDecodeError:
            return jsonify({
                'success': True,
                'summary': text_response,
                'tasks': []
            })
        
        # Aggiungi attività al CRM se richiesto
        tasks_added = []
        if add_tasks_to_crm and analysis.get('tasks'):
            for task in analysis['tasks']:
                try:
                    due_date = None
                    if task.get('due_date'):
                        try:
                            due_date = datetime.strptime(task['due_date'], '%Y-%m-%d').date()
                        except:
                            pass
                    
                    new_task = OpportunityTask(
                        opportunity_id=None,  # Task globale
                        account_id=None,
                        task_type=task.get('task_type', 'other'),
                        description=task.get('description', ''),
                        assigned_to_id=user_id,
                        due_date=due_date
                    )
                    db.session.add(new_task)
                    db.session.commit()
                    
                    tasks_added.append({
                        **task,
                        'added_to_crm': True,
                        'task_id': new_task.id
                    })
                except Exception as e:
                    tasks_added.append({
                        **task,
                        'added_to_crm': False,
                        'error': str(e)
                    })
        
        return jsonify({
            'success': True,
            'summary': analysis.get('summary', ''),
            'tasks': tasks_added if tasks_added else analysis.get('tasks', []),
            'emails_analyzed': len(emails)
        })
        
    except Exception as e:
        return jsonify({'error': f'Errore durante l\'analisi: {str(e)}'}), 500


@app.route('/api/email/test', methods=['POST'])
@app.route('/api/email/test-send', methods=['POST'])
def test_email():
    """Invia email di test con sostituzione placeholder"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        data = request.json
        to_email = data.get('to_email')
        subject = data.get('subject', 'Test Email')
        html_content = data.get('html_content') or data.get('html_body') or '<p>Questa è un\'email di test</p>'
        
        if not to_email:
            return jsonify({'error': 'Email destinatario richiesta'}), 400
        
        # Ottieni nome normalizzato
        first_name, last_name, full_name = get_recipient_name(to_email)
        
        # Sostituisci placeholder nel contenuto
        html_content = html_content.replace("%nome%", full_name)
        html_content = html_content.replace("%nome_completo%", full_name)
        html_content = html_content.replace("%nome_destinatario%", full_name)
        html_content = html_content.replace("%first_name%", first_name)
        html_content = html_content.replace("%last_name%", last_name)
        
        # Aggiungi firma email
        from email_signature import add_signature_to_email
        
        html_with_signature, text_with_signature = add_signature_to_email(
            html_content=html_content,
            text_content=None
        )
        
        from mailgun_service import MailgunService
        mailgun = MailgunService()
        
        result = mailgun.send_email(
            to_email=to_email,
            subject=subject,
            html_content=html_with_signature,
            text_content=text_with_signature
        )
        
        if result.get('success'):
            return jsonify({'success': True, 'message': 'Email di test inviata con successo', 'message_id': result.get('message_id')})
        else:
            return jsonify({'error': result.get('error', 'Errore sconosciuto')}), 500
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/email/lists/<int:list_id>/subscribers', methods=['GET', 'DELETE'])
def manage_list_subscribers(list_id):
    """Gestisci iscritti a una lista"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    email_list = EmailList.query.get_or_404(list_id)
    
    if request.method == 'GET':
        subscribers = EmailListSubscription.query.filter_by(list_id=list_id).all()
        return jsonify({
            'subscribers': [{
                'id': s.id,
                'email': s.email,
                'first_name': s.first_name,
                'last_name': s.last_name,
                'contact_id': s.contact_id,
                'lead_id': s.lead_id,
                'status': s.status,
                'subscribed_at': s.subscribed_at.isoformat() if s.subscribed_at else None
            } for s in subscribers]
        })
    
    elif request.method == 'DELETE':
        data = request.json
        subscription_id = data.get('subscription_id')
        
        if subscription_id:
            subscription = EmailListSubscription.query.get(subscription_id)
            if subscription and subscription.list_id == list_id:
                subscription.status = 'unsubscribed'
                subscription.unsubscribed_at = datetime.utcnow()
                db.session.commit()
                return jsonify({'success': True})
        
        return jsonify({'error': 'Subscription non trovata'}), 404


@app.route('/unsubscribe', methods=['GET', 'POST'])
def unsubscribe_email():
    """Endpoint pubblico per disiscrizioni email"""
    email = request.args.get('email') or (request.form.get('email') if request.method == 'POST' else None)
    
    if not email:
        return render_template('unsubscribe.html', 
                             success=False, 
                             message='Indirizzo email non fornito')
    
    try:
        # Cerca tutte le subscription per questa email
        subscriptions = EmailListSubscription.query.filter_by(email=email, status='subscribed').all()
        
        if not subscriptions:
            return render_template('unsubscribe.html', 
                                 success=False, 
                                 message=f'L\'indirizzo email {email} non risulta iscritto a nessuna newsletter')
        
        # Disiscrivi da tutte le liste
        for subscription in subscriptions:
            subscription.status = 'unsubscribed'
            subscription.unsubscribed_at = datetime.utcnow()
        
        db.session.commit()
        
        return render_template('unsubscribe.html', 
                             success=True, 
                             message=f'Sei stato disiscritto con successo dalla newsletter. L\'indirizzo {email} non riceverà più comunicazioni da parte nostra.')
    
    except Exception as e:
        db.session.rollback()
        return render_template('unsubscribe.html', 
                             success=False, 
                             message=f'Si è verificato un errore durante la disiscrizione: {str(e)}')


@app.route('/api/email/contacts-available', methods=['GET'])
def get_available_contacts_for_lists():
    """Ottieni contatti e leads disponibili con email per aggiungerli alle liste"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Ottieni tutti i contatti con email
    contacts = Contact.query.filter(Contact.email.isnot(None), Contact.email != '').all()
    contacts_data = [{
        'id': c.id,
        'type': 'contact',
        'email': c.email,
        'first_name': c.first_name,
        'last_name': c.last_name,
        'name': f"{c.first_name} {c.last_name}".strip(),
        'company': c.account.name if c.account else None
    } for c in contacts]
    
    # Ottieni tutti i leads con email
    leads = Lead.query.filter(Lead.email.isnot(None), Lead.email != '').all()
    leads_data = [{
        'id': l.id,
        'type': 'lead',
        'email': l.email,
        'first_name': l.contact_first_name,
        'last_name': l.contact_last_name,
        'name': f"{l.contact_first_name or ''} {l.contact_last_name or ''}".strip() or l.company_name,
        'company': l.company_name
    } for l in leads]
    
    return jsonify({
        'contacts': contacts_data,
        'leads': leads_data,
        'all': contacts_data + leads_data
    })


@app.route('/api/email/lists/unique-recipients', methods=['POST'])
def get_unique_recipients():
    """Calcola destinatari unici da più liste (elimina duplicati)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    list_ids = data.get('list_ids', [])
    
    if not list_ids:
        return jsonify({'success': True, 'unique_count': 0, 'total_with_duplicates': 0})
    
    # Recupera tutte le subscription dalle liste selezionate
    subscriptions = EmailListSubscription.query.filter(
        EmailListSubscription.list_id.in_(list_ids),
        EmailListSubscription.status == 'subscribed'
    ).all()
    
    # Conta totali (con duplicati)
    total_with_duplicates = len(subscriptions)
    
    # Deduplica per email (case-insensitive)
    unique_emails = set()
    for sub in subscriptions:
        if sub.email:
            unique_emails.add(sub.email.lower().strip())
    
    unique_count = len(unique_emails)
    duplicates = total_with_duplicates - unique_count
    
    return jsonify({
        'success': True,
        'unique_count': unique_count,
        'total_with_duplicates': total_with_duplicates,
        'duplicates': duplicates,
        'list_count': len(list_ids)
    })


@app.route('/api/email/templates/init-defaults', methods=['POST'])
def init_default_email_templates():
    """Inizializza template email predefiniti"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    try:
        from email_template_loader import get_default_templates, create_sample_content
        
        templates = get_default_templates()
        sample_content = create_sample_content()
        
        created = []
        for key, template_data in templates.items():
            # Verifica se esiste già
            existing = EmailTemplate.query.filter_by(name=template_data['name']).first()
            if existing:
                continue
            
            # Sostituisci placeholder content
            html_content = template_data['html_content']
            if key in sample_content:
                html_content = html_content.replace('{{content}}', sample_content[key])
            
            template = EmailTemplate(
                name=template_data['name'],
                subject=template_data['subject'],
                html_content=html_content,
                category=template_data['category'],
                owner_id=user_id
            )
            db.session.add(template)
            created.append(template_data['name'])
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Creati {len(created)} template predefiniti',
            'created': created
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ==================== API TICKET ====================

def generate_ticket_number():
    """Genera numero ticket univoco: TKT-YYYY-NNN"""
    from datetime import datetime
    year = datetime.now().year
    last_ticket = Ticket.query.filter(Ticket.ticket_number.like(f'TKT-{year}-%')).order_by(Ticket.ticket_number.desc()).first()
    if last_ticket:
        try:
            last_num = int(last_ticket.ticket_number.split('-')[-1])
            new_num = last_num + 1
        except:
            new_num = 1
    else:
        new_num = 1
    return f'TKT-{year}-{new_num:03d}'


@app.route('/api/tickets', methods=['GET', 'POST'])
def tickets():
    """Gestisci ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    # Lista operatori autorizzati
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    if request.method == 'GET':
        # Filtri
        status_filter = request.args.get('status')
        tab_filter = request.args.get('tab', 'all')  # all, my, open, closed
        
        query = Ticket.query
        
        # Clienti vedono solo i loro ticket
        if not is_operatore:
            query = query.filter_by(customer_id=user_id)
        # Operatori vedono tutti, ma possono filtrare
        elif tab_filter == 'my':
            query = query.filter_by(owner_id=user_id)
        elif tab_filter == 'open':
            query = query.filter(Ticket.status.in_(['Aperto', 'In Lavorazione', 'In Attesa Cliente']))
        elif tab_filter == 'closed':
            query = query.filter_by(status='Chiuso')
        
        if status_filter:
            query = query.filter_by(status=status_filter)
        
        tickets_list = query.order_by(Ticket.created_at.desc()).all()
        
        return jsonify({
            'tickets': [{
                'id': t.id,
                'ticket_number': t.ticket_number,
                'title': t.title,
                'description': t.description,
                'status': t.status,
                'priority': t.priority,
                'category': t.category,
                'product_category': t.product_category,
                'customer_id': t.customer_id,
                'customer_name': f"{t.customer.first_name} {t.customer.last_name}" if t.customer else 'Unknown',
                'owner_id': t.owner_id,
                'owner_name': f"{t.owner.first_name} {t.owner.last_name}" if t.owner else None,
                'is_my_ticket': t.customer_id == user_id,
                'is_my_ownership': t.owner_id == user_id,
                'created_at': t.created_at.isoformat() if t.created_at else None,
                'updated_at': t.updated_at.isoformat() if t.updated_at else None,
                'comments_count': len(t.comments)
            } for t in tickets_list],
            'is_operatore': is_operatore
        })
    
    elif request.method == 'POST':
        try:
            data = request.json
            if not data:
                return jsonify({'error': 'Dati mancanti'}), 400
            
            ticket_number = generate_ticket_number()
            
            ticket = Ticket(
                ticket_number=ticket_number,
                title=data.get('title', ''),
                description=data.get('description', ''),
                status='Aperto',
                priority=data.get('priority', 'Media'),
                category=data.get('category'),
                product_category=data.get('product_category'),
                customer_id=user_id
            )
            
            db.session.add(ticket)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'ticket': {
                    'id': ticket.id,
                    'ticket_number': ticket.ticket_number,
                    'title': ticket.title
                }
            })
        except Exception as e:
            import traceback
            traceback.print_exc()
            db.session.rollback()
            return jsonify({'error': f'Errore creazione ticket: {str(e)}'}), 500


@app.route('/api/tickets/<int:ticket_id>', methods=['GET', 'PUT'])
def get_or_update_ticket(ticket_id):
    """Ottieni o aggiorna ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    
    # Verifica permessi: clienti vedono solo i loro ticket
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        # Filtra commenti: clienti vedono solo quelli pubblici
        all_comments = ticket.comments
        if not is_operatore:
            all_comments = [c for c in all_comments if not c.is_internal]
        
        return jsonify({
            'ticket': {
                'id': ticket.id,
                'ticket_number': ticket.ticket_number,
                'title': ticket.title,
                'description': ticket.description,
                'status': ticket.status,
                'priority': ticket.priority,
                'category': ticket.category,
                'product_category': ticket.product_category,
                'customer_id': ticket.customer_id,
                'customer_name': f"{ticket.customer.first_name} {ticket.customer.last_name}" if ticket.customer else 'Unknown',
                'customer_email': ticket.customer.email if ticket.customer else None,
                'owner_id': ticket.owner_id,
                'owner_name': f"{ticket.owner.first_name} {ticket.owner.last_name}" if ticket.owner else None,
                'created_at': ticket.created_at.isoformat() if ticket.created_at else None,
                'updated_at': ticket.updated_at.isoformat() if ticket.updated_at else None,
                'comments': [{
                    'id': c.id,
                    'message': c.message,
                    'is_internal': c.is_internal,
                    'user_id': c.user_id,
                    'user_name': f"{c.user.first_name} {c.user.last_name}" if c.user else 'Unknown',
                    'created_at': c.created_at.isoformat() if c.created_at else None
                } for c in ticket.comments if not c.is_internal or is_operatore]
            },
            'is_operatore': is_operatore
        })
    
    elif request.method == 'PUT':
        data = request.json
        
        # Solo operatori possono modificare status, priority, owner
        if is_operatore:
            if 'status' in data:
                ticket.status = data['status']
            if 'priority' in data:
                ticket.priority = data['priority']
            if 'owner_id' in data:
                ticket.owner_id = data['owner_id']
            if 'category' in data:
                ticket.category = data['category']
            if 'product_category' in data:
                ticket.product_category = data['product_category']
        
        # Tutti possono modificare title e description
        if 'title' in data:
            ticket.title = data['title']
        if 'description' in data:
            ticket.description = data['description']
        
        db.session.commit()
        
        return jsonify({'success': True, 'ticket_id': ticket.id})


@app.route('/api/tickets/<int:ticket_id>/take-ownership', methods=['POST'])
def take_ticket_ownership(ticket_id):
    """Prendi ownership di un ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    if not is_operatore:
        return jsonify({'error': 'Solo operatori possono prendere ownership'}), 403
    
    ticket = Ticket.query.get_or_404(ticket_id)
    old_owner_id = ticket.owner_id
    ticket.owner_id = user_id
    if ticket.status == 'Aperto':
        ticket.status = 'In Lavorazione'
    
    # Crea notifica per il nuovo owner (se diverso dal vecchio)
    if old_owner_id != user_id:
        create_notification(
            user_id=user_id,
            title=f'Ticket {ticket.ticket_number} assegnato',
            message=f'Ti è stato assegnato il ticket "{ticket.title}"',
            type='info',
            related_ticket_id=ticket_id
        )
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': ticket.owner_id,
        'owner_name': f"{ticket.owner.first_name} {ticket.owner.last_name}" if ticket.owner else 'Unknown'
    })


@app.route('/api/tickets/<int:ticket_id>/assign-ownership', methods=['POST'])
def assign_ticket_ownership(ticket_id):
    """Assegna un ticket a un operatore"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    if not is_operatore:
        return jsonify({'error': 'Solo operatori possono assegnare ticket'}), 403
    
    data = request.get_json()
    new_owner_id = data.get('owner_id')
    
    if not new_owner_id:
        return jsonify({'error': 'owner_id richiesto'}), 400
    
    ticket = Ticket.query.get_or_404(ticket_id)
    old_owner_id = ticket.owner_id
    ticket.owner_id = new_owner_id
    if ticket.status == 'Aperto':
        ticket.status = 'In Lavorazione'
    
    # Crea notifica per il nuovo owner (se diverso dal vecchio)
    if old_owner_id != new_owner_id:
        new_owner = User.query.get(new_owner_id)
        create_notification(
            user_id=new_owner_id,
            title=f'Ticket {ticket.ticket_number} assegnato',
            message=f'Ti è stato assegnato il ticket "{ticket.title}" da {user.first_name} {user.last_name}',
            type='info',
            related_ticket_id=ticket_id
        )
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': ticket.owner_id,
        'owner_name': f"{ticket.owner.first_name} {ticket.owner.last_name}" if ticket.owner else 'Unknown'
    })


@app.route('/api/users/operatori', methods=['GET'])
def get_operatori():
    """Ottieni lista operatori"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    operatori_usernames = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo']
    
    # Cerca utenti con ruolo operatore o admin, o username che contiene i nomi degli operatori
    operatori = User.query.filter(
        (User.role == 'admin') | 
        (User.role == 'operatore') |
        db.or_(*[User.username.ilike(f'%{op}%') for op in operatori_usernames])
    ).all()
    
    return jsonify({
        'operatori': [{
            'id': u.id,
            'username': u.username,
            'first_name': u.first_name,
            'last_name': u.last_name,
            'full_name': f"{u.first_name} {u.last_name}"
        } for u in operatori]
    })


@app.route('/api/tickets/<int:ticket_id>/comments', methods=['GET', 'POST'])
def ticket_comments(ticket_id):
    """Gestisci commenti su ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    
    # Verifica permessi
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        comments = ticket.comments
        # Clienti non vedono commenti interni
        if not is_operatore:
            comments = [c for c in comments if not c.is_internal]
        
        return jsonify({
            'comments': [{
                'id': c.id,
                'message': c.message,
                'is_internal': c.is_internal,
                'user_id': c.user_id,
                'user_name': f"{c.user.first_name} {c.user.last_name}" if c.user else 'Unknown',
                'created_at': c.created_at.isoformat() if c.created_at else None
            } for c in comments]
        })
    
    elif request.method == 'POST':
        data = request.json
        comment = TicketComment(
            ticket_id=ticket_id,
            user_id=user_id,
            message=data['message'],
            is_internal=data.get('is_internal', False)
        )
        
        db.session.add(comment)
        
        # Aggiorna status se cliente risponde
        if not is_operatore and ticket.status == 'In Attesa Cliente':
            ticket.status = 'In Lavorazione'
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'comment': {
                'id': comment.id,
                'message': comment.message,
                'user_name': f"{comment.user.first_name} {comment.user.last_name}" if comment.user else 'Unknown',
                'created_at': comment.created_at.isoformat() if comment.created_at else None
            }
        })


# ==================== GESTIONE ALLEGATI TICKET ====================

def get_storage_usage_mb():
    """Calcola spazio utilizzato in MB"""
    total_size = 0
    for root, dirs, files in os.walk(app.config['UPLOAD_FOLDER']):
        for file in files:
            filepath = os.path.join(root, file)
            if os.path.exists(filepath):
                total_size += os.path.getsize(filepath)
    return total_size / (1024 * 1024)  # Converti in MB


def get_client_storage_usage_mb(user_id):
    """Calcola spazio utilizzato da un cliente in MB"""
    attachments = TicketAttachment.query.filter_by(user_id=user_id).all()
    total_size = sum(att.file_size or 0 for att in attachments)
    return total_size / (1024 * 1024)  # Converti in MB


def compress_image(input_path, output_path, quality=85, max_size=(1920, 1080)):
    """Comprimi immagine usando Pillow"""
    try:
        from PIL import Image
        img = Image.open(input_path)
        
        # Converti RGBA in RGB se necessario
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        
        # Ridimensiona se troppo grande
        if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Salva compressa
        img.save(output_path, 'JPEG', quality=quality, optimize=True)
        return True
    except Exception as e:
        print(f"Errore compressione immagine: {e}")
        return False


def compress_video(input_path, output_path):
    """Comprimi video usando ffmpeg"""
    try:
        import subprocess
        # Comprimi video: riduci risoluzione e bitrate
        cmd = [
            'ffmpeg', '-i', input_path,
            '-vf', 'scale=1280:-1',  # Max width 1280px
            '-c:v', 'libx264',
            '-preset', 'medium',
            '-crf', '28',  # Qualità compressa
            '-c:a', 'aac',
            '-b:a', '128k',
            '-y',  # Sovrascrivi
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        return result.returncode == 0
    except Exception as e:
        print(f"Errore compressione video: {e}")
        return False


def cleanup_old_files():
    """Elimina file vecchi se lo spazio è limitato"""
    total_usage_mb = get_storage_usage_mb()
    total_storage_mb = app.config['TOTAL_STORAGE_GB'] * 1024
    
    # Se uso più dell'80% dello spazio, elimina file più vecchi
    if total_usage_mb > (total_storage_mb * 0.8):
        # Ordina allegati per data (più vecchi prima)
        old_attachments = TicketAttachment.query.order_by(TicketAttachment.created_at.asc()).limit(50).all()
        
        deleted_count = 0
        for att in old_attachments:
            if os.path.exists(att.file_path):
                try:
                    os.remove(att.file_path)
                    db.session.delete(att)
                    deleted_count += 1
                except Exception as e:
                    print(f"Errore eliminazione file {att.file_path}: {e}")
        
        db.session.commit()
        return deleted_count
    
    return 0


@app.route('/api/tickets/<int:ticket_id>/attachments', methods=['GET', 'POST', 'DELETE'])
def ticket_attachments(ticket_id):
    """Gestisci allegati ai ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    
    # Verifica permessi
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        attachments = ticket.attachments
        return jsonify({
            'attachments': [{
                'id': att.id,
                'filename': att.filename if hasattr(att, 'filename') else (att.original_filename if hasattr(att, 'original_filename') else 'file'),
                'file_type': att.file_type,
                'file_size': att.file_size,
                'mime_type': getattr(att, 'mime_type', None),
                'is_compressed': att.is_compressed,
                'created_at': att.created_at.isoformat() if att.created_at else None,
                'user_name': f"{att.user.first_name} {att.user.last_name}" if att.user else 'Unknown'
            } for att in attachments]
        })
    
    elif request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({'error': 'Nessun file caricato'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nessun file selezionato'}), 400
        
        # Verifica quota cliente
        if user_role == 'cliente':
            client_usage_mb = get_client_storage_usage_mb(user_id)
            if client_usage_mb >= app.config['CLIENT_QUOTA_MB']:
                return jsonify({'error': f'Quota esaurita. Spazio utilizzato: {client_usage_mb:.2f}MB / {app.config["CLIENT_QUOTA_MB"]}MB'}), 400
        
        # Verifica spazio totale
        cleanup_old_files()  # Pulisci file vecchi se necessario
        
        # Salva file originale
        original_filename = secure_filename(file.filename)
        file_ext = os.path.splitext(original_filename)[1].lower()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{ticket_id}_{timestamp}_{original_filename}"
        original_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        
        file.save(original_path)
        original_size = os.path.getsize(original_path)
        
        # Determina tipo file
        is_image = file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']
        is_video = file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.webm']
        
        if not (is_image or is_video):
            os.remove(original_path)
            return jsonify({'error': 'Formato file non supportato. Usa immagini (JPG, PNG) o video (MP4)'}), 400
        
        file_type = 'image' if is_image else 'video'
        compressed_path = original_path
        is_compressed = False
        
        # Comprimi se necessario
        if is_image and original_size > 2 * 1024 * 1024:  # > 2MB
            compressed_filename = f"compressed_{safe_filename}"
            compressed_path = os.path.join(app.config['UPLOAD_FOLDER'], compressed_filename)
            if compress_image(original_path, compressed_path):
                os.remove(original_path)
                original_path = compressed_path
                is_compressed = True
        elif is_video and original_size > 10 * 1024 * 1024:  # > 10MB
            compressed_filename = f"compressed_{safe_filename}"
            compressed_path = os.path.join(app.config['UPLOAD_FOLDER'], compressed_filename)
            if compress_video(original_path, compressed_path):
                os.remove(original_path)
                original_path = compressed_path
                is_compressed = True
        
        final_size = os.path.getsize(original_path)
        
        # Crea record allegato
        attachment = TicketAttachment(
            ticket_id=ticket_id,
            user_id=user_id,
            filename=safe_filename,
            original_filename=original_filename,
            file_path=original_path,
            file_type=file_type,
            file_size=final_size,
            mime_type=file.content_type,
            is_compressed=is_compressed
        )
        
        db.session.add(attachment)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'attachment': {
                'id': attachment.id,
                'filename': attachment.original_filename,
                'file_size': attachment.file_size,
                'is_compressed': attachment.is_compressed
            }
        })
    
    elif request.method == 'DELETE':
        attachment_id = request.json.get('attachment_id')
        if not attachment_id:
            return jsonify({'error': 'ID allegato mancante'}), 400
        
        attachment = TicketAttachment.query.get_or_404(attachment_id)
        
        # Solo il proprietario o operatore può eliminare
        if not is_operatore and attachment.user_id != user_id:
            return jsonify({'error': 'Non autorizzato'}), 403
        
        # Elimina file
        if os.path.exists(attachment.file_path):
            try:
                os.remove(attachment.file_path)
            except Exception as e:
                print(f"Errore eliminazione file: {e}")
        
        db.session.delete(attachment)
        db.session.commit()
        
        return jsonify({'success': True})


@app.route('/api/tickets/<int:ticket_id>/attachments/<int:attachment_id>/download')
def download_attachment(ticket_id, attachment_id):
    """Scarica allegato"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    attachment = TicketAttachment.query.get_or_404(attachment_id)
    
    # Verifica permessi
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if not os.path.exists(attachment.file_path):
        return jsonify({'error': 'File non trovato'}), 404
    
    return send_file(attachment.file_path, as_attachment=True, download_name=attachment.original_filename)


@app.route('/api/storage/usage')
def storage_usage():
    """Ottieni informazioni spazio utilizzato"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    total_usage_mb = get_storage_usage_mb()
    total_storage_mb = app.config['TOTAL_STORAGE_GB'] * 1024
    
    client_usage_mb = 0
    client_quota_mb = 0
    if user_role == 'cliente':
        client_usage_mb = get_client_storage_usage_mb(user_id)
        client_quota_mb = app.config['CLIENT_QUOTA_MB']
    
    return jsonify({
        'total_usage_mb': round(total_usage_mb, 2),
        'total_storage_mb': total_storage_mb,
        'total_usage_percent': round((total_usage_mb / total_storage_mb) * 100, 2),
        'client_usage_mb': round(client_usage_mb, 2),
        'client_quota_mb': client_quota_mb,
        'client_usage_percent': round((client_usage_mb / client_quota_mb * 100) if client_quota_mb > 0 else 0, 2)
    })


def create_notification(user_id, title, message, type='info', related_ticket_id=None, related_opportunity_id=None):
    """Crea una nuova notifica"""
    try:
        notification = Notification(
            user_id=user_id,
            title=title,
            message=message,
            type=type,
            related_ticket_id=related_ticket_id,
            related_opportunity_id=related_opportunity_id
        )
        db.session.add(notification)
        db.session.commit()
        return notification
    except Exception as e:
        db.session.rollback()
        print(f"Errore creazione notifica: {e}")
        return None


@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    """Ottieni notifiche dell'utente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    unread_only = request.args.get('unread_only', 'false').lower() == 'true'
    
    query = Notification.query.filter_by(user_id=user_id)
    if unread_only:
        query = query.filter_by(is_read=False)
    
    notifications = query.order_by(Notification.created_at.desc()).limit(50).all()
    
    return jsonify({
        'notifications': [{
            'id': n.id,
            'title': n.title,
            'message': n.message,
            'type': n.type,
            'is_read': n.is_read,
            'related_ticket_id': n.related_ticket_id,
            'related_opportunity_id': n.related_opportunity_id,
            'created_at': n.created_at.isoformat() if n.created_at else None
        } for n in notifications]
    })


@app.route('/api/notifications/<int:notification_id>/read', methods=['POST'])
def mark_notification_read(notification_id):
    """Marca una notifica come letta"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    notification = Notification.query.filter_by(id=notification_id, user_id=user_id).first()
    
    if not notification:
        return jsonify({'error': 'Notifica non trovata'}), 404
    
    notification.is_read = True
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/api/notifications/read-all', methods=['POST'])
def mark_all_notifications_read():
    """Marca tutte le notifiche come lette"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    Notification.query.filter_by(user_id=user_id, is_read=False).update({'is_read': True})
    db.session.commit()
    
    return jsonify({'success': True})


def init_db():
    """Inizializza database"""
    with app.app_context():
        db.create_all()
        
        # Crea admin
        if not User.query.filter_by(username='admin').first():
            admin = User(
                username='admin',
                password=generate_password_hash('admin123'),
                first_name='Admin',
                last_name='System',
                role='admin'
            )
            db.session.add(admin)
        
        # Crea commerciale demo
        if not User.query.filter_by(username='demo').first():
            demo = User(
                username='demo',
                password=generate_password_hash('demo123'),
                first_name='Demo',
                last_name='User',
                role='commerciale'
            )
            db.session.add(demo)
        
        db.session.commit()
        print("✓ Database inizializzato")


if __name__ == '__main__':
    init_db()
    app.run(host='0.0.0.0', port=8000, debug=False)

