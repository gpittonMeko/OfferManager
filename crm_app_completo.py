#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MekoCRM - Sistema CRM completo stile Salesforce
"""
import sys
import os
from pathlib import Path

# Carica variabili d'ambiente da file .env se esiste
def load_env_file():
    """Carica variabili d'ambiente da file .env"""
    env_file = Path('.env')
    if env_file.exists():
        try:
            # Prova con python-dotenv se disponibile
            try:
                from dotenv import load_dotenv
                load_dotenv()
                return
            except ImportError:
                pass
            
            # Fallback: leggi manualmente il file .env
            with open(env_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        key, value = line.split('=', 1)
                        os.environ[key.strip()] = value.strip()
        except Exception as e:
            print(f"Warning: Errore caricamento .env: {e}")

load_env_file()

if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date, timedelta
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import csv
import io
import os
from sqlalchemy import func, text
import json
import os
import shutil
from pathlib import Path
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

app = Flask(__name__)
app.config['SECRET_KEY'] = 'mekocrm-2025-secret-key'

# Database originale: Flask cerca automaticamente in instance/ quando si usa sqlite:///mekocrm.db
# Questo è il percorso CORRETTO che era già configurato
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///mekocrm.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)


# ==================== MODELLI DATABASE SALESFORCE-LIKE ====================

class TimestampMixin:
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class User(TimestampMixin, db.Model):
    __tablename__ = "users"
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    first_name = db.Column(db.String(100))
    last_name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    role = db.Column(db.String(20), default="commerciale")
    
    # Configurazione server per accesso file (IP server condiviso da tutti i PC della subnet)
    local_pc_ip = db.Column(db.String(50), default="192.168.10.240")  # IP del server (default: 192.168.10.240)
    local_pc_share = db.Column(db.String(100), default="commerciale")  # Nome condivisione (default: "commerciale")
    local_pc_base_path = db.Column(db.String(500))  # Percorso base sul server (es. "OFFERTE - K")
    local_pc_username = db.Column(db.String(100))  # Username opzionale per SMB
    local_pc_password = db.Column(db.String(200))  # Password opzionale per SMB (criptata)
    
    leads = db.relationship("Lead", backref="owner", lazy=True)
    accounts = db.relationship("Account", backref="owner", lazy=True)
    contacts = db.relationship("Contact", backref="owner", lazy=True)
    opportunities = db.relationship("Opportunity", backref="owner", lazy=True)
    offers = db.relationship("OfferDocument", backref="owner", lazy=True)


class Lead(TimestampMixin, db.Model):
    __tablename__ = "leads"
    id = db.Column(db.Integer, primary_key=True)
    company_name = db.Column(db.String(200), nullable=False)
    contact_first_name = db.Column(db.String(100))
    contact_last_name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    vat_number = db.Column(db.String(20))
    address = db.Column(db.Text)
    industry = db.Column(db.String(100))
    interest = db.Column(db.String(50))  # MIR, UR, Unitree, ROEQ, Servizi
    source = db.Column(db.String(50))
    rating = db.Column(db.String(20), default="Warm")
    status = db.Column(db.String(20), default="New")  # New, Contacted, Qualified, Converted, Lost
    notes = db.Column(db.Text)
    last_contact_at = db.Column(db.DateTime)
    
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    converted_account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"))
    converted_contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"))
    converted_opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"))


class Account(TimestampMixin, db.Model):
    __tablename__ = "accounts"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), unique=True, nullable=False)
    industry = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    website = db.Column(db.String(200))
    billing_address = db.Column(db.Text)
    shipping_address = db.Column(db.Text)
    description = db.Column(db.Text)
    image_url = db.Column(db.String(500))  # URL immagine account (fittizia)
    
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    contacts = db.relationship("Contact", backref="account", lazy=True)
    opportunities = db.relationship("Opportunity", backref="account", lazy=True)


class Contact(TimestampMixin, db.Model):
    __tablename__ = "contacts"
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(100), nullable=False)
    last_name = db.Column(db.String(100), nullable=False)
    title = db.Column(db.String(100))
    email = db.Column(db.String(100))
    phone = db.Column(db.String(50))
    department = db.Column(db.String(100))
    description = db.Column(db.Text)
    
    account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"))
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)


class Opportunity(TimestampMixin, db.Model):
    __tablename__ = "opportunities"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    stage = db.Column(db.String(40), default="Qualification")
    amount = db.Column(db.Float, default=0.0)
    probability = db.Column(db.Integer, default=10)
    close_date = db.Column(db.Date)
    opportunity_type = db.Column(db.String(50))
    description = db.Column(db.Text)
    supplier_category = db.Column(db.String(50))  # Universal Robots, Unitree, MiR, Altri Fornitori, Servizi
    folder_path = db.Column(db.String(400))  # Cartella per documenti dell'opportunità
    heat_level = db.Column(db.String(30), nullable=True)  # calda, tiepida, fredda_speranza, fredda_gelo (nullable perché potrebbe non esistere nel DB)
    lost_reason = db.Column(db.Text)  # Motivazione quando l'opportunità è chiusa come persa
    
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"))
    contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"))
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"))
    
    line_items = db.relationship("OpportunityLineItem", backref="opportunity", lazy=True, cascade="all, delete-orphan")
    offers = db.relationship("OfferDocument", backref="opportunity", lazy=True)
    chatter_posts = db.relationship("OpportunityChatter", backref="opportunity", lazy=True, cascade="all, delete-orphan", order_by="OpportunityChatter.created_at.desc()")
    tasks = db.relationship("OpportunityTask", backref="opportunity", lazy=True, cascade="all, delete-orphan", order_by="OpportunityTask.created_at.desc()")


class OpportunityChatter(TimestampMixin, db.Model):
    __tablename__ = "opportunity_chatter"
    id = db.Column(db.Integer, primary_key=True)
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    message = db.Column(db.Text, nullable=False)
    
    user = db.relationship("User", backref="chatter_posts")


class OpportunityTask(TimestampMixin, db.Model):
    __tablename__ = "opportunity_tasks"
    id = db.Column(db.Integer, primary_key=True)
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=True)  # NULL = task globale
    account_id = db.Column(db.Integer, db.ForeignKey("accounts.id"), nullable=True)  # Collegamento diretto ad account
    task_type = db.Column(db.String(50), nullable=False)  # 'send_offer', 'call', 'recall', 'other'
    description = db.Column(db.Text)
    assigned_to_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)  # Chi deve farlo
    due_date = db.Column(db.Date)


class AssistantLog(TimestampMixin, db.Model):
    """Log delle richieste all'assistente AI"""
    __tablename__ = "assistant_logs"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    user_message = db.Column(db.Text, nullable=False)
    ai_reply = db.Column(db.Text)
    ai_response_raw = db.Column(db.Text)  # Risposta completa JSON dall'AI
    actions_executed = db.Column(db.Text)  # JSON delle azioni eseguite
    errors = db.Column(db.Text)  # JSON degli errori
    tokens_used = db.Column(db.Integer)  # Token utilizzati (se disponibile)
    model_used = db.Column(db.String(50))  # Modello AI usato
    
    user = db.relationship("User", backref="assistant_logs")


class OpportunityLineItem(TimestampMixin, db.Model):
    __tablename__ = "opportunity_line_items"
    id = db.Column(db.Integer, primary_key=True)
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=False)
    product_id = db.Column(db.Integer, db.ForeignKey("products.id"), nullable=False)
    quantity = db.Column(db.Integer, default=1)
    unit_price = db.Column(db.Float, nullable=False)
    discount = db.Column(db.Float, default=0.0)
    total_price = db.Column(db.Float, nullable=False)
    
    product = db.relationship("Product", backref="line_items")


class Product(TimestampMixin, db.Model):
    __tablename__ = "products"
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(80), unique=True, nullable=False)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    category = db.Column(db.String(80))  # MIR, UR, Unitree, ROEQ, Servizi
    family = db.Column(db.String(80))
    is_active = db.Column(db.Boolean, default=True)
    
    prices = db.relationship("PriceBookEntry", backref="product", lazy=True, cascade="all, delete-orphan")


class PriceBookEntry(TimestampMixin, db.Model):
    __tablename__ = "pricebook_entries"
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey("products.id"), nullable=False)
    price_level = db.Column(db.String(80), nullable=False)  # bronze, silver, gold, ur_acquisto, ur_si, mir_distributor, etc.
    currency = db.Column(db.String(8), default="EUR")
    amount = db.Column(db.Float, nullable=False)
    source = db.Column(db.String(80), default="official")
    effective_date = db.Column(db.Date, default=date.today)
    is_override = db.Column(db.Boolean, default=False)
    is_active = db.Column(db.Boolean, default=True)


class OfferDocument(TimestampMixin, db.Model):
    __tablename__ = "offers"
    id = db.Column(db.Integer, primary_key=True)
    number = db.Column(db.String(32), unique=True, nullable=False)
    status = db.Column(db.String(20), default="draft")  # draft, sent, accepted, rejected, expired
    amount = db.Column(db.Float, default=0.0)  # Valore offerta iniziale
    closed_amount = db.Column(db.Float, default=0.0)  # Valore a cui è stata chiusa l'offerta
    margin = db.Column(db.Float, default=0.0)  # Margine percentuale quando chiusa
    description = db.Column(db.Text)
    docx_path = db.Column(db.String(400))
    pdf_path = db.Column(db.String(400))
    folder_path = db.Column(db.String(400))
    generated_from_nlp = db.Column(db.Boolean, default=False)
    
    opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"))
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"))
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"))


# ==================== MODELLI EMAIL MARKETING ====================

class EmailList(TimestampMixin, db.Model):
    """Lista email per campagne marketing"""
    __tablename__ = "email_lists"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    is_active = db.Column(db.Boolean, default=True)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    
    subscriptions = db.relationship("EmailListSubscription", backref="email_list", lazy=True, cascade="all, delete-orphan")
    campaigns = db.relationship("EmailCampaign", backref="email_list", lazy=True)


class EmailListSubscription(TimestampMixin, db.Model):
    """Iscrizioni a liste email"""
    __tablename__ = "email_list_subscriptions"
    id = db.Column(db.Integer, primary_key=True)
    list_id = db.Column(db.Integer, db.ForeignKey("email_lists.id"), nullable=False)
    contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"), nullable=True)
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"), nullable=True)
    email = db.Column(db.String(200), nullable=False)
    first_name = db.Column(db.String(100))
    last_name = db.Column(db.String(100))
    status = db.Column(db.String(20), default="subscribed")  # subscribed, unsubscribed, bounced
    subscribed_at = db.Column(db.DateTime, default=datetime.utcnow)
    unsubscribed_at = db.Column(db.DateTime)


class EmailTemplate(TimestampMixin, db.Model):
    """Template email HTML"""
    __tablename__ = "email_templates"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    subject = db.Column(db.String(500), nullable=False)
    html_content = db.Column(db.Text, nullable=False)
    text_content = db.Column(db.Text)
    category = db.Column(db.String(100))  # newsletter, promotional, transactional, etc.
    is_active = db.Column(db.Boolean, default=True)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    
    campaigns = db.relationship("EmailCampaign", backref="template", lazy=True)


class EmailCampaign(TimestampMixin, db.Model):
    """Campagna email"""
    __tablename__ = "email_campaigns"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    subject = db.Column(db.String(500), nullable=False)
    html_content = db.Column(db.Text, nullable=False)
    text_content = db.Column(db.Text)
    status = db.Column(db.String(20), default="draft")  # draft, scheduled, sending, sent, paused, cancelled
    list_id = db.Column(db.Integer, db.ForeignKey("email_lists.id"), nullable=False)
    template_id = db.Column(db.Integer, db.ForeignKey("email_templates.id"), nullable=True)
    scheduled_at = db.Column(db.DateTime)
    sent_at = db.Column(db.DateTime)
    total_recipients = db.Column(db.Integer, default=0)
    total_sent = db.Column(db.Integer, default=0)
    total_delivered = db.Column(db.Integer, default=0)
    total_opened = db.Column(db.Integer, default=0)
    total_clicked = db.Column(db.Integer, default=0)
    total_bounced = db.Column(db.Integer, default=0)
    total_unsubscribed = db.Column(db.Integer, default=0)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    
    sends = db.relationship("EmailCampaignSend", backref="campaign", lazy=True, cascade="all, delete-orphan")


class EmailCampaignSend(TimestampMixin, db.Model):
    """Singola email inviata in una campagna"""
    __tablename__ = "email_campaign_sends"
    id = db.Column(db.Integer, primary_key=True)
    campaign_id = db.Column(db.Integer, db.ForeignKey("email_campaigns.id"), nullable=False)
    email = db.Column(db.String(200), nullable=False)
    contact_id = db.Column(db.Integer, db.ForeignKey("contacts.id"), nullable=True)
    lead_id = db.Column(db.Integer, db.ForeignKey("leads.id"), nullable=True)


class Ticket(TimestampMixin, db.Model):
    """Ticket di supporto/richiesta clienti"""
    __tablename__ = "tickets"
    id = db.Column(db.Integer, primary_key=True)
    ticket_number = db.Column(db.String(20), unique=True, nullable=False)  # TKT-2025-001
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(20), default="Aperto")  # Aperto, In Lavorazione, In Attesa Cliente, Chiuso
    priority = db.Column(db.String(20), default="Media")  # Bassa, Media, Alta, Urgente
    category = db.Column(db.String(50))  # Supporto, Richiesta, Bug, Altro
    product_category = db.Column(db.String(100))  # MiR, Universal Robots, Unitree, ecc.
    
    customer_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)  # Cliente che ha creato il ticket
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)  # Operatore che ha preso ownership
    
    customer = db.relationship("User", foreign_keys=[customer_id], backref="tickets_as_customer")
    owner = db.relationship("User", foreign_keys=[owner_id], backref="tickets_as_owner")
    comments = db.relationship("TicketComment", backref="ticket", lazy=True, cascade="all, delete-orphan", order_by="TicketComment.created_at.asc()")


class TicketComment(TimestampMixin, db.Model):
    """Commenti su ticket"""
    __tablename__ = "ticket_comments"
    id = db.Column(db.Integer, primary_key=True)
    ticket_id = db.Column(db.Integer, db.ForeignKey("tickets.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    message = db.Column(db.Text, nullable=False)
    is_internal = db.Column(db.Boolean, default=False)  # True = solo operatori vedono
    
    user = db.relationship("User", backref="ticket_comments")


class TicketAttachment(TimestampMixin, db.Model):
    """Allegati ai ticket (immagini, video)"""
    __tablename__ = "ticket_attachments"
    id = db.Column(db.Integer, primary_key=True)
    ticket_id = db.Column(db.Integer, db.ForeignKey("tickets.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_type = db.Column(db.String(50))  # image, video, pdf
    file_size = db.Column(db.Integer)  # in bytes
    mime_type = db.Column(db.String(100))
    is_compressed = db.Column(db.Boolean, default=False)
    
    user = db.relationship("User", backref="ticket_attachments")
    ticket = db.relationship("Ticket", backref="attachments")


class Notification(TimestampMixin, db.Model):
    """Notifiche per gli utenti"""
    __tablename__ = "notifications"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    message = db.Column(db.Text, nullable=False)
    type = db.Column(db.String(50), default="info")  # info, success, warning, error
    is_read = db.Column(db.Boolean, default=False)
    related_ticket_id = db.Column(db.Integer, db.ForeignKey("tickets.id"), nullable=True)
    related_opportunity_id = db.Column(db.Integer, db.ForeignKey("opportunities.id"), nullable=True)
    
    user = db.relationship("User", backref="notifications")
    ticket = db.relationship("Ticket", backref="notifications")
    opportunity = db.relationship("Opportunity", backref="notifications")


def load_opportunity_safe(opportunity_id):
    """Carica opportunità usando SQL diretto per evitare problemi con colonne mancanti (heat_level)"""
    from sqlalchemy import text
    from datetime import datetime
    
    # Prova prima con heat_level, poi senza
    try:
        opp_result = db.session.execute(text("""
            SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                   description, supplier_category, folder_path, lost_reason, owner_id, 
                   account_id, contact_id, lead_id, created_at, updated_at, heat_level
            FROM opportunities 
            WHERE id = :opp_id
        """), {"opp_id": opportunity_id}).fetchone()
        has_heat_level = True
    except Exception:
        opp_result = db.session.execute(text("""
            SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                   description, supplier_category, folder_path, lost_reason, owner_id, 
                   account_id, contact_id, lead_id, created_at, updated_at
            FROM opportunities 
            WHERE id = :opp_id
        """), {"opp_id": opportunity_id}).fetchone()
        has_heat_level = False
    
    if not opp_result:
        return None
    
    # Costruisci oggetto Opportunity
    opp = Opportunity()
    opp.id = opp_result[0]
    opp.name = opp_result[1]
    opp.stage = opp_result[2]
    opp.amount = opp_result[3]
    opp.probability = opp_result[4]
    opp.close_date = opp_result[5]
    opp.opportunity_type = opp_result[6]
    opp.description = opp_result[7]
    opp.supplier_category = opp_result[8]
    opp.folder_path = opp_result[9]
    opp.lost_reason = opp_result[10]
    opp.owner_id = opp_result[11]
    opp.account_id = opp_result[12]
    opp.contact_id = opp_result[13]
    opp.lead_id = opp_result[14]
    # Converti date
    if isinstance(opp_result[15], str):
        opp.created_at = datetime.fromisoformat(opp_result[15].replace('Z', '+00:00')) if opp_result[15] else None
    else:
        opp.created_at = opp_result[15]
    if isinstance(opp_result[16], str):
        opp.updated_at = datetime.fromisoformat(opp_result[16].replace('Z', '+00:00')) if opp_result[16] else None
    else:
        opp.updated_at = opp_result[16]
    # Carica heat_level se disponibile
    if has_heat_level and len(opp_result) > 17:
        opp.heat_level = opp_result[17]
    else:
        opp.heat_level = None
    # Carica relazioni
    opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
    opp.account = db.session.query(Account).filter_by(id=opp.account_id).first() if opp.account_id else None
    opp.owner = db.session.query(User).filter_by(id=opp.owner_id).first() if opp.owner_id else None
    opp.contact = db.session.query(Contact).filter_by(id=opp.contact_id).first() if opp.contact_id else None
    opp.line_items = db.session.query(OpportunityLineItem).filter_by(opportunity_id=opp.id).all()
    
    return opp


def extract_date_from_folder_name(folder_path):
    """Estrae la data dal nome della cartella dell'offerta"""
    import re
    from datetime import datetime
    
    if not folder_path:
        return None
    
    folder_name = os.path.basename(folder_path)
    
    # Cerca pattern di date comuni nel nome della cartella
    # Formato: DD-MM-YYYY, DD-MM-YY, DD/MM/YYYY, DD/MM/YY, YYYY-MM-DD, DDMMYYYY
    date_patterns = [
        r'(\d{2})-(\d{2})-(\d{4})',  # DD-MM-YYYY
        r'(\d{2})-(\d{2})-(\d{2})',  # DD-MM-YY (es. 30-01-25)
        r'(\d{2})/(\d{2})/(\d{4})',  # DD/MM/YYYY
        r'(\d{2})/(\d{2})/(\d{2})',  # DD/MM/YY
        r'(\d{4})-(\d{2})-(\d{2})',  # YYYY-MM-DD
        r'(\d{2})(\d{2})(\d{4})',    # DDMMYYYY
    ]
    
    for pattern in date_patterns:
        match = re.search(pattern, folder_name)
        if match:
            try:
                if len(match.groups()) == 3:
                    if pattern.startswith(r'(\d{4})'):  # YYYY-MM-DD
                        year, month, day = match.groups()
                        return datetime(int(year), int(month), int(day)).date()
                    else:  # DD-MM-YYYY, DD-MM-YY, DD/MM/YYYY, DD/MM/YY
                        day, month, year = match.groups()
                        # Se l'anno ha solo 2 cifre, assumi 20XX
                        if len(year) == 2:
                            year = '20' + year
                        return datetime(int(year), int(month), int(day)).date()
                elif len(match.groups()) == 3 and pattern == r'(\d{2})(\d{2})(\d{4})':  # DDMMYYYY
                    day, month, year = match.groups()
                    return datetime(int(year), int(month), int(day)).date()
            except (ValueError, TypeError):
                continue
    
    # Se non trova una data nel nome, prova a estrarre dalla data di modifica della cartella
    try:
        if os.path.exists(folder_path):
            mod_time = os.path.getmtime(folder_path)
            return datetime.fromtimestamp(mod_time).date()
    except:
        pass
    
    return None


def extract_offer_amount(offer):
    """Estrae il valore da un'offerta (dal campo amount o dal PDF se disponibile)"""
    if offer.amount and offer.amount > 0:
        return offer.amount
    
    # Se non c'è amount, prova a estrarre dal PDF
    if offer.pdf_path and os.path.exists(offer.pdf_path):
        return extract_offer_amount_from_file(offer.pdf_path)
    
    return 0.0


def extract_offer_amount_from_file(file_path):
    """Estrae il valore da un file PDF"""
    if not file_path or not file_path.endswith('.pdf') or not os.path.exists(file_path):
        return 0.0
    
    try:
        import pdfplumber
        import re
        with pdfplumber.open(file_path) as pdf:
            all_text = ""
            for page in pdf.pages:
                all_text += page.extract_text() or ""
            
            # Cerca "TOTALE" o "TOTALE FORNITURA"
            patterns = [
                r'TOTALE\s*(?:FORNITURA)?\s*[:\s]*€?\s*([\d.,]+)',
                r'TOTALE\s*[:\s]*€?\s*([\d.,]+)',
            ]
            
            max_amount = 0.0
            for pattern in patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.replace('.', '').replace(',', '.')
                    try:
                        amount = float(amount_str)
                        if amount > max_amount:
                            max_amount = amount
                    except ValueError:
                        continue
            
            if max_amount > 0:
                return max_amount
            
            # Cerca il numero più alto
            numbers = re.findall(r'€\s*([\d.,]+)', all_text)
            amounts = []
            for num in numbers:
                try:
                    amount_str = num.replace('.', '').replace(',', '.')
                    amount = float(amount_str)
                    if amount > 100:
                        amounts.append(amount)
                except ValueError:
                    continue
            
            if amounts:
                return max(amounts)
    except Exception:
        pass
    
    return 0.0


# ==================== ROUTES ====================

@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('app_route'))
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        try:
            data = request.json
            if not data:
                return jsonify({'success': False, 'error': 'Dati mancanti'}), 400
            
            # Prova prima con username, poi con email
            user = User.query.filter_by(username=data.get('username')).first()
            if not user:
                user = User.query.filter_by(email=data.get('username')).first()
            
            if user and check_password_hash(user.password, data.get('password', '')):
                session['user_id'] = user.id
                session['username'] = user.username
                session['user_role'] = user.role
                return jsonify({'success': True, 'role': user.role})
            return jsonify({'success': False, 'error': 'Credenziali errate'}), 401
        except Exception as e:
            return jsonify({'success': False, 'error': f'Errore: {str(e)}'}), 500
    return render_template('login.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    """Registrazione clienti"""
    if request.method == 'POST':
        try:
            data = request.json
            if not data:
                return jsonify({'success': False, 'error': 'Dati mancanti'}), 400
            
            # Verifica campi obbligatori
            if not data.get('email') or not data.get('password') or not data.get('first_name'):
                return jsonify({'success': False, 'error': 'Email, password e nome sono obbligatori'}), 400
            
            # Verifica se email già esistente
            existing = User.query.filter_by(email=data.get('email')).first()
            if existing:
                return jsonify({'success': False, 'error': 'Email già registrata'}), 400
            
            # Verifica se username già esistente
            username = data.get('username') or data.get('email').split('@')[0]
            existing_username = User.query.filter_by(username=username).first()
            if existing_username:
                # Aggiungi numero se username già esiste
                counter = 1
                while User.query.filter_by(username=f"{username}{counter}").first():
                    counter += 1
                username = f"{username}{counter}"
            
            # Crea nuovo utente cliente
            user = User(
                username=username,
                password=generate_password_hash(data.get('password')),
                email=data.get('email'),
                first_name=data.get('first_name'),
                last_name=data.get('last_name', ''),
                role='cliente'  # Ruolo cliente per distinguerli dagli operatori
            )
            
            db.session.add(user)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Registrazione completata! Ora puoi accedere.',
                'username': username
            })
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': f'Errore: {str(e)}'}), 500
    
    # GET: mostra pagina registrazione
    return render_template('register.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/dashboard')
def dashboard():
    """Redirect a /app per compatibilità"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('app_route'))


@app.route('/app')
def app_route():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    user = User.query.get(session['user_id'])
    if user and user.role == 'cliente':
        # Cliente vede solo la vista ticket
        return render_template('app_cliente.html')
    elif user:
        # Operatori tecnici vedono solo ticketing
        # NOTA: filippo e mariuzzo sono commerciali, non operatori tecnici!
        operatori_tecnici = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'mattia', 'scevola']
        is_operatore_tecnico = any(op in user.username.lower() for op in operatori_tecnici)
        
        if is_operatore_tecnico and user.role != 'admin':
            # Operatori tecnici vedono solo ticketing (usa template cliente ma con permessi operatori)
            return render_template('app_cliente.html')
        else:
            # Commerciali (Giovanni, Mauro, Davide, Filippo) e admin vedono il CRM completo
            return render_template('app_mekocrm_completo.html')
    else:
        return redirect(url_for('login'))


@app.route('/api/user/info')
def user_info():
    """Ottieni informazioni utente corrente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user:
        return jsonify({'error': 'Utente non trovato'}), 404
    
    return jsonify({
        'user': {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'role': user.role
        }
    })


@app.route('/api/user/local-pc-config', methods=['GET', 'PUT'])
def user_local_pc_config():
    """Gestisce la configurazione del PC locale per l'utente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user:
        return jsonify({'error': 'Utente non trovato'}), 404
    
    if request.method == 'GET':
        # Restituisci valori di default se vuoti - MAI restituire la password in chiaro
        return jsonify({
            'local_pc_ip': user.local_pc_ip or DEFAULT_SERVER_IP,
            'local_pc_share': user.local_pc_share or DEFAULT_SERVER_SHARE,
            'local_pc_base_path': user.local_pc_base_path or '',
            'local_pc_username': user.local_pc_username or '',
            'has_password': bool(user.local_pc_password),  # Solo indicatore booleano, mai la password
            'default_ip': DEFAULT_SERVER_IP,
            'default_share': DEFAULT_SERVER_SHARE,
            'is_using_default': not (user.local_pc_ip and user.local_pc_share)
        })
    
    elif request.method == 'PUT':
        data = request.get_json()
        
        user.local_pc_ip = data.get('local_pc_ip', '').strip()
        user.local_pc_share = data.get('local_pc_share', '').strip()
        user.local_pc_base_path = data.get('local_pc_base_path', '').strip()
        user.local_pc_username = data.get('local_pc_username', '').strip()
        
        # Aggiorna password solo se fornita - usa hash sicuro
        if 'local_pc_password' in data and data['local_pc_password']:
            from werkzeug.security import generate_password_hash
            # Genera hash sicuro della password SMB (non mostrare mai in chiaro)
            user.local_pc_password = generate_password_hash(data['local_pc_password'])
            # Log dell'aggiornamento (senza password)
            app.logger.info(f"Password SMB aggiornata per utente {user.id} ({user.username})")
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Configurazione PC locale salvata'
        })


@app.route('/favicon.ico')
def favicon():
    """Route per favicon per evitare errori 404"""
    return '', 204  # No Content


# ==================== API ROUTES ====================

def extract_offer_amount(offer):
    """Estrae il valore dell'offerta dal PDF o DOCX se non presente nel database"""
    if offer.amount and offer.amount > 0:
        return offer.amount
    
    # Prova a estrarre dal PDF
    if offer.pdf_path and os.path.exists(offer.pdf_path):
        try:
            import pdfplumber
            import re
            with pdfplumber.open(offer.pdf_path) as pdf:
                all_text = ""
                for page in pdf.pages:
                    all_text += page.extract_text() or ""
                
                # Cerca "TOTALE" o "TOTALE FORNITURA" seguito da un numero
                patterns = [
                    r'TOTALE\s*(?:FORNITURA)?\s*[:\s]*€?\s*([\d.,]+)',
                    r'TOTALE\s*[:\s]*€?\s*([\d.,]+)',
                    r'€\s*([\d.,]+)\s*(?:TOTALE|TOTALE FORNITURA)',
                ]
                
                max_amount = 0.0
                for pattern in patterns:
                    matches = re.findall(pattern, all_text, re.IGNORECASE)
                    for match in matches:
                        # Converti il numero (gestisce sia . che , come separatore decimale)
                        amount_str = match.replace('.', '').replace(',', '.')
                        try:
                            amount = float(amount_str)
                            if amount > max_amount:
                                max_amount = amount
                        except ValueError:
                            continue
                
                if max_amount > 0:
                    return max_amount
                
                # Cerca TUTTI i numeri nel documento (non solo quelli con €)
                # Pattern più completo: numeri con separatori italiani o internazionali
                number_patterns = [
                    r'€\s*([\d.,]+)',  # € seguito da numero
                    r'([\d]{1,3}(?:\.\d{3})*(?:,\d{2})?)',  # Formato italiano: 1.234,56
                    r'([\d]{1,3}(?:,\d{3})*(?:\.\d{2})?)',  # Formato internazionale: 1,234.56
                    r'([\d]+[.,]\d{2})',  # Numeri con decimali
                ]
                
                all_amounts = []
                for pattern in number_patterns:
                    matches = re.findall(pattern, all_text)
                    for match in matches:
                        try:
                            # Normalizza: rimuovi punti (migliaia) e sostituisci virgola con punto
                            amount_str = match.replace('.', '').replace(',', '.')
                            amount = float(amount_str)
                            # Filtra numeri ragionevoli (tra 100 e 10 milioni)
                            if 100 <= amount <= 10000000:
                                all_amounts.append(amount)
                        except ValueError:
                            continue
                
                # Restituisci il numero più alto trovato
                if all_amounts:
                    return max(all_amounts)
        except Exception as e:
            pass  # Ignora errori di estrazione
    
    # Prova a estrarre dal DOCX
    if offer.docx_path and os.path.exists(offer.docx_path):
        try:
            from docx import Document
            import re
            doc = Document(offer.docx_path)
            all_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Cerca "TOTALE" o "TOTALE FORNITURA"
            patterns = [
                r'TOTALE\s*(?:FORNITURA)?\s*[:\s]*€?\s*([\d.,]+)',
                r'TOTALE\s*[:\s]*€?\s*([\d.,]+)',
            ]
            
            max_amount = 0.0
            for pattern in patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.replace('.', '').replace(',', '.')
                    try:
                        amount = float(amount_str)
                        if amount > max_amount:
                            max_amount = amount
                    except ValueError:
                        continue
            
            if max_amount > 0:
                return max_amount
            
            # Cerca il numero più alto
            numbers = re.findall(r'€\s*([\d.,]+)', all_text)
            amounts = []
            for num in numbers:
                try:
                    amount_str = num.replace('.', '').replace(',', '.')
                    amount = float(amount_str)
                    if amount > 100:
                        amounts.append(amount)
                except ValueError:
                    continue
            
            if amounts:
                return max(amounts)
        except Exception as e:
            pass  # Ignora errori di estrazione
    
    return 0.0


@app.route('/api/user/me', methods=['GET'])
def get_current_user():
    """Ottieni informazioni sull'utente corrente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user:
        return jsonify({'error': 'Utente non trovato'}), 404
    
    return jsonify({
        'user': {
            'id': user.id,
            'username': user.username,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email,
            'role': user.role
        }
    })


@app.route('/api/dashboard/summary')
def dashboard_summary():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    # Metriche base per l'utente corrente
    # Leads e Accounts condivisi - conta tutti
    leads_count = Lead.query.count()
    accounts_count = Account.query.count()
    # Carica opportunità con offerte usando joinedload per performance
    # Usa SQL diretto per evitare problemi con colonne mancanti come heat_level
    from sqlalchemy import text
    from sqlalchemy.orm import joinedload
    
    # Carica opportunità - Commerciali vedono TUTTE, altri solo le loro
    user = User.query.get(user_id)
    is_commerciale = user and user.role == 'commerciale'
    
    if is_commerciale:
        # Commerciali vedono TUTTE le opportunità (per dashboard completa)
        opps_result = db.session.execute(text("""
            SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                   description, supplier_category, folder_path, lost_reason, owner_id, 
                   account_id, contact_id, lead_id, created_at, updated_at, heat_level
            FROM opportunities
        """)).fetchall()
    else:
        # Altri utenti vedono solo le loro opportunità
        try:
            opps_result = db.session.execute(text("""
                SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                       description, supplier_category, folder_path, lost_reason, owner_id, 
                       account_id, contact_id, lead_id, created_at, updated_at, heat_level
                FROM opportunities 
                WHERE owner_id = :user_id
            """), {"user_id": user_id}).fetchall()
        except Exception:
            # Fallback se heat_level non esiste
            opps_result = db.session.execute(text("""
                SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                       description, supplier_category, folder_path, lost_reason, owner_id, 
                       account_id, contact_id, lead_id, created_at, updated_at
                FROM opportunities 
                WHERE owner_id = :user_id
            """), {"user_id": user_id}).fetchall()
    
    # Costruisci oggetti Opportunity manualmente senza heat_level
    from datetime import datetime
    opps = []
    for row in opps_result:
        opp = Opportunity()
        opp.id = row[0]
        opp.name = row[1]
        opp.stage = row[2]
        opp.amount = row[3]
        opp.probability = row[4]
        opp.close_date = row[5]
        opp.opportunity_type = row[6]
        opp.description = row[7]
        opp.supplier_category = row[8]
        opp.folder_path = row[9]
        opp.lost_reason = row[10]
        opp.owner_id = row[11]
        opp.account_id = row[12]
        opp.contact_id = row[13]
        opp.lead_id = row[14]
        # Converti date da stringhe a datetime se necessario
        if isinstance(row[15], str):
            opp.created_at = datetime.fromisoformat(row[15].replace('Z', '+00:00')) if row[15] else None
        else:
            opp.created_at = row[15]
        if isinstance(row[16], str):
            opp.updated_at = datetime.fromisoformat(row[16].replace('Z', '+00:00')) if row[16] else None
        else:
            opp.updated_at = row[16]
        
        # Carica heat_level se disponibile (può essere nella colonna 17 o non esistere)
        if len(row) > 17:
            opp.heat_level = row[17]
        else:
            # Prova a caricarlo dalla tabella se esiste
            try:
                heat_result = db.session.execute(text("""
                    SELECT heat_level FROM opportunities WHERE id = :opp_id
                """), {"opp_id": opp.id}).fetchone()
                opp.heat_level = heat_result[0] if heat_result else None
            except:
                opp.heat_level = None
        
        # Carica offerte per questa opportunità
        opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
        opps.append(opp)
    opportunities_count = len(opps)
    
    # Calcola pipeline includendo valori offerte
    pipeline = 0.0
    for opp in opps:
        if opp.stage not in ['Closed Won', 'Closed Lost']:
            # Calcola valore opportunità considerando la probabilità
            opp_value = 0.0
            if opp.amount and opp.amount > 0:
                opp_value = opp.amount
            else:
                # Carica le offerte - accedi direttamente per triggerare lazy loading
                try:
                    offers_list = list(opp.offers)  # Questo triggera il lazy loading
                    if offers_list:
                        opp_value = sum(
                            (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                            for o in offers_list
                        )
                except Exception as e:
                    # Se fallisce, continua senza aggiungere alla pipeline
                    pass
            
            # Determina probabilità basandosi su heat_level invece che su probability
            # Fredde gelo: 0%, Fredde con speranza: 5%, Tiepide: 15%, Calde: 50%
            heat_level = getattr(opp, 'heat_level', None) or 'da_categorizzare'
            if heat_level == 'calda':
                probability = 50
            elif heat_level == 'tiepida':
                probability = 15
            elif heat_level == 'fredda_speranza':
                probability = 5
            elif heat_level == 'fredda_gelo':
                probability = 0
            else:
                # Default per da_categorizzare o valori non riconosciuti
                probability = 0
            pipeline += opp_value * (probability / 100.0)
    
    # Calcola chiusi includendo valori offerte
    closed_won = 0.0
    for opp in opps:
        if opp.stage == 'Closed Won':
            if opp.amount and opp.amount > 0:
                closed_won += opp.amount
            else:
                try:
                    offers_list = list(opp.offers)  # Triggera lazy loading
                    if offers_list:
                        opp_value = sum(
                            (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                            for o in offers_list
                        )
                        if opp_value > 0:
                            closed_won += opp_value
                except Exception as e:
                    pass
    
    closed_won_count = len([o for o in opps if o.stage == 'Closed Won'])
    
    # Metriche globali (tutte le opportunità - leads UR condivisi)
    # Usa sempre SQL diretto per evitare problemi con colonne mancanti
    from sqlalchemy import text
    all_opps_result = db.session.execute(text("""
        SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
               description, supplier_category, folder_path, lost_reason, owner_id, 
               account_id, contact_id, lead_id, created_at, updated_at
        FROM opportunities
    """)).fetchall()
    
    from datetime import datetime
    all_opps = []
    for row in all_opps_result:
        opp = Opportunity()
        opp.id = row[0]
        opp.name = row[1]
        opp.stage = row[2]
        opp.amount = row[3]
        opp.probability = row[4]
        opp.close_date = row[5]
        opp.opportunity_type = row[6]
        opp.description = row[7]
        opp.supplier_category = row[8]
        opp.folder_path = row[9]
        opp.lost_reason = row[10]
        opp.owner_id = row[11]
        opp.account_id = row[12]
        opp.contact_id = row[13]
        opp.lead_id = row[14]
        # Converti date da stringhe a datetime se necessario
        if isinstance(row[15], str):
            opp.created_at = datetime.fromisoformat(row[15].replace('Z', '+00:00')) if row[15] else None
        else:
            opp.created_at = row[15]
        if isinstance(row[16], str):
            opp.updated_at = datetime.fromisoformat(row[16].replace('Z', '+00:00')) if row[16] else None
        else:
            opp.updated_at = row[16]
        opp.heat_level = None
        
        # Carica offerte per questa opportunità
        opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
        all_opps.append(opp)
    all_opportunities_count = len(all_opps)
    
    # Calcola pipeline globale includendo valori offerte
    all_pipeline = 0.0
    for opp in all_opps:
        if opp.stage not in ['Closed Won', 'Closed Lost']:
            # Calcola valore opportunità considerando la probabilità
            opp_value = 0.0
            if opp.amount and opp.amount > 0:
                opp_value = opp.amount
            else:
                try:
                    offers_list = list(opp.offers)  # Triggera lazy loading
                    if offers_list:
                        opp_value = sum(
                            (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                            for o in offers_list
                        )
                except Exception as e:
                    pass
            
            # Applica probabilità (default 10% se non specificata)
            probability = opp.probability if opp.probability and opp.probability > 0 else 10
            all_pipeline += opp_value * (probability / 100.0)
    
    # Calcola chiusi globali includendo valori offerte
    all_closed_won = 0.0
    for opp in all_opps:
        if opp.stage == 'Closed Won':
            if opp.amount and opp.amount > 0:
                all_closed_won += opp.amount
            else:
                try:
                    offers_list = list(opp.offers)  # Triggera lazy loading
                    if offers_list:
                        opp_value = sum(
                            (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                            for o in offers_list
                        )
                        if opp_value > 0:
                            all_closed_won += opp_value
                except Exception as e:
                    pass
    
    all_closed_count = len([o for o in all_opps if o.stage == 'Closed Won'])
    
    # Leaderboard - metriche per tutti gli utenti commerciali (senza duplicati)
    commercial_users = User.query.filter_by(role='commerciale').all()
    # Rimuovi duplicati per nome completo (first_name + last_name) o email
    # Mantieni l'utente con più opportunità o quello principale
    seen_names = {}
    unique_users = []
    
    for user in commercial_users:
        # Crea chiave univoca: nome completo normalizzato o email
        full_name = f"{user.first_name or ''} {user.last_name or ''}".strip().lower()
        email_key = (user.email or '').lower()
        
        # Usa email se disponibile, altrimenti nome
        key = email_key if email_key else full_name
        
        if not key:
            # Se non c'è né nome né email, usa username
            key = user.username.lower()
        
        if key not in seen_names:
            seen_names[key] = user
            unique_users.append(user)
        else:
            # Se esiste già, mantieni quello con più opportunità o quello principale
            existing_user = seen_names[key]
            existing_opps = Opportunity.query.filter_by(owner_id=existing_user.id).count()
            current_opps = Opportunity.query.filter_by(owner_id=user.id).count()
            
            # Se il nuovo utente ha più opportunità, sostituisci
            # Oppure se il nuovo ha un username più "principale" (più corto o senza numeri)
            if current_opps > existing_opps or (current_opps == existing_opps and len(user.username) < len(existing_user.username)):
                # Rimuovi il vecchio e aggiungi il nuovo
                unique_users.remove(existing_user)
                seen_names[key] = user
                unique_users.append(user)
    
    leaderboard = []
    for user in unique_users:
        # Usa sempre SQL diretto per evitare problemi con colonne mancanti
        from sqlalchemy import text
        user_opps_result = db.session.execute(text("""
            SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                   description, supplier_category, folder_path, lost_reason, owner_id, 
                   account_id, contact_id, lead_id, created_at, updated_at
            FROM opportunities 
            WHERE owner_id = :user_id
        """), {"user_id": user.id}).fetchall()
        
        from datetime import datetime
        user_opps = []
        for row in user_opps_result:
            opp = Opportunity()
            opp.id = row[0]
            opp.name = row[1]
            opp.stage = row[2]
            opp.amount = row[3]
            opp.probability = row[4]
            opp.close_date = row[5]
            opp.opportunity_type = row[6]
            opp.description = row[7]
            opp.supplier_category = row[8]
            opp.folder_path = row[9]
            opp.lost_reason = row[10]
            opp.owner_id = row[11]
            opp.account_id = row[12]
            opp.contact_id = row[13]
            opp.lead_id = row[14]
            # Converti date da stringhe a datetime se necessario
            if isinstance(row[15], str):
                opp.created_at = datetime.fromisoformat(row[15].replace('Z', '+00:00')) if row[15] else None
            else:
                opp.created_at = row[15]
            if isinstance(row[16], str):
                opp.updated_at = datetime.fromisoformat(row[16].replace('Z', '+00:00')) if row[16] else None
            else:
                opp.updated_at = row[16]
            opp.heat_level = None
            
            # Carica offerte per questa opportunità
            opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
            user_opps.append(opp)
        
        # Calcola pipeline includendo valori offerte
        user_pipeline = 0.0
        for opp in user_opps:
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                # Usa amount opportunità o somma delle offerte
                if opp.amount and opp.amount > 0:
                    user_pipeline += opp.amount
                else:
                    # Calcola dalla somma delle offerte
                    try:
                        offers_list = list(opp.offers)  # Triggera lazy loading
                        if offers_list:
                            opp_value = sum(
                                (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                                for o in offers_list
                            )
                            if opp_value > 0:
                                user_pipeline += opp_value
                    except Exception as e:
                        pass
        
        # Calcola chiusi includendo valori offerte
        user_closed_won = 0.0
        for opp in user_opps:
            if opp.stage == 'Closed Won':
                if opp.amount and opp.amount > 0:
                    user_closed_won += opp.amount
                else:
                    try:
                        offers_list = list(opp.offers)  # Triggera lazy loading
                        if offers_list:
                            opp_value = sum(
                                (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                                for o in offers_list
                            )
                            if opp_value > 0:
                                user_closed_won += opp_value
                    except Exception as e:
                        pass
        
        user_closed_count = len([o for o in user_opps if o.stage == 'Closed Won'])
        user_opps_count = len(user_opps)
        
        leaderboard.append({
            'user_id': user.id,
            'username': user.username,
            'name': f"{user.first_name} {user.last_name}",
            'opportunities': user_opps_count,
            'pipeline': user_pipeline,
            'closed_won': user_closed_won,
            'closed_count': user_closed_count
        })
    
    # Ordina per pipeline (decrescente)
    leaderboard.sort(key=lambda x: x['pipeline'], reverse=True)
    
    # Statistiche per mese/giorno (ultimi 30 giorni)
    from datetime import timedelta
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    
    # Opportunità chiuse per giorno (ultimi 30 giorni)
    daily_closed = {}
    daily_pipeline = {}
    daily_offers = {}  # Valori offerte per giorno
    monthly_stats = {}
    monthly_offers = {}  # Valori offerte per mese
    
    for opp in all_opps:
        if opp.created_at and opp.created_at >= thirty_days_ago:
            day_key = opp.created_at.date().isoformat()
            month_key = opp.created_at.strftime('%Y-%m')
            
            # Pipeline giornaliera
            if day_key not in daily_pipeline:
                daily_pipeline[day_key] = {'count': 0, 'amount': 0}
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                daily_pipeline[day_key]['count'] += 1
                daily_pipeline[day_key]['amount'] += opp.amount
            
            # Chiusi giornalieri
            if opp.stage == 'Closed Won':
                closed_day = opp.close_date.isoformat() if opp.close_date else day_key
                if closed_day not in daily_closed:
                    daily_closed[closed_day] = {'count': 0, 'amount': 0}
                daily_closed[closed_day]['count'] += 1
                daily_closed[closed_day]['amount'] += opp.amount
            
            # Statistiche mensili
            if month_key not in monthly_stats:
                monthly_stats[month_key] = {'pipeline': 0, 'closed': 0, 'count': 0}
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                monthly_stats[month_key]['pipeline'] += opp.amount
            elif opp.stage == 'Closed Won':
                monthly_stats[month_key]['closed'] += opp.amount
                monthly_stats[month_key]['count'] += 1
        
        # Calcola valori offerte per questa opportunità
        opp_offers_value = 0.0
        for offer in opp.offers:
            offer_amount = offer.amount if offer.amount and offer.amount > 0 else extract_offer_amount(offer)
            opp_offers_value += offer_amount
        
        if opp_offers_value > 0:
            # Aggiungi ai totali giornalieri
            if opp.created_at and opp.created_at >= thirty_days_ago:
                day_key = opp.created_at.date().isoformat()
                if day_key not in daily_offers:
                    daily_offers[day_key] = {'count': 0, 'amount': 0.0}
                daily_offers[day_key]['count'] += len([o for o in opp.offers if (o.amount and o.amount > 0) or extract_offer_amount(o) > 0])
                daily_offers[day_key]['amount'] += opp_offers_value
                
                # Aggiungi ai totali mensili
                month_key = opp.created_at.strftime('%Y-%m')
                if month_key not in monthly_offers:
                    monthly_offers[month_key] = {'count': 0, 'amount': 0.0}
                monthly_offers[month_key]['count'] += len([o for o in opp.offers if (o.amount and o.amount > 0) or extract_offer_amount(o) > 0])
                monthly_offers[month_key]['amount'] += opp_offers_value
    
    # Ultima sincronizzazione
    last_sync = PriceBookEntry.query.order_by(PriceBookEntry.created_at.desc()).first()
    last_sync_date = last_sync.created_at if last_sync else None
    
    # Calcola valore totale offerte per tutte le opportunità
    total_offers_value = 0.0
    for opp in all_opps:
        for offer in opp.offers:
            offer_amount = offer.amount if offer.amount and offer.amount > 0 else extract_offer_amount(offer)
            total_offers_value += offer_amount
    
    return jsonify({
        'metrics': {
            'leads': leads_count,
            'accounts': accounts_count,
            'opportunities': opportunities_count,
            'pipeline': pipeline,
            'closed_won': closed_won,
            'closed_count': closed_won_count
        },
        'global_metrics': {
            'opportunities': all_opportunities_count,
            'pipeline': all_pipeline,
            'closed_won': all_closed_won,
            'closed_count': all_closed_count,
            'total_offers_value': total_offers_value
        },
        'leaderboard': leaderboard,
        'daily_pipeline': daily_pipeline,
        'daily_closed': daily_closed,
        'daily_offers': daily_offers,
        'monthly_stats': monthly_stats,
        'monthly_offers': monthly_offers,
        'last_sync': last_sync_date.isoformat() if last_sync_date else None
    })


@app.route('/api/leads', methods=['GET', 'POST'])
def leads():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Leads condivisi per tutti (come opportunità UR)
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        
        query = Lead.query.order_by(Lead.created_at.desc())
        
        # Se per_page è molto grande, carica tutti senza paginazione
        if per_page >= 1000:
            leads_list = query.all()
            pagination = type('obj', (object,), {
                'total': len(leads_list),
                'pages': 1,
                'has_next': False,
                'has_prev': False,
                'next_num': None,
                'prev_num': None
            })()
        else:
            # Paginazione normale
            pagination = query.paginate(page=page, per_page=per_page, error_out=False)
            leads_list = pagination.items
        
        return jsonify({
            'leads': [{
                'id': l.id,
                'company_name': l.company_name,
                'contact_first_name': l.contact_first_name,
                'contact_last_name': l.contact_last_name,
                'email': l.email,
                'phone': l.phone,
                'interest': l.interest,
                'source': l.source,
                'status': l.status,
                'rating': l.rating,
                'owner_id': l.owner_id,
                'owner_name': f"{l.owner.first_name} {l.owner.last_name}" if l.owner else 'Unknown',
                'is_my_lead': l.owner_id == user_id,
                'created_at': l.created_at.isoformat() if l.created_at else None,
                'last_contact_at': l.last_contact_at.isoformat() if l.last_contact_at else None,
                'address': l.address,
                'has_pending_tasks': False,  # Leads non hanno attività dirette, solo tramite account convertito
                'pending_tasks_count': 0
            } for l in leads_list],
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': pagination.total,
                'pages': pagination.pages,
                'has_next': pagination.has_next,
                'has_prev': pagination.has_prev,
                'next_page': pagination.next_num if pagination.has_next else None,
                'prev_page': pagination.prev_num if pagination.has_prev else None
            }
        })
    
    elif request.method == 'POST':
        data = request.json
        lead = Lead(
            company_name=data['company_name'],
            contact_first_name=data.get('contact_first_name'),
            contact_last_name=data.get('contact_last_name'),
            email=data.get('email'),
            phone=data.get('phone'),
            vat_number=data.get('vat_number'),
            address=data.get('address'),
            industry=data.get('industry'),
            interest=data.get('interest'),
            source=data.get('source'),
            rating=data.get('rating', 'Warm'),
            notes=data.get('notes'),
            owner_id=user_id
        )
        db.session.add(lead)
        db.session.commit()
        return jsonify({'success': True, 'lead_id': lead.id})


@app.route('/api/leads/<int:lead_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_lead(lead_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead = Lead.query.get_or_404(lead_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and lead.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        db.session.delete(lead)
        db.session.commit()
        return jsonify({'success': True})
    
    # Permetti modifica anche agli admin o se non sei owner (leads condivisi)
    user_id = session['user_id']
    user = User.query.get(user_id)
    is_admin = user and user.role == 'admin'
    is_owner = lead.owner_id == user_id
    
    if request.method == 'GET':
        return jsonify({
            'lead': {
                'id': lead.id,
                'company_name': lead.company_name,
                'contact_first_name': lead.contact_first_name,
                'contact_last_name': lead.contact_last_name,
                'email': lead.email,
                'phone': lead.phone,
                'vat_number': lead.vat_number,
                'address': lead.address,
                'industry': lead.industry,
                'interest': lead.interest,
                'source': lead.source,
                'rating': lead.rating,
                'status': lead.status,
                'notes': lead.notes,
                'created_at': lead.created_at.isoformat() if lead.created_at else None
            }
        })
    
    # PUT method - permetti modifica se admin o owner
    if not is_admin and not is_owner:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    data = request.json
    if 'company_name' in data:
        lead.company_name = data['company_name']
    if 'contact_first_name' in data:
        lead.contact_first_name = data['contact_first_name']
    if 'contact_last_name' in data:
        lead.contact_last_name = data['contact_last_name']
    if 'email' in data:
        lead.email = data['email']
    if 'phone' in data:
        lead.phone = data['phone']
    if 'vat_number' in data:
        lead.vat_number = data['vat_number']
    if 'address' in data:
        lead.address = data['address']
    if 'industry' in data:
        lead.industry = data['industry']
    if 'interest' in data:
        lead.interest = data['interest']
    if 'source' in data:
        lead.source = data['source']
    if 'rating' in data:
        lead.rating = data['rating']
    if 'status' in data:
        lead.status = data['status']
    if 'notes' in data:
        lead.notes = data['notes']
    if 'created_at' in data:
        if data['created_at']:
            lead.created_at = datetime.fromisoformat(data['created_at'].replace('Z', '+00:00'))
        else:
            lead.created_at = datetime.utcnow()
    if 'last_contact_at' in data:
        if data['last_contact_at']:
            lead.last_contact_at = datetime.fromisoformat(data['last_contact_at'].replace('Z', '+00:00'))
        else:
            lead.last_contact_at = None
    
    db.session.commit()
    return jsonify({'success': True, 'lead_id': lead.id})


@app.route('/api/leads/<int:lead_id>/convert', methods=['POST'])
def convert_lead(lead_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead = Lead.query.get_or_404(lead_id)
    # Tutti possono convertire leads (leads condivisi)
    
    # Crea Account
    account = Account(
        name=lead.company_name,
        industry=lead.industry,
        phone=lead.phone,
        billing_address=lead.address,
        owner_id=lead.owner_id
    )
    db.session.add(account)
    db.session.flush()
    
    # Crea Contact
    contact = Contact(
        first_name=lead.contact_first_name or 'Contact',
        last_name=lead.contact_last_name or 'Unknown',
        email=lead.email,
        phone=lead.phone,
        account_id=account.id,
        owner_id=lead.owner_id
    )
    db.session.add(contact)
    db.session.flush()
    
    # Crea Opportunity
    opportunity = Opportunity(
        name=f"Opportunity {lead.company_name}",
        stage="Qualification",
        account_id=account.id,
        contact_id=contact.id,
        lead_id=lead.id,
        owner_id=lead.owner_id
    )
    db.session.add(opportunity)
    db.session.flush()
    
    # Aggiorna Lead
    lead.status = 'Converted'
    lead.converted_account_id = account.id
    lead.converted_contact_id = contact.id
    lead.converted_opportunity_id = opportunity.id
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/accounts', methods=['GET', 'POST'])
def accounts():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Accounts condivisi per tutti (come opportunità UR)
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        
        query = Account.query.order_by(Account.created_at.desc())
        
        # Paginazione
        try:
            pagination = query.paginate(page=page, per_page=per_page, error_out=False)
            accounts_list = pagination.items
        except Exception as e:
            return jsonify({'error': f'Errore paginazione: {str(e)}'}), 500
        
        accounts_data = []
        for a in accounts_list:
            try:
                account_dict = {
                    'id': a.id,
                    'name': a.name or '',
                    'industry': a.industry or '',
                    'phone': a.phone or '',
                    'website': a.website or '',
                    'billing_address': a.billing_address or '',
                    'shipping_address': a.shipping_address or '',
                    'description': a.description or '',
                    'image_url': a.image_url or '',
                    'owner_id': a.owner_id,
                    'owner_name': f"{a.owner.first_name} {a.owner.last_name}" if a.owner else 'Unknown',
                    'is_my_account': a.owner_id == user_id,
                    'created_at': a.created_at.isoformat() if a.created_at else None,
                    'has_pending_tasks': False,
                    'pending_tasks_count': 0,
                    'opportunities': []
                }
                
                # Aggiungi opportunità in modo sicuro
                try:
                    if hasattr(a, 'opportunities') and a.opportunities:
                        for o in a.opportunities:
                            try:
                                opp_amount = o.amount if o.amount and o.amount > 0 else 0
                                if not opp_amount and hasattr(o, 'offers') and o.offers:
                                    try:
                                        opp_amount = sum(
                                            (off.amount if off.amount and off.amount > 0 else extract_offer_amount(off))
                                            for off in o.offers
                                        )
                                    except:
                                        pass
                                
                                account_dict['opportunities'].append({
                                    'id': o.id, 
                                    'name': o.name or '', 
                                    'stage': o.stage or '', 
                                    'amount': opp_amount,
                                    'created_at': o.created_at.isoformat() if o.created_at else None
                                })
                            except Exception:
                                continue
                except Exception:
                    pass
                
                accounts_data.append(account_dict)
            except Exception as e:
                # Se fallisce un account, continua con gli altri
                continue
        
        return jsonify({
            'accounts': accounts_data,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': pagination.total,
                'pages': pagination.pages,
                'has_next': pagination.has_next,
                'has_prev': pagination.has_prev,
                'next_page': pagination.next_num if pagination.has_next else None,
                'prev_page': pagination.prev_num if pagination.has_prev else None
            }
        })
    
    elif request.method == 'POST':
        data = request.json
        account = Account(
            name=data['name'],
            industry=data.get('industry'),
            phone=data.get('phone'),
            website=data.get('website'),
            billing_address=data.get('billing_address'),
            shipping_address=data.get('shipping_address'),
            description=data.get('description'),
            owner_id=user_id
        )
        db.session.add(account)
        db.session.commit()
        return jsonify({'success': True, 'account_id': account.id})


@app.route('/api/map-data', methods=['GET'])
def map_data():
    """Endpoint ottimizzato per caricare solo account e leads con indirizzi validi per la mappa"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        # Carica solo account con indirizzo valido
        accounts_list = Account.query.filter(
            db.or_(
                Account.billing_address.isnot(None),
                Account.shipping_address.isnot(None)
            )
        ).all()
        
        accounts_data = []
        for a in accounts_list:
            try:
                # Determina il miglior indirizzo
                address = a.billing_address if a.billing_address and len(a.billing_address.strip()) > 3 else (a.shipping_address if a.shipping_address and len(a.shipping_address.strip()) > 3 else None)
                if not address:
                    continue
                
                # Ottieni gli ID delle opportunità usando una query separata (più sicuro)
                opp_ids = db.session.query(Opportunity.id).filter_by(account_id=a.id).all()
                opp_id_list = [opp_id[0] for opp_id in opp_ids] if opp_ids else []
                
                # Conta task pendenti (sia diretti all'account che alle opportunità)
                pending_tasks = 0
                if opp_id_list:
                    pending_tasks = OpportunityTask.query.filter(
                        db.or_(
                            OpportunityTask.account_id == a.id,
                            OpportunityTask.opportunity_id.in_(opp_id_list)
                        )
                    ).filter_by(
                        is_completed=False
                    ).count()
                else:
                    # Se non ci sono opportunità, conta solo i task diretti all'account
                    pending_tasks = OpportunityTask.query.filter_by(
                        account_id=a.id,
                        is_completed=False
                    ).count()
                
                # Determina categoria dalla prima opportunità (se presente)
                category = None
                first_opp = Opportunity.query.filter_by(account_id=a.id).first()
                if first_opp and first_opp.supplier_category:
                    category_lower = first_opp.supplier_category.lower()
                    if 'mir' in category_lower:
                        category = 'MiR'
                    elif 'ur' in category_lower or 'universal' in category_lower:
                        category = 'UR'
                    elif 'unitree' in category_lower:
                        category = 'Unitree'
                
                # Conta le opportunità
                opportunities_count = len(opp_id_list) if opp_id_list else 0
                
                accounts_data.append({
                    'id': a.id,
                    'name': a.name or '',
                    'address': address,
                    'billing_address': a.billing_address or '',
                    'shipping_address': a.shipping_address or '',
                    'has_pending_tasks': pending_tasks > 0,
                    'pending_tasks_count': pending_tasks,
                    'category': category,
                    'opportunities_count': opportunities_count
                })
            except Exception as e:
                # Se c'è un errore con un account, continua con gli altri
                print(f"Errore processando account {a.id}: {str(e)}")
                continue
        
        # Carica solo leads con indirizzo valido
        leads_list = Lead.query.filter(
            Lead.address.isnot(None)
        ).filter(
            db.func.length(Lead.address) > 3
        ).all()
        
        leads_data = []
        for l in leads_list:
            if not l.address or len(l.address.strip()) <= 3:
                continue
            
            # Determina categoria dal lead interest
            category = None
            if l.interest:
                interest_lower = l.interest.lower()
                if 'mir' in interest_lower:
                    category = 'MiR'
                elif 'ur' in interest_lower or 'universal' in interest_lower:
                    category = 'UR'
                elif 'unitree' in interest_lower:
                    category = 'Unitree'
            
            leads_data.append({
                'id': l.id,
                'name': l.company_name or '',
                'address': l.address,
                'category': category
            })
        
        return jsonify({
            'accounts': accounts_data,
            'leads': leads_data,
            'total': len(accounts_data) + len(leads_data)
        })
    except Exception as e:
        return jsonify({'error': f'Errore: {str(e)}'}), 500


@app.route('/api/accounts/<int:account_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_account(account_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    account = Account.query.get_or_404(account_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and account.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        # Verifica se ci sono opportunità collegate
        opp_count = Opportunity.query.filter_by(account_id=account_id).count()
        if opp_count > 0:
            return jsonify({
                'error': f'Impossibile eliminare: ci sono {opp_count} opportunità collegate',
                'opportunities_count': opp_count
            }), 400
        db.session.delete(account)
        db.session.commit()
        return jsonify({'success': True})
    
    # Rimuovo controllo owner - tutti possono vedere tutti gli accounts
    
    if request.method == 'GET':
        return jsonify({
            'account': {
                'id': account.id,
                'name': account.name,
                'industry': account.industry,
                'phone': account.phone,
                'website': account.website,
                'billing_address': account.billing_address,
                'shipping_address': account.shipping_address,
                'description': account.description
            }
        })
    
    # PUT method
    data = request.json
    if 'name' in data:
        account.name = data['name']
    if 'industry' in data:
        account.industry = data['industry']
    if 'phone' in data:
        account.phone = data['phone']
    if 'website' in data:
        account.website = data['website']
    if 'billing_address' in data:
        account.billing_address = data['billing_address']
    if 'shipping_address' in data:
        account.shipping_address = data['shipping_address']
    if 'description' in data:
        account.description = data['description']
    if 'created_at' in data:
        if data['created_at']:
            account.created_at = datetime.fromisoformat(data['created_at'].replace('Z', '+00:00'))
        else:
            account.created_at = datetime.utcnow()
    
    db.session.commit()
    return jsonify({'success': True, 'account_id': account.id})


@app.route('/api/opportunities', methods=['GET', 'POST'])
def opportunities():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Mostra tutte le opportunità (leads UR condivisi per tutti)
        # Filtri opzionali
        category_filter = request.args.get('category')
        account_id_filter = request.args.get('account_id')
        limit = request.args.get('limit', type=int)
        sort = request.args.get('sort', 'created_at')  # created_at, amount, stage
        
        # Usa SQL diretto per evitare problemi con colonne mancanti (heat_level)
        from sqlalchemy import text
        from datetime import datetime
        
        # Costruisci query SQL
        where_clauses = []
        params = {}
        
        if category_filter:
            where_clauses.append("supplier_category = :category_filter")
            params['category_filter'] = category_filter
        if account_id_filter:
            where_clauses.append("account_id = :account_id_filter")
            params['account_id_filter'] = account_id_filter
        
        where_sql = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        
        # Ordinamento
        order_by = "ORDER BY "
        if sort == 'amount':
            order_by += "amount DESC"
        elif sort == 'stage':
            order_by += "stage"
        else:  # default: created_at
            order_by += "created_at DESC"
        
        # Paginazione
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        offset = (page - 1) * per_page
        
        # Conta totale
        count_result = db.session.execute(text(f"""
            SELECT COUNT(*) FROM opportunities {where_sql}
        """), params).fetchone()
        total = count_result[0] if count_result else 0
        
        # Carica opportunità (include heat_level)
        opps_result = db.session.execute(text(f"""
            SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                   description, supplier_category, folder_path, lost_reason, owner_id, 
                   account_id, contact_id, lead_id, created_at, updated_at, heat_level
            FROM opportunities 
            {where_sql}
            {order_by}
            LIMIT :limit OFFSET :offset
        """), {**params, 'limit': per_page, 'offset': offset}).fetchall()
        
        # Costruisci oggetti Opportunity
        opps = []
        for row in opps_result:
            opp = Opportunity()
            opp.id = row[0]
            opp.name = row[1]
            opp.stage = row[2]
            opp.amount = row[3]
            opp.probability = row[4]
            opp.close_date = row[5]
            opp.opportunity_type = row[6]
            opp.description = row[7]
            opp.supplier_category = row[8]
            opp.folder_path = row[9]
            opp.lost_reason = row[10]
            opp.owner_id = row[11]
            opp.account_id = row[12]
            opp.contact_id = row[13]
            opp.lead_id = row[14]
            # Converti date
            if isinstance(row[15], str):
                opp.created_at = datetime.fromisoformat(row[15].replace('Z', '+00:00')) if row[15] else None
            else:
                opp.created_at = row[15]
            if isinstance(row[16], str):
                opp.updated_at = datetime.fromisoformat(row[16].replace('Z', '+00:00')) if row[16] else None
            else:
                opp.updated_at = row[16]
            # Carica heat_level (colonna 17)
            opp.heat_level = row[17] if len(row) > 17 else 'da_categorizzare'
            # Carica relazioni
            opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
            opp.account = db.session.query(Account).filter_by(id=opp.account_id).first() if opp.account_id else None
            opp.owner = db.session.query(User).filter_by(id=opp.owner_id).first() if opp.owner_id else None
            opps.append(opp)
        
        # Crea oggetto pagination fittizio
        class FakePagination:
            def __init__(self, total, page, per_page):
                self.total = total
                self.page = page
                self.per_page = per_page
                self.pages = (total + per_page - 1) // per_page if per_page > 0 else 1
                self.has_next = page < self.pages
                self.has_prev = page > 1
                self.next_num = page + 1 if self.has_next else None
                self.prev_num = page - 1 if self.has_prev else None
        
        pagination = FakePagination(total, page, per_page)
        
        return jsonify({
            'opportunities': [{
                'id': o.id,
                'name': o.name,
                'stage': o.stage,
                'amount': o.amount if o.amount and o.amount > 0 else sum(
                    (off.amount if off.amount and off.amount > 0 else extract_offer_amount(off))
                    for off in o.offers
                ),
                'created_at': o.created_at.isoformat() if o.created_at else None,
                'probability': o.probability,
                'close_date': o.close_date.isoformat() if o.close_date else None,
                'opportunity_type': o.opportunity_type,
                'description': o.description,
                'supplier_category': o.supplier_category,
                'heat_level': o.heat_level or 'da_categorizzare',
                'lost_reason': o.lost_reason,
                'account_id': o.account_id,
                'account_name': o.account.name if o.account else None,
                'account_image': o.account.image_url if o.account and o.account.image_url else None,
                'account_address': o.account.billing_address if o.account and o.account.billing_address else None,
                'contact_id': o.contact_id,
                'owner_id': o.owner_id,
                'owner_name': f"{o.owner.first_name} {o.owner.last_name}" if o.owner else 'Unknown',
                'is_my_opportunity': o.owner_id == user_id,
                'folder_path': o.folder_path,
                'offers': [{
                    'id': off.id,
                    'number': off.number,
                    'amount': off.amount if off.amount and off.amount > 0 else extract_offer_amount(off),
                    'pdf_path': off.pdf_path,
                    'docx_path': off.docx_path,
                    'status': off.status
                } for off in o.offers]
            } for o in opps],
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': pagination.total,
                'pages': pagination.pages,
                'has_next': pagination.has_next,
                'has_prev': pagination.has_prev,
                'next_page': pagination.next_num if pagination.has_next else None,
                'prev_page': pagination.prev_num if pagination.has_prev else None
            }
        })
    
    elif request.method == 'POST':
        data = request.json
        close_date = None
        if data.get('close_date'):
            close_date = datetime.fromisoformat(data['close_date'].replace('Z', '+00:00')).date()
        
        # Crea cartella per l'opportunità
        import platform
        if platform.system() == 'Windows':
            base_dir = r"K:\OFFERTE - K\OPPORTUNITA"
        else:
            base_dir = os.environ.get('OPPORTUNITIES_FOLDER', '/mnt/k/OFFERTE - K/OPPORTUNITA')
        
        if not os.path.exists(base_dir):
            base_dir = os.path.join(os.path.dirname(__file__), 'opportunita')
        
        os.makedirs(base_dir, exist_ok=True)
        
        # Nome cartella: OPP-{ID}-{NOME_OPPORTUNITA}
        opp_name_clean = "".join(c for c in data['name'] if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
        folder_name = f"OPP-{datetime.now().strftime('%Y%m%d')}-{opp_name_clean}"
        folder_path = os.path.join(base_dir, folder_name)
        
        # Crea la cartella
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        opp = Opportunity(
            name=data['name'],
            stage=data.get('stage', 'Qualification'),
            amount=float(data.get('amount', 0)),
            probability=int(data.get('probability', 10)),
            close_date=close_date,
            opportunity_type=data.get('opportunity_type'),
            description=data.get('description'),
            account_id=data.get('account_id'),
            contact_id=data.get('contact_id'),
            folder_path=folder_path,
            owner_id=user_id
        )
        db.session.add(opp)
        db.session.commit()
        return jsonify({'success': True, 'opportunity_id': opp.id, 'folder_path': folder_path})


@app.route('/api/opportunities/<int:opportunity_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_opportunity(opportunity_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = load_opportunity_safe(opportunity_id)
    if not opp:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and opp.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        db.session.delete(opp)
        db.session.commit()
        return jsonify({'success': True})
    
    # Rimuovo il controllo owner - tutti possono vedere tutte le opportunità
    
    if request.method == 'GET':
        return jsonify({
            'opportunity': {
                'id': opp.id,
                'name': opp.name,
                'stage': opp.stage,
                'amount': opp.amount,
                'probability': opp.probability,
                'close_date': opp.close_date.isoformat() if opp.close_date else None,
                'opportunity_type': opp.opportunity_type,
                'description': opp.description,
                'supplier_category': opp.supplier_category,
                'heat_level': opp.heat_level or 'da_categorizzare',
                'lost_reason': opp.lost_reason,
                'account_id': opp.account_id,
                'account_name': opp.account.name if opp.account else None,
                'account_image': opp.account.image_url if opp.account and opp.account.image_url else None,
                'contact_id': opp.contact_id,
                'owner_id': opp.owner_id,
                'owner_name': f"{opp.owner.first_name} {opp.owner.last_name}" if opp.owner else 'Unknown',
                'folder_path': opp.folder_path,
                'offers': [{
                    'id': off.id,
                    'number': off.number,
                    'amount': off.amount if off.amount and off.amount > 0 else extract_offer_amount(off),
                    'closed_amount': getattr(off, 'closed_amount', 0.0) or 0.0,
                    'margin': getattr(off, 'margin', 0.0) or 0.0,
                    'pdf_path': off.pdf_path,
                    'docx_path': off.docx_path,
                    'status': off.status,
                    'description': off.description,
                    'folder_path': off.folder_path
                } for off in opp.offers]
            }
        })
    
    # PUT method - Usa SQL diretto per aggiornare perché load_opportunity_safe restituisce oggetto non collegato al DB
    data = request.json
    from sqlalchemy import text
    
    # Costruisci query UPDATE dinamica
    update_fields = []
    update_params = {"opp_id": opportunity_id}
    
    if 'name' in data:
        update_fields.append("name = :name")
        update_params['name'] = data['name']
    if 'stage' in data:
        update_fields.append("stage = :stage")
        update_params['stage'] = data['stage']
    if 'amount' in data:
        update_fields.append("amount = :amount")
        update_params['amount'] = float(data['amount'])
    if 'probability' in data:
        update_fields.append("probability = :probability")
        update_params['probability'] = int(data['probability'])
    if 'close_date' in data:
        if data['close_date']:
            update_fields.append("close_date = :close_date")
            update_params['close_date'] = datetime.fromisoformat(data['close_date'].replace('Z', '+00:00')).date()
        else:
            update_fields.append("close_date = NULL")
    if 'opportunity_type' in data:
        update_fields.append("opportunity_type = :opportunity_type")
        update_params['opportunity_type'] = data['opportunity_type']
    if 'description' in data:
        update_fields.append("description = :description")
        update_params['description'] = data['description']
    if 'account_id' in data:
        update_fields.append("account_id = :account_id")
        update_params['account_id'] = data['account_id']
    if 'contact_id' in data:
        update_fields.append("contact_id = :contact_id")
        update_params['contact_id'] = data['contact_id']
    if 'supplier_category' in data:
        update_fields.append("supplier_category = :supplier_category")
        update_params['supplier_category'] = data['supplier_category']
    if 'heat_level' in data:
        # Verifica se la colonna heat_level esiste
        try:
            columns_check = db.session.execute(text("PRAGMA table_info(opportunities)")).fetchall()
            column_names = [col[1] for col in columns_check]
            if 'heat_level' in column_names:
                new_heat_level = data['heat_level']
                update_fields.append("heat_level = :heat_level")
                update_params['heat_level'] = new_heat_level
                
                # Aggiorna automaticamente la probabilità in base al heat_level
                # Fredde gelo: 0%, Fredde con speranza: 5%, Tiepide: 15%, Calde: 50%
                if new_heat_level == 'calda':
                    probability_value = 50
                elif new_heat_level == 'tiepida':
                    probability_value = 15
                elif new_heat_level == 'fredda_speranza':
                    probability_value = 5
                elif new_heat_level == 'fredda_gelo':
                    probability_value = 0
                else:  # da_categorizzare o altri
                    probability_value = 0
                
                # Aggiorna anche probability
                update_fields.append("probability = :probability")
                update_params['probability'] = probability_value
        except:
            pass  # Se non esiste, salta
    if 'folder_path' in data:
        update_fields.append("folder_path = :folder_path")
        update_params['folder_path'] = data['folder_path']
    if 'owner_id' in data:
        old_owner_id = opp.owner_id
        new_owner_id = data['owner_id']
        update_fields.append("owner_id = :owner_id")
        update_params['owner_id'] = new_owner_id
        
        # Crea notifica per il vecchio owner se diverso
        if old_owner_id and old_owner_id != new_owner_id:
            user = User.query.get(session['user_id'])
            create_notification(
                user_id=old_owner_id,
                title=f'Ownership opportunità cambiata',
                message=f'L\'opportunità "{opp.name}" è stata assegnata a {user.first_name if user else "altro utente"} {user.last_name if user else ""}',
                type='info',
                related_opportunity_id=opportunity_id
            )
    
    # Aggiorna sempre updated_at
    update_fields.append("updated_at = :updated_at")
    update_params['updated_at'] = datetime.utcnow()
    
    if update_fields:
        # Esegui UPDATE
        update_query = f"UPDATE opportunities SET {', '.join(update_fields)} WHERE id = :opp_id"
        db.session.execute(text(update_query), update_params)
        db.session.commit()
    
    return jsonify({'success': True, 'opportunity_id': opportunity_id})


@app.route('/api/leads/<int:lead_id>/take-ownership', methods=['POST'])
def take_lead_ownership(lead_id):
    """Prendi ownership di un lead"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead = Lead.query.get_or_404(lead_id)
    user_id = session['user_id']
    
    lead.owner_id = user_id
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': lead.owner_id,
        'owner_name': f"{lead.owner.first_name} {lead.owner.last_name}" if lead.owner else 'Unknown'
    })


@app.route('/api/accounts/<int:account_id>/take-ownership', methods=['POST'])
def take_account_ownership(account_id):
    """Prendi ownership di un account"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    account = Account.query.get_or_404(account_id)
    user_id = session['user_id']
    
    account.owner_id = user_id
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': account.owner_id,
        'owner_name': f"{account.owner.first_name} {account.owner.last_name}" if account.owner else 'Unknown'
    })


@app.route('/api/contacts', methods=['GET', 'POST'])
def contacts():
    """Gestisci contatti"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Contatti condivisi per tutti
        contacts_list = Contact.query.order_by(Contact.created_at.desc()).all()
        return jsonify({
            'contacts': [{
                'id': c.id,
                'first_name': c.first_name,
                'last_name': c.last_name,
                'title': c.title,
                'email': c.email,
                'phone': c.phone,
                'department': c.department,
                'description': c.description,
                'account_id': c.account_id,
                'account_name': c.account.name if c.account else None,
                'owner_id': c.owner_id,
                'owner_name': f"{c.owner.first_name} {c.owner.last_name}" if c.owner else 'Unknown'
            } for c in contacts_list]
        })
    
    elif request.method == 'POST':
        data = request.json
        contact = Contact(
            first_name=data['first_name'],
            last_name=data['last_name'],
            title=data.get('title'),
            email=data.get('email'),
            phone=data.get('phone'),
            department=data.get('department'),
            description=data.get('description'),
            account_id=data.get('account_id') if data.get('account_id') else None,
            owner_id=user_id
        )
        db.session.add(contact)
        db.session.commit()
        return jsonify({'success': True, 'contact_id': contact.id})


@app.route('/api/contacts/<int:contact_id>', methods=['GET', 'PUT', 'DELETE'])
def get_or_update_contact(contact_id):
    """Ottieni, aggiorna o elimina un contatto"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    contact = Contact.query.get_or_404(contact_id)
    
    if request.method == 'DELETE':
        # Solo admin o owner può eliminare
        if session.get('user_role') != 'admin' and contact.owner_id != session['user_id']:
            return jsonify({'error': 'Non autorizzato'}), 403
        db.session.delete(contact)
        db.session.commit()
        return jsonify({'success': True})
    
    if request.method == 'GET':
        return jsonify({
            'contact': {
                'id': contact.id,
                'first_name': contact.first_name,
                'last_name': contact.last_name,
                'title': contact.title,
                'email': contact.email,
                'phone': contact.phone,
                'department': contact.department,
                'description': contact.description,
                'account_id': contact.account_id,
                'account_name': contact.account.name if contact.account else None,
                'owner_id': contact.owner_id
            }
        })
    
    # PUT method
    data = request.json
    if 'first_name' in data:
        contact.first_name = data['first_name']
    if 'last_name' in data:
        contact.last_name = data['last_name']
    if 'title' in data:
        contact.title = data['title']
    if 'email' in data:
        contact.email = data['email']
    if 'phone' in data:
        contact.phone = data['phone']
    if 'department' in data:
        contact.department = data['department']
    if 'description' in data:
        contact.description = data['description']
    if 'account_id' in data:
        contact.account_id = data['account_id'] if data['account_id'] else None
    
    db.session.commit()
    return jsonify({'success': True, 'contact_id': contact.id})


@app.route('/api/accounts/search', methods=['GET'])
def search_accounts():
    """Cerca accounts per autocomplete"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    if len(query) < 2:
        return jsonify({'accounts': []})
    
    accounts = Account.query.filter(
        Account.name.ilike(f'%{query}%')
    ).limit(10).all()
    
    return jsonify({
        'accounts': [{
            'id': a.id,
            'name': a.name,
            'industry': a.industry
        } for a in accounts]
    })


@app.route('/api/accounts/search-all', methods=['GET'])
def search_all_accounts():
    """Cerca accounts su TUTTE le pagine nel database"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 50, type=int)
    
    # Costruisci query di ricerca su più campi
    search_query = Account.query
    
    if query:
        search_query = search_query.filter(
            db.or_(
                Account.name.ilike(f'%{query}%'),
                Account.industry.ilike(f'%{query}%'),
                Account.phone.ilike(f'%{query}%'),
                Account.email.ilike(f'%{query}%'),
                Account.billing_address.ilike(f'%{query}%'),
                Account.shipping_address.ilike(f'%{query}%')
            )
        )
    
    # Ordina per created_at desc (come nella lista normale)
    search_query = search_query.order_by(Account.created_at.desc())
    
    # Paginazione
    pagination = search_query.paginate(page=page, per_page=per_page, error_out=False)
    accounts_list = pagination.items
    
    # Prepara dati come nella lista normale
    accounts_data = []
    for a in accounts_list:
        # Conta opportunità
        opps_count = Opportunity.query.filter_by(account_id=a.id).count()
        
        accounts_data.append({
            'id': a.id,
            'name': a.name,
            'industry': a.industry,
            'phone': a.phone,
            'email': a.email,
            'billing_address': a.billing_address,
            'shipping_address': a.shipping_address,
            'created_at': a.created_at.isoformat() if a.created_at else None,
            'opportunities_count': opps_count,
            'opportunities': []  # Non carichiamo tutte le opportunità per performance
        })
    
    return jsonify({
        'accounts': accounts_data,
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': pagination.total,
            'pages': pagination.pages,
            'has_next': pagination.has_next,
            'has_prev': pagination.has_prev,
            'next_page': pagination.next_num if pagination.has_next else None,
            'prev_page': pagination.prev_num if pagination.has_prev else None
        }
    })


@app.route('/api/leads/search-all', methods=['GET'])
def search_all_leads():
    """Cerca leads su TUTTE le pagine nel database"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    query = request.args.get('q', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 50, type=int)
    
    # Costruisci query di ricerca su più campi
    search_query = Lead.query
    
    if query:
        search_query = search_query.filter(
            db.or_(
                Lead.company_name.ilike(f'%{query}%'),
                Lead.contact_first_name.ilike(f'%{query}%'),
                Lead.contact_last_name.ilike(f'%{query}%'),
                Lead.email.ilike(f'%{query}%'),
                Lead.phone.ilike(f'%{query}%'),
                Lead.address.ilike(f'%{query}%'),
                Lead.interest.ilike(f'%{query}%'),
                Lead.source.ilike(f'%{query}%')
            )
        )
    
    # Ordina per created_at desc (come nella lista normale)
    search_query = search_query.order_by(Lead.created_at.desc())
    
    # Paginazione
    pagination = search_query.paginate(page=page, per_page=per_page, error_out=False)
    leads_list = pagination.items
    
    # Prepara dati come nella lista normale
    leads_data = []
    for l in leads_list:
        owner = User.query.get(l.owner_id) if l.owner_id else None
        leads_data.append({
            'id': l.id,
            'company_name': l.company_name,
            'contact_first_name': l.contact_first_name,
            'contact_last_name': l.contact_last_name,
            'email': l.email,
            'phone': l.phone,
            'interest': l.interest,
            'source': l.source,
            'status': l.status,
            'rating': l.rating,
            'owner_id': l.owner_id,
            'owner_name': f"{owner.first_name} {owner.last_name}" if owner else 'Unknown',
            'is_my_lead': l.owner_id == user_id,
            'created_at': l.created_at.isoformat() if l.created_at else None,
            'last_contact_at': l.last_contact_at.isoformat() if l.last_contact_at else None,
            'address': l.address,
            'has_pending_tasks': False,
            'pending_tasks_count': 0
        })
    
    return jsonify({
        'leads': leads_data,
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': pagination.total,
            'pages': pagination.pages,
            'has_next': pagination.has_next,
            'has_prev': pagination.has_prev,
            'next_page': pagination.next_num if pagination.has_next else None,
            'prev_page': pagination.prev_num if pagination.has_prev else None
        }
    })


@app.route('/api/accounts/find-page', methods=['GET'])
def find_account_page():
    """Trova la pagina dove si trova un account specifico"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    account_id = request.args.get('id', type=int)
    per_page = request.args.get('per_page', 50, type=int)
    
    if not account_id:
        return jsonify({'error': 'ID mancante'}), 400
    
    # Conta quanti account ci sono prima di questo (ordinati per created_at desc)
    account = Account.query.get(account_id)
    if not account:
        return jsonify({'error': 'Account non trovato'}), 404
    
    # Conta gli account creati dopo questo (perché l'ordine è DESC)
    count_before = Account.query.filter(Account.created_at > account.created_at).count()
    # Se stesso timestamp, conta quelli con ID maggiore
    count_before += Account.query.filter(
        Account.created_at == account.created_at,
        Account.id > account_id
    ).count()
    
    page = (count_before // per_page) + 1
    
    return jsonify({
        'page': page,
        'account': {
            'id': account.id,
            'name': account.name
        }
    })


@app.route('/api/leads/find-page', methods=['GET'])
def find_lead_page():
    """Trova la pagina dove si trova un lead specifico"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    lead_id = request.args.get('id', type=int)
    per_page = request.args.get('per_page', 50, type=int)
    
    if not lead_id:
        return jsonify({'error': 'ID mancante'}), 400
    
    # Conta quanti lead ci sono prima di questo (ordinati per created_at desc)
    lead = Lead.query.get(lead_id)
    if not lead:
        return jsonify({'error': 'Lead non trovato'}), 404
    
    # Conta i lead creati dopo questo (perché l'ordine è DESC)
    count_before = Lead.query.filter(Lead.created_at > lead.created_at).count()
    # Se stesso timestamp, conta quelli con ID maggiore
    count_before += Lead.query.filter(
        Lead.created_at == lead.created_at,
        Lead.id > lead_id
    ).count()
    
    page = (count_before // per_page) + 1
    
    return jsonify({
        'page': page,
        'lead': {
            'id': lead.id,
            'company_name': lead.company_name
        }
    })


@app.route('/api/opportunities/search', methods=['GET'])
def search_opportunities():
    """Cerca opportunità per autocomplete"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    account_id = request.args.get('account_id')
    
    if len(query) < 2:
        return jsonify({'opportunities': []})
    
    opps_query = Opportunity.query.filter(
        Opportunity.name.ilike(f'%{query}%')
    )
    
    if account_id:
        opps_query = opps_query.filter(Opportunity.account_id == account_id)
    
    opps = opps_query.limit(10).all()
    
    return jsonify({
        'opportunities': [{
            'id': o.id,
            'name': o.name,
            'account_id': o.account_id,
            'account_name': o.account.name if o.account else None,
            'stage': o.stage
        } for o in opps]
    })


@app.route('/api/products/search', methods=['GET'])
def search_products():
    """Cerca prodotti per autocomplete"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    query = request.args.get('q', '').strip()
    if len(query) < 2:
        return jsonify({'products': []})
    
    products = Product.query.filter(
        Product.is_active == True
    ).filter(
        (Product.name.ilike(f'%{query}%')) | 
        (Product.code.ilike(f'%{query}%'))
    ).limit(20).all()
    
    return jsonify({
        'products': [{
            'id': p.id,
            'code': p.code,
            'name': p.name,
            'category': p.category
        } for p in products]
    })


@app.route('/api/opportunities/<int:opportunity_id>/chatter', methods=['GET', 'POST'])
def opportunity_chatter(opportunity_id):
    """API per chatter delle opportunità"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = load_opportunity_safe(opportunity_id)
    if not opp:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Ottieni tutti i post del chatter per questa opportunità
        posts = OpportunityChatter.query.filter_by(opportunity_id=opportunity_id).order_by(OpportunityChatter.created_at.desc()).all()
        return jsonify({
            'posts': [{
                'id': p.id,
                'message': p.message,
                'user_id': p.user_id,
                'user_name': f"{p.user.first_name} {p.user.last_name}" if p.user else 'Unknown',
                'created_at': p.created_at.isoformat() if p.created_at else None
            } for p in posts]
        })
    
    elif request.method == 'POST':
        # Crea un nuovo post
        data = request.json
        if not data.get('message'):
            return jsonify({'error': 'Messaggio richiesto'}), 400
        
        post = OpportunityChatter(
            opportunity_id=opportunity_id,
            user_id=user_id,
            message=data['message']
        )
        db.session.add(post)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'post': {
                'id': post.id,
                'message': post.message,
                'user_id': post.user_id,
                'user_name': f"{post.user.first_name} {post.user.last_name}" if post.user else 'Unknown',
                'created_at': post.created_at.isoformat() if post.created_at else None
            }
        })


@app.route('/api/opportunities/<int:opportunity_id>/tasks', methods=['GET', 'POST'])
def opportunity_tasks(opportunity_id):
    """API per tasks delle opportunità"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = load_opportunity_safe(opportunity_id)
    if not opp:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Ottieni tutti i task per questa opportunità
        tasks = OpportunityTask.query.filter_by(opportunity_id=opportunity_id).order_by(OpportunityTask.created_at.desc()).all()
        return jsonify({
            'tasks': [{
                'id': t.id,
                'task_type': t.task_type,
                'description': t.description,
                'assigned_to_id': t.assigned_to_id,
                'assigned_to_name': f"{t.assigned_to.first_name} {t.assigned_to.last_name}" if t.assigned_to else None,
                'due_date': t.due_date.isoformat() if t.due_date else None,
                'is_completed': t.is_completed,
                'completed_at': t.completed_at.isoformat() if t.completed_at else None,
                'completed_by_name': f"{t.completed_by.first_name} {t.completed_by.last_name}" if t.completed_by else None,
                'created_at': t.created_at.isoformat() if t.created_at else None
            } for t in tasks]
        })
    
    elif request.method == 'POST':
        # Crea un nuovo task
        data = request.json
        if not data.get('task_type'):
            return jsonify({'error': 'Tipo task richiesto'}), 400
        
        due_date = None
        if data.get('due_date'):
            due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
        
        # Ottieni informazioni opportunità per email
        opportunity = Opportunity.query.get(opportunity_id)
        opportunity_name = opportunity.name if opportunity else "Opportunità"
        
        task = OpportunityTask(
            opportunity_id=opportunity_id,
            task_type=data['task_type'],
            description=data.get('description', ''),
            assigned_to_id=data.get('assigned_to_id'),
            due_date=due_date
        )
        db.session.add(task)
        db.session.commit()
        
        # Invia email agli account commerciali se il task è creato manualmente
        try:
            send_task_notification_email(task, opportunity_name, opportunity)
        except Exception as e:
            print(f"Errore invio email task: {str(e)}")
            # Non bloccare la creazione del task se l'email fallisce
        
        return jsonify({
            'success': True,
            'task': {
                'id': task.id,
                'task_type': task.task_type,
                'description': task.description,
                'assigned_to_id': task.assigned_to_id,
                'assigned_to_name': f"{task.assigned_to.first_name} {task.assigned_to.last_name}" if task.assigned_to else None,
                'due_date': task.due_date.isoformat() if task.due_date else None,
                'is_completed': task.is_completed
            }
        })


@app.route('/api/opportunities/<int:opportunity_id>/tasks/<int:task_id>', methods=['PUT', 'DELETE'])
def update_or_delete_task(opportunity_id, task_id):
    """Aggiorna o elimina un task"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    task = OpportunityTask.query.filter_by(id=task_id, opportunity_id=opportunity_id).first_or_404()
    user_id = session['user_id']
    
    if request.method == 'PUT':
        data = request.json
        if 'is_completed' in data:
            task.is_completed = data['is_completed']
            if data['is_completed']:
                task.completed_at = datetime.utcnow()
                task.completed_by_id = user_id
            else:
                task.completed_at = None
                task.completed_by_id = None
        if 'assigned_to_id' in data:
            task.assigned_to_id = data['assigned_to_id']
        if 'description' in data:
            task.description = data['description']
        if 'due_date' in data:
            if data['due_date']:
                task.due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
            else:
                task.due_date = None
        
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        db.session.delete(task)
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/my-tasks', methods=['GET'])
def get_my_tasks():
    """Ottieni tutte le task dell'utente corrente (globali e legate a opportunità)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    # Task globali (opportunity_id IS NULL) assegnate all'utente
    # Usa query raw per evitare errori se account_id non esiste nel database
    try:
        global_tasks = db.session.execute(
            text("""
                SELECT id, opportunity_id, task_type, description, assigned_to_id, 
                       due_date, is_completed, completed_at, completed_by_id, 
                       created_at, updated_at
                FROM opportunity_tasks
                WHERE opportunity_id IS NULL AND assigned_to_id = :user_id
                ORDER BY due_date ASC, created_at DESC
            """),
            {'user_id': user_id}
        ).fetchall()
        # Converti i risultati in oggetti OpportunityTask
        global_tasks = [OpportunityTask.query.get(row[0]) for row in global_tasks if row[0]]
    except Exception as e:
        # Se fallisce, prova query normale (potrebbe funzionare se account_id esiste)
        try:
            global_tasks = OpportunityTask.query.filter(
                OpportunityTask.opportunity_id.is_(None),
                OpportunityTask.assigned_to_id == user_id
            ).order_by(OpportunityTask.due_date.asc(), OpportunityTask.created_at.desc()).all()
        except:
            global_tasks = []
    
    # Task legate a opportunità assegnate all'utente
    try:
        opp_tasks = db.session.execute(
            text("""
                SELECT id, opportunity_id, task_type, description, assigned_to_id, 
                       due_date, is_completed, completed_at, completed_by_id, 
                       created_at, updated_at
                FROM opportunity_tasks
                WHERE opportunity_id IS NOT NULL AND assigned_to_id = :user_id
                ORDER BY due_date ASC, created_at DESC
            """),
            {'user_id': user_id}
        ).fetchall()
        # Converti i risultati in oggetti OpportunityTask
        opp_tasks = [OpportunityTask.query.get(row[0]) for row in opp_tasks if row[0]]
    except Exception as e:
        # Se fallisce, prova query normale
        try:
            opp_tasks = OpportunityTask.query.filter(
                OpportunityTask.opportunity_id.isnot(None),
                OpportunityTask.assigned_to_id == user_id
            ).order_by(OpportunityTask.due_date.asc(), OpportunityTask.created_at.desc()).all()
        except:
            opp_tasks = []
    
    def task_to_dict(t):
        # Gestisci account_id se non esiste nel database
        account_id = None
        account_name = None
        account_image_url = None
        
        try:
            account_id = t.account_id if hasattr(t, 'account_id') else None
            if hasattr(t, 'account') and t.account:
                account_name = t.account.name
                account_image_url = getattr(t.account, 'image_url', None)
        except:
            pass
        
        # Fallback su opportunity.account
        if not account_name and t.opportunity and t.opportunity.account:
            account_name = t.opportunity.account.name
            account_image_url = getattr(t.opportunity.account, 'image_url', None)
        
        return {
            'id': t.id,
            'task_type': t.task_type,
            'description': t.description,
            'assigned_to_id': t.assigned_to_id,
            'assigned_to_name': f"{t.assigned_to.first_name} {t.assigned_to.last_name}" if t.assigned_to else None,
            'due_date': t.due_date.isoformat() if t.due_date else None,
            'is_completed': t.is_completed,
            'completed_at': t.completed_at.isoformat() if t.completed_at else None,
            'completed_by_name': f"{t.completed_by.first_name} {t.completed_by.last_name}" if t.completed_by else None,
            'created_at': t.created_at.isoformat() if t.created_at else None,
            'opportunity_id': t.opportunity_id,
            'opportunity_name': t.opportunity.name if t.opportunity else None,
            'account_id': account_id,
            'account_name': account_name,
            'account_image_url': account_image_url,
            'opportunity_offers': [{'id': o.id, 'number': o.number, 'pdf_path': o.pdf_path, 'docx_path': o.docx_path, 'amount': (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))} for o in t.opportunity.offers] if t.opportunity and t.opportunity.offers else []
        }
    
    return jsonify({
        'global_tasks': [task_to_dict(t) for t in global_tasks],
        'opportunity_tasks': [task_to_dict(t) for t in opp_tasks],
        'total': len(global_tasks) + len(opp_tasks)
    })


@app.route('/api/tasks', methods=['POST'])
def create_global_task():
    """Crea una task globale (non legata a opportunità)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    if not data.get('task_type'):
        return jsonify({'error': 'Tipo task richiesto'}), 400
    
    due_date = None
    if data.get('due_date'):
        due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
    
    # Ottieni informazioni opportunità se presente
    opportunity_name = None
    opportunity = None
    if data.get('opportunity_id'):
        opportunity = Opportunity.query.get(data['opportunity_id'])
        opportunity_name = opportunity.name if opportunity else None
    
    task = OpportunityTask(
        opportunity_id=data.get('opportunity_id'),  # Può essere None per task globale
        account_id=data.get('account_id'),  # Collegamento diretto ad account
        task_type=data['task_type'],
        description=data.get('description', ''),
        assigned_to_id=data.get('assigned_to_id'),
        due_date=due_date
    )
    db.session.add(task)
    db.session.commit()
    
    # Invia email agli account commerciali se il task è creato manualmente
    try:
        send_task_notification_email(task, opportunity_name or "Task globale", opportunity)
    except Exception as e:
        print(f"Errore invio email task: {str(e)}")
        # Non bloccare la creazione del task se l'email fallisce
    
    return jsonify({
        'success': True,
        'task': {
            'id': task.id,
            'task_type': task.task_type,
            'description': task.description,
            'assigned_to_id': task.assigned_to_id,
            'assigned_to_name': f"{task.assigned_to.first_name} {task.assigned_to.last_name}" if task.assigned_to else None,
            'due_date': task.due_date.isoformat() if task.due_date else None,
            'is_completed': task.is_completed
        }
    })


@app.route('/api/tasks/all', methods=['GET'])
def get_all_tasks():
    """Ottieni tutte le task di tutti gli operatori (per vista per operatore)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Solo admin può vedere tutte le attività
    if session.get('user_role') != 'admin':
        return jsonify({'error': 'Non autorizzato'}), 403
    
    all_tasks = OpportunityTask.query.order_by(
        OpportunityTask.assigned_to_id,
        OpportunityTask.due_date.asc(),
        OpportunityTask.created_at.desc()
    ).all()
    
    def task_to_dict(t):
        return {
            'id': t.id,
            'task_type': t.task_type,
            'description': t.description,
            'assigned_to_id': t.assigned_to_id,
            'assigned_to_name': f"{t.assigned_to.first_name} {t.assigned_to.last_name}" if t.assigned_to else 'Non assegnato',
            'due_date': t.due_date.isoformat() if t.due_date else None,
            'is_completed': t.is_completed,
            'completed_at': t.completed_at.isoformat() if t.completed_at else None,
            'opportunity_id': t.opportunity_id,
            'opportunity_name': t.opportunity.name if t.opportunity else None,
            'account_id': t.account_id,
            'account_name': (t.account.name if t.account else None) or (t.opportunity.account.name if t.opportunity and t.opportunity.account else None),
            'account_image_url': (t.account.image_url if t.account and t.account.image_url else None) or (t.opportunity.account.image_url if t.opportunity and t.opportunity.account and t.opportunity.account.image_url else None),
            'opportunity_offers': [{'id': o.id, 'number': o.number, 'pdf_path': o.pdf_path, 'docx_path': o.docx_path, 'amount': (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))} for o in t.opportunity.offers] if t.opportunity and t.opportunity.offers else []
        }
    
    return jsonify({
        'tasks': [task_to_dict(t) for t in all_tasks]
    })


@app.route('/api/tasks/<int:task_id>', methods=['PUT', 'DELETE'])
def update_or_delete_global_task(task_id):
    """Aggiorna o elimina una task (globale o legata a opportunità)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    task = OpportunityTask.query.get_or_404(task_id)
    user_id = session['user_id']
    
    if request.method == 'PUT':
        data = request.json
        if 'is_completed' in data:
            task.is_completed = data['is_completed']
            if data['is_completed']:
                task.completed_at = datetime.utcnow()
                task.completed_by_id = user_id
            else:
                task.completed_at = None
                task.completed_by_id = None
        if 'assigned_to_id' in data:
            task.assigned_to_id = data['assigned_to_id']
        if 'description' in data:
            task.description = data['description']
        if 'due_date' in data:
            if data['due_date']:
                task.due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00')).date()
            else:
                task.due_date = None
        
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        db.session.delete(task)
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/opportunities/<int:opportunity_id>/open-folder', methods=['POST'])
def open_opportunity_folder(opportunity_id):
    """Restituisce percorso SMB per aprire la cartella dell'opportunità su Windows"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opp = load_opportunity_safe(opportunity_id)
    if not opp:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    
    # Recupera configurazione server per accesso file
    DEFAULT_SERVER_IP = "192.168.10.240"
    DEFAULT_SERVER_SHARE = "commerciale"
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_local_pc_ip = user.local_pc_ip if user and user.local_pc_ip else DEFAULT_SERVER_IP
    user_local_pc_share = user.local_pc_share if user and user.local_pc_share else DEFAULT_SERVER_SHARE
    user_local_pc_base_path = user.local_pc_base_path if user and user.local_pc_base_path else 'OFFERTE - K'
    
    # Cerca cartella nelle offerte associate (priorità)
    folder_path = None
    if opp.offers:
        for offer in opp.offers:
            if offer.folder_path:
                folder_path = offer.folder_path
                break
    
    # Se non trovata nelle offerte, usa folder_path dell'opportunità
    if not folder_path:
        folder_path = opp.folder_path
    
    if not folder_path:
        return jsonify({'error': 'Nessuna cartella associata a questa opportunità'}), 404
    
    # Converti percorso Windows (K:\) in percorso UNC/SMB
    relative_path = folder_path
    if folder_path.startswith('K:'):
        relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
    
    # Costruisci percorso completo
    base_path = user_local_pc_base_path.replace('\\', '/').strip('/')
    relative_path = relative_path.replace('\\', '/').strip('/')
    
    if not relative_path.startswith(base_path):
        full_path = f"{base_path}/{relative_path}" if base_path else relative_path
    else:
        full_path = relative_path
    
    full_path = full_path.replace('//', '/').replace('\\', '/')
    
    # Genera link SMB per Windows
    smb_path = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}\\{full_path}"
    smb_path = smb_path.replace('/', '\\')
    
    return jsonify({
        'success': True,
        'folder_path': folder_path,
        'smb_path': smb_path
    })

@app.route('/api/opportunities/<int:opportunity_id>/files', methods=['GET'])
def get_opportunity_files(opportunity_id):
    """Ottiene la lista dei file nella cartella dell'opportunità e nelle cartelle offerte associate"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Usa query SQL diretta per evitare problemi con colonne mancanti
    from sqlalchemy import text
    
    # Recupera opportunità usando SQL diretto
    opp_result = db.session.execute(text("""
        SELECT id, name, folder_path, account_id 
        FROM opportunities 
        WHERE id = :opp_id
    """), {"opp_id": opportunity_id}).fetchone()
    
    if not opp_result:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    
    opp_id, opp_name, opp_folder_path, opp_account_id = opp_result
    
    # Inizializza variabili per debug e file
    files_list = []
    folders_checked = []
    debug_info = []
    
    # Crea oggetto fittizio per compatibilità
    opp = type('Opportunity', (), {
        'id': opp_id,
        'name': opp_name,
        'folder_path': opp_folder_path,
        'account_id': opp_account_id,
        'offers': []
    })()
    
    # Recupera offerte associate usando SQL diretto
    offers_result = db.session.execute(text("""
        SELECT id, number, folder_path, pdf_path, docx_path
        FROM offers 
        WHERE opportunity_id = :opp_id
    """), {"opp_id": opportunity_id}).fetchall()
    
    # Crea oggetti fittizi per le offerte
    offers_list = []
    for offer_row in offers_result:
        offer_obj = type('OfferDocument', (), {
            'id': offer_row[0],
            'number': offer_row[1],
            'folder_path': offer_row[2],
            'pdf_path': offer_row[3],
            'docx_path': offer_row[4]
        })()
        offers_list.append(offer_obj)
    
    # Se non ci sono offerte associate, cerca automaticamente l'offerta con il numero nel nome
    if not offers_list and opp_name:
        import re
        match = re.search(r'K(\d{4})-(\d{2})', opp_name)
        if match:
            offer_number = match.group(0)  # es. "K0800-25"
            # Cerca offerta con questo numero
            matching_offer = db.session.execute(text("""
                SELECT id, number, folder_path, pdf_path, docx_path, opportunity_id
                FROM offers 
                WHERE number = :offer_number
                LIMIT 1
            """), {"offer_number": offer_number}).fetchone()
            
            if matching_offer:
                # Se l'offerta esiste ma è associata a un'altra opportunità, associamola a questa
                if matching_offer[5] != opportunity_id:
                    db.session.execute(text("""
                        UPDATE offers 
                        SET opportunity_id = :opp_id
                        WHERE id = :offer_id
                    """), {
                        "opp_id": opportunity_id,
                        "offer_id": matching_offer[0]
                    })
                    db.session.commit()
                
                # Aggiungi l'offerta alla lista
                offer_obj = type('OfferDocument', (), {
                    'id': matching_offer[0],
                    'number': matching_offer[1],
                    'folder_path': matching_offer[2],
                    'pdf_path': matching_offer[3],
                    'docx_path': matching_offer[4]
                })()
                offers_list.append(offer_obj)
    
    opp.offers = offers_list
    
    # IMPORTANTE: NON cercare file nella cartella dell'opportunità se è la cartella principale!
    # I file devono essere solo nelle cartelle specifiche delle offerte associate
    # La cartella dell'opportunità è solo per documenti generali dell'opportunità, non per le offerte
    # Se l'opportunità ha una cartella, verifichiamo che non sia la cartella principale delle offerte
    if opp.folder_path:
        folder_path = opp.folder_path
        folder_basename = os.path.basename(folder_path)
        
        # SKIP se è la cartella principale delle offerte
        if folder_basename not in ['OFFERTE 2025 - K', 'OFFERTE 2026 - K']:
            # Normalizza il percorso
            if not os.path.isabs(folder_path):
                # Se è relativo, prova a costruirlo
                import platform
                if platform.system() == 'Windows':
                    base_dir = r"K:\OFFERTE - K\OPPORTUNITA"
                else:
                    base_dir = os.environ.get('OPPORTUNITIES_FOLDER', '/mnt/k/OFFERTE - K/OPPORTUNITA')
                folder_path = os.path.join(base_dir, os.path.basename(folder_path))
            
            if os.path.exists(folder_path):
                folders_checked.append(folder_path)
                try:
                    file_count_opp = 0
                    for item in os.listdir(folder_path):
                        item_path = os.path.join(folder_path, item)
                        if os.path.isfile(item_path) and item.lower().endswith(('.pdf', '.docx', '.doc')):
                            file_size = os.path.getsize(item_path)
                            file_ext = os.path.splitext(item)[1].lower()
                            files_list.append({
                                'name': item,
                                'path': item_path,
                                'size': file_size,
                                'extension': file_ext,
                                'type': 'opportunity_folder'
                            })
                            file_count_opp += 1
                except Exception as e:
                    pass
    
    # Recupera configurazione server per accesso file (IP server condiviso)
    # Valori di default: IP=192.168.10.240, Share=commerciale (stessi per tutti i PC della subnet)
    DEFAULT_SERVER_IP = "192.168.10.240"
    DEFAULT_SERVER_SHARE = "commerciale"
    
    user_has_local_config = False
    user_local_pc_ip = None
    user_local_pc_share = None
    user_local_pc_base_path = None
    
    user_id_from_session = session.get('user_id')
    
    try:
        # Prima verifica se le colonne esistono
        columns_check = db.session.execute(text("PRAGMA table_info(users)")).fetchall()
        column_names = [col[1] for col in columns_check]
        has_local_pc_columns = 'local_pc_ip' in column_names and 'local_pc_share' in column_names
        
        if not has_local_pc_columns:
            user_local_pc_ip = DEFAULT_SERVER_IP
            user_local_pc_share = DEFAULT_SERVER_SHARE
            user_has_local_config = True
        else:
            user_result = db.session.execute(text("""
                SELECT id, local_pc_ip, local_pc_share, local_pc_base_path
                FROM users 
                WHERE id = :user_id
            """), {"user_id": user_id_from_session}).fetchone()
            
            if user_result:
                user_id_db, user_local_pc_ip, user_local_pc_share, user_local_pc_base_path = user_result
                if not user_local_pc_ip or not user_local_pc_share:
                    user_local_pc_ip = user_local_pc_ip or DEFAULT_SERVER_IP
                    user_local_pc_share = user_local_pc_share or DEFAULT_SERVER_SHARE
                user_has_local_config = True
            else:
                user_local_pc_ip = DEFAULT_SERVER_IP
                user_local_pc_share = DEFAULT_SERVER_SHARE
                user_has_local_config = True
    except Exception as e:
        user_local_pc_ip = DEFAULT_SERVER_IP
        user_local_pc_share = DEFAULT_SERVER_SHARE
        user_has_local_config = True
    
    if not user_has_local_config:
        user_local_pc_ip = DEFAULT_SERVER_IP
        user_local_pc_share = DEFAULT_SERVER_SHARE
        user_has_local_config = True
    
    # Per ogni offerta associata, cerca SOLO nella sua cartella specifica
    for offer in offers_list:
        if not offer.folder_path:
            continue
        
        folder_path = offer.folder_path
        
        # IMPORTANTE: Se la cartella è la cartella principale, NON mostrare file!
        folder_basename = os.path.basename(folder_path)
        if folder_basename in ['OFFERTE 2025 - K', 'OFFERTE 2026 - K']:
            continue
        
        # Costruisci il percorso usando la configurazione del PC locale dell'utente
        import platform
        is_windows = platform.system() == 'Windows'
        actual_folder_path = None
        
        if folder_path.startswith('K:') or folder_path.startswith('K:\\'):
            # Percorso Windows K:\... 
            if is_windows:
                if os.path.exists(folder_path):
                    actual_folder_path = folder_path.replace('/', '\\')
                elif user_has_local_config:
                    relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                    share_name = user_local_pc_share
                    actual_folder_path = f"\\\\{user_local_pc_ip}\\{share_name}\\{relative_path}".replace('/', '\\')
            else:
                # Su Linux (AWS), prova più percorsi possibili
                relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                folder_name = os.path.basename(relative_path)  # Nome cartella specifica (es: "K1650-25 - VRS GAS SRL...")
                
                # Lista di percorsi possibili da provare sulla macchina AWS
                possible_paths = []
                
                # 1. Se l'utente ha configurazione PC locale, prova quella
                if user_has_local_config:
                    possible_paths.append(f"/mnt/{user_local_pc_share}/{relative_path}")
                
                # 2. Prova cartelle locali sulla macchina AWS
                base_dir = os.path.dirname(os.path.abspath(__file__))
                
                # Cartelle principali da cercare
                main_folders = [
                    os.path.join(base_dir, 'offerte_2025_backup'),
                    os.path.join(base_dir, 'offerte_2026_backup'),
                    os.path.join(base_dir, 'offerte_generate'),
                    '/mnt/k/OFFERTE - K/OFFERTE 2025 - K',
                    '/mnt/k/OFFERTE - K/OFFERTE 2026 - K',
                    os.environ.get('OFFERS_FOLDER', ''),
                    '/mnt/k',
                    '/mnt/share'
                ]
                
                # Rimuovi percorsi vuoti
                main_folders = [f for f in main_folders if f]
                
                # Per ogni cartella principale, prova:
                # a) Percorso completo relativo
                # b) Cerca ricorsivamente la cartella specifica per nome
                for main_folder in main_folders:
                    if not os.path.exists(main_folder):
                        continue
                    
                    # Prova percorso completo
                    full_path = os.path.join(main_folder, relative_path)
                    if os.path.exists(full_path):
                        possible_paths.append(full_path)
                    
                    # Cerca ricorsivamente la cartella specifica (aumenta profondità per trovare meglio)
                    try:
                        # Prima cerca direttamente nella cartella principale
                        if folder_name:
                            direct_path = os.path.join(main_folder, folder_name)
                            if os.path.exists(direct_path):
                                possible_paths.append(direct_path)
                        
                        # Cerca anche usando solo il numero offerta (es. "K1470-25")
                        if offer.number:
                            # Rimuovi eventuali prefissi/suffissi dal numero
                            clean_number = offer.number.replace('FILE-', '').strip()
                            if clean_number:
                                # Cerca cartelle che contengono questo numero
                                if os.path.isdir(main_folder):
                                    try:
                                        for item in os.listdir(main_folder):
                                            item_path = os.path.join(main_folder, item)
                                            if os.path.isdir(item_path):
                                                # Verifica se il nome contiene il numero offerta
                                                item_upper = item.upper()
                                                if clean_number.upper() in item_upper or folder_name and folder_name.lower() in item.lower():
                                                    possible_paths.append(item_path)
                                                
                                                # Cerca anche nelle sottocartelle (2 livelli per trovare meglio)
                                                try:
                                                    for sub_item in os.listdir(item_path):
                                                        sub_item_path = os.path.join(item_path, sub_item)
                                                        if os.path.isdir(sub_item_path):
                                                            sub_item_upper = sub_item.upper()
                                                            if clean_number.upper() in sub_item_upper or (folder_name and folder_name.lower() in sub_item.lower()):
                                                                possible_paths.append(sub_item_path)
                                                            
                                                            # Cerca anche nel terzo livello
                                                            try:
                                                                for sub_sub_item in os.listdir(sub_item_path):
                                                                    sub_sub_item_path = os.path.join(sub_item_path, sub_sub_item)
                                                                    if os.path.isdir(sub_sub_item_path):
                                                                        sub_sub_item_upper = sub_sub_item.upper()
                                                                        if clean_number.upper() in sub_sub_item_upper or (folder_name and folder_name.lower() in sub_sub_item.lower()):
                                                                            possible_paths.append(sub_sub_item_path)
                                                            except Exception:
                                                                pass
                                                except Exception:
                                                    pass
                                    except Exception:
                                        pass
                    except Exception:
                        pass
                
                # Rimuovi duplicati mantenendo ordine
                seen = set()
                unique_paths = []
                for p in possible_paths:
                    if p not in seen:
                        seen.add(p)
                        unique_paths.append(p)
                possible_paths = unique_paths
                
                # Prova ogni percorso fino a trovarne uno valido
                actual_folder_path = None
                for test_path in possible_paths:
                    if test_path and os.path.exists(test_path) and os.path.isdir(test_path):
                        actual_folder_path = test_path
                        break
        else:
            # Percorso già assoluto o relativo - usa così com'è
            actual_folder_path = folder_path
        
        # Verifica se la cartella esiste e leggi i file
        if actual_folder_path and os.path.exists(actual_folder_path):
            folders_checked.append(actual_folder_path)
            try:
                for item in os.listdir(actual_folder_path):
                    item_path = os.path.join(actual_folder_path, item)
                    if os.path.isfile(item_path) and item.lower().endswith(('.pdf', '.docx', '.doc')):
                        if not any(f['name'] == item and f['path'] == item_path for f in files_list):
                            file_size = os.path.getsize(item_path)
                            file_ext = os.path.splitext(item)[1].lower()
                            
                            # IMPORTANTE: Su Linux, genera SEMPRE link SMB anche se troviamo file locali
                            # perché i file reali sono sul server Windows
                            file_entry = {
                                'name': item,
                                'path': item_path,
                                'size': file_size,
                                'extension': file_ext,
                                'type': 'offer_folder',
                                'offer_number': offer.number,
                                'folder': os.path.basename(actual_folder_path)
                            }
                            
                            # Se siamo su Linux e il folder_path inizia con K:, genera link SMB
                            if not is_windows and folder_path.startswith('K:'):
                                relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                                smb_base = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}"
                                file_relative = os.path.join(relative_path, item).replace('\\', '/')
                                file_entry['smb_path'] = f"{smb_base}\\{file_relative}".replace('/', '\\')
                            
                            files_list.append(file_entry)
            except Exception:
                pass
    
    # IMPORTANTE: Su Linux (AWS), genera SEMPRE link SMB anche se troviamo file locali,
    # perché i file reali sono sul server Windows e devono essere aperti via SMB
    # Su Windows, genera link SMB solo se non troviamo file fisicamente
    import platform
    is_windows = platform.system() == 'Windows'
    generate_smb_links = not is_windows or (not files_list and offers_list)
    
    # Se non abbiamo trovato file fisicamente ma abbiamo offerte associate con folder_path,
    # OPPURE siamo su Linux (AWS) e abbiamo offerte associate,
    # genera link SMB usando i folder_path delle offerte per aprire i file dal PC locale
    if generate_smb_links and offers_list:
        import platform
        is_windows = platform.system() == 'Windows'
        
        for offer in offers_list:
            if not offer.folder_path:
                continue
            
            folder_path = offer.folder_path
            
            # Skip cartelle principali
            folder_basename = os.path.basename(folder_path)
            if folder_basename in ['OFFERTE 2025 - K', 'OFFERTE 2026 - K']:
                continue
            
            # Se il percorso inizia con K:, genera link SMB
            if folder_path.startswith('K:') or folder_path.startswith('K:\\'):
                relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                
                # Costruisci percorso SMB base
                smb_base = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}"
                
                # Prova a leggere file dalla cartella via UNC (solo su Windows)
                if is_windows:
                    try:
                        unc_path = f"{smb_base}\\{relative_path}".replace('/', '\\')
                        if os.path.exists(unc_path):
                            for item in os.listdir(unc_path):
                                item_path = os.path.join(unc_path, item)
                                if os.path.isfile(item_path) and item.lower().endswith(('.pdf', '.docx', '.doc')):
                                    file_size = os.path.getsize(item_path)
                                    file_ext = os.path.splitext(item)[1].lower()
                                    files_list.append({
                                        'name': item,
                                        'path': item_path,
                                        'size': file_size,
                                        'extension': file_ext,
                                        'type': 'offer_folder',
                                        'offer_number': offer.number,
                                        'folder': os.path.basename(unc_path),
                                        'smb_path': f"{smb_base}\\{relative_path}\\{item}".replace('/', '\\')
                                    })
                    except Exception:
                        pass
                
                # IMPORTANTE: Su Linux, aggiungi SEMPRE link SMB anche se abbiamo già file locali
                # Per Windows, aggiungi solo se non abbiamo file
                files_for_offer = [f for f in files_list if f.get('offer_number') == offer.number]
                should_add_smb = not is_windows or len(files_for_offer) == 0
                
                # Funzione helper per verificare se un file esiste già nella lista
                def file_exists_in_list(file_name):
                    base_name = os.path.basename(file_name).lower()
                    for f in files_for_offer:
                        existing_name = os.path.basename(f.get('name', '')).lower()
                        if base_name == existing_name:
                            return True
                    return False
                
                # Se ancora non abbiamo file OPPURE siamo su Linux, aggiungi i file PDF/DOCX direttamente collegati
                if should_add_smb:
                    # Aggiungi PDF se presente e non già nella lista
                    if offer.pdf_path:
                        pdf_name = os.path.basename(offer.pdf_path)
                        if not file_exists_in_list(pdf_name):
                            pdf_relative = offer.pdf_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                            pdf_smb = f"{smb_base}\\{pdf_relative}".replace('/', '\\')
                            files_list.append({
                                'name': pdf_name,
                                'path': pdf_smb,
                                'size': 0,
                                'extension': '.pdf',
                                'type': 'offer_file',
                                'offer_number': offer.number,
                                'smb_path': pdf_smb
                            })
                    
                    # Aggiungi DOCX se presente e non già nella lista
                    if offer.docx_path:
                        docx_name = os.path.basename(offer.docx_path)
                        if not file_exists_in_list(docx_name):
                            docx_relative = offer.docx_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                            docx_smb = f"{smb_base}\\{docx_relative}".replace('/', '\\')
                            files_list.append({
                                'name': docx_name,
                                'path': docx_smb,
                                'size': 0,
                                'extension': '.docx',
                                'type': 'offer_file',
                                'offer_number': offer.number,
                                'smb_path': docx_smb
                            })
                    
                    # Se ancora non abbiamo file, genera link alla cartella
                    # IMPORTANTE: Su Linux, genera sempre link anche se abbiamo trovato file locali
                    if not is_windows or not any(f.get('offer_number') == offer.number for f in files_list):
                        folder_smb = f"{smb_base}\\{relative_path}".replace('/', '\\')
                        # Se siamo su Linux e abbiamo già file locali, aggiungi comunque il link SMB
                        # ma con un nome diverso per distinguerlo
                        if not is_windows and any(f.get('offer_number') == offer.number for f in files_list):
                            link_name = f'🔗 Apri cartella SMB {offer.number}'
                        else:
                            link_name = f'📁 Apri cartella {offer.number}'
                        
                        files_list.append({
                            'name': link_name,
                            'path': folder_smb,
                            'size': 0,
                            'extension': '',
                            'type': 'offer_folder_link',
                            'offer_number': offer.number,
                            'folder': os.path.basename(folder_path),
                            'smb_path': folder_smb,
                            'is_folder_link': True
                        })
    
    # Rimuovi duplicati: stesso nome file (senza percorso) e stesso smb_path
    seen_files = set()
    unique_files = []
    for f in files_list:
        # Estrai solo il nome file (senza percorso)
        file_name = os.path.basename(f.get('name', ''))
        smb_path = f.get('smb_path', '')
        # Se smb_path è presente, estrai anche il nome file da quello
        if smb_path:
            smb_file_name = os.path.basename(smb_path)
        else:
            smb_file_name = ''
        
        # Crea chiave univoca: nome file + smb_path (se presente)
        # Se smb_path è presente, usa quello come chiave principale
        if smb_path:
            key = (smb_file_name.lower(), smb_path.lower())
        else:
            key = (file_name.lower(), f.get('path', '').lower())
        
        if key not in seen_files:
            seen_files.add(key)
            unique_files.append(f)
    files_list = unique_files
    
    # Ordina per nome
    files_list.sort(key=lambda x: x['name'])
    
    # Recupera nome account se disponibile
    account_name = None
    if opp_account_id:
        account_result = db.session.execute(text("""
            SELECT name FROM accounts WHERE id = :account_id
        """), {"account_id": opp_account_id}).fetchone()
        if account_result:
            account_name = account_result[0]
    
    return jsonify({
        'files': files_list,
        'folder_path': opp_folder_path,
        'account_name': account_name,
        'offers_count': len(offers_list),
        'files_found': len(files_list)
    })

@app.route('/api/opportunities/<int:opportunity_id>/files/<path:filename>', methods=['GET'])
def download_opportunity_file(opportunity_id, filename):
    """Genera link SMB per aprire file sul PC locale o scarica se non configurato"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    from urllib.parse import unquote, quote
    filename = unquote(filename)
    
    opp = load_opportunity_safe(opportunity_id)
    if not opp:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    # Carica offerte
    opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
    
    # Recupera configurazione server per accesso file (valori di default se non configurati)
    DEFAULT_SERVER_IP = "192.168.10.240"
    DEFAULT_SERVER_SHARE = "commerciale"
    
    user = User.query.get(session['user_id'])
    if user:
        user_local_pc_ip = user.local_pc_ip or DEFAULT_SERVER_IP
        user_local_pc_share = user.local_pc_share or DEFAULT_SERVER_SHARE
        user_has_local_config = True
    else:
        user_local_pc_ip = DEFAULT_SERVER_IP
        user_local_pc_share = DEFAULT_SERVER_SHARE
        user_has_local_config = True
    
    # Cerca il file SOLO nelle cartelle specifiche delle offerte associate
    # NON cercare in tutte le cartelle principali!
    
    # Prima cerca nella cartella dell'opportunità
    if opp.folder_path:
        folder_path = opp.folder_path
        if user_has_local_config and folder_path.startswith('K:'):
            # Converti percorso K:\ usando configurazione server
            relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
            import platform
            if platform.system() == 'Windows':
                folder_path = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}\\{relative_path}".replace('/', '\\')
            else:
                folder_path = f"/mnt/{user_local_pc_share}/{relative_path}"
        
        if os.path.exists(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.exists(file_path) and os.path.isfile(file_path):
                return send_file(file_path, as_attachment=True)
    
    # Poi cerca nelle cartelle specifiche delle offerte associate
    offers_list = list(opp.offers)
    for offer in offers_list:
        if not offer.folder_path:
            continue  # Salta offerte senza cartella
        
        folder_path = offer.folder_path
        
        # Converti percorso usando configurazione server
        if user_has_local_config and folder_path.startswith('K:'):
            relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
            import platform
            if platform.system() == 'Windows':
                folder_path = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}\\{relative_path}".replace('/', '\\')
            else:
                folder_path = f"/mnt/{user_local_pc_share}/{relative_path}"
        
        if os.path.exists(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.exists(file_path) and os.path.isfile(file_path):
                return send_file(file_path, as_attachment=True)
        
        # Prova anche nei file PDF/DOCX direttamente collegati (se nella stessa cartella)
        if offer.pdf_path and os.path.basename(offer.pdf_path) == filename:
            pdf_path = offer.pdf_path
            if user_has_local_config and pdf_path.startswith('K:'):
                relative_path = pdf_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                import platform
                if platform.system() == 'Windows':
                    pdf_path = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}\\{relative_path}".replace('/', '\\')
                else:
                    pdf_path = f"/mnt/{user_local_pc_share}/{relative_path}"
            
            if os.path.exists(pdf_path):
                return send_file(pdf_path, as_attachment=True)
        
        if offer.docx_path and os.path.basename(offer.docx_path) == filename:
            docx_path = offer.docx_path
            if user_has_local_config and docx_path.startswith('K:'):
                relative_path = docx_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                import platform
                if platform.system() == 'Windows':
                    docx_path = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}\\{relative_path}".replace('/', '\\')
                else:
                    docx_path = f"/mnt/{user_local_pc_share}/{relative_path}"
            
            if os.path.exists(docx_path):
                return send_file(docx_path, as_attachment=True)
    
    # Se l'utente ha configurato il PC locale, genera link SMB per aprire sul PC locale
    if user_has_local_config:
        # Trova il percorso relativo del file dalle offerte associate
        relative_path = None
        
        # Cerca nelle offerte associate
        for offer in opp.offers:
            if offer.folder_path:
                folder_path = offer.folder_path
                # Estrai percorso relativo
                if folder_path.startswith('K:'):
                    relative_path = folder_path.replace('K:\\', '').replace('K:/', '').replace('\\', '/')
                
                if relative_path:
                    # Aggiungi il nome del file
                    relative_path = relative_path.replace('\\', '/')
                    if not relative_path.endswith('/'):
                        relative_path += '/'
                    relative_path += filename
                    break
        
        # Se non trovato nelle offerte, costruisci percorso basandosi sul nome file
        if not relative_path:
            # Prova a costruire il percorso basandosi sul nome del file
            # Esempio: se il file è "K1020-25 - BWS SRL - UR12E E UR20 - 28-05-2025.pdf"
            # potrebbe essere in "OFFERTE 2025 - K/K1020-25 - BWS SRL..."
            base_path = 'OFFERTE - K'  # Valore di default
            # Prova entrambe le cartelle anno
            for year_folder in ['OFFERTE 2025 - K', 'OFFERTE 2026 - K']:
                # Cerca cartelle che contengono parte del nome file
                if 'K' in filename and len(filename) > 10:
                    # Estrai numero offerta (es. K1020-25)
                    import re
                    match = re.search(r'K\d{4}-\d{2}', filename)
                    if match:
                        offer_num = match.group(0)
                        # Costruisci percorso probabile
                        relative_path = f"{year_folder}/{offer_num} - */{filename}"
                        break
        
        if relative_path:
            # Costruisci percorso completo
            # user_local_pc_base_path è definito nella funzione get_opportunity_files, qui usiamo un fallback
            base_path = 'OFFERTE - K'  # Valore di default per base_path
            base_path = base_path.replace('\\', '/').strip('/')
            relative_path = relative_path.replace('\\', '/').strip('/')
            
            # Se il percorso relativo inizia già con la base path, non duplicarlo
            if not relative_path.startswith(base_path):
                full_path = f"{base_path}/{relative_path}" if base_path else relative_path
            else:
                full_path = relative_path
            
            full_path = full_path.replace('//', '/').replace('\\', '/')
            
            # Genera link SMB per Windows
            # Formato: \\IP\share\path\to\file
            smb_path = f"\\\\{user_local_pc_ip}\\{user_local_pc_share}\\{full_path}"
            smb_path = smb_path.replace('/', '\\')
            
            # Per browser, crea una pagina HTML che copia il percorso negli appunti e mostra istruzioni
            # Prepara le variabili fuori dall'f-string per evitare problemi con backslash
            smb_path_normalized = smb_path.replace('\\', '/')
            smb_path_escaped = smb_path.replace('\\', '\\\\').replace("'", "\\'").replace('"', '&quot;')
            
            # Crea pagina HTML che copia automaticamente il percorso e mostra istruzioni
            html_content = f'''<!DOCTYPE html>
<html>
<head>
    <title>Apertura File SMB</title>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            padding: 40px;
            text-align: center;
            background: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 30px;
            border-radius: 8px;
            max-width: 600px;
            margin: 0 auto;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .path-box {{
            background: #f0f0f0;
            padding: 16px;
            border-radius: 8px;
            margin: 20px 0;
            font-family: monospace;
            word-break: break-all;
            font-size: 14px;
            border: 2px solid #ddd;
        }}
        button {{
            padding: 12px 24px;
            background: #667eea;
            color: white;
            border: none;
            border-radius: 6px;
            cursor: pointer;
            font-size: 16px;
            margin: 10px;
        }}
        button:hover {{
            background: #5568d3;
        }}
        .instructions {{
            margin-top: 20px;
            padding: 15px;
            background: #e3f2fd;
            border-radius: 6px;
            text-align: left;
        }}
        .instructions ol {{
            margin: 10px 0;
            padding-left: 20px;
        }}
        .instructions li {{
            margin: 8px 0;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h2>📁 Apertura File di Rete</h2>
        <p>Il percorso è stato copiato automaticamente negli appunti.</p>
        <div class="path-box" id="path-display">{smb_path}</div>
        <div>
            <button onclick="copyPath()">📋 Copia di Nuovo</button>
            <button onclick="openExplorer()">🔍 Apri Esplora File</button>
            <button onclick="openFolderDirect() || copyPath()">🚀 Apri Direttamente</button>
        </div>
        <div class="instructions">
            <strong>Istruzioni:</strong>
            <ol>
                <li>Apri Esplora File (premi <kbd>Win + E</kbd>)</li>
                <li>Incolla il percorso nella barra degli indirizzi (<kbd>Ctrl + V</kbd>)</li>
                <li>Premi <kbd>Invio</kbd></li>
                <li>Se richiesto, inserisci le credenziali di accesso al server</li>
            </ol>
            <p style="margin-top: 10px; font-size: 12px; color: #666;">
                <strong>Nota:</strong> Se Windows chiede credenziali, usa quelle configurate nelle Impostazioni del CRM.
            </p>
        </div>
    </div>
    <script>
        const path = '{smb_path_escaped}';
        
        // Copia automaticamente il percorso negli appunti
        function copyPath() {{
            navigator.clipboard.writeText(path).then(() => {{
                alert('✅ Percorso copiato negli appunti!');
            }}).catch(() => {{
                // Fallback per browser vecchi
                const textArea = document.createElement('textarea');
                textArea.value = path;
                textArea.style.position = 'fixed';
                textArea.style.opacity = '0';
                document.body.appendChild(textArea);
                textArea.select();
                document.execCommand('copy');
                document.body.removeChild(textArea);
                alert('✅ Percorso copiato negli appunti!');
            }});
        }}
        
        // Prova ad aprire Esplora File (funziona solo su Windows)
        function openExplorer() {{
            // Su Windows, prova ad aprire Esplora File con il percorso UNC
            try {{
                // Metodo 1: Prova con protocollo file:// (può essere bloccato dal browser)
                const link = document.createElement('a');
                link.href = 'file:///' + path.replace(/\\\\/g, '/');
                link.style.display = 'none';
                document.body.appendChild(link);
                link.click();
                setTimeout(() => document.body.removeChild(link), 100);
                
                // Metodo 2: Prova con shell: (solo Windows)
                setTimeout(() => {{
                    try {{
                        // Crea un link con protocollo shell per aprire Esplora File
                        const shellLink = document.createElement('a');
                        shellLink.href = 'shell:AppsFolder\\Microsoft.Windows.Explorer_8wekyb3d8bbwe!App';
                        shellLink.style.display = 'none';
                        document.body.appendChild(shellLink);
                        shellLink.click();
                        setTimeout(() => document.body.removeChild(shellLink), 100);
                    }} catch (e) {{
                        console.log('Metodo shell non disponibile');
                    }}
                }}, 200);
            }} catch (e) {{
                console.log('Apertura automatica non supportata');
            }}
            copyPath();
        }}
        
        // Funzione per aprire direttamente la cartella usando ActiveX (solo IE/Edge legacy)
        function openFolderDirect() {{
            try {{
                if (window.ActiveXObject) {{
                    const shell = new ActiveXObject('Shell.Application');
                    shell.Open(path);
                    return true;
                }}
            }} catch (e) {{
                // ActiveX non disponibile (browser moderno)
            }}
            return false;
        }}
        
        // Copia automaticamente all'apertura della pagina
        window.onload = function() {{
            copyPath();
        }};
    </script>
</body>
</html>'''
            
            return html_content, 200, {'Content-Type': 'text/html; charset=utf-8'}
    
    # RIMOSSO: Non cerchiamo più in tutte le cartelle principali
    # Cerchiamo SOLO nelle cartelle specifiche delle offerte usando la configurazione del PC locale
    
    return jsonify({'error': 'File non trovato'}), 404


@app.route('/api/opportunities/<int:opportunity_id>/take-ownership', methods=['POST'])
def take_opportunity_ownership(opportunity_id):
    """Prendi ownership di un'opportunità - Solo per Giovanni, Mauro, Davide, Filippo"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    
    # Solo questi utenti possono prendere ownership delle opportunità
    allowed_users = ['giovanni', 'mauro', 'davide', 'filippo', 'admin']
    if not user or user.username.lower() not in allowed_users:
        return jsonify({'error': 'Solo utenti commerciali autorizzati possono prendere ownership delle opportunità'}), 403
    
    opp = load_opportunity_safe(opportunity_id)
    if not opp:
        return jsonify({'error': 'Opportunità non trovata'}), 404
    old_owner_id = opp.owner_id
    
    opp.owner_id = user_id
    db.session.commit()
    
    # Crea notifica per il vecchio owner se diverso
    if old_owner_id and old_owner_id != user_id:
        create_notification(
            user_id=old_owner_id,
            title=f'Ownership opportunità cambiata',
            message=f'L\'opportunità "{opp.name}" è stata assegnata a {user.first_name} {user.last_name}',
            type='info',
            related_opportunity_id=opportunity_id
        )
    
    return jsonify({
        'success': True,
        'owner_id': opp.owner_id,
        'owner_name': f"{opp.owner.first_name} {opp.owner.last_name}" if opp.owner else 'Unknown'
    })


@app.route('/api/users/commerciali', methods=['GET'])
def get_commercial_users():
    """Ottieni lista utenti commerciali"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    users = User.query.filter_by(role='commerciale').all()
    return jsonify({
        'users': [{
            'id': u.id,
            'username': u.username,
            'name': f"{u.first_name} {u.last_name}",
            'email': u.email
        } for u in users]
    })


@app.route('/api/products', methods=['GET'])
def products():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    products_list = Product.query.filter_by(is_active=True).all()
    result = []
    
    for p in products_list:
        prices_dict = {}
        for price_entry in p.prices:
            if price_entry.is_active:
                prices_dict[price_entry.price_level] = price_entry.amount
        
        # Compatibilità bronze/silver/gold
        primary_prices = {}
        if 'bronze' in prices_dict:
            primary_prices['bronze'] = prices_dict['bronze']
        if 'silver' in prices_dict:
            primary_prices['silver'] = prices_dict['silver']
        if 'gold' in prices_dict:
            primary_prices['gold'] = prices_dict['gold']
        
        result.append({
            'id': p.id,
            'code': p.code,
            'name': p.name,
            'description': p.description,
            'category': p.category,
            'family': p.family,
            'primary_prices': primary_prices,
            'prices': [{
                'price_level': pe.price_level,
                'amount': pe.amount,
                'currency': pe.currency,
                'source': pe.source,
                'is_override': pe.is_override
            } for pe in p.prices if pe.is_active]
        })
    
    return jsonify({
        'products': result,
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': pagination.total,
            'pages': pagination.pages,
            'has_next': pagination.has_next,
            'has_prev': pagination.has_prev,
            'next_page': pagination.next_num if pagination.has_next else None,
            'prev_page': pagination.prev_num if pagination.has_prev else None
        }
    })


@app.route('/api/products/<int:product_id>/prices', methods=['PUT'])
def update_product_price(product_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    product = Product.query.get_or_404(product_id)
    
    price_level = data['price_level']
    amount = float(data['amount'])
    
    # Cerca entry esistente o crea nuova
    entry = PriceBookEntry.query.filter_by(
        product_id=product_id,
        price_level=price_level
    ).first()
    
    if entry:
        entry.amount = amount
        entry.is_override = True
        entry.is_active = True
    else:
        entry = PriceBookEntry(
            product_id=product_id,
            price_level=price_level,
            amount=amount,
            currency='EUR',
            source='manual',
            is_override=True,
            is_active=True
        )
        db.session.add(entry)
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/pricebook/sync', methods=['POST'])
def sync_pricebook():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        from price_loader import load_price_catalog
        
        catalog = load_price_catalog()
        
        # Sincronizza nel database
        for code, entry in catalog.items():
            # Crea o aggiorna prodotto
            product = Product.query.filter_by(code=code).first()
            if not product:
                product = Product(
                    code=code,
                    name=entry.name,
                    description=entry.description,
                    category=entry.category,
                    family=entry.family,
                    is_active=True
                )
                db.session.add(product)
                db.session.flush()
            
            # Aggiorna prezzi
            for level, amount in entry.price_levels.items():
                price_entry = PriceBookEntry.query.filter_by(
                    product_id=product.id,
                    price_level=level
                ).first()
                
                if price_entry:
                    if not price_entry.is_override:  # Non sovrascrivere override manuali
                        price_entry.amount = amount
                        price_entry.source = entry.sources.get(level, 'official')
                        price_entry.is_active = True
                else:
                    price_entry = PriceBookEntry(
                        product_id=product.id,
                        price_level=level,
                        amount=amount,
                        currency='EUR',
                        source=entry.sources.get(level, 'official'),
                        is_active=True
                    )
                    db.session.add(price_entry)
        
        db.session.commit()
        return jsonify({'success': True, 'message': f'Sincronizzati {len(catalog)} prodotti'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/offers', methods=['GET'])
def offers():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Mostra tutte le offerte (condivise)
    offers_list = OfferDocument.query.order_by(OfferDocument.created_at.desc()).limit(100).all()
    
    offers_data = []
    for o in offers_list:
        # Determina nome cliente
        client_name = 'N/A'
        if o.opportunity and o.opportunity.account:
            client_name = o.opportunity.account.name
        elif o.owner:
            client_name = f"{o.owner.first_name} {o.owner.last_name}"
        
        offers_data.append({
            'id': o.id,
            'number': o.number,
            'status': o.status,
            'amount': o.amount,
            'description': o.description,
            'client': client_name,
            'opportunity_id': o.opportunity_id,
            'opportunity_name': o.opportunity.name if o.opportunity else None,
            'docx_path': o.docx_path,
            'pdf_path': o.pdf_path,
            'created_at': o.created_at.isoformat() if o.created_at else None
        })
    
    return jsonify({'offers': offers_data})


@app.route('/api/offers/<int:offer_id>/pdf', methods=['GET'])
def download_offer_pdf(offer_id):
    """Apri il PDF di un'offerta nel browser"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    offer = OfferDocument.query.get_or_404(offer_id)
    
    if not offer.pdf_path or not os.path.exists(offer.pdf_path):
        return jsonify({'error': 'PDF non trovato'}), 404
    
    # Apri nel browser invece di forzare il download
    return send_file(offer.pdf_path, as_attachment=False, mimetype='application/pdf')


@app.route('/api/offers/<int:offer_id>/docx', methods=['GET'])
def download_offer_docx(offer_id):
    """Scarica il DOCX di un'offerta"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    offer = OfferDocument.query.get_or_404(offer_id)
    
    if not offer.docx_path or not os.path.exists(offer.docx_path):
        return jsonify({'error': 'DOCX non trovato'}), 404
    
    return send_file(offer.docx_path, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/offers/generate', methods=['POST'])
def generate_offer():
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        data = request.json
        user_id = session['user_id']
        
        # Estrai dati
        request_text = data.get('request', '').strip()
        account_id = data.get('account_id')
        opportunity_id = data.get('opportunity_id')
        client_name = data.get('client_name', '').strip()
        
        # Se account_id o opportunity_id sono forniti, usa quelli
        if account_id:
            account = Account.query.get(account_id)
            if account:
                client_name = account.name
        
        if not client_name and not account_id:
            return jsonify({'error': 'Specifica un account o un nome cliente'}), 400
        
        if not request_text:
            return jsonify({'error': 'Specifica i prodotti richiesti'}), 400
        
        # Importa generatore
        try:
            from offer_generator import OfferGenerator
            import tempfile
            import platform
            
            # Determina directory offerte
            if platform.system() == 'Windows':
                offers_dir = r"K:\OFFERTE - K\OFFERTE 2025 - K"
            else:
                offers_dir = os.environ.get('OFFERS_FOLDER', '/mnt/k/OFFERTE - K/OFFERTE 2025 - K')
            
            # Fallback se K:\ non disponibile
            if not os.path.exists(offers_dir):
                offers_dir = os.path.join(os.path.dirname(__file__), 'offerte_generate')
                os.makedirs(offers_dir, exist_ok=True)
            
            generator = OfferGenerator(offers_directory=offers_dir)
            
            # Carica listini se disponibili
            pricelist_dir = os.path.join(os.path.dirname(__file__), 'listini')
            if os.path.exists(pricelist_dir):
                pdf_files = [os.path.join(pricelist_dir, f) for f in os.listdir(pricelist_dir) if f.endswith('.pdf')]
                if pdf_files:
                    generator.load_pricelists(pdf_files)
            
            # Crea offerta
            try:
                doc_handler, offer_info = generator.create_offer(
                    request=request_text,
                    company_name=client_name,
                    search_company_online=False
                )
            except Exception as e:
                import traceback
                error_trace = traceback.format_exc()
                logger.error(f"Errore nella creazione offerta: {e}\n{error_trace}")
                return jsonify({
                    'error': f'Errore nella generazione offerta: {str(e)}',
                    'trace': error_trace
                }), 500
            
            if not doc_handler:
                # Verifica se ci sono prodotti disponibili
                available_products = generator.parser.get_all_products()
                if not available_products:
                    return jsonify({
                        'error': 'Nessun prodotto disponibile. Carica prima i listini PDF nella sezione Listini.',
                        'hint': 'Vai su "Listini" e clicca "Sincronizza Listini" per caricare i prodotti.'
                    }), 400
                return jsonify({
                    'error': 'Nessun prodotto trovato nella richiesta. Verifica che i nomi dei prodotti siano corretti.',
                    'hint': f'Prodotti disponibili: {len(available_products)}. Prova a cercare prodotti nella sezione Listini.'
                }), 400
            
            # Salva offerta
            docx_path, pdf_path = generator.save_offer(doc_handler, offer_info, generate_pdf=True)
            
            if not docx_path:
                return jsonify({'error': 'Errore nel salvataggio offerta'}), 500
            
            # Salva offerta nel database
            offer_doc = OfferDocument(
                number=offer_info['offer_number'],
                status='draft',
                amount=offer_info.get('total', 0.0),
                description=f"Prodotti: {', '.join([p.get('name', '') for p in offer_info.get('products', [])[:3]])}",
                docx_path=docx_path,
                pdf_path=pdf_path,
                folder_path=os.path.dirname(docx_path) if docx_path else None,
                generated_from_nlp=True,
                opportunity_id=opportunity_id if opportunity_id else None,
                owner_id=user_id
            )
            db.session.add(offer_doc)
            
            # Se opportunity_id è fornito, aggiorna l'opportunità
            if opportunity_id:
                opp = load_opportunity_safe(opportunity_id)
                if opp:
                    opp.description = f"Offerta generata: {offer_info['offer_number']}\n{opp.description or ''}"
                    # Aggiorna anche l'amount se non è impostato
                    if opp.amount == 0.0 and offer_info.get('total', 0) > 0:
                        opp.amount = offer_info['total']
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'offer_number': offer_info['offer_number'],
                'client': offer_info['client_name'],
                'total': offer_info['total'],
                'products_count': len(offer_info['products']),
                'docx_path': docx_path,
                'pdf_path': pdf_path,
                'offer_id': offer_doc.id
            })
            
        except ImportError as e:
            return jsonify({'error': f'Modulo generatore non disponibile: {e}'}), 500
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Errore generazione offerta: {str(e)}'}), 500
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/offers/import-past', methods=['POST'])
def import_past_offers():
    """Importa offerte passate dalla cartella"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        from analyze_past_offers import analyze_offers_folder, import_offers_to_crm
        offers_data = analyze_offers_folder()
        imported = import_offers_to_crm(offers_data, db.session, session['user_id'])
        return jsonify({
            'success': True,
            'imported': imported,
            'total_found': len(offers_data)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ur/sync', methods=['POST'])
def ur_sync_route():
    """Endpoint per sincronizzazione dati CRM dal portale UR con Selenium"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session.get('user_id', 1)
    
    try:
        # Usa la funzione di sincronizzazione con logging completo
        from ur_portal_integration import sync_ur_portal
        
        # Esegui sincronizzazione (include già tutto il logging)
        stats = sync_ur_portal(db.session, user_id)
        
        if stats.get('errors'):
            return jsonify({
                'success': False,
                'error': '\n'.join(stats['errors']),
                'stats': stats,
                'message': f"Sincronizzazione completata con {len(stats['errors'])} errori"
            }), 500
        
        return jsonify({
            'success': True,
            'message': 'Sincronizzazione completata con successo',
            'stats': stats
        })
        
    except ImportError as e:
        return jsonify({
            'success': False,
            'error': f'Moduli integrazione UR non trovati: {str(e)}. Installa: pip install selenium webdriver-manager'
        }), 500
    except ConnectionError as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'Errore connessione al portale UR: {str(e)}. Verifica la connessione internet e che il dominio sia raggiungibile.',
            'error_type': 'ConnectionError'
        }), 500
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'Errore sincronizzazione UR: {str(e)}',
            'error_type': type(e).__name__
        }), 500


@app.route('/api/wix/sync', methods=['POST'])
def wix_sync_route():
    """Endpoint per sincronizzare form Wix Studio"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session.get('user_id', 1)
    
    try:
        from utils.wix_forms_import import import_wix_submissions_from_json
        
        # Cerca file JSON con submissions Wix
        import os
        wix_json_path = os.path.join(os.path.dirname(__file__), 'data', 'wix_submissions.json')
        
        if not os.path.exists(wix_json_path):
            return jsonify({
                'success': False,
                'error': f'File {wix_json_path} non trovato. Crea il file JSON con i dati del form Wix.'
            }), 404
        
        imported, skipped = import_wix_submissions_from_json(wix_json_path, user_id)
        
        return jsonify({
            'success': True,
            'message': f'Sincronizzazione completata: {imported} importate, {skipped} saltate',
            'stats': {
                'imported': imported,
                'skipped': skipped
            }
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'Errore sincronizzazione Wix: {str(e)}',
            'error_type': type(e).__name__
        }), 500


@app.route('/api/sync/all', methods=['POST'])
def sync_all_portals_route():
    """Endpoint per sincronizzare tutti i portali (UR, Offerte, Wix)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        import subprocess
        import os
        
        script_path = os.path.join(os.path.dirname(__file__), 'scripts', 'sync_all_portals.py')
        
        if not os.path.exists(script_path):
            return jsonify({
                'success': False,
                'error': f'Script {script_path} non trovato'
            }), 404
        
        # Esegui script di sincronizzazione
        result = subprocess.run(
            ['python3', script_path],
            capture_output=True,
            text=True,
            cwd=os.path.dirname(__file__),
            timeout=600  # 10 minuti timeout
        )
        
        if result.returncode == 0:
            return jsonify({
                'success': True,
                'message': 'Sincronizzazione completata',
                'output': result.stdout
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Errore durante sincronizzazione',
                'output': result.stdout,
                'error_output': result.stderr
            }), 500
            
    except subprocess.TimeoutExpired:
        return jsonify({
            'success': False,
            'error': 'Timeout durante sincronizzazione (oltre 10 minuti)'
        }), 500
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Errore sincronizzazione: {str(e)}',
            'error_type': type(e).__name__
        }), 500


@app.route('/api/import/leads/csv', methods=['POST'])
def import_leads_csv():
    """Importa Lead da file CSV"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session.get('user_id', 1)
    
    if 'file' not in request.files:
        return jsonify({'error': 'Nessun file caricato'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nessun file selezionato'}), 400
    
    try:
        # Leggi il file CSV
        stream = io.StringIO(file.stream.read().decode("UTF8"), newline=None)
        csv_input = csv.DictReader(stream, delimiter=',')
        
        imported = 0
        updated = 0
        skipped = 0
        errors = []
        
        for row_num, row in enumerate(csv_input, 1):
            try:
                email = row.get('Email', '').strip()
                if not email:
                    skipped += 1
                    continue
                
                company_name = row.get('Azienda', '').strip() or 'Azienda non specificata'
                first_name = row.get('Nome', '').strip()
                last_name = row.get('Cognome', '').strip()
                phone = row.get('Telefono', '').strip()
                address = row.get('Indirizzo', '').strip()
                interest = row.get('Interesse', '').strip() or 'MIR'
                source = row.get('Source', '').strip() or 'CSV Import'
                rating = row.get('Rating', '').strip() or 'Cold'
                notes = row.get('Note', '').strip()
                
                # Verifica se esiste già
                existing = db.session.execute(text("""
                    SELECT id FROM leads WHERE email = :email LIMIT 1
                """), {"email": email}).fetchone()
                
                if existing:
                    # Aggiorna Lead esistente
                    db.session.execute(text("""
                        UPDATE leads 
                        SET company_name = :company_name,
                            contact_first_name = :first_name,
                            contact_last_name = :last_name,
                            phone = :phone,
                            address = :address,
                            source = :source,
                            interest = :interest,
                            rating = :rating,
                            notes = :notes,
                            updated_at = :updated_at
                        WHERE id = :lead_id
                    """), {
                        "company_name": company_name[:200],
                        "first_name": first_name[:100],
                        "last_name": last_name[:100],
                        "phone": phone[:50] if phone else None,
                        "address": address[:500] if address else None,
                        "source": source[:50],
                        "interest": interest[:50],
                        "rating": rating[:20],
                        "notes": notes[:500] if notes else None,
                        "updated_at": datetime.utcnow(),
                        "lead_id": existing[0]
                    })
                    updated += 1
                else:
                    # Crea nuovo Lead
                    db.session.execute(text("""
                        INSERT INTO leads 
                        (company_name, contact_first_name, contact_last_name, email, phone, address,
                         source, interest, rating, status, notes, owner_id, created_at, updated_at)
                        VALUES (:company_name, :first_name, :last_name, :email, :phone, :address,
                                :source, :interest, :rating, :status, :notes, :owner_id, :created_at, :updated_at)
                    """), {
                        "company_name": company_name[:200],
                        "first_name": first_name[:100],
                        "last_name": last_name[:100],
                        "email": email[:100],
                        "phone": phone[:50] if phone else None,
                        "address": address[:500] if address else None,
                        "source": source[:50],
                        "interest": interest[:50],
                        "rating": rating[:20],
                        "status": "New",
                        "notes": notes[:500] if notes else None,
                        "owner_id": user_id,
                        "created_at": datetime.utcnow(),
                        "updated_at": datetime.utcnow()
                    })
                    imported += 1
                
                if row_num % 50 == 0:
                    db.session.commit()
            
            except Exception as e:
                errors.append(f"Riga {row_num}: {str(e)}")
                continue
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'imported': imported,
            'updated': updated,
            'skipped': skipped,
            'errors': errors[:10]  # Limita a 10 errori
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Errore importazione: {str(e)}'}), 500


@app.route('/api/import/opportunities/csv', methods=['POST'])
def import_opportunities_csv():
    """Importa Opportunità da file CSV"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session.get('user_id', 1)
    
    if 'file' not in request.files:
        return jsonify({'error': 'Nessun file caricato'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nessun file selezionato'}), 400
    
    try:
        # Leggi il file CSV
        stream = io.StringIO(file.stream.read().decode("UTF8"), newline=None)
        csv_input = csv.DictReader(stream, delimiter=',')
        
        imported = 0
        skipped = 0
        errors = []
        
        for row_num, row in enumerate(csv_input, 1):
            try:
                opp_name = row.get('Nome Opportunità', '').strip()
                company_name = row.get('Azienda', '').strip()
                contact_email = row.get('Contatto Email', '').strip()
                
                if not opp_name or not company_name:
                    skipped += 1
                    continue
                
                # Cerca o crea Account
                account = db.session.execute(text("""
                    SELECT id FROM accounts WHERE name = :name LIMIT 1
                """), {"name": company_name}).fetchone()
                
                if account:
                    account_id = account[0]
                else:
                    result = db.session.execute(text("""
                        INSERT INTO accounts (name, owner_id, created_at, updated_at)
                        VALUES (:name, :owner_id, :created_at, :updated_at)
                    """), {
                        "name": company_name[:200],
                        "owner_id": user_id,
                        "created_at": datetime.utcnow(),
                        "updated_at": datetime.utcnow()
                    })
                    db.session.commit()
                    account_id = result.lastrowid
                
                # Cerca o crea Contact
                contact_id = None
                if contact_email:
                    contact = db.session.execute(text("""
                        SELECT id FROM contacts WHERE email = :email LIMIT 1
                    """), {"email": contact_email}).fetchone()
                    
                    if contact:
                        contact_id = contact[0]
                    else:
                        # Crea contatto base
                        result = db.session.execute(text("""
                            INSERT INTO contacts (first_name, last_name, email, account_id, owner_id, created_at, updated_at)
                            VALUES (:first_name, :last_name, :email, :account_id, :owner_id, :created_at, :updated_at)
                        """), {
                            "first_name": 'Contatto',
                            "last_name": company_name[:50],
                            "email": contact_email[:100],
                            "account_id": account_id,
                            "owner_id": user_id,
                            "created_at": datetime.utcnow(),
                            "updated_at": datetime.utcnow()
                        })
                        db.session.commit()
                        contact_id = result.lastrowid
                
                # Estrai altri dati
                product = row.get('Prodotto', '').strip() or 'Richiesta informazioni'
                amount_str = row.get('Valore', '').strip().replace('€', '').replace(',', '').replace('.', '')
                try:
                    amount = float(amount_str) if amount_str else 0.0
                except:
                    amount = 0.0
                
                probability_str = row.get('Probabilità', '').strip().replace('%', '')
                try:
                    probability = int(probability_str) if probability_str else 5
                except:
                    probability = 5
                
                stage = row.get('Stage', '').strip() or 'Qualification'
                heat_level_map = {
                    'Fredde gelo': 'fredde_gelo',
                    'Fredde con speranza': 'fredda_speranza',
                    'Tiepide': 'tiepide',
                    'Calde': 'calde'
                }
                heat_level_raw = row.get('Heat Level', '').strip()
                heat_level = heat_level_map.get(heat_level_raw, 'fredda_speranza')
                
                description = row.get('Descrizione', '').strip()
                
                # Determina supplier_category
                product_lower = product.lower()
                if 'unitree' in product_lower or 'g1' in product_lower or 'go2' in product_lower:
                    supplier_category = 'Unitree'
                elif 'ur' in product_lower or 'universal' in product_lower or 'cobot' in product_lower:
                    supplier_category = 'Universal Robots'
                elif 'mir' in product_lower or 'pallet' in product_lower:
                    supplier_category = 'MIR'
                else:
                    supplier_category = 'Altro'
                
                # Crea Opportunità
                db.session.execute(text("""
                    INSERT INTO opportunities 
                    (name, stage, amount, probability, supplier_category, account_id, contact_id,
                     owner_id, description, heat_level, created_at, updated_at)
                    VALUES (:name, :stage, :amount, :probability, :supplier_category, :account_id, :contact_id,
                            :owner_id, :description, :heat_level, :created_at, :updated_at)
                """), {
                    "name": opp_name[:200],
                    "stage": stage[:50],
                    "amount": amount,
                    "probability": probability,
                    "supplier_category": supplier_category[:50],
                    "account_id": account_id,
                    "contact_id": contact_id,
                    "owner_id": user_id,
                    "description": description[:1000] if description else None,
                    "heat_level": heat_level[:50],
                    "created_at": datetime.utcnow(),
                    "updated_at": datetime.utcnow()
                })
                imported += 1
                
                if row_num % 50 == 0:
                    db.session.commit()
            
            except Exception as e:
                errors.append(f"Riga {row_num}: {str(e)}")
                continue
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'imported': imported,
            'skipped': skipped,
            'errors': errors[:10]
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Errore importazione: {str(e)}'}), 500


@app.route('/api/import/template/<type>', methods=['GET'])
def download_import_template(type):
    """Scarica template CSV per importazione"""
    if type == 'leads':
        return send_file('templates/import_leads_template.csv', as_attachment=True, download_name='template_import_leads.csv')
    elif type == 'opportunities':
        return send_file('templates/import_opportunities_template.csv', as_attachment=True, download_name='template_import_opportunities.csv')
    else:
        return jsonify({'error': 'Tipo non valido'}), 400


@app.route('/api/accounts/clean-names', methods=['POST'])
def clean_account_names():
    """Endpoint per pulire i nomi degli account"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Solo admin può eseguire questa operazione
    if session.get('user_role') != 'admin':
        return jsonify({'error': 'Solo admin può eseguire questa operazione'}), 403
    
    try:
        from pulisci_nomi_account import pulisci_nomi_account
        
        # Esegui la pulizia
        with app.app_context():
            accounts = Account.query.all()
            updated_count = 0
            skipped_count = 0
            
            import re
            def clean_account_name(name):
                if not name:
                    return ""
                cleaned = re.sub(r'^[~\$\s]+', '', name)
                cleaned = re.sub(r'[^\w\s\-\.\,\'\(\)]', '', cleaned)
                cleaned = re.sub(r'\s+', ' ', cleaned)
                cleaned = cleaned.strip()
                if len(cleaned) < 2:
                    return ""
                words = cleaned.split()
                cleaned_words = []
                for word in words:
                    if word.isupper() or any(c.isdigit() for c in word):
                        cleaned_words.append(word)
                    else:
                        cleaned_words.append(word.capitalize())
                return ' '.join(cleaned_words)
            
            for account in accounts:
                original_name = account.name
                cleaned_name = clean_account_name(original_name)
                
                if cleaned_name != original_name and cleaned_name:
                    existing = Account.query.filter_by(name=cleaned_name).first()
                    if existing and existing.id != account.id:
                        skipped_count += 1
                    else:
                        account.name = cleaned_name
                        updated_count += 1
            
            db.session.commit()
        
        return jsonify({
            'success': True,
            'updated': updated_count,
            'skipped': skipped_count,
            'message': f'Aggiornati {updated_count} account, saltati {skipped_count} (duplicati)'
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Errore pulizia nomi account: {error_trace}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/setup/associate-addresses', methods=['POST'])
def associate_addresses():
    """Associa indirizzi ai contatti cercando nelle offerte e nei siti web"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        return jsonify({'error': 'Solo admin può eseguire questa operazione'}), 403
    
    try:
        # Import dinamico per evitare errori se il file non esiste
        try:
            import sys
            import os
            sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
            from associa_indirizzi_contatti import associa_indirizzi_contatti
        except ImportError:
            return jsonify({
                'success': False,
                'error': 'Modulo associa_indirizzi_contatti non trovato. Assicurati che il file esista.'
            }), 500
        updated_count = associa_indirizzi_contatti()
        return jsonify({
            'success': True,
            'updated': updated_count,
            'message': f'Aggiornati {updated_count} account con indirizzi'
        })
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

@app.route('/api/setup/backup', methods=['POST'])
def backup_database():
    """Crea un backup del database"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        return jsonify({'error': 'Solo admin può creare backup'}), 403
    
    try:
        import shutil
        from datetime import datetime
        
        db_path = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        if not os.path.exists(db_path):
            return jsonify({'error': 'Database non trovato'}), 404
        
        # Crea directory backup se non esiste
        backup_dir = os.path.join(os.path.dirname(db_path), 'backups')
        os.makedirs(backup_dir, exist_ok=True)
        
        # Nome file backup con timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_filename = f'mekocrm_backup_{timestamp}.db'
        backup_path = os.path.join(backup_dir, backup_filename)
        
        # Copia database
        shutil.copy2(db_path, backup_path)
        
        return jsonify({
            'success': True,
            'filename': backup_filename,
            'path': backup_path,
            'size': os.path.getsize(backup_path)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/setup/optimize', methods=['POST'])
def optimize_database():
    """Ottimizza il database SQLite"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        return jsonify({'error': 'Solo admin può ottimizzare il database'}), 403
    
    try:
        # Esegui ANALYZE per ottimizzare le statistiche
        db.session.execute(text('ANALYZE'))
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Database ottimizzato'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/setup/vacuum', methods=['POST'])
def vacuum_database():
    """Esegue VACUUM sul database SQLite"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        return jsonify({'error': 'Solo admin può eseguire VACUUM'}), 403
    
    try:
        # Esegui VACUUM per recuperare spazio
        db.session.execute(text('VACUUM'))
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'VACUUM completato'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/setup/stats', methods=['GET'])
def get_database_stats():
    """Restituisce statistiche sul database"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        db_path = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        db_size = os.path.getsize(db_path) if os.path.exists(db_path) else 0
        
        stats = {
            'accounts': Account.query.count(),
            'leads': Lead.query.count(),
            'opportunities': Opportunity.query.count(),
            'users': User.query.count(),
            'tasks': OpportunityTask.query.count(),
            'db_size': f"{db_size / (1024 * 1024):.2f} MB"
        }
        
        return jsonify({'success': True, 'stats': stats})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/offers/sync', methods=['POST'])
def sync_offers_route():
    """Endpoint per sincronizzare offerte dalle cartelle 2025 e 2026"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        # Importa e esegui sync dalle cartelle 2025 e 2026
        from sync_offers_2025_2026 import sync_all_offers
        import logging
        import io
        import sys
        
        # Cattura output
        log_capture = io.StringIO()
        handler = logging.StreamHandler(log_capture)
        handler.setLevel(logging.INFO)
        
        # Esegui sync
        result_msg = []
        try:
            sync_all_offers()
            result_msg.append("✅ Sincronizzazione completata")
            
            # Dopo la sincronizzazione, aggiorna tutti i valori delle offerte
            update_all_offer_values()
            result_msg.append("✅ Valori offerte aggiornati")
        except Exception as sync_error:
            error_msg = f"❌ Errore durante sync: {str(sync_error)}"
            result_msg.append(error_msg)
            import traceback
            result_msg.append(traceback.format_exc())
            raise
        
        return jsonify({
            'success': True,
            'message': '\n'.join(result_msg) if result_msg else 'Sincronizzazione offerte completata. Opportunità aggiornate.'
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Errore sync offerte: {error_trace}")
        return jsonify({
            'success': False,
            'error': str(e),
            'trace': error_trace
        }), 500


def update_all_offer_values():
    """Aggiorna tutti i valori delle offerte estraendoli dai PDF"""
    # Aggiorna tutte le offerte con PDF disponibile
    all_offers = OfferDocument.query.filter(OfferDocument.pdf_path != None).all()
    
    updated_count = 0
    for offer in all_offers:
        if offer.pdf_path and os.path.exists(offer.pdf_path):
            try:
                amount = extract_offer_amount_from_file(offer.pdf_path)
                if amount > 0:
                    # Aggiorna sempre se il valore estratto è maggiore o se non c'è valore
                    old_amount = offer.amount or 0.0
                    if abs(amount - old_amount) > 0.01:
                        offer.amount = amount
                        updated_count += 1
                        if old_amount > 0:
                            print(f"  💰 Valore aggiornato per {offer.number}: €{old_amount:,.2f} → €{amount:,.2f}")
                        else:
                            print(f"  💰 Valore estratto per {offer.number}: €{amount:,.2f}")
            except Exception as e:
                print(f"  ⚠️  Errore estrazione valore per {offer.number}: {e}")
    
    # Aggiorna anche i valori delle opportunità
    opportunities = Opportunity.query.all()
    opp_updated = 0
    for opp in opportunities:
        total_value = 0.0
        for offer in opp.offers:
            if offer.pdf_path and os.path.exists(offer.pdf_path):
                amount = offer.amount if offer.amount and offer.amount > 0 else extract_offer_amount_from_file(offer.pdf_path)
                total_value += amount
            elif offer.amount and offer.amount > 0:
                total_value += offer.amount
        
        if total_value > 0:
            old_amount = opp.amount or 0.0
            if abs(old_amount - total_value) > 0.01:
                opp.amount = total_value
                opp_updated += 1
                print(f"  💰 Valore opportunità {opp.name}: €{old_amount:,.2f} → €{total_value:,.2f}")
    
    db.session.commit()
    print(f"✅ Aggiornati {updated_count} valori offerte e {opp_updated} valori opportunità")


@app.route('/api/offers/update-values', methods=['POST'])
def update_offer_values_route():
    """Endpoint per aggiornare tutti i valori delle offerte dai PDF"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        update_all_offer_values()
        return jsonify({
            'success': True,
            'message': 'Valori offerte aggiornati con successo'
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Errore aggiornamento valori: {error_trace}")
        return jsonify({
            'success': False,
            'error': str(e),
            'trace': error_trace
        }), 500


@app.route('/api/offers/<int:offer_id>', methods=['GET', 'PUT'])
def get_or_update_offer(offer_id):
    """Ottiene o aggiorna un'offerta"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    from sqlalchemy import text
    
    # GET - Recupera offerta
    if request.method == 'GET':
        # Usa query che gestisce colonne opzionali
        try:
            offer_result = db.session.execute(text("""
                SELECT id, number, status, amount, 
                       COALESCE(closed_amount, 0.0) as closed_amount, 
                       COALESCE(margin, 0.0) as margin, 
                       description, folder_path, pdf_path, docx_path, opportunity_id
                FROM offers 
                WHERE id = :offer_id
            """), {"offer_id": offer_id}).fetchone()
        except Exception:
            # Fallback se le colonne non esistono ancora
            offer_result = db.session.execute(text("""
                SELECT id, number, status, amount, 0.0 as closed_amount, 0.0 as margin,
                       description, folder_path, pdf_path, docx_path, opportunity_id
                FROM offers 
                WHERE id = :offer_id
            """), {"offer_id": offer_id}).fetchone()
        
        if not offer_result:
            return jsonify({'error': 'Offerta non trovata'}), 404
        
        return jsonify({
            'success': True,
            'offer': {
                'id': offer_result[0],
                'number': offer_result[1],
                'status': offer_result[2],
                'amount': offer_result[3] or 0.0,
                'closed_amount': offer_result[4] or 0.0,
                'margin': offer_result[5] or 0.0,
                'description': offer_result[6],
                'folder_path': offer_result[7],
                'pdf_path': offer_result[8],
                'docx_path': offer_result[9],
                'opportunity_id': offer_result[10]
            }
        })
    
    # PUT - Aggiorna offerta
    data = request.json
    
    # Costruisci query UPDATE dinamica
    updates = []
    params = {"offer_id": offer_id}
    
    if 'folder_path' in data:
        updates.append("folder_path = :folder_path")
        params['folder_path'] = data['folder_path']
    
    if 'status' in data:
        updates.append("status = :status")
        params['status'] = data['status']
        
        # Se l'offerta viene chiusa come "accepted" o "closed_won", aggiorna anche closed_amount e margin
        if data['status'] in ['accepted', 'closed_won']:
            if 'closed_amount' in data:
                try:
                    # Verifica se la colonna esiste
                    db.session.execute(text("SELECT closed_amount FROM offers WHERE id = :offer_id LIMIT 1"), {"offer_id": offer_id})
                    updates.append("closed_amount = :closed_amount")
                    params['closed_amount'] = float(data['closed_amount']) if data['closed_amount'] else 0.0
                except Exception:
                    pass  # Colonna non esiste ancora
            
            if 'margin' in data:
                try:
                    db.session.execute(text("SELECT margin FROM offers WHERE id = :offer_id LIMIT 1"), {"offer_id": offer_id})
                    updates.append("margin = :margin")
                    params['margin'] = float(data['margin']) if data['margin'] else 0.0
                except Exception:
                    pass  # Colonna non esiste ancora
            elif 'closed_amount' in data:
                # Calcola margine automaticamente se non fornito
                try:
                    # Recupera amount attuale
                    current_offer = db.session.execute(text("SELECT amount FROM offers WHERE id = :offer_id"), {"offer_id": offer_id}).fetchone()
                    if current_offer:
                        amount = float(current_offer[0] or 0)
                        closed_amount = float(data['closed_amount'])
                        if amount > 0:
                            margin = ((closed_amount - amount) / amount) * 100
                            db.session.execute(text("SELECT margin FROM offers WHERE id = :offer_id LIMIT 1"), {"offer_id": offer_id})
                            updates.append("margin = :margin")
                            params['margin'] = margin
                except Exception:
                    pass  # Colonna non esiste ancora
    
    if 'amount' in data:
        updates.append("amount = :amount")
        params['amount'] = float(data['amount']) if data['amount'] else 0.0
    
    if 'closed_amount' in data and 'status' not in data:
        # Permetti aggiornamento closed_amount anche senza cambiare status
        try:
            db.session.execute(text("SELECT closed_amount FROM offers WHERE id = :offer_id LIMIT 1"), {"offer_id": offer_id})
            updates.append("closed_amount = :closed_amount")
            params['closed_amount'] = float(data['closed_amount']) if data['closed_amount'] else 0.0
        except Exception:
            pass  # Colonna non esiste ancora
    
    if 'margin' in data and 'status' not in data:
        try:
            db.session.execute(text("SELECT margin FROM offers WHERE id = :offer_id LIMIT 1"), {"offer_id": offer_id})
            updates.append("margin = :margin")
            params['margin'] = float(data['margin']) if data['margin'] else 0.0
        except Exception:
            pass  # Colonna non esiste ancora
    
    if 'description' in data:
        updates.append("description = :description")
        params['description'] = data['description']
    
    if not updates:
        return jsonify({'error': 'Nessun campo da aggiornare'}), 400
    
    # Esegui UPDATE
    update_query = f"UPDATE offers SET {', '.join(updates)}, updated_at = CURRENT_TIMESTAMP WHERE id = :offer_id"
    db.session.execute(text(update_query), params)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Offerta aggiornata con successo'})


@app.route('/api/offers/<int:offer_id>/search-folders', methods=['POST'])
def search_offer_folders(offer_id):
    """Cerca cartelle disponibili per un'offerta sulla macchina AWS"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    from sqlalchemy import text
    import platform
    
    # Recupera offerta
    offer_result = db.session.execute(text("""
        SELECT number, folder_path FROM offers WHERE id = :offer_id
    """), {"offer_id": offer_id}).fetchone()
    
    if not offer_result:
        return jsonify({'error': 'Offerta non trovata'}), 404
    
    offer_number = offer_result[0]
    current_folder_path = offer_result[1]
    
    # Cerca cartelle possibili
    possible_folders = []
    is_windows = platform.system() == 'Windows'
    
    if is_windows:
        # Su Windows, cerca in K:\
        base_paths = [
            r"K:\OFFERTE - K\OFFERTE 2025 - K",
            r"K:\OFFERTE - K\OFFERTE 2026 - K"
        ]
        
        for base_path in base_paths:
            if os.path.exists(base_path):
                try:
                    for item in os.listdir(base_path):
                        item_path = os.path.join(base_path, item)
                        if os.path.isdir(item_path):
                            # Verifica se il numero offerta è nel nome della cartella
                            if offer_number and offer_number.upper() in item.upper():
                                possible_folders.append({
                                    'path': item_path,
                                    'name': item,
                                    'exists': True
                                })
                except Exception as e:
                    pass
    else:
        # Su Linux (AWS), cerca in vari percorsi possibili
        base_dir = os.path.dirname(os.path.abspath(__file__))
        search_paths = [
            os.path.join(base_dir, 'offerte_2025_backup'),
            os.path.join(base_dir, 'offerte_2026_backup'),
            os.path.join(base_dir, 'offerte_generate'),
            '/mnt/k/OFFERTE - K/OFFERTE 2025 - K',
            '/mnt/k/OFFERTE - K/OFFERTE 2026 - K',
            os.environ.get('OFFERS_FOLDER', '')
        ]
        
        for base_path in search_paths:
            if base_path and os.path.exists(base_path):
                try:
                    for item in os.listdir(base_path):
                        item_path = os.path.join(base_path, item)
                        if os.path.isdir(item_path):
                            # Verifica se il numero offerta è nel nome della cartella
                            if offer_number and offer_number.upper() in item.upper():
                                possible_folders.append({
                                    'path': item_path,
                                    'name': item,
                                    'exists': True
                                })
                except Exception as e:
                    pass
    
    return jsonify({
        'success': True,
        'current_folder': current_folder_path,
        'possible_folders': possible_folders[:20]  # Limita a 20 risultati
    })


@app.route('/api/ur/logs', methods=['GET'])
def ur_sync_logs():
    """Endpoint per leggere i log della sincronizzazione UR"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        log_file = 'ur_portal_sync.log'
        lines = request.args.get('lines', 100, type=int)  # Ultime N righe
        
        if not os.path.exists(log_file):
            return jsonify({
                'success': True,
                'logs': [],
                'message': 'File log non ancora creato'
            })
        
        with open(log_file, 'r', encoding='utf-8') as f:
            all_lines = f.readlines()
            # Prendi le ultime N righe
            recent_lines = all_lines[-lines:] if len(all_lines) > lines else all_lines
        
        return jsonify({
            'success': True,
            'logs': [line.strip() for line in recent_lines],
            'total_lines': len(all_lines),
            'returned_lines': len(recent_lines)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Errore lettura log: {str(e)}'
        }), 500


@app.route('/api/analytics/pipeline', methods=['GET'])
def pipeline_analytics():
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Non autorizzato'}), 401
        
        user_id = session['user_id']
        period = request.args.get('period', 'month')  # day, week, month
        
        # Calcola data di inizio (ultimi 12 mesi per includere tutte le offerte)
        from datetime import datetime, timedelta
        end_date = datetime.utcnow()
        
        # Calcola start_date (12 mesi fa) - usa logica manuale per evitare dipendenze
        start_date = end_date
        for _ in range(12):
            # Sottrai un mese (approssimato a 30 giorni)
            start_date = start_date - timedelta(days=30)
        
        # Query opportunità - Carica TUTTE le opportunità per i commerciali (non solo quelle dell'utente)
        # Questo permette a tutti i commerciali di vedere i grafici completi
        from sqlalchemy import text
        from datetime import datetime
        
        user = User.query.get(user_id)
        is_commerciale = user and user.role == 'commerciale'
        
        # Carica anche heat_level se esiste nella tabella
        try:
            if is_commerciale:
                # Commerciali vedono TUTTE le opportunità (per grafici completi)
                opps_result = db.session.execute(text("""
                    SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                           description, supplier_category, folder_path, lost_reason, owner_id, 
                           account_id, contact_id, lead_id, created_at, updated_at, heat_level
                    FROM opportunities
                """)).fetchall()
            else:
                # Altri utenti vedono solo le loro opportunità
                opps_result = db.session.execute(text("""
                    SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                           description, supplier_category, folder_path, lost_reason, owner_id, 
                           account_id, contact_id, lead_id, created_at, updated_at, heat_level
                    FROM opportunities 
                    WHERE owner_id = :user_id
                """), {
                    "user_id": user_id
                }).fetchall()
            has_heat_level = True
        except Exception:
            # Se heat_level non esiste, usa query senza quella colonna
            if is_commerciale:
                opps_result = db.session.execute(text("""
                    SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                           description, supplier_category, folder_path, lost_reason, owner_id, 
                           account_id, contact_id, lead_id, created_at, updated_at
                    FROM opportunities
                """)).fetchall()
            else:
                opps_result = db.session.execute(text("""
                    SELECT id, name, stage, amount, probability, close_date, opportunity_type, 
                           description, supplier_category, folder_path, lost_reason, owner_id, 
                           account_id, contact_id, lead_id, created_at, updated_at
                    FROM opportunities 
                    WHERE owner_id = :user_id
                """), {
                    "user_id": user_id
                }).fetchall()
            has_heat_level = False
        
        # Costruisci oggetti Opportunity manualmente
        opps = []
        for row in opps_result:
            opp = Opportunity()
            opp.id = row[0]
            opp.name = row[1]
            opp.stage = row[2]
            opp.amount = row[3]
            opp.probability = row[4]
            opp.close_date = row[5]
            opp.opportunity_type = row[6]
            opp.description = row[7]
            opp.supplier_category = row[8]
            opp.folder_path = row[9]
            opp.lost_reason = row[10]
            opp.owner_id = row[11]
            opp.account_id = row[12]
            opp.contact_id = row[13]
            opp.lead_id = row[14]
            # Converti date
            if isinstance(row[15], str):
                opp.created_at = datetime.fromisoformat(row[15].replace('Z', '+00:00')) if row[15] else None
            else:
                opp.created_at = row[15]
            if isinstance(row[16], str):
                opp.updated_at = datetime.fromisoformat(row[16].replace('Z', '+00:00')) if row[16] else None
            else:
                opp.updated_at = row[16]
            
            # Carica heat_level se disponibile
            if has_heat_level and len(row) > 17:
                opp.heat_level = row[17]
            else:
                # Se non c'è heat_level nella query, prova a caricarlo separatamente
                try:
                    heat_result = db.session.execute(text("SELECT heat_level FROM opportunities WHERE id = :opp_id"), {"opp_id": opp.id}).fetchone()
                    opp.heat_level = heat_result[0] if heat_result else None
                except Exception:
                    opp.heat_level = None
            
            # Carica offerte
            opp.offers = db.session.query(OfferDocument).filter_by(opportunity_id=opp.id).all()
            # Carica line_items se necessario
            opp.line_items = db.session.query(OpportunityLineItem).filter_by(opportunity_id=opp.id).all()
            opps.append(opp)
        
        # Prima raccogli tutte le date per determinare il range reale
        # Poi genera i periodi necessari
        all_periods = set()
        
        # Raccogli periodi da tutte le opportunità usando le date estratte dalle cartelle
        for opp in opps:
            opp_date = None
            
            # Prova a estrarre la data dalla cartella dell'offerta più recente
            if opp.offers:
                for offer in opp.offers:
                    if offer.folder_path:
                        folder_date = extract_date_from_folder_name(offer.folder_path)
                        if folder_date:
                            from datetime import date as date_type
                            if isinstance(folder_date, date_type):
                                folder_datetime = datetime.combine(folder_date, datetime.min.time())
                            else:
                                folder_datetime = folder_date
                            if not opp_date or folder_datetime > opp_date:
                                opp_date = folder_datetime
                            break
            
            # Se non c'è data dalla cartella, usa created_at
            if not opp_date:
                opp_date = opp.created_at
            
            if opp_date:
                if period == 'day':
                    period_key = opp_date.strftime('%Y-%m-%d')
                elif period == 'week':
                    period_key = opp_date.strftime('%Y-W%W')
                else:
                    period_key = opp_date.strftime('%Y-%m')
                all_periods.add(period_key)
        
        # Genera struttura dati per tutti i periodi trovati
        data = {}
        for period_key in sorted(all_periods):
            data[period_key] = {
            'period': period_key,
            'MIR': {'open': 0, 'closed': 0},
            'UR': {'open': 0, 'closed': 0},
            'Unitree': {'open': 0, 'closed': 0},
            'ROEQ': {'open': 0, 'closed': 0},
            'Servizi': {'open': 0, 'closed': 0},
                'Altro': {'open': 0, 'closed': 0}
            }
            
        # Popola con i dati reali
        for opp in opps:
            # Determina periodo - PRIORITÀ ALLE DATE ESTRATTE DALLE CARTELLE
            opp_date = None
            
            # Prova a estrarre la data dalla cartella dell'offerta più recente
            if opp.offers:
                for offer in opp.offers:
                    if offer.folder_path:
                        folder_date = extract_date_from_folder_name(offer.folder_path)
                        if folder_date:
                            # Converti date in datetime per confronto
                            from datetime import date as date_type
                            if isinstance(folder_date, date_type):
                                folder_datetime = datetime.combine(folder_date, datetime.min.time())
                            else:
                                folder_datetime = folder_date
                            # Usa la data più recente tra le offerte (se ci sono più offerte)
                            if not opp_date or folder_datetime > opp_date:
                                opp_date = folder_datetime
                            break
            
            # Se non c'è data dalla cartella, usa created_at come fallback
            if not opp_date:
                opp_date = opp.created_at
            
            if not opp_date:
                continue  # Skip se non c'è nessuna data
            
            # NON filtrare per data qui - includiamo tutte le opportunità aperte
            # Il filtro per periodo viene fatto dopo quando assegniamo al periodo corretto
            
            # Determina periodo usando la data estratta
            if period == 'day':
                period_key = opp_date.strftime('%Y-%m-%d')
            elif period == 'week':
                period_key = opp_date.strftime('%Y-W%W')
            else:
                period_key = opp_date.strftime('%Y-%m')
            
            # Se il periodo non esiste, crealo
            if period_key not in data:
                data[period_key] = {
                    'period': period_key,
                    'MIR': {'open': 0, 'closed': 0},
                    'UR': {'open': 0, 'closed': 0},
                    'Unitree': {'open': 0, 'closed': 0},
                    'ROEQ': {'open': 0, 'closed': 0},
                    'Servizi': {'open': 0, 'closed': 0},
                    'Altro': {'open': 0, 'closed': 0}
                }
            
            # Calcola valore opportunità - somma TUTTE le offerte
            opp_value = 0.0
            if opp.amount and opp.amount > 0:
                # Se c'è un amount sull'opportunità, usa quello
                opp_value = opp.amount
            else:
                # Altrimenti somma tutte le offerte
                if opp.offers:
                    for o in opp.offers:
                        if o.amount and o.amount > 0:
                            opp_value += o.amount
                        else:
                            # Prova a estrarre dal PDF se disponibile
                            try:
                                extracted = extract_offer_amount(o)
                                if extracted > 0:
                                    opp_value += extracted
                            except Exception:
                                pass
            
            # Determina probabilità basandosi su heat_level invece che su probability
            # Fredde gelo: 0%, Fredde con speranza: 5%, Tiepide: 15%, Calde: 40%
            heat_level = opp.heat_level or 'da_categorizzare'
            if heat_level == 'calda':
                probability = 40
            elif heat_level == 'tiepida':
                probability = 15
            elif heat_level == 'fredda_speranza':
                probability = 5
            elif heat_level == 'fredda_gelo':
                probability = 0
            else:
                # Default per da_categorizzare o valori non riconosciuti
                probability = 0
            
            weighted_value = opp_value * (probability / 100.0)
            
            # Determina categoria dai prodotti o supplier_category
            category = 'Altro'
            if opp.supplier_category:
                # Mappa supplier_category alle categorie
                if 'UR' in opp.supplier_category or 'Universal Robot' in opp.supplier_category:
                    category = 'UR'
                elif 'MiR' in opp.supplier_category or 'MIR' in opp.supplier_category:
                    category = 'MIR'
                elif 'Unitree' in opp.supplier_category:
                    category = 'Unitree'
                else:
                    category = opp.supplier_category
            elif opp.line_items:
                # Prendi categoria dal primo prodotto
                first_product = opp.line_items[0].product
                if first_product:
                    category = first_product.category or 'Altro'
            
            # Normalizza categoria
            if category not in data[period_key]:
                category = 'Altro'
            
            # Calcola valore (usa amount o somma offerte) - già calcolato sopra con probabilità
            # weighted_value è già stato calcolato sopra, lo usiamo direttamente
            
            # Determina se chiusa
            is_closed = opp.stage in ['Closed Won', 'Closed Lost']
            
            if is_closed:
                # Per le chiuse, usa il valore pieno senza probabilità
                opp_value = opp.amount if opp.amount and opp.amount > 0 else 0
                if opp_value == 0 and opp.offers:
                    opp_value = sum(
                        (o.amount if o.amount and o.amount > 0 else extract_offer_amount(o))
                        for o in opp.offers
                    )
                data[period_key][category]['closed'] += opp_value
            else:
                # Per le aperte, usa il valore pesato per probabilità
                data[period_key][category]['open'] += weighted_value
        
        # Ordina per periodo
        sorted_data = sorted(data.values(), key=lambda x: x['period'])
        
        # Conta opportunità per stato di calore
        heat_counts = {
            'calda': 0,
            'tiepida': 0,
            'fredda_speranza': 0,
            'fredda_gelo': 0,
            'da_categorizzare': 0
        }
        
        for opp in opps:
            if opp.stage not in ['Closed Won', 'Closed Lost']:
                heat_level = opp.heat_level or 'da_categorizzare'
                if heat_level in heat_counts:
                    heat_counts[heat_level] += 1
                else:
                    heat_counts['da_categorizzare'] += 1
        
        return jsonify({
            'data': sorted_data,
            'heat_counts': heat_counts,
            'total_in_pipeline': sum(heat_counts.values())
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Errore pipeline_analytics: {e}\n{error_trace}")
        return jsonify({
            'error': str(e),
            'trace': error_trace
        }), 500


# ==================== EMAIL MARKETING API ROUTES ====================

@app.route('/api/email/lists', methods=['GET', 'POST'])
def email_lists():
    """Gestisci liste email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        lists = EmailList.query.filter_by(is_active=True).order_by(EmailList.created_at.desc()).all()
        return jsonify({
            'lists': [{
                'id': l.id,
                'name': l.name,
                'description': l.description,
                'subscriber_count': len([s for s in l.subscriptions if s.status == 'subscribed']),
                'subscribers_count': len([s for s in l.subscriptions if s.status == 'subscribed']),  # Mantenuto per retrocompatibilità
                'owner_id': l.owner_id,
                'created_at': l.created_at.isoformat() if l.created_at else None
            } for l in lists]
        })
    
    elif request.method == 'POST':
        data = request.json
        email_list = EmailList(
            name=data['name'],
            description=data.get('description'),
            owner_id=user_id
        )
        db.session.add(email_list)
        db.session.commit()
        return jsonify({'success': True, 'list_id': email_list.id})


@app.route('/api/email/lists/<int:list_id>', methods=['GET', 'PUT', 'DELETE'])
def email_list_detail(list_id):
    """Dettagli, aggiorna o elimina lista"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    email_list = EmailList.query.get_or_404(list_id)
    
    if request.method == 'GET':
        subscribers = EmailListSubscription.query.filter_by(
            list_id=list_id,
            status='subscribed'
        ).all()
        
        return jsonify({
            'list': {
                'id': email_list.id,
                'name': email_list.name,
                'description': email_list.description,
                'subscribers': [{
                    'id': s.id,
                    'email': s.email,
                    'first_name': s.first_name,
                    'last_name': s.last_name,
                    'contact_id': s.contact_id,
                    'lead_id': s.lead_id
                } for s in subscribers],
                'subscriber_count': len(subscribers),
                'subscribers_count': len(subscribers)  # Mantenuto per retrocompatibilità
            }
        })
    
    elif request.method == 'PUT':
        data = request.json
        if 'name' in data:
            email_list.name = data['name']
        if 'description' in data:
            email_list.description = data['description']
        if 'is_active' in data:
            email_list.is_active = data['is_active']
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        email_list.is_active = False
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/email/lists/<int:list_id>/subscribe', methods=['POST'])
def subscribe_to_list(list_id):
    """Iscrivi contatti/leads a una lista"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    email_list = EmailList.query.get_or_404(list_id)
    data = request.json
    
    contact_ids = data.get('contact_ids', [])
    lead_ids = data.get('lead_ids', [])
    emails = data.get('emails', [])  # Email manuali
    
    subscribed = []
    
    # Iscrivi contatti
    for contact_id in contact_ids:
        contact = Contact.query.get(contact_id)
        if contact and contact.email:
            sub = EmailListSubscription.query.filter_by(
                list_id=list_id,
                contact_id=contact_id
            ).first()
            
            if not sub:
                sub = EmailListSubscription(
                    list_id=list_id,
                    contact_id=contact_id,
                    email=contact.email,
                    first_name=contact.first_name,
                    last_name=contact.last_name,
                    status='subscribed'
                )
                db.session.add(sub)
                subscribed.append(contact.email)
    
    # Iscrivi leads
    for lead_id in lead_ids:
        lead = Lead.query.get(lead_id)
        if lead and lead.email:
            sub = EmailListSubscription.query.filter_by(
                list_id=list_id,
                lead_id=lead_id
            ).first()
            
            if not sub:
                sub = EmailListSubscription(
                    list_id=list_id,
                    lead_id=lead_id,
                    email=lead.email,
                    first_name=lead.contact_first_name,
                    last_name=lead.contact_last_name,
                    status='subscribed'
                )
                db.session.add(sub)
                subscribed.append(lead.email)
    
    # Iscrivi email manuali
    for email_data in emails:
        if isinstance(email_data, str):
            email = email_data
            name = ""
        else:
            email = email_data.get('email')
            name = email_data.get('name', '')
        
        if email:
            # Controlla se esiste già (per email o per contatto/lead collegato)
            sub = EmailListSubscription.query.filter_by(
                list_id=list_id,
                email=email,
                status='subscribed'
            ).first()
            
            if not sub:
                first_name, last_name = (name.split(' ', 1) + [''])[:2] if name else ('', '')
                sub = EmailListSubscription(
                    list_id=list_id,
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    status='subscribed'
                )
                db.session.add(sub)
                subscribed.append(email)
            else:
                # Se esiste già, aggiorna nome se fornito
                if name and (not sub.first_name or not sub.last_name):
                    first_name, last_name = (name.split(' ', 1) + [''])[:2]
                    sub.first_name = sub.first_name or first_name
                    sub.last_name = sub.last_name or last_name
                    db.session.commit()
    
    db.session.commit()
    return jsonify({'success': True, 'subscribed': subscribed, 'count': len(subscribed)})


@app.route('/api/email/templates', methods=['GET', 'POST'])
def email_templates():
    """Gestisci template email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        templates = EmailTemplate.query.filter_by(is_active=True).order_by(EmailTemplate.created_at.desc()).all()
        return jsonify({
            'templates': [{
                'id': t.id,
                'name': t.name,
                'subject': t.subject,
                'category': t.category,
                'template_type': t.category,  # Aggiunto per retrocompatibilità con frontend
                'owner_id': t.owner_id,
                'created_at': t.created_at.isoformat() if t.created_at else None
            } for t in templates]
        })
    
    elif request.method == 'POST':
        data = request.json
        template = EmailTemplate(
            name=data['name'],
            subject=data['subject'],
            html_content=data['html_content'],
            text_content=data.get('text_content'),
            category=data.get('category', 'general'),
            owner_id=user_id
        )
        db.session.add(template)
        db.session.commit()
        return jsonify({'success': True, 'template_id': template.id})


@app.route('/api/email/templates/<int:template_id>', methods=['GET', 'PUT', 'DELETE'])
def email_template_detail(template_id):
    """Dettagli, aggiorna o elimina template"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    template = EmailTemplate.query.get_or_404(template_id)
    
    if request.method == 'GET':
        return jsonify({
            'template': {
                'id': template.id,
                'name': template.name,
                'subject': template.subject,
                'html_content': template.html_content,
                'text_content': template.text_content,
                'category': template.category
            }
        })
    
    elif request.method == 'PUT':
        data = request.json
        if 'name' in data:
            template.name = data['name']
        if 'subject' in data:
            template.subject = data['subject']
        if 'html_content' in data:
            template.html_content = data['html_content']
        if 'text_content' in data:
            template.text_content = data['text_content']
        if 'category' in data:
            template.category = data['category']
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        template.is_active = False
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/email/campaigns', methods=['GET', 'POST'])
def email_campaigns():
    """Gestisci campagne email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    if request.method == 'GET':
        # Carica TUTTE le campagne (non solo le ultime 100) per vedere anche quelle passate
        campaigns = EmailCampaign.query.order_by(EmailCampaign.created_at.desc()).all()
        result_campaigns = []
        for c in campaigns:
            # Calcola numero destinatari dalla lista (sempre aggiornato)
            subscriber_count = 0
            if c.list_id:
                subscriber_count = EmailListSubscription.query.filter_by(
                    list_id=c.list_id,
                    status='subscribed'
                ).count()
            
            # Usa sempre il conteggio aggiornato degli iscritti per campagne non ancora inviate
            # Per campagne già inviate, mostra il numero effettivamente inviato
            if c.status == 'sent' and c.total_recipients > 0:
                recipients = c.total_recipients  # Numero effettivamente inviato
            else:
                recipients = subscriber_count  # Conteggio aggiornato dalla lista
            
            result_campaigns.append({
                'id': c.id,
                'name': c.name,
                'subject': c.subject,
                'status': c.status,
                'list_id': c.list_id,
                'list_name': c.email_list.name if c.email_list else None,
                'total_recipients': recipients,
                'list_subscriber_count': subscriber_count,
                'total_sent': c.total_sent,
                'total_delivered': c.total_delivered,
                'total_opened': c.total_opened,
                'total_clicked': c.total_clicked,
                'scheduled_at': c.scheduled_at.isoformat() if c.scheduled_at else None,
                'sent_at': c.sent_at.isoformat() if c.sent_at else None,
                'created_at': c.created_at.isoformat() if c.created_at else None
            })
        
        return jsonify({'campaigns': result_campaigns})
    
    elif request.method == 'POST':
        data = request.json
        scheduled_at = None
        if data.get('scheduled_at'):
            scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00'))
        
        # Supporta sia list_id (singolo) che list_ids (array)
        list_ids = []
        if 'list_ids' in data and isinstance(data['list_ids'], list):
            list_ids = data['list_ids']
        elif 'list_id' in data:
            list_ids = [data['list_id']]
        else:
            return jsonify({'success': False, 'error': 'list_id o list_ids richiesto'}), 400
        
        # Se più liste, crea una campagna unica che invierà a tutte le liste (con deduplicazione)
        # Se una sola lista, crea campagna normale
        if len(list_ids) == 1:
            campaign = EmailCampaign(
                name=data['name'],
                subject=data['subject'],
                html_content=data['html_content'],
                text_content=data.get('text_content'),
                list_id=list_ids[0],
                template_id=data.get('template_id'),
                scheduled_at=scheduled_at,
                status='draft',
                owner_id=user_id
            )
            db.session.add(campaign)
            db.session.commit()
            return jsonify({'success': True, 'campaign_id': campaign.id})
        else:
            # Per più liste, salva tutte le liste come JSON nel campo text_content (se vuoto) o in un commento HTML
            import json
            list_ids_json = json.dumps(list_ids)
            
            # Salva le liste come commento JSON nell'HTML (non visibile nell'email)
            html_with_metadata = data['html_content'] + f'\n<!-- CAMPAIGN_LIST_IDS: {list_ids_json} -->'
            
            campaign = EmailCampaign(
                name=f"{data['name']} ({len(list_ids)} liste)",
                subject=data['subject'],
                html_content=html_with_metadata,
                text_content=data.get('text_content') or f'LIST_IDS:{list_ids_json}',  # Backup nel text_content
                list_id=list_ids[0],  # Prima lista come riferimento
                template_id=data.get('template_id'),
                scheduled_at=scheduled_at,
                status='draft',
                owner_id=user_id
            )
            db.session.add(campaign)
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'campaign_id': campaign.id,
                'list_ids': list_ids,
                'message': f'Campagna creata per {len(list_ids)} liste. I destinatari verranno deduplicati all\'invio.'
            })


@app.route('/api/email/campaigns/<int:campaign_id>', methods=['GET', 'PUT', 'DELETE'])
def email_campaign_detail(campaign_id):
    """Dettagli, aggiorna o elimina campagna"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    campaign = EmailCampaign.query.get_or_404(campaign_id)
    
    if request.method == 'GET':
        return jsonify({
            'campaign': {
                'id': campaign.id,
                'name': campaign.name,
                'subject': campaign.subject,
                'html_content': campaign.html_content,
                'text_content': campaign.text_content,
                'status': campaign.status,
                'list_id': campaign.list_id,
                'template_id': campaign.template_id,
                'scheduled_at': campaign.scheduled_at.isoformat() if campaign.scheduled_at else None,
                'sent_at': campaign.sent_at.isoformat() if campaign.sent_at else None,
                'total_recipients': campaign.total_recipients,
                'total_sent': campaign.total_sent,
                'total_delivered': campaign.total_delivered,
                'total_opened': campaign.total_opened,
                'total_clicked': campaign.total_clicked,
                'total_bounced': campaign.total_bounced,
                'total_unsubscribed': campaign.total_unsubscribed
            }
        })
    
    elif request.method == 'PUT':
        data = request.json
        if 'name' in data:
            campaign.name = data['name']
        if 'subject' in data:
            campaign.subject = data['subject']
        if 'html_content' in data:
            campaign.html_content = data['html_content']
        if 'text_content' in data:
            campaign.text_content = data['text_content']
        if 'list_id' in data:
            campaign.list_id = data['list_id']
        if 'template_id' in data:
            campaign.template_id = data['template_id']
        if 'status' in data:
            campaign.status = data['status']
        if 'scheduled_at' in data:
            campaign.scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00')) if data['scheduled_at'] else None
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        db.session.delete(campaign)
        db.session.commit()
        return jsonify({'success': True})


@app.route('/api/email/campaigns/<int:campaign_id>/send', methods=['POST'])
def send_email_campaign(campaign_id):
    """Invia una campagna email"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        campaign = EmailCampaign.query.get_or_404(campaign_id)
        data = request.json or {}
        
        # Determina le liste da usare
        # Se vengono passate liste nell'endpoint, usale (per supportare più liste)
        # Altrimenti, recupera le liste salvate nella campagna (dal commento HTML o text_content)
        if 'list_ids' in data and isinstance(data['list_ids'], list) and len(data['list_ids']) > 0:
            list_ids = data['list_ids']
        else:
            # Prova a recuperare le liste dal commento HTML o text_content
            import json
            import re
            list_ids = [campaign.list_id]  # Default: solo la lista principale
            
            # Cerca nel commento HTML
            html_match = re.search(r'<!--\s*CAMPAIGN_LIST_IDS:\s*(\[.*?\])\s*-->', campaign.html_content)
            if html_match:
                try:
                    list_ids = json.loads(html_match.group(1))
                except:
                    pass
            
            # Se non trovato, cerca nel text_content
            if len(list_ids) == 1 and campaign.text_content and campaign.text_content.startswith('LIST_IDS:'):
                try:
                    list_ids_json = campaign.text_content.replace('LIST_IDS:', '').strip()
                    list_ids = json.loads(list_ids_json)
                except:
                    pass
        
        # Se il nome contiene "(X liste)", significa che ci sono più liste
        # Per ora, usiamo solo la lista principale, ma in futuro possiamo salvare le liste in un campo JSON
        # TODO: Aggiungere campo list_ids JSON al modello EmailCampaign
        
        # Prepara destinatari da tutte le liste (con deduplicazione)
        subscriptions = EmailListSubscription.query.filter(
            EmailListSubscription.list_id.in_(list_ids),
            EmailListSubscription.status == 'subscribed'
        ).all()
        
        if not subscriptions:
            return jsonify({'error': 'Nessun destinatario trovato nelle liste'}), 400
        
        # Deduplica per email (case-insensitive) mantenendo la prima occorrenza
        seen_emails = set()
        recipients = []
        for sub in subscriptions:
            email_lower = sub.email.lower().strip() if sub.email else None
            if not email_lower or email_lower in seen_emails:
                continue  # Salta duplicati
            
            seen_emails.add(email_lower)
            
            # Ottieni nome normalizzato
            first_name, last_name, full_name = get_recipient_name(sub.email)
            
            recipient = {
                'email': sub.email,
                'name': full_name,
                'first_name': first_name,
                'last_name': last_name,
                'variables': {}
            }
            recipients.append(recipient)
        
        # Inizializza Mailgun
        try:
            from mailgun_service import MailgunService
            mailgun = MailgunService()
        except Exception as e:
            return jsonify({'error': f'Errore configurazione Mailgun: {str(e)}'}), 500
        
        # Aggiorna campagna
        campaign.status = 'sending'
        campaign.total_recipients = len(recipients)
        db.session.commit()
        
        # Rimuovi commenti HTML (come CAMPAIGN_LIST_IDS) prima di inviare
        import re
        html_clean = re.sub(r'<!--\s*CAMPAIGN_LIST_IDS:.*?-->\s*', '', campaign.html_content, flags=re.DOTALL)
        text_clean = campaign.text_content
        if text_clean and text_clean.startswith('LIST_IDS:'):
            text_clean = None  # Rimuovi se era solo metadata
        
        # Aggiungi firma email a tutti i contenuti
        from email_signature import add_signature_to_email
        
        html_with_signature, text_with_signature = add_signature_to_email(
            html_content=html_clean,
            text_content=text_clean
        )
        
        # Invia email
        campaign_data = {
            'name': campaign.name,
            'subject': campaign.subject,
            'html_content': html_with_signature,
            'text_content': text_with_signature
        }
        
        stats = mailgun.send_campaign_email(campaign_data, recipients)
        
        # Aggiorna statistiche campagna
        campaign.total_sent = stats['sent']
        campaign.status = 'sent' if stats['failed'] == 0 else 'partial'
        campaign.sent_at = datetime.utcnow()
        db.session.commit()
        
        # Salva dettagli invii
        for result in stats['results']:
            recipient = result.get('recipient', {})
            send = EmailCampaignSend(
                campaign_id=campaign_id,
                email=recipient.get('email', ''),
                status='sent' if result.get('success') else 'failed',
                mailgun_message_id=result.get('message_id'),
                sent_at=datetime.utcnow() if result.get('success') else None,
                error_message=result.get('error') if not result.get('success') else None
            )
            db.session.add(send)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'stats': {
                'total': stats['total'],
                'sent': stats['sent'],
                'failed': stats['failed']
            }
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        if campaign:
            campaign.status = 'failed'
            db.session.commit()
        return jsonify({'error': str(e)}), 500


@app.route('/api/email/campaigns/<int:campaign_id>/recipients', methods=['GET'])
def campaign_recipients(campaign_id):
    """Lista destinatari di una campagna"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    campaign = EmailCampaign.query.get_or_404(campaign_id)
    
    # Ottieni tutti i destinatari dalla lista
    subscriptions = EmailListSubscription.query.filter_by(
        list_id=campaign.list_id,
        status='subscribed'
    ).all()
    
    recipients = []
    for sub in subscriptions:
        recipients.append({
            'email': sub.email,
            'first_name': sub.first_name or '',
            'last_name': sub.last_name or '',
            'name': f"{sub.first_name or ''} {sub.last_name or ''}".strip() or sub.email,
            'subscribed_at': sub.subscribed_at.isoformat() if sub.subscribed_at else None
        })
    
    return jsonify({
        'success': True,
        'campaign_id': campaign_id,
        'campaign_name': campaign.name,
        'total': len(recipients),
        'recipients': recipients
    })


def normalize_name(name_str):
    """Normalizza un nome: rimuove spazi extra, capitalizza correttamente"""
    if not name_str:
        return ""
    # Rimuovi spazi multipli e trim
    name_str = " ".join(name_str.split())
    # Capitalizza ogni parola (gestisce anche apostrofi e trattini)
    parts = []
    for part in name_str.split():
        if part:
            # Gestisci apostrofi e trattini
            if "'" in part:
                subparts = part.split("'")
                parts.append("'".join([subparts[0].capitalize()] + [s.lower() for s in subparts[1:]]))
            elif "-" in part:
                subparts = part.split("-")
                parts.append("-".join([s.capitalize() for s in subparts]))
            else:
                parts.append(part.capitalize())
    return " ".join(parts)


def extract_name_from_email(email):
    """Estrae e formatta un nome dall'indirizzo email"""
    if not email or '@' not in email:
        return "", "", ""
    
    email_local = email.split('@')[0]
    
    # Prova diversi pattern comuni
    # Pattern 1: nome.cognome
    if '.' in email_local:
        parts = email_local.split('.')
        if len(parts) >= 2:
            first = normalize_name(parts[0])
            last = normalize_name('.'.join(parts[1:]))
            full = f"{first} {last}".strip()
            return first, last, full
    
    # Pattern 2: nome-cognome
    if '-' in email_local:
        parts = email_local.split('-')
        if len(parts) >= 2:
            first = normalize_name(parts[0])
            last = normalize_name('-'.join(parts[1:]))
            full = f"{first} {last}".strip()
            return first, last, full
    
    # Pattern 3: nomecognome (senza separatori)
    if len(email_local) > 3:
        # Prova a dividere a metà (euristico)
        mid = len(email_local) // 2
        first = normalize_name(email_local[:mid])
        last = normalize_name(email_local[mid:])
        full = f"{first} {last}".strip()
        return first, last, full
    
    # Fallback: usa tutto come nome
    name = normalize_name(email_local)
    return name, "", name


def get_recipient_name(email):
    """Ottiene il nome del destinatario dal database o dall'email, normalizzato"""
    first_name = ""
    last_name = ""
    full_name = ""
    
    # Cerca prima nelle subscription (priorità più alta per email marketing)
    subscription = EmailListSubscription.query.filter_by(email=email).first()
    if subscription:
        first_name = normalize_name(subscription.first_name or "")
        last_name = normalize_name(subscription.last_name or "")
        if first_name or last_name:
            full_name = f"{first_name} {last_name}".strip()
    
    # Se non trovato, cerca nei contatti
    if not full_name:
        contact = Contact.query.filter_by(email=email).first()
        if contact:
            first_name = normalize_name(contact.first_name or "")
            last_name = normalize_name(contact.last_name or "")
            if first_name or last_name:
                full_name = f"{first_name} {last_name}".strip()
    
    # Se non trovato, cerca nei lead
    if not full_name:
        lead = Lead.query.filter_by(email=email).first()
        if lead:
            first_name = normalize_name(lead.contact_first_name or "")
            last_name = normalize_name(lead.contact_last_name or "")
            if first_name or last_name:
                full_name = f"{first_name} {last_name}".strip()
    
    # Se ancora non trovato, estrai dall'email
    if not full_name:
        first_name, last_name, full_name = extract_name_from_email(email)
    
    # Fallback finale
    if not full_name:
        full_name = "Cliente"
        first_name = "Cliente"
    
    return first_name, last_name, full_name


@app.route('/api/email/analyze-outlook', methods=['POST'])
def analyze_outlook_emails():
    """Analizza email Outlook con Claude API e aggiunge attività al CRM"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    add_tasks_to_crm = data.get('add_tasks_to_crm', True)
    user_id = session['user_id']
    
    try:
        import requests
        import re
        import time
        from datetime import timedelta
        
        # Claude API Key - usa variabile d'ambiente per sicurezza
        CLAUDE_API_KEY = os.getenv('ANTHROPIC_API_KEY', 'YOUR_ANTHROPIC_API_KEY_HERE')
        CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"
        
        # Microsoft Graph API - Device Code Flow
        # Usa un'app pubblica Microsoft per autenticazione
        CLIENT_ID = "14d82eec-204b-4c2f-b7e8-296a70dab67e"  # Microsoft Graph Explorer App ID (pubblico)
        TENANT_ID = "common"
        SCOPE = "https://graph.microsoft.com/Mail.Read"
        
        # Step 1: Ottieni device code
        device_url = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/devicecode"
        device_data = {
            "client_id": CLIENT_ID,
            "scope": SCOPE
        }
        
        device_response = requests.post(device_url, data=device_data)
        if device_response.status_code != 200:
            return jsonify({
                'error': 'Errore autenticazione Microsoft',
                'details': 'Impossibile ottenere device code. Usa il metodo alternativo con email/password.'
            }), 500
        
        device_info = device_response.json()
        
        # Ritorna il device code al frontend per mostrare all'utente
        return jsonify({
            'success': False,
            'auth_required': True,
            'verification_uri': device_info['verification_uri'],
            'user_code': device_info['user_code'],
            'device_code': device_info['device_code'],
            'expires_in': device_info['expires_in'],
            'interval': device_info['interval'],
            'message': f"Apri {device_info['verification_uri']} e inserisci il codice: {device_info['user_code']}"
        })
        
    except Exception as e:
        return jsonify({'error': f'Errore inizializzazione: {str(e)}'}), 500


@app.route('/api/email/analyze-outlook-complete', methods=['POST'])
def analyze_outlook_emails_complete():
    """Completa l'analisi email dopo l'autenticazione"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    device_code = data.get('device_code')
    add_tasks_to_crm = data.get('add_tasks_to_crm', True)
    user_id = session['user_id']
    
    if not device_code:
        return jsonify({'error': 'Device code richiesto'}), 400
    
    try:
        import requests
        import re
        import time
        from datetime import timedelta
        
        # Claude API Key - usa variabile d'ambiente per sicurezza
        CLAUDE_API_KEY = os.getenv('ANTHROPIC_API_KEY', 'YOUR_ANTHROPIC_API_KEY_HERE')
        CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"
        
        # Microsoft Graph API
        CLIENT_ID = "14d82eec-204b-4c2f-b7e8-296a70dab67e"
        TENANT_ID = "common"
        
        # Step 2: Ottieni access token
        token_url = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"
        token_data = {
            "grant_type": "urn:ietf:params:oauth:grant-type:device_code",
            "client_id": CLIENT_ID,
            "device_code": device_code
        }
        
        # Prova a ottenere il token (l'utente deve aver autorizzato)
        token_response = requests.post(token_url, data=token_data)
        token_result = token_response.json()
        
        if "access_token" not in token_result:
            if token_result.get("error") == "authorization_pending":
                return jsonify({
                    'success': False,
                    'auth_pending': True,
                    'message': 'Attendi che autorizzi l\'app nel browser...'
                })
            else:
                return jsonify({
                    'error': 'Errore autenticazione',
                    'details': token_result.get('error_description', token_result.get('error'))
                }), 500
        
        access_token = token_result["access_token"]
        
        # Step 3: Leggi email da Microsoft Graph
        graph_url = "https://graph.microsoft.com/v1.0/me/messages"
        filter_date = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%dT%H:%M:%SZ')
        
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        params = {
            "$filter": f"receivedDateTime ge {filter_date}",
            "$orderby": "receivedDateTime desc",
            "$top": 50,
            "$select": "subject,from,receivedDateTime,bodyPreview,body"
        }
        
        graph_response = requests.get(graph_url, headers=headers, params=params)
        
        if graph_response.status_code != 200:
            return jsonify({
                'error': f'Errore lettura email: {graph_response.status_code}',
                'details': graph_response.text
            }), 500
        
        graph_data = graph_response.json()
        emails = []
        
        for msg in graph_data.get('value', []):
            body_content = msg.get('bodyPreview', '') or msg.get('body', {}).get('content', '')
            # Rimuovi HTML se presente
            body_content = re.sub(r'<[^>]+>', '', body_content)
            
            emails.append({
                'subject': msg.get('subject', '(Nessun oggetto)'),
                'from': msg.get('from', {}).get('emailAddress', {}).get('address', ''),
                'date': msg.get('receivedDateTime', ''),
                'body': body_content[:5000]
            })
        
        if not emails:
            return jsonify({
                'success': True,
                'summary': 'Nessuna email trovata negli ultimi 7 giorni',
                'tasks': []
            })
        
        # Analizza con Claude
        email_text = ""
        for i, email_data in enumerate(emails, 1):
            email_text += f"\n--- Email {i} ---\n"
            email_text += f"Da: {email_data['from']}\n"
            email_text += f"Oggetto: {email_data['subject']}\n"
            email_text += f"Data: {email_data['date']}\n"
            email_text += f"Contenuto: {email_data['body']}\n\n"
        
        prompt = f"""Analizza le seguenti email e fornisci:
1. Un riassunto generale di tutte le email
2. Una lista di attività da fare basate sulle email

Per ogni attività, indica:
- Tipo attività (send_offer, call, recall, other)
- Descrizione chiara dell'attività
- Data scadenza se presente o suggerita (formato YYYY-MM-DD o null)
- Priorità se evidente (high, medium, low)

Rispondi SOLO in formato JSON valido con questa struttura:
{{
    "summary": "Riassunto generale delle email",
    "tasks": [
        {{
            "task_type": "send_offer|call|recall|other",
            "description": "Descrizione dettagliata dell'attività",
            "due_date": "YYYY-MM-DD o null",
            "priority": "high|medium|low"
        }}
    ]
}}

Email da analizzare:
{email_text}
"""
        
        response = requests.post(
            CLAUDE_API_URL,
            headers={
                "x-api-key": CLAUDE_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json"
            },
            json={
                "model": "claude-3-5-sonnet-20241022",
                "max_tokens": 4000,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=120
        )
        
        if response.status_code != 200:
            return jsonify({'error': f'Errore API Claude: {response.status_code}'}), 500
        
        result = response.json()
        content = result.get('content', [])
        if not content or len(content) == 0:
            return jsonify({'error': 'Risposta Claude vuota'}), 500
        
        text_response = content[0].get('text', '')
        
        # Estrai JSON dalla risposta
        json_match = re.search(r'\{.*\}', text_response, re.DOTALL)
        if not json_match:
            return jsonify({
                'success': True,
                'summary': text_response,
                'tasks': []
            })
        
        try:
            analysis = json.loads(json_match.group())
        except json.JSONDecodeError:
            return jsonify({
                'success': True,
                'summary': text_response,
                'tasks': []
            })
        
        # Aggiungi attività al CRM se richiesto
        tasks_added = []
        if add_tasks_to_crm and analysis.get('tasks'):
            for task in analysis['tasks']:
                try:
                    due_date = None
                    if task.get('due_date'):
                        try:
                            due_date = datetime.strptime(task['due_date'], '%Y-%m-%d').date()
                        except:
                            pass
                    
                    new_task = OpportunityTask(
                        opportunity_id=None,  # Task globale
                        account_id=None,
                        task_type=task.get('task_type', 'other'),
                        description=task.get('description', ''),
                        assigned_to_id=user_id,
                        due_date=due_date
                    )
                    db.session.add(new_task)
                    db.session.commit()
                    
                    tasks_added.append({
                        **task,
                        'added_to_crm': True,
                        'task_id': new_task.id
                    })
                except Exception as e:
                    tasks_added.append({
                        **task,
                        'added_to_crm': False,
                        'error': str(e)
                    })
        
        return jsonify({
            'success': True,
            'summary': analysis.get('summary', ''),
            'tasks': tasks_added if tasks_added else analysis.get('tasks', []),
            'emails_analyzed': len(emails)
        })
        
    except Exception as e:
        return jsonify({'error': f'Errore durante l\'analisi: {str(e)}'}), 500


@app.route('/api/email/test', methods=['POST'])
@app.route('/api/email/test-send', methods=['POST'])
def test_email():
    """Invia email di test con sostituzione placeholder"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    try:
        data = request.json
        to_email = data.get('to_email')
        subject = data.get('subject', 'Test Email')
        html_content = data.get('html_content') or data.get('html_body') or '<p>Questa è un\'email di test</p>'
        
        if not to_email:
            return jsonify({'error': 'Email destinatario richiesta'}), 400
        
        # Ottieni nome normalizzato
        first_name, last_name, full_name = get_recipient_name(to_email)
        
        # Sostituisci placeholder nel contenuto
        html_content = html_content.replace("%nome%", full_name)
        html_content = html_content.replace("%nome_completo%", full_name)
        html_content = html_content.replace("%nome_destinatario%", full_name)
        html_content = html_content.replace("%first_name%", first_name)
        html_content = html_content.replace("%last_name%", last_name)
        
        # Aggiungi firma email
        from email_signature import add_signature_to_email
        
        html_with_signature, text_with_signature = add_signature_to_email(
            html_content=html_content,
            text_content=None
        )
        
        from mailgun_service import MailgunService
        mailgun = MailgunService()
        
        result = mailgun.send_email(
            to_email=to_email,
            subject=subject,
            html_content=html_with_signature,
            text_content=text_with_signature
        )
        
        if result.get('success'):
            return jsonify({'success': True, 'message': 'Email di test inviata con successo', 'message_id': result.get('message_id')})
        else:
            return jsonify({'error': result.get('error', 'Errore sconosciuto')}), 500
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/email/lists/<int:list_id>/subscribers', methods=['GET', 'DELETE'])
def manage_list_subscribers(list_id):
    """Gestisci iscritti a una lista"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    email_list = EmailList.query.get_or_404(list_id)
    
    if request.method == 'GET':
        subscribers = EmailListSubscription.query.filter_by(list_id=list_id).all()
        return jsonify({
            'subscribers': [{
                'id': s.id,
                'email': s.email,
                'first_name': s.first_name,
                'last_name': s.last_name,
                'contact_id': s.contact_id,
                'lead_id': s.lead_id,
                'status': s.status,
                'subscribed_at': s.subscribed_at.isoformat() if s.subscribed_at else None
            } for s in subscribers]
        })
    
    elif request.method == 'DELETE':
        data = request.json
        subscription_id = data.get('subscription_id')
        
        if subscription_id:
            subscription = EmailListSubscription.query.get(subscription_id)
            if subscription and subscription.list_id == list_id:
                subscription.status = 'unsubscribed'
                subscription.unsubscribed_at = datetime.utcnow()
                db.session.commit()
                return jsonify({'success': True})
        
        return jsonify({'error': 'Subscription non trovata'}), 404


@app.route('/unsubscribe', methods=['GET', 'POST'])
def unsubscribe_email():
    """Endpoint pubblico per disiscrizioni email"""
    email = request.args.get('email') or (request.form.get('email') if request.method == 'POST' else None)
    
    if not email:
        return render_template('unsubscribe.html', 
                             success=False, 
                             message='Indirizzo email non fornito')
    
    try:
        # Cerca tutte le subscription per questa email
        subscriptions = EmailListSubscription.query.filter_by(email=email, status='subscribed').all()
        
        if not subscriptions:
            return render_template('unsubscribe.html', 
                                 success=False, 
                                 message=f'L\'indirizzo email {email} non risulta iscritto a nessuna newsletter')
        
        # Disiscrivi da tutte le liste
        for subscription in subscriptions:
            subscription.status = 'unsubscribed'
            subscription.unsubscribed_at = datetime.utcnow()
        
        db.session.commit()
        
        return render_template('unsubscribe.html', 
                             success=True, 
                             message=f'Sei stato disiscritto con successo dalla newsletter. L\'indirizzo {email} non riceverà più comunicazioni da parte nostra.')
    
    except Exception as e:
        db.session.rollback()
        return render_template('unsubscribe.html', 
                             success=False, 
                             message=f'Si è verificato un errore durante la disiscrizione: {str(e)}')


@app.route('/api/email/contacts-available', methods=['GET'])
def get_available_contacts_for_lists():
    """Ottieni contatti e leads disponibili con email per aggiungerli alle liste"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    # Ottieni tutti i contatti con email
    contacts = Contact.query.filter(Contact.email.isnot(None), Contact.email != '').all()
    contacts_data = [{
        'id': c.id,
        'type': 'contact',
        'email': c.email,
        'first_name': c.first_name,
        'last_name': c.last_name,
        'name': f"{c.first_name} {c.last_name}".strip(),
        'company': c.account.name if c.account else None
    } for c in contacts]
    
    # Ottieni tutti i leads con email
    leads = Lead.query.filter(Lead.email.isnot(None), Lead.email != '').all()
    leads_data = [{
        'id': l.id,
        'type': 'lead',
        'email': l.email,
        'first_name': l.contact_first_name,
        'last_name': l.contact_last_name,
        'name': f"{l.contact_first_name or ''} {l.contact_last_name or ''}".strip() or l.company_name,
        'company': l.company_name
    } for l in leads]
    
    return jsonify({
        'contacts': contacts_data,
        'leads': leads_data,
        'all': contacts_data + leads_data
    })


@app.route('/api/email/lists/unique-recipients', methods=['POST'])
def get_unique_recipients():
    """Calcola destinatari unici da più liste (elimina duplicati)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    data = request.json
    list_ids = data.get('list_ids', [])
    
    if not list_ids:
        return jsonify({'success': True, 'unique_count': 0, 'total_with_duplicates': 0})
    
    # Recupera tutte le subscription dalle liste selezionate
    subscriptions = EmailListSubscription.query.filter(
        EmailListSubscription.list_id.in_(list_ids),
        EmailListSubscription.status == 'subscribed'
    ).all()
    
    # Conta totali (con duplicati)
    total_with_duplicates = len(subscriptions)
    
    # Deduplica per email (case-insensitive)
    unique_emails = set()
    for sub in subscriptions:
        if sub.email:
            unique_emails.add(sub.email.lower().strip())
    
    unique_count = len(unique_emails)
    duplicates = total_with_duplicates - unique_count
    
    return jsonify({
        'success': True,
        'unique_count': unique_count,
        'total_with_duplicates': total_with_duplicates,
        'duplicates': duplicates,
        'list_count': len(list_ids)
    })


@app.route('/api/email/templates/init-defaults', methods=['POST'])
def init_default_email_templates():
    """Inizializza template email predefiniti"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    
    try:
        from email_template_loader import get_default_templates, create_sample_content
        
        templates = get_default_templates()
        sample_content = create_sample_content()
        
        created = []
        for key, template_data in templates.items():
            # Verifica se esiste già
            existing = EmailTemplate.query.filter_by(name=template_data['name']).first()
            if existing:
                continue
            
            # Sostituisci placeholder content
            html_content = template_data['html_content']
            if key in sample_content:
                html_content = html_content.replace('{{content}}', sample_content[key])
            
            template = EmailTemplate(
                name=template_data['name'],
                subject=template_data['subject'],
                html_content=html_content,
                category=template_data['category'],
                owner_id=user_id
            )
            db.session.add(template)
            created.append(template_data['name'])
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Creati {len(created)} template predefiniti',
            'created': created
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ==================== API TICKET ====================

def generate_ticket_number():
    """Genera numero ticket univoco: TKT-YYYY-NNN"""
    from datetime import datetime
    year = datetime.now().year
    last_ticket = Ticket.query.filter(Ticket.ticket_number.like(f'TKT-{year}-%')).order_by(Ticket.ticket_number.desc()).first()
    if last_ticket:
        try:
            last_num = int(last_ticket.ticket_number.split('-')[-1])
            new_num = last_num + 1
        except:
            new_num = 1
    else:
        new_num = 1
    return f'TKT-{year}-{new_num:03d}'


@app.route('/api/tickets', methods=['GET', 'POST'])
def tickets():
    """Gestisci ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    # Lista operatori autorizzati
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    if request.method == 'GET':
        # Filtri
        status_filter = request.args.get('status')
        tab_filter = request.args.get('tab', 'all')  # all, my, open, closed
        
        query = Ticket.query
        
        # Clienti vedono solo i loro ticket
        if not is_operatore:
            query = query.filter_by(customer_id=user_id)
        # Operatori vedono tutti, ma possono filtrare
        elif tab_filter == 'my':
            query = query.filter_by(owner_id=user_id)
        elif tab_filter == 'open':
            query = query.filter(Ticket.status.in_(['Aperto', 'In Lavorazione', 'In Attesa Cliente']))
        elif tab_filter == 'closed':
            query = query.filter_by(status='Chiuso')
        
        if status_filter:
            query = query.filter_by(status=status_filter)
        
        tickets_list = query.order_by(Ticket.created_at.desc()).all()
        
        return jsonify({
            'tickets': [{
                'id': t.id,
                'ticket_number': t.ticket_number,
                'title': t.title,
                'description': t.description,
                'status': t.status,
                'priority': t.priority,
                'category': t.category,
                'product_category': t.product_category,
                'customer_id': t.customer_id,
                'customer_name': f"{t.customer.first_name} {t.customer.last_name}" if t.customer else 'Unknown',
                'owner_id': t.owner_id,
                'owner_name': f"{t.owner.first_name} {t.owner.last_name}" if t.owner else None,
                'is_my_ticket': t.customer_id == user_id,
                'is_my_ownership': t.owner_id == user_id,
                'created_at': t.created_at.isoformat() if t.created_at else None,
                'updated_at': t.updated_at.isoformat() if t.updated_at else None,
                'comments_count': len(t.comments)
            } for t in tickets_list],
            'is_operatore': is_operatore
        })
    
    elif request.method == 'POST':
        try:
            data = request.json
            if not data:
                return jsonify({'error': 'Dati mancanti'}), 400
            
            ticket_number = generate_ticket_number()
            
            ticket = Ticket(
                ticket_number=ticket_number,
                title=data.get('title', ''),
                description=data.get('description', ''),
                status='Aperto',
                priority=data.get('priority', 'Media'),
                category=data.get('category'),
                product_category=data.get('product_category'),
                customer_id=user_id
            )
            
            db.session.add(ticket)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'ticket': {
                    'id': ticket.id,
                    'ticket_number': ticket.ticket_number,
                    'title': ticket.title
                }
            })
        except Exception as e:
            import traceback
            traceback.print_exc()
            db.session.rollback()
            return jsonify({'error': f'Errore creazione ticket: {str(e)}'}), 500


@app.route('/api/tickets/<int:ticket_id>', methods=['GET', 'PUT'])
def get_or_update_ticket(ticket_id):
    """Ottieni o aggiorna ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    
    # Verifica permessi: clienti vedono solo i loro ticket
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        # Filtra commenti: clienti vedono solo quelli pubblici
        all_comments = ticket.comments
        if not is_operatore:
            all_comments = [c for c in all_comments if not c.is_internal]
        
        return jsonify({
            'ticket': {
                'id': ticket.id,
                'ticket_number': ticket.ticket_number,
                'title': ticket.title,
                'description': ticket.description,
                'status': ticket.status,
                'priority': ticket.priority,
                'category': ticket.category,
                'product_category': ticket.product_category,
                'customer_id': ticket.customer_id,
                'customer_name': f"{ticket.customer.first_name} {ticket.customer.last_name}" if ticket.customer else 'Unknown',
                'customer_email': ticket.customer.email if ticket.customer else None,
                'owner_id': ticket.owner_id,
                'owner_name': f"{ticket.owner.first_name} {ticket.owner.last_name}" if ticket.owner else None,
                'created_at': ticket.created_at.isoformat() if ticket.created_at else None,
                'updated_at': ticket.updated_at.isoformat() if ticket.updated_at else None,
                'comments': [{
                    'id': c.id,
                    'message': c.message,
                    'is_internal': c.is_internal,
                    'user_id': c.user_id,
                    'user_name': f"{c.user.first_name} {c.user.last_name}" if c.user else 'Unknown',
                    'created_at': c.created_at.isoformat() if c.created_at else None
                } for c in ticket.comments if not c.is_internal or is_operatore]
            },
            'is_operatore': is_operatore
        })
    
    elif request.method == 'PUT':
        data = request.json
        
        # Solo operatori possono modificare status, priority, owner
        if is_operatore:
            if 'status' in data:
                ticket.status = data['status']
            if 'priority' in data:
                ticket.priority = data['priority']
            if 'owner_id' in data:
                ticket.owner_id = data['owner_id']
            if 'category' in data:
                ticket.category = data['category']
            if 'product_category' in data:
                ticket.product_category = data['product_category']
        
        # Tutti possono modificare title e description
        if 'title' in data:
            ticket.title = data['title']
        if 'description' in data:
            ticket.description = data['description']
        
        db.session.commit()
        
        return jsonify({'success': True, 'ticket_id': ticket.id})


@app.route('/api/tickets/<int:ticket_id>/take-ownership', methods=['POST'])
def take_ticket_ownership(ticket_id):
    """Prendi ownership di un ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    if not is_operatore:
        return jsonify({'error': 'Solo operatori possono prendere ownership'}), 403
    
    ticket = Ticket.query.get_or_404(ticket_id)
    old_owner_id = ticket.owner_id
    ticket.owner_id = user_id
    if ticket.status == 'Aperto':
        ticket.status = 'In Lavorazione'
    
    # Crea notifica per il nuovo owner (se diverso dal vecchio)
    if old_owner_id != user_id:
        create_notification(
            user_id=user_id,
            title=f'Ticket {ticket.ticket_number} assegnato',
            message=f'Ti è stato assegnato il ticket "{ticket.title}"',
            type='info',
            related_ticket_id=ticket_id
        )
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': ticket.owner_id,
        'owner_name': f"{ticket.owner.first_name} {ticket.owner.last_name}" if ticket.owner else 'Unknown'
    })


@app.route('/api/tickets/<int:ticket_id>/assign-ownership', methods=['POST'])
def assign_ticket_ownership(ticket_id):
    """Assegna un ticket a un operatore"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    if not is_operatore:
        return jsonify({'error': 'Solo operatori possono assegnare ticket'}), 403
    
    data = request.get_json()
    new_owner_id = data.get('owner_id')
    
    if not new_owner_id:
        return jsonify({'error': 'owner_id richiesto'}), 400
    
    ticket = Ticket.query.get_or_404(ticket_id)
    old_owner_id = ticket.owner_id
    ticket.owner_id = new_owner_id
    if ticket.status == 'Aperto':
        ticket.status = 'In Lavorazione'
    
    # Crea notifica per il nuovo owner (se diverso dal vecchio)
    if old_owner_id != new_owner_id:
        new_owner = User.query.get(new_owner_id)
        create_notification(
            user_id=new_owner_id,
            title=f'Ticket {ticket.ticket_number} assegnato',
            message=f'Ti è stato assegnato il ticket "{ticket.title}" da {user.first_name} {user.last_name}',
            type='info',
            related_ticket_id=ticket_id
        )
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'owner_id': ticket.owner_id,
        'owner_name': f"{ticket.owner.first_name} {ticket.owner.last_name}" if ticket.owner else 'Unknown'
    })


@app.route('/api/users/operatori', methods=['GET'])
def get_operatori():
    """Ottieni lista operatori"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    operatori_usernames = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo']
    
    # Cerca utenti con ruolo operatore o admin, o username che contiene i nomi degli operatori
    operatori = User.query.filter(
        (User.role == 'admin') | 
        (User.role == 'operatore') |
        db.or_(*[User.username.ilike(f'%{op}%') for op in operatori_usernames])
    ).all()
    
    return jsonify({
        'operatori': [{
            'id': u.id,
            'username': u.username,
            'first_name': u.first_name,
            'last_name': u.last_name,
            'full_name': f"{u.first_name} {u.last_name}"
        } for u in operatori]
    })


@app.route('/api/tickets/<int:ticket_id>/comments', methods=['GET', 'POST'])
def ticket_comments(ticket_id):
    """Gestisci commenti su ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    
    # Verifica permessi
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        comments = ticket.comments
        # Clienti non vedono commenti interni
        if not is_operatore:
            comments = [c for c in comments if not c.is_internal]
        
        return jsonify({
            'comments': [{
                'id': c.id,
                'message': c.message,
                'is_internal': c.is_internal,
                'user_id': c.user_id,
                'user_name': f"{c.user.first_name} {c.user.last_name}" if c.user else 'Unknown',
                'created_at': c.created_at.isoformat() if c.created_at else None
            } for c in comments]
        })
    
    elif request.method == 'POST':
        data = request.json
        comment = TicketComment(
            ticket_id=ticket_id,
            user_id=user_id,
            message=data['message'],
            is_internal=data.get('is_internal', False)
        )
        
        db.session.add(comment)
        
        # Aggiorna status se cliente risponde
        if not is_operatore and ticket.status == 'In Attesa Cliente':
            ticket.status = 'In Lavorazione'
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'comment': {
                'id': comment.id,
                'message': comment.message,
                'user_name': f"{comment.user.first_name} {comment.user.last_name}" if comment.user else 'Unknown',
                'created_at': comment.created_at.isoformat() if comment.created_at else None
            }
        })


# ==================== GESTIONE ALLEGATI TICKET ====================

def get_storage_usage_mb():
    """Calcola spazio utilizzato in MB"""
    total_size = 0
    for root, dirs, files in os.walk(app.config['UPLOAD_FOLDER']):
        for file in files:
            filepath = os.path.join(root, file)
            if os.path.exists(filepath):
                total_size += os.path.getsize(filepath)
    return total_size / (1024 * 1024)  # Converti in MB


def get_client_storage_usage_mb(user_id):
    """Calcola spazio utilizzato da un cliente in MB"""
    attachments = TicketAttachment.query.filter_by(user_id=user_id).all()
    total_size = sum(att.file_size or 0 for att in attachments)
    return total_size / (1024 * 1024)  # Converti in MB


def compress_image(input_path, output_path, quality=85, max_size=(1920, 1080)):
    """Comprimi immagine usando Pillow"""
    try:
        from PIL import Image
        img = Image.open(input_path)
        
        # Converti RGBA in RGB se necessario
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        
        # Ridimensiona se troppo grande
        if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Salva compressa
        img.save(output_path, 'JPEG', quality=quality, optimize=True)
        return True
    except Exception as e:
        print(f"Errore compressione immagine: {e}")
        return False


def compress_video(input_path, output_path):
    """Comprimi video usando ffmpeg"""
    try:
        import subprocess
        # Comprimi video: riduci risoluzione e bitrate
        cmd = [
            'ffmpeg', '-i', input_path,
            '-vf', 'scale=1280:-1',  # Max width 1280px
            '-c:v', 'libx264',
            '-preset', 'medium',
            '-crf', '28',  # Qualità compressa
            '-c:a', 'aac',
            '-b:a', '128k',
            '-y',  # Sovrascrivi
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        return result.returncode == 0
    except Exception as e:
        print(f"Errore compressione video: {e}")
        return False


def cleanup_old_files():
    """Elimina file vecchi se lo spazio è limitato"""
    total_usage_mb = get_storage_usage_mb()
    total_storage_mb = app.config['TOTAL_STORAGE_GB'] * 1024
    
    # Se uso più dell'80% dello spazio, elimina file più vecchi
    if total_usage_mb > (total_storage_mb * 0.8):
        # Ordina allegati per data (più vecchi prima)
        old_attachments = TicketAttachment.query.order_by(TicketAttachment.created_at.asc()).limit(50).all()
        
        deleted_count = 0
        for att in old_attachments:
            if os.path.exists(att.file_path):
                try:
                    os.remove(att.file_path)
                    db.session.delete(att)
                    deleted_count += 1
                except Exception as e:
                    print(f"Errore eliminazione file {att.file_path}: {e}")
        
        db.session.commit()
        return deleted_count
    
    return 0


@app.route('/api/tickets/<int:ticket_id>/attachments', methods=['GET', 'POST', 'DELETE'])
def ticket_attachments(ticket_id):
    """Gestisci allegati ai ticket"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    
    # Verifica permessi
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if request.method == 'GET':
        attachments = ticket.attachments
        return jsonify({
            'attachments': [{
                'id': att.id,
                'filename': att.filename if hasattr(att, 'filename') else (att.original_filename if hasattr(att, 'original_filename') else 'file'),
                'file_type': att.file_type,
                'file_size': att.file_size,
                'mime_type': getattr(att, 'mime_type', None),
                'is_compressed': att.is_compressed,
                'created_at': att.created_at.isoformat() if att.created_at else None,
                'user_name': f"{att.user.first_name} {att.user.last_name}" if att.user else 'Unknown'
            } for att in attachments]
        })
    
    elif request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({'error': 'Nessun file caricato'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nessun file selezionato'}), 400
        
        # Verifica quota cliente
        if user_role == 'cliente':
            client_usage_mb = get_client_storage_usage_mb(user_id)
            if client_usage_mb >= app.config['CLIENT_QUOTA_MB']:
                return jsonify({'error': f'Quota esaurita. Spazio utilizzato: {client_usage_mb:.2f}MB / {app.config["CLIENT_QUOTA_MB"]}MB'}), 400
        
        # Verifica spazio totale
        cleanup_old_files()  # Pulisci file vecchi se necessario
        
        # Salva file originale
        original_filename = secure_filename(file.filename)
        file_ext = os.path.splitext(original_filename)[1].lower()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{ticket_id}_{timestamp}_{original_filename}"
        original_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        
        file.save(original_path)
        original_size = os.path.getsize(original_path)
        
        # Determina tipo file
        is_image = file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']
        is_video = file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.webm']
        is_pdf = file_ext == '.pdf'
        
        if not (is_image or is_video or is_pdf):
            os.remove(original_path)
            return jsonify({'error': 'Formato file non supportato. Usa immagini (JPG, PNG), video (MP4) o PDF'}), 400
        
        file_type = 'image' if is_image else ('video' if is_video else 'pdf')
        compressed_path = original_path
        is_compressed = False
        
        # Comprimi se necessario
        if is_image and original_size > 2 * 1024 * 1024:  # > 2MB
            compressed_filename = f"compressed_{safe_filename}"
            compressed_path = os.path.join(app.config['UPLOAD_FOLDER'], compressed_filename)
            if compress_image(original_path, compressed_path):
                os.remove(original_path)
                original_path = compressed_path
                is_compressed = True
        elif is_video and original_size > 10 * 1024 * 1024:  # > 10MB
            compressed_filename = f"compressed_{safe_filename}"
            compressed_path = os.path.join(app.config['UPLOAD_FOLDER'], compressed_filename)
            if compress_video(original_path, compressed_path):
                os.remove(original_path)
                original_path = compressed_path
                is_compressed = True
        
        final_size = os.path.getsize(original_path)
        
        # Crea record allegato
        attachment = TicketAttachment(
            ticket_id=ticket_id,
            user_id=user_id,
            filename=safe_filename,
            original_filename=original_filename,
            file_path=original_path,
            file_type=file_type,
            file_size=final_size,
            mime_type=file.content_type,
            is_compressed=is_compressed
        )
        
        db.session.add(attachment)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'attachment': {
                'id': attachment.id,
                'filename': attachment.original_filename,
                'file_size': attachment.file_size,
                'is_compressed': attachment.is_compressed
            }
        })
    
    elif request.method == 'DELETE':
        attachment_id = request.json.get('attachment_id')
        if not attachment_id:
            return jsonify({'error': 'ID allegato mancante'}), 400
        
        attachment = TicketAttachment.query.get_or_404(attachment_id)
        
        # Solo il proprietario o operatore può eliminare
        if not is_operatore and attachment.user_id != user_id:
            return jsonify({'error': 'Non autorizzato'}), 403
        
        # Elimina file
        if os.path.exists(attachment.file_path):
            try:
                os.remove(attachment.file_path)
            except Exception as e:
                print(f"Errore eliminazione file: {e}")
        
        db.session.delete(attachment)
        db.session.commit()
        
        return jsonify({'success': True})


@app.route('/api/tickets/<int:ticket_id>/attachments/<int:attachment_id>/download')
def download_attachment(ticket_id, attachment_id):
    """Scarica allegato"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    operatori = ['sandy', 'turchetto', 'luca', 'furlan', 'leonardo', 'giovanni', 'pitton', 'mattia', 'scevola', 'filippo', 'mariuzzo', 'admin']
    is_operatore = user_role == 'admin' or any(op in user.username.lower() for op in operatori)
    
    ticket = Ticket.query.get_or_404(ticket_id)
    attachment = TicketAttachment.query.get_or_404(attachment_id)
    
    # Verifica permessi
    if not is_operatore and ticket.customer_id != user_id:
        return jsonify({'error': 'Non autorizzato'}), 403
    
    if not os.path.exists(attachment.file_path):
        return jsonify({'error': 'File non trovato'}), 404
    
    # Per PDF, permettere visualizzazione inline per anteprima
    as_attachment = request.args.get('download', 'false').lower() == 'true'
    mimetype = 'application/pdf' if attachment.file_type == 'pdf' else None
    
    return send_file(attachment.file_path, as_attachment=as_attachment, download_name=attachment.original_filename, mimetype=mimetype)


@app.route('/api/storage/usage')
def storage_usage():
    """Ottieni informazioni spazio utilizzato"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    user = User.query.get(user_id)
    user_role = user.role if user else 'commerciale'
    
    total_usage_mb = get_storage_usage_mb()
    total_storage_mb = app.config['TOTAL_STORAGE_GB'] * 1024
    
    client_usage_mb = 0
    client_quota_mb = 0
    if user_role == 'cliente':
        client_usage_mb = get_client_storage_usage_mb(user_id)
        client_quota_mb = app.config['CLIENT_QUOTA_MB']
    
    return jsonify({
        'total_usage_mb': round(total_usage_mb, 2),
        'total_storage_mb': total_storage_mb,
        'total_usage_percent': round((total_usage_mb / total_storage_mb) * 100, 2),
        'client_usage_mb': round(client_usage_mb, 2),
        'client_quota_mb': client_quota_mb,
        'client_usage_percent': round((client_usage_mb / client_quota_mb * 100) if client_quota_mb > 0 else 0, 2)
    })


def create_notification(user_id, title, message, type='info', related_ticket_id=None, related_opportunity_id=None):
    """Crea una nuova notifica"""
    try:
        notification = Notification(
            user_id=user_id,
            title=title,
            message=message,
            type=type,
            related_ticket_id=related_ticket_id,
            related_opportunity_id=related_opportunity_id
        )
        db.session.add(notification)
        db.session.commit()
        return notification
    except Exception as e:
        db.session.rollback()
        print(f"Errore creazione notifica: {e}")
        return None


@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    """Ottieni notifiche dell'utente"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    unread_only = request.args.get('unread_only', 'false').lower() == 'true'
    
    query = Notification.query.filter_by(user_id=user_id)
    if unread_only:
        query = query.filter_by(is_read=False)
    
    notifications = query.order_by(Notification.created_at.desc()).limit(50).all()
    
    return jsonify({
        'notifications': [{
            'id': n.id,
            'title': n.title,
            'message': n.message,
            'type': n.type,
            'is_read': n.is_read,
            'related_ticket_id': n.related_ticket_id,
            'related_opportunity_id': n.related_opportunity_id,
            'created_at': n.created_at.isoformat() if n.created_at else None
        } for n in notifications]
    })


@app.route('/api/notifications/<int:notification_id>/read', methods=['POST'])
def mark_notification_read(notification_id):
    """Marca una notifica come letta"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    notification = Notification.query.filter_by(id=notification_id, user_id=user_id).first()
    
    if not notification:
        return jsonify({'error': 'Notifica non trovata'}), 404
    
    notification.is_read = True
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/api/notifications/read-all', methods=['POST'])
def mark_all_notifications_read():
    """Marca tutte le notifiche come lette"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user_id = session['user_id']
    Notification.query.filter_by(user_id=user_id, is_read=False).update({'is_read': True})
    db.session.commit()
    
    return jsonify({'success': True})


def init_db():
    """Inizializza database"""
    with app.app_context():
        db.create_all()
        
        # Crea admin
        if not User.query.filter_by(username='admin').first():
            admin = User(
                username='admin',
                password=generate_password_hash('admin123'),
                first_name='Admin',
                last_name='System',
                role='admin'
            )
            db.session.add(admin)
        
        # Crea commerciale demo
        if not User.query.filter_by(username='demo').first():
            demo = User(
                username='demo',
                password=generate_password_hash('demo123'),
                first_name='Demo',
                last_name='User',
                role='commerciale'
            )
            db.session.add(demo)
        
        db.session.commit()
        print("✓ Database inizializzato")


def send_task_notification_email(task, opportunity_name, opportunity=None):
    """Invia email alla persona assegnata quando viene creata un'attività manualmente"""
    # Verifica che ci sia una persona assegnata
    if not task.assigned_to:
        print("⚠️  Task senza persona assegnata, email non inviata")
        return
    
    # Mappa username a email commerciali
    username_email_map = {
        'giovanni': 'giovanni.pitton@mekosrl.it',
        'pitton': 'giovanni.pitton@mekosrl.it',
        'mauro': 'mauro.zanet@mekosrl.it',
        'zanet': 'mauro.zanet@mekosrl.it',
        'davide': 'davide.cal@mekosrl.it',
        'cal': 'davide.cal@mekosrl.it',
        'filippo': 'filippo.mariuzzo@mekosrl.it',
        'mariuzzo': 'filippo.mariuzzo@mekosrl.it'
    }
    
    # Ottieni email dalla persona assegnata
    assigned_user = task.assigned_to
    recipient_email = None
    
    # Prima prova a usare l'email dell'utente se presente
    if assigned_user.email:
        recipient_email = assigned_user.email
    else:
        # Altrimenti mappa username a email
        username_lower = assigned_user.username.lower()
        for key, email in username_email_map.items():
            if key in username_lower:
                recipient_email = email
                break
    
    if not recipient_email:
        print(f"⚠️  Nessuna email trovata per utente {assigned_user.username}, email non inviata")
        return
    
    # Ottieni informazioni task
    task_type_map = {
        'call': 'Chiamare',
        'send_offer': 'Inviare Offerta',
        'recall': 'Richiamare',
        'meeting': 'Riunione',
        'other': 'Altro'
    }
    task_type_label = task_type_map.get(task.task_type, task.task_type)
    
    assigned_to_name = f"{assigned_user.first_name} {assigned_user.last_name}"
    due_date_str = task.due_date.strftime('%d/%m/%Y') if task.due_date else "Non specificata"
    
    # Costruisci corpo email
    account_name = "N/A"
    contact_name = "N/A"
    opportunity_amount = 0.0
    
    if opportunity:
        account_name = opportunity.account.name if opportunity.account else "N/A"
        if opportunity.contact:
            contact_name = f"{opportunity.contact.first_name} {opportunity.contact.last_name}"
        opportunity_amount = float(opportunity.amount) if opportunity.amount else 0.0
    
    subject = f"📋 Nuova Attività: {task_type_label} - {opportunity_name}"
    
    body = f"""
Ciao,

È stata creata una nuova attività nel CRM:

📋 Tipo Attività: {task_type_label}
📝 Descrizione: {task.description or 'Nessuna descrizione'}
👤 Assegnata a: {assigned_to_name}
📅 Scadenza: {due_date_str}

💼 Opportunità: {opportunity_name}
🏢 Account: {account_name}
👥 Contatto: {contact_name}
💰 Valore: €{opportunity_amount:,.2f}

Puoi visualizzare l'attività nel CRM: http://crm.infomekosrl.it:8000

Cordiali saluti,
Sistema CRM ME.KO.
    """
    
    body_html = f"""
<html>
<body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
    <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
        <h2 style="color: #2563eb;">Nuova Attività Assegnata</h2>
        <p>Ciao {assigned_user.first_name},</p>
        <p>È stata creata una nuova attività nel CRM e ti è stata assegnata:</p>
        
        <div style="background: #f3f4f6; padding: 16px; border-radius: 8px; margin: 20px 0;">
            <p style="margin: 8px 0;"><strong>📋 Tipo Attività:</strong> {task_type_label}</p>
            <p style="margin: 8px 0;"><strong>📝 Descrizione:</strong> {task.description or 'Nessuna descrizione'}</p>
            <p style="margin: 8px 0;"><strong>👤 Assegnata a:</strong> {assigned_to_name}</p>
            <p style="margin: 8px 0;"><strong>📅 Scadenza:</strong> {due_date_str}</p>
        </div>
        
        <div style="background: #eff6ff; padding: 16px; border-radius: 8px; margin: 20px 0;">
            <h3 style="margin-top: 0; color: #1e40af;">Dettagli Opportunità</h3>
            <p style="margin: 8px 0;"><strong>💼 Opportunità:</strong> {opportunity_name}</p>
            <p style="margin: 8px 0;"><strong>🏢 Account:</strong> {account_name}</p>
            <p style="margin: 8px 0;"><strong>👥 Contatto:</strong> {contact_name}</p>
            <p style="margin: 8px 0;"><strong>💰 Valore:</strong> €{opportunity_amount:,.2f}</p>
        </div>
        
        <div style="margin-top: 30px; text-align: center;">
            <a href="http://crm.infomekosrl.it:8000" 
               style="background: #2563eb; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block;">
                Vai al CRM
            </a>
        </div>
        
        <p style="margin-top: 30px; color: #6b7280; font-size: 12px;">
            Cordiali saluti,<br>
            Sistema CRM ME.KO.
        </p>
    </div>
</body>
</html>
    """
    
    body_text = f"""
Ciao {assigned_user.first_name},

È stata creata una nuova attività nel CRM e ti è stata assegnata:

📋 Tipo Attività: {task_type_label}
📝 Descrizione: {task.description or 'Nessuna descrizione'}
👤 Assegnata a: {assigned_to_name}
📅 Scadenza: {due_date_str}

💼 Opportunità: {opportunity_name}
🏢 Account: {account_name}
👥 Contatto: {contact_name}
💰 Valore: €{opportunity_amount:,.2f}

Puoi visualizzare l'attività nel CRM: http://crm.infomekosrl.it:8000

Cordiali saluti,
Sistema CRM ME.KO.
    """
    
    # Usa Mailgun come per l'email marketing
    try:
        from mailgun_service import MailgunService
        mailgun = MailgunService()
        
        # Prepara dati email
        email_data = {
            'subject': subject,
            'html_content': body_html,
            'text_content': body_text
        }
        
        # Invia email singola
        result = mailgun.send_single_email(
            to_email=recipient_email,
            to_name=assigned_to_name,
            email_data=email_data
        )
        
        if result.get('success'):
            print(f"✅ Email inviata a {recipient_email} per task {task.id}")
        else:
            print(f"⚠️  Errore invio email: {result.get('error', 'Unknown error')}")
    
    except ImportError:
        # Fallback: usa SMTP diretto se Mailgun non disponibile
        print("⚠️  Mailgun non disponibile, uso SMTP diretto")
        try:
            smtp_server = os.getenv('SMTP_SERVER', 'smtp.office365.com')
            smtp_port = int(os.getenv('SMTP_PORT', '587'))
            smtp_user = os.getenv('SMTP_USER', 'noreply@mekosrl.it')
            smtp_password = os.getenv('SMTP_PASSWORD', '')
            
            if not smtp_password:
                print("⚠️  SMTP_PASSWORD non configurato, email non inviata")
                return
            
            msg = MIMEMultipart('alternative')
            msg['From'] = smtp_user
            msg['To'] = recipient_email
            msg['Subject'] = subject
            
            msg.attach(MIMEText(body_text, 'plain', 'utf-8'))
            msg.attach(MIMEText(body_html, 'html', 'utf-8'))
            
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            server.quit()
            
            print(f"✅ Email inviata via SMTP a {recipient_email} per task {task.id}")
        except Exception as e:
            print(f"❌ Errore invio email SMTP: {str(e)}")
    
    except Exception as e:
        print(f"❌ Errore invio email: {str(e)}")
        # Non bloccare la creazione del task se l'email fallisce


@app.route('/api/tasks/morning-summary', methods=['GET'])
def morning_task_summary():
    """Riepilogo mattutino delle attività da fare"""
    try:
        if 'user_id' not in session:
            return jsonify({'success': False, 'error': 'Non autorizzato'}), 401
        
        today = date.today()
        
        # Trova tutte le attività non completate con scadenza oggi o passata
        try:
            pending_tasks = db.session.execute(text("""
                SELECT t.id, t.task_type, t.description, t.due_date, t.opportunity_id,
                       o.name as opportunity_name, o.amount,
                       a.name as account_name,
                       u.first_name || ' ' || u.last_name as assigned_to_name,
                       u.email as assigned_to_email
                FROM opportunity_tasks t
                LEFT JOIN opportunities o ON t.opportunity_id = o.id
                LEFT JOIN accounts a ON o.account_id = a.id
                LEFT JOIN users u ON t.assigned_to_id = u.id
                WHERE t.is_completed = 0
                AND (t.due_date IS NULL OR t.due_date <= :today)
                ORDER BY t.due_date ASC, t.created_at DESC
            """), {"today": today}).fetchall()
        except Exception as e:
            # Se la query fallisce, prova con una query più semplice
            print(f"Errore query morning-summary: {e}")
            pending_tasks = []
        
        tasks_list = []
        for task in pending_tasks:
            try:
                tasks_list.append({
                    'id': task[0] if len(task) > 0 else None,
                    'task_type': task[1] if len(task) > 1 else 'other',
                    'description': task[2] if len(task) > 2 else '',
                    'due_date': task[3].isoformat() if len(task) > 3 and task[3] else None,
                    'opportunity_id': task[4] if len(task) > 4 else None,
                    'opportunity_name': task[5] if len(task) > 5 and task[5] else 'N/A',
                    'amount': float(task[6]) if len(task) > 6 and task[6] else 0,
                    'account_name': task[7] if len(task) > 7 and task[7] else 'N/A',
                    'assigned_to_name': task[8] if len(task) > 8 and task[8] else 'Non assegnato',
                    'assigned_to_email': task[9] if len(task) > 9 else None
                })
            except Exception as e:
                print(f"Errore processamento task: {e}")
                continue
        
        return jsonify({
            'success': True,
            'date': today.isoformat(),
            'total_pending': len(tasks_list),
            'tasks': tasks_list
        })
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Errore morning_task_summary: {e}\n{error_trace}")
        return jsonify({
            'success': False,
            'error': str(e),
            'trace': error_trace
        }), 500


@app.route('/api/opportunities/<int:opportunity_id>/update-category', methods=['PUT'])
def update_opportunity_category(opportunity_id):
    """Aggiorna la categoria (supplier_category) di un'opportunità"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    opportunity = Opportunity.query.get_or_404(opportunity_id)
    data = request.json
    
    if 'supplier_category' not in data:
        return jsonify({'error': 'supplier_category richiesto'}), 400
    
    opportunity.supplier_category = data['supplier_category']
    opportunity.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({
        'success': True,
        'opportunity': {
            'id': opportunity.id,
            'name': opportunity.name,
            'supplier_category': opportunity.supplier_category
        }
    })


@app.route('/api/geocode', methods=['GET'])
def geocode_proxy():
    """Proxy per geocoding Nominatim per evitare problemi CORS e rate limiting"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    address = request.args.get('q', '')
    if not address:
        return jsonify({'error': 'Indirizzo mancante'}), 400
    
    try:
        import requests
        import time
        
        # Usa Nominatim con User-Agent appropriato
        url = 'https://nominatim.openstreetmap.org/search'
        params = {
            'format': 'json',
            'q': address,
            'limit': 1,
            'addressdetails': 1
        }
        headers = {
            'User-Agent': 'ME.KO. CRM (contact@mekosrl.it)'
        }
        
        # Rate limiting: aspetta 1 secondo tra le richieste
        time.sleep(1)
        
        response = requests.get(url, params=params, headers=headers, timeout=10)
        response.raise_for_status()
        
        data = response.json()
        return jsonify(data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/assistant', methods=['POST'])
def assistant():
    """Natural-language assistant for CRM actions."""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401

    data = request.json or {}
    user_message = (data.get('message') or '').strip()
    if not user_message:
        return jsonify({'error': 'Messaggio vuoto'}), 400

    api_key = os.environ.get('OPENAI_API_KEY') or os.environ.get('OPENAI_KEY')
    if not api_key:
        return jsonify({'error': 'OPENAI_API_KEY non configurata'}), 500

    model = os.environ.get('OPENAI_MODEL', 'gpt-4o-mini')

    system_prompt = (
        "You are a CRM assistant. Reply ONLY in JSON with keys: reply (string), actions (array). "
        "Allowed action types: update_opportunity_stage, update_opportunity_category, update_opportunity_heat_level, "
        "update_opportunity_amount, update_opportunity_close_date, create_task. "
        "For update_opportunity_stage: provide opportunity_id or opportunity_name, and stage. "
        "For update_opportunity_category: provide opportunity_id or opportunity_name, and category (UR, MiR, Unitree, Altri). "
        "For update_opportunity_heat_level: provide opportunity_id or opportunity_name, and heat_level (calda, tiepida, fredda_speranza, fredda_gelo, da_categorizzare). "
        "For update_opportunity_amount: provide opportunity_id or opportunity_name, and amount (number). "
        "For update_opportunity_close_date: provide opportunity_id or opportunity_name, and close_date (YYYY-MM-DD). "
        "For create_task: provide task_type (send_offer, call, recall, other), description, "
        "optional due_date (YYYY-MM-DD), optional assigned_to_id (user ID) or assigned_to_name (username like 'giovanni', 'mauro', 'davide', 'filippo'), "
        "and opportunity_id/opportunity_name or account_id/account_name. "
        "If assigned_to_name is provided, try to match it to a user (e.g., 'giovanni' or 'pitton' for Giovanni Pitton). "
        "Use standard stage values: Qualification, Needs Analysis, Value Proposition, Id. Decision Makers, "
        "Perception Analysis, Proposal/Price Quote, Negotiation/Review, Closed Won, Closed Lost. "
        "For names, be tolerant of typos and partial matches - use the closest match you can find."
    )

    # Log della richiesta
    log_entry = AssistantLog(
        user_id=session['user_id'],
        user_message=user_message,
        model_used=model
    )
    
    try:
        import requests

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            "temperature": 0.2,
            "response_format": {"type": "json_object"}
        }
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}"},
            json=payload,
            timeout=25
        )
        response.raise_for_status()
        result = response.json()
        content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
        
        # Salva token utilizzati se disponibili
        usage = result.get('usage', {})
        log_entry.tokens_used = usage.get('total_tokens')
        log_entry.ai_response_raw = json.dumps(result)
    except Exception as e:
        log_entry.errors = json.dumps({'error': str(e)})
        db.session.add(log_entry)
        db.session.commit()
        return jsonify({'error': f'Errore chiamata AI: {str(e)}'}), 500

    def parse_json(text_value):
        try:
            return json.loads(text_value)
        except Exception:
            try:
                start = text_value.find('{')
                end = text_value.rfind('}')
                if start != -1 and end != -1 and end > start:
                    return json.loads(text_value[start:end + 1])
            except Exception:
                return None
        return None

    parsed = parse_json(content)
    if not parsed:
        return jsonify({'error': 'Risposta AI non valida'}), 500

    reply = parsed.get('reply') or ''
    actions = parsed.get('actions') or []

    def normalize_stage(value):
        if not value:
            return None
        normalized = value.strip().lower()
        mapping = {
            'chiusa vinta': 'Closed Won',
            'vinta': 'Closed Won',
            'chiusa persa': 'Closed Lost',
            'persa': 'Closed Lost',
            'qualificazione': 'Qualification',
            'analisi necessita': 'Needs Analysis',
            'analisi necessit?': 'Needs Analysis',
            'proposta di valore': 'Value Proposition',
            'identificazione decision maker': 'Id. Decision Makers',
            'analisi percezione': 'Perception Analysis',
            'offerta': 'Proposal/Price Quote',
            'preventivo': 'Proposal/Price Quote',
            'negoziazione': 'Negotiation/Review'
        }
        return mapping.get(normalized, value)

    def normalize_task_type(value):
        if not value:
            return 'other'
        normalized = value.strip().lower()
        if normalized in ['call', 'chiamata', 'chiama']:
            return 'call'
        if normalized in ['recall', 'richiama']:
            return 'recall'
        if normalized in ['send_offer', 'invia offerta', 'offerta']:
            return 'send_offer'
        return 'other'

    def parse_date(value):
        if not value:
            return None
        try:
            return datetime.fromisoformat(value).date()
        except Exception:
            return None

    def fuzzy_match_name(search_name, names_list, threshold=0.6):
        """Fuzzy matching per trovare il nome più simile"""
        try:
            from difflib import SequenceMatcher
            search_lower = search_name.lower().strip()
            best_match = None
            best_score = 0
            
            for name in names_list:
                name_lower = name.lower().strip()
                # Calcola similarità
                score = SequenceMatcher(None, search_lower, name_lower).ratio()
                # Bonus se contiene la stringa di ricerca
                if search_lower in name_lower:
                    score += 0.2
                # Bonus se inizia con la stringa di ricerca
                if name_lower.startswith(search_lower):
                    score += 0.1
                
                if score > best_score:
                    best_score = score
                    best_match = name
            
            return best_match if best_score >= threshold else None
        except:
            return None

    def resolve_opportunity(action):
        opp_id = action.get('opportunity_id')
        if opp_id:
            return Opportunity.query.get(opp_id)
        name = (action.get('opportunity_name') or '').strip()
        if not name:
            return None
        
        # Prima prova match esatto/parziale
        matches = Opportunity.query.filter(Opportunity.name.ilike(f"%{name}%")).all()
        if len(matches) == 1:
            return matches[0]
        elif len(matches) > 1:
            # Se ci sono più match, usa fuzzy matching
            names = [m.name for m in matches]
            best_name = fuzzy_match_name(name, names, threshold=0.5)
            if best_name:
                return next((m for m in matches if m.name == best_name), None)
            # Se non trova match fuzzy, restituisci il primo
            return matches[0]
        
        # Se non trova nulla, prova ricerca più ampia (senza case sensitivity)
        all_opps = Opportunity.query.all()
        if all_opps:
            names = [o.name for o in all_opps]
            best_name = fuzzy_match_name(name, names, threshold=0.4)
            if best_name:
                return next((o for o in all_opps if o.name == best_name), None)
        
        return None

    def resolve_account(action):
        account_id = action.get('account_id')
        if account_id:
            return Account.query.get(account_id)
        name = (action.get('account_name') or '').strip()
        if not name:
            return None
        
        # Prima prova match esatto/parziale
        matches = Account.query.filter(Account.name.ilike(f"%{name}%")).all()
        if len(matches) == 1:
            return matches[0]
        elif len(matches) > 1:
            # Se ci sono più match, usa fuzzy matching
            names = [a.name for a in matches]
            best_name = fuzzy_match_name(name, names, threshold=0.5)
            if best_name:
                return next((a for a in matches if a.name == best_name), None)
            # Se non trova match fuzzy, restituisci il primo
            return matches[0]
        
        # Se non trova nulla, prova ricerca più ampia
        all_accounts = Account.query.all()
        if all_accounts:
            names = [a.name for a in all_accounts]
            best_name = fuzzy_match_name(name, names, threshold=0.4)
            if best_name:
                return next((a for a in all_accounts if a.name == best_name), None)
        
        return None

    results = []
    errors = []

    for action in actions:
        action_type = (action.get('type') or '').strip()
        if action_type == 'update_opportunity_stage':
            opp = resolve_opportunity(action)
            stage = normalize_stage(action.get('stage'))
            if not opp:
                errors.append({'type': action_type, 'error': 'Opportunita non trovata'})
                continue
            if not stage:
                errors.append({'type': action_type, 'error': 'Stage mancante'})
                continue
            opp.stage = stage
            opp.updated_at = datetime.utcnow()
            db.session.commit()
            results.append({'type': action_type, 'opportunity_id': opp.id, 'stage': stage})
        elif action_type == 'update_opportunity_category':
            opp = resolve_opportunity(action)
            category = (action.get('category') or '').strip()
            if not opp:
                errors.append({'type': action_type, 'error': 'Opportunita non trovata'})
                continue
            if not category:
                errors.append({'type': action_type, 'error': 'Categoria mancante'})
                continue
            opp.supplier_category = category
            opp.updated_at = datetime.utcnow()
            db.session.commit()
            results.append({'type': action_type, 'opportunity_id': opp.id, 'category': category})
        elif action_type == 'create_task':
            task_type = normalize_task_type(action.get('task_type'))
            description = (action.get('description') or '').strip()
            due_date = parse_date(action.get('due_date'))
            opp = resolve_opportunity(action)
            account = resolve_account(action)
            
            # Risolvi assigned_to_id
            assigned_to_id = None
            if action.get('assigned_to_id'):
                assigned_to_id = action.get('assigned_to_id')
            elif action.get('assigned_to_name'):
                # Cerca utente per nome (fuzzy matching)
                assigned_name = action.get('assigned_to_name').strip().lower()
                users = User.query.all()
                for user in users:
                    username_lower = user.username.lower()
                    first_name_lower = (user.first_name or '').lower()
                    last_name_lower = (user.last_name or '').lower()
                    full_name_lower = f"{first_name_lower} {last_name_lower}".strip()
                    
                    if (assigned_name in username_lower or 
                        assigned_name in first_name_lower or 
                        assigned_name in last_name_lower or
                        assigned_name in full_name_lower):
                        assigned_to_id = user.id
                        break
            else:
                assigned_to_id = session['user_id']  # Default: utente corrente
            
            if not description:
                errors.append({'type': action_type, 'error': 'Descrizione mancante'})
                continue
            
            task = OpportunityTask(
                opportunity_id=opp.id if opp else None,
                account_id=account.id if account else None,
                task_type=task_type,
                description=description,
                assigned_to_id=assigned_to_id,
                due_date=due_date
            )
            db.session.add(task)
            db.session.commit()
            
            # Invia email alla persona assegnata
            try:
                opp_name = opp.name if opp else (account.name if account else "Task globale")
                send_task_notification_email(task, opp_name, opp)
            except Exception as e:
                print(f"Errore invio email task (assistente): {str(e)}")
                # Non bloccare se l'email fallisce
            
            results.append({
                'type': action_type,
                'task_id': task.id,
                'task_type': task.task_type,
                'opportunity_id': task.opportunity_id,
                'account_id': task.account_id,
                'assigned_to_id': task.assigned_to_id
            })
        elif action_type == 'update_opportunity_heat_level':
            opp = resolve_opportunity(action)
            heat_level = (action.get('heat_level') or '').strip().lower()
            valid_heat_levels = ['calda', 'tiepida', 'fredda_speranza', 'fredda_gelo', 'da_categorizzare']
            if not opp:
                errors.append({'type': action_type, 'error': 'Opportunita non trovata'})
                continue
            if heat_level not in valid_heat_levels:
                errors.append({'type': action_type, 'error': f'Heat level non valido. Usa: {", ".join(valid_heat_levels)}'})
                continue
            
            # Aggiorna heat_level e probability automaticamente
            opp.heat_level = heat_level
            if heat_level == 'calda':
                opp.probability = 50
            elif heat_level == 'tiepida':
                opp.probability = 15
            elif heat_level == 'fredda_speranza':
                opp.probability = 5
            elif heat_level == 'fredda_gelo':
                opp.probability = 0
            else:
                opp.probability = 0
            opp.updated_at = datetime.utcnow()
            db.session.commit()
            results.append({'type': action_type, 'opportunity_id': opp.id, 'heat_level': heat_level, 'probability': opp.probability})
        elif action_type == 'update_opportunity_amount':
            opp = resolve_opportunity(action)
            amount = action.get('amount')
            if not opp:
                errors.append({'type': action_type, 'error': 'Opportunita non trovata'})
                continue
            try:
                amount_float = float(amount) if amount else 0.0
                opp.amount = amount_float
                opp.updated_at = datetime.utcnow()
                db.session.commit()
                results.append({'type': action_type, 'opportunity_id': opp.id, 'amount': amount_float})
            except (ValueError, TypeError):
                errors.append({'type': action_type, 'error': 'Importo non valido'})
        elif action_type == 'update_opportunity_close_date':
            opp = resolve_opportunity(action)
            close_date_str = action.get('close_date')
            if not opp:
                errors.append({'type': action_type, 'error': 'Opportunita non trovata'})
                continue
            try:
                if close_date_str:
                    close_date = datetime.fromisoformat(close_date_str.replace('Z', '+00:00')).date()
                else:
                    close_date = None
                opp.close_date = close_date
                opp.updated_at = datetime.utcnow()
                db.session.commit()
                results.append({'type': action_type, 'opportunity_id': opp.id, 'close_date': close_date.isoformat() if close_date else None})
            except (ValueError, TypeError) as e:
                errors.append({'type': action_type, 'error': f'Data non valida: {str(e)}'})
        else:
            errors.append({'type': action_type or 'unknown', 'error': 'Tipo azione non supportato'})

    # Salva log con risultati
    log_entry.ai_reply = reply
    log_entry.actions_executed = json.dumps(results)
    log_entry.errors = json.dumps(errors) if errors else None
    db.session.add(log_entry)
    db.session.commit()

    return jsonify({
        'reply': reply,
        'actions': results,
        'errors': errors
    })


@app.route('/api/users/all', methods=['GET'])
def get_all_users():
    """Ottieni lista completa di tutti gli utenti (solo per admin)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        return jsonify({'error': 'Solo admin può vedere tutti gli utenti'}), 403
    
    users = User.query.all()
    return jsonify({
        'users': [{
            'id': u.id,
            'username': u.username,
            'first_name': u.first_name,
            'last_name': u.last_name,
            'email': u.email,
            'role': u.role,
            'created_at': u.created_at.isoformat() if u.created_at else None
        } for u in users]
    })


@app.route('/api/assistant/logs', methods=['GET'])
def get_assistant_logs():
    """Ottieni log delle richieste all'assistente (solo admin)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Non autorizzato'}), 401
    
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        return jsonify({'error': 'Solo admin può vedere i log'}), 403
    
    page = request.args.get('page', 1, type=int)
    per_page = min(request.args.get('per_page', 50, type=int), 100)
    
    logs = AssistantLog.query.order_by(AssistantLog.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    return jsonify({
        'logs': [{
            'id': log.id,
            'user_id': log.user_id,
            'user_name': f"{log.user.first_name} {log.user.last_name}" if log.user else 'Unknown',
            'user_message': log.user_message,
            'ai_reply': log.ai_reply,
            'actions_executed': json.loads(log.actions_executed) if log.actions_executed else [],
            'errors': json.loads(log.errors) if log.errors else [],
            'tokens_used': log.tokens_used,
            'model_used': log.model_used,
            'created_at': log.created_at.isoformat() if log.created_at else None
        } for log in logs.items],
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': logs.total,
            'pages': logs.pages
        }
    })

if __name__ == '__main__':
    init_db()
    app.run(host='0.0.0.0', port=8000, debug=False)

