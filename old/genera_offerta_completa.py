#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generatore offerte COMPLETO
- Crea cartelle in K:\OFFERTE - K\OFFERTE 2025 - K
- Genera DOCX e PDF in locale
- Usa template professionale
"""
import os
import re
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def trova_ultima_offerta_k_drive():
    """Trova ultimo numero offerta in K:\OFFERTE - K\OFFERTE 2025 - K"""
    
    base_paths = [
        r"K:\OFFERTE - K\OFFERTE 2025 - K",
        r"K:\OFFERTE - K",
        r"K:\OFFERTE"
    ]
    
    for base_path in base_paths:
        if os.path.exists(base_path):
            try:
                folders = [f for f in os.listdir(base_path) 
                          if os.path.isdir(os.path.join(base_path, f))]
                
                # Trova numeri K####-YY
                pattern = r'^K(\d{4})-\d{2}'
                numeri = []
                
                for folder in folders:
                    match = re.match(pattern, folder)
                    if match:
                        numeri.append(int(match.group(1)))
                
                if numeri:
                    ultimo = max(numeri)
                    print(f"Ultima offerta in K:\\ trovata: K{ultimo:04d}")
                    return ultimo
                    
            except Exception as e:
                print(f"Errore lettura K:\\: {e}")
    
    # Default: parte da 2000 (come nelle immagini)
    print("K:\\ non accessibile, parto da K2000")
    return 2000


def crea_cartella_offerta_k_drive(numero_offerta, cliente, prodotto, data=None):
    """
    Crea cartella vuota in K:\OFFERTE - K\OFFERTE 2025 - K
    Formato: K####-YY - CLIENTE - PRODOTTO - DD-MM-YYYY
    """
    
    if data is None:
        data = datetime.now()
    
    # Path base
    base_path = r"K:\OFFERTE - K\OFFERTE 2025 - K"
    
    # Se non esiste, crea le directory
    if not os.path.exists(base_path):
        # Prova percorsi alternativi
        alt_paths = [
            r"K:\OFFERTE",
            r"./offerte_k_drive"  # Fallback locale per test
        ]
        
        for path in alt_paths:
            if os.path.exists(os.path.dirname(path)) or path.startswith('.'):
                base_path = path
                if not os.path.exists(base_path):
                    os.makedirs(base_path)
                break
    
    # Pulisci nomi
    def clean(text):
        invalid = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        for char in invalid:
            text = text.replace(char, '')
        return ' '.join(text.split())[:50]
    
    cliente_clean = clean(cliente)
    prodotto_clean = clean(prodotto)
    data_str = data.strftime('%d-%m-%Y')
    
    # Nome cartella
    folder_name = f"{numero_offerta} - {cliente_clean} - {prodotto_clean} - {data_str}"
    folder_path = os.path.join(base_path, folder_name)
    
    # Crea cartella
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"✓ Cartella creata: {folder_name}")
    
    return folder_path


def genera_docx_offerta(numero, cliente, prodotti_lista, valore_totale, folder_path):
    """Genera documento Word dell'offerta"""
    
    doc = Document()
    
    # Intestazione
    heading = doc.add_heading(f'OFFERTA COMMERCIALE N. {numero}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Info offerta
    p = doc.add_paragraph()
    p.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n').bold = True
    p.add_run(f'Validità: {(datetime.now() + timedelta(days=30)).strftime("%d/%m/%Y")}\n').bold = True
    
    # Cliente
    doc.add_heading('Cliente:', level=2)
    doc.add_paragraph(cliente)
    
    # Prodotti
    doc.add_heading('Prodotti:', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Codice'
    hdr_cells[1].text = 'Descrizione'
    hdr_cells[2].text = 'Q.tà'
    hdr_cells[3].text = 'Valore €'
    
    for prod in prodotti_lista:
        row_cells = table.add_row().cells
        row_cells[0].text = prod['code']
        row_cells[1].text = prod['nome']
        row_cells[2].text = str(prod['quantita'])
        row_cells[3].text = f"€{prod['valore']:,.2f}"
    
    # Totale
    doc.add_paragraph()
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_total.add_run(f'TOTALE: €{valore_totale:,.2f}')
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Salva
    docx_file = os.path.join(folder_path, f"{os.path.basename(folder_path)}.docx")
    doc.save(docx_file)
    
    return docx_file


def genera_pdf_da_docx(docx_path):
    """Converte DOCX in PDF"""
    
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"✓ PDF generato: {os.path.basename(pdf_path)}")
        return pdf_path
    except:
        print("⚠️ PDF non generato (installa: pip install docx2pdf)")
        return None


if __name__ == '__main__':
    # Test
    from datetime import timedelta
    
    ultimo = trova_ultima_offerta_k_drive()
    prossimo = f"K{ultimo + 1:04d}-25"
    
    print(f"\nProssima offerta: {prossimo}")
    
    folder = crea_cartella_offerta_k_drive(
        prossimo,
        "CLIENTE TEST SRL",
        "ROBOT UR10e E PROGRAMMAZIONE"
    )
    
    prodotti = [
        {'code': 'UR10e', 'nome': 'UR10e Cobot', 'quantita': 1, 'valore': 35000}
    ]
    
    docx = genera_docx_offerta(prossimo, "CLIENTE TEST SRL", prodotti, 35000, folder)
    pdf = genera_pdf_da_docx(docx)
    
    print(f"\n✓ Offerta generata:")
    print(f"  Cartella: {folder}")
    print(f"  DOCX: {docx}")
    if pdf:
        print(f"  PDF: {pdf}")

