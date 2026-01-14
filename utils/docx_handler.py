"""
Gestore per documenti Word (DOCX) - Template e generazione offerte
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from typing import Dict, List, Optional
import os


class DOCXHandler:
    """Gestisce la lettura e scrittura di documenti Word"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Inizializza il gestore
        
        Args:
            template_path: Percorso al template DOCX (opzionale)
        """
        self.template_path = template_path
        self.document = None
        
        if template_path and os.path.exists(template_path):
            self.load_template(template_path)
        else:
            self.document = Document()
    
    def load_template(self, template_path: str):
        """Carica un template esistente"""
        try:
            self.document = Document(template_path)
            self.template_path = template_path
            print(f"Template caricato: {template_path}")
        except Exception as e:
            print(f"Errore nel caricamento del template: {e}")
            self.document = Document()
    
    def create_new_document(self):
        """Crea un nuovo documento vuoto"""
        self.document = Document()
    
    def set_company_header(self, company_info: Dict[str, str]):
        """
        Imposta l'intestazione aziendale
        
        Args:
            company_info: Dizionario con info azienda
                - name: Nome azienda
                - address: Indirizzo
                - vat: Partita IVA
                - email: Email
                - phone: Telefono
        """
        if not self.document:
            self.create_new_document()
        
        # Aggiungi intestazione
        header = self.document.add_heading(company_info.get('name', ''), level=0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info azienda
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if 'address' in company_info:
            p.add_run(company_info['address'] + '\n')
        if 'vat' in company_info:
            p.add_run(f"P.IVA: {company_info['vat']}\n")
        if 'email' in company_info:
            p.add_run(f"Email: {company_info['email']}\n")
        if 'phone' in company_info:
            p.add_run(f"Tel: {company_info['phone']}\n")
        
        # Linea separatrice
        self.document.add_paragraph('_' * 80)
    
    def add_offer_header(self, 
                        offer_number: str,
                        client_name: str,
                        client_info: Optional[Dict[str, str]] = None,
                        validity_days: int = 30):
        """
        Aggiunge l'intestazione dell'offerta
        
        Args:
            offer_number: Numero offerta (es. K0001-25)
            client_name: Nome cliente
            client_info: Info aggiuntive cliente (opzionale)
            validity_days: Giorni di validità dell'offerta
        """
        if not self.document:
            self.create_new_document()
        
        # Titolo offerta
        title = self.document.add_heading(f'OFFERTA N. {offer_number}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data e validità
        today = datetime.now()
        validity_date = today + timedelta(days=validity_days)
        
        p = self.document.add_paragraph()
        p.add_run(f'Data: {today.strftime("%d/%m/%Y")}\n').bold = True
        p.add_run(f'Validità: {validity_date.strftime("%d/%m/%Y")}\n').bold = True
        
        # Cliente
        self.document.add_heading('Cliente:', level=2)
        p = self.document.add_paragraph(client_name)
        
        if client_info:
            if 'address' in client_info:
                p.add_run(f"\n{client_info['address']}")
            if 'vat' in client_info:
                p.add_run(f"\nP.IVA: {client_info['vat']}")
            if 'contact' in client_info:
                p.add_run(f"\nRif: {client_info['contact']}")
        
        self.document.add_paragraph()
    
    def add_products_table(self, products: List[Dict]):
        """
        Aggiunge una tabella con i prodotti dell'offerta
        
        Args:
            products: Lista di dizionari con info prodotti:
                - name: Nome prodotto
                - code: Codice prodotto
                - description: Descrizione
                - quantity: Quantità
                - unit_price: Prezzo unitario
                - total: Totale
        """
        if not self.document:
            self.create_new_document()
        
        self.document.add_heading('Prodotti:', level=2)
        
        # Crea tabella
        num_rows = len(products) + 1  # +1 per header
        table = self.document.add_table(rows=num_rows, cols=5)
        try:
            table.style = 'Light Grid Accent 1'
        except KeyError:
            table.style = 'Table Grid'
        
        # Header
        headers = ['Codice', 'Descrizione', 'Q.tà', 'Prezzo Unit.', 'Totale']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Grassetto per l'header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Dati prodotti
        total_amount = 0.0
        for idx, product in enumerate(products, start=1):
            row = table.rows[idx]
            row.cells[0].text = product.get('code', '')
            row.cells[1].text = f"{product.get('name', '')}\n{product.get('description', '')}"
            row.cells[2].text = str(product.get('quantity', 1))
            row.cells[3].text = f"€ {product.get('unit_price', 0.0):.2f}"
            
            total = product.get('total', product.get('unit_price', 0.0) * product.get('quantity', 1))
            row.cells[4].text = f"€ {total:.2f}"
            total_amount += total
        
        # Totale
        self.document.add_paragraph()
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f'TOTALE: € {total_amount:.2f}')
        run.font.size = Pt(14)
        run.font.bold = True
        
        return total_amount
    
    def add_image(self, image_path: str, width_inches: float = 3.0):
        """
        Aggiunge un'immagine al documento
        
        Args:
            image_path: Percorso all'immagine
            width_inches: Larghezza in pollici
        """
        if not self.document:
            self.create_new_document()
        
        if os.path.exists(image_path):
            try:
                self.document.add_picture(image_path, width=Inches(width_inches))
                last_paragraph = self.document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Errore nell'aggiunta dell'immagine: {e}")
    
    def add_notes(self, notes: str):
        """Aggiunge note all'offerta"""
        if not self.document:
            self.create_new_document()
        
        self.document.add_heading('Note:', level=2)
        self.document.add_paragraph(notes)
    
    def add_terms_and_conditions(self, terms: List[str]):
        """
        Aggiunge termini e condizioni
        
        Args:
            terms: Lista di termini/condizioni
        """
        if not self.document:
            self.create_new_document()
        
        self.document.add_page_break()
        self.document.add_heading('Termini e Condizioni', level=2)
        
        for term in terms:
            self.document.add_paragraph(term, style='List Bullet')
    
    def add_signature_section(self):
        """Aggiunge la sezione per le firme"""
        if not self.document:
            self.create_new_document()
        
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Due colonne per le firme
        table = self.document.add_table(rows=2, cols=2)
        
        # Fornitore
        table.rows[0].cells[0].text = 'Il Fornitore'
        table.rows[1].cells[0].text = '\n\n_______________________'
        
        # Cliente
        table.rows[0].cells[1].text = 'Il Cliente'
        table.rows[1].cells[1].text = '\n\n_______________________'
    
    def save(self, output_path: str):
        """
        Salva il documento
        
        Args:
            output_path: Percorso di salvataggio
        """
        if not self.document:
            raise ValueError("Nessun documento da salvare")
        
        try:
            self.document.save(output_path)
            print(f"Documento salvato: {output_path}")
            return True
        except Exception as e:
            print(f"Errore nel salvataggio: {e}")
            return False
    
    def replace_text(self, old_text: str, new_text: str):
        """
        Sostituisce del testo nel documento
        
        Args:
            old_text: Testo da sostituire
            new_text: Nuovo testo
        """
        if not self.document:
            return
        
        for paragraph in self.document.paragraphs:
            if old_text in paragraph.text:
                # Sostituisci nel testo
                inline = paragraph.runs
                for run in inline:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
        
        # Cerca anche nelle tabelle
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
    
    def extract_text(self) -> str:
        """Estrae tutto il testo dal documento"""
        if not self.document:
            return ""
        
        full_text = []
        for para in self.document.paragraphs:
            full_text.append(para.text)
        
        return '\n'.join(full_text)


if __name__ == '__main__':
    # Test del gestore DOCX
    handler = DOCXHandler()
    
    # Test creazione offerta
    handler.set_company_header({
        'name': 'ME.KO. Srl',
        'address': 'Via Example 123, 00100 Roma',
        'vat': '12345678901',
        'email': 'info@meko.it',
        'phone': '+39 06 1234567'
    })
    
    handler.add_offer_header(
        offer_number='K0001-25',
        client_name='Cliente Test Srl',
        validity_days=30
    )
    
    # Aggiungi prodotti di esempio
    products = [
        {
            'code': 'G1-U6',
            'name': 'Unitree G1 U6',
            'description': 'Robot umanoide di nuova generazione',
            'quantity': 2,
            'unit_price': 15000.00
        },
        {
            'code': 'GO2-EDU',
            'name': 'Unitree GO2 EDU',
            'description': 'Robot quadrupede educativo',
            'quantity': 1,
            'unit_price': 3500.00
        }
    ]
    
    handler.add_products_table(products)
    
    handler.add_notes("Consegna prevista entro 30 giorni dalla conferma d'ordine.")
    
    handler.add_terms_and_conditions([
        "Pagamento: 50% all'ordine, 50% alla consegna",
        "Garanzia: 12 mesi dalla data di consegna",
        "Trasporto incluso nel prezzo"
    ])
    
    handler.add_signature_section()
    
    # Salva
    output_file = 'test_offerta.docx'
    if handler.save(output_file):
        print(f"\nOfferta di test creata: {output_file}")

